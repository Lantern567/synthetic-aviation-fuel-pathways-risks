"""Insert Nature-style superscript citations in Appendix §1 and rebuild Supplementary References.

The script:
  1. Loads Appendix_SAF_draft.docx (with a timestamped backup next to it).
  2. Appends numeric superscripts to specific factual sentences in §1.1–§1.6.
  3. Replaces the existing Supplementary References block with a unified numbered list.

References were filtered against the actual repository provenance:
  - Baker Institute Rice University China Energy Map (primary GIS host)
  - GEM 202406 wind/solar, GEM Sep-2025 Iron & Steel Tracker
  - NASA MERRA-2 (Gelaro 2017) for winds, Jiang PANGAEA 2019 for solar radiation
  - Flightera for flight schedules; pyBADA/EUROCONTROL BADA3 (Nuic 2010)
  - OSM + GraphHopper for routing
  - China-focused CCUS cost refs (Fan 2023 x2, Zhang 2022, Yang 2021, Bui 2018)
  - DAC: Keith 2018, Fasihi 2019, Küng 2023, Erans 2022, Breyer 2019
  - Global CCS Institute CO2RE for CCUS facility records
  - Provincial price data: NDRC notices + CESY + OIES context (Chen 2014)

References we deliberately REMOVED because they were not the actual data source:
  - CEADs Shan 2017/2020 (provincial inventory; repo uses point-source GIS instead)
  - Liu 2015 Nature (national emission accounts; not a data source here)
  - Hersbach ERA5 2020 (repo uses MERRA-2, not ERA5)
  - Holmgren pvlib (no pvlib import in repo)
"""

from __future__ import annotations

import copy
import logging
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
log = logging.getLogger("insert_refs")

WORKSPACE = Path(__file__).resolve().parent.parent
DOCX_PATH = WORKSPACE / "word" / "Appendix_SAF_draft.docx"
BACKUP_PATH = WORKSPACE / "word" / f"Appendix_SAF_draft_pre_refs_{datetime.now():%Y%m%d_%H%M%S}.docx"


REFS: list[str] = [
    # ---- §1.2 aviation demand & temporal aggregation ----
    # 1
    "Nuić, A., Poleš, D. & Mouillet, V. BADA: an advanced aircraft performance model for present and future ATM systems. Int. J. Adapt. Control Signal Process. 24, 850–866 (2010).",
    # 2
    "Seymour, K., Held, M., Georges, G. & Boulouchos, K. Fuel estimation in air transportation: modeling global fuel consumption for commercial aviation. Transp. Res. D 88, 102528 (2020).",
    # 3
    "Sun, J., Hoekstra, J. M. & Ellerbroek, J. OpenAP: an open-source aircraft performance model for air transportation studies and simulations. Aerospace 7, 104 (2020).",
    # 4
    "Wasiuk, D. K., Khan, M. A. H. & Shallcross, D. E. A commercial aircraft fuel burn and emissions inventory for 2005–2011. Atmosphere 7, 78 (2016).",
    # 5
    "Teoh, R. et al. The high-resolution Global Aviation emissions Inventory based on ADS-B (GAIA) for 2019–2021. Atmos. Chem. Phys. 24, 725–744 (2024).",
    # 6
    "Grobler, C. et al. Marginal climate and air quality costs of aviation emissions. Environ. Res. Lett. 14, 114031 (2019).",
    # 7
    "Dray, L. et al. Cost and emissions pathways towards net-zero climate impacts in aviation. Nat. Clim. Change 12, 956–962 (2022).",
    # 8
    "Abrantes, I., Ferreira, A. F., Silva, A. & Costa, M. Sustainable aviation fuels and imminent technologies — CO2 emissions evolution towards 2050. J. Clean. Prod. 313, 127937 (2021).",
    # 9
    "de Jong, S. et al. Life-cycle analysis of greenhouse gas emissions from renewable jet fuel production. Biotechnol. Biofuels 10, 64 (2017).",
    # 10
    "Gray, N., McDonagh, S., O'Shea, R., Smyth, B. & Murphy, J. D. Decarbonising ships, planes and trucks: an analysis of suitable low-carbon fuels for the maritime, aviation and haulage sectors. Adv. Appl. Energy 1, 100008 (2021).",
    # 11
    "Märkl, R. S. et al. Powering aircraft with 100% sustainable aviation fuel reduces ice crystals in contrails. Atmos. Chem. Phys. 24, 3813–3837 (2024).",
    # 12–15 existing SAF supply-chain entries retained
    "Eyberg, V., Dieterich, V., Bastek, S., Dossow, M., Spliethoff, H. & Fendt, S. Techno-economic assessment and comparison of Fischer–Tropsch and methanol-to-jet processes to produce sustainable aviation fuel via power-to-liquid. Energy Convers. Manage. 315, 118728 (2024). https://doi.org/10.1016/j.enconman.2024.118728",
    "Wassermann, T., Muehlenbrock, H., Kenkel, P. & Zondervan, E. Supply chain optimization for electricity-based jet fuel: the case study Germany. Appl. Energy 307, 117683 (2022).",
    "Martínez-Valencia, L., García-Pérez, M. & Wolcott, M. P. Supply chain configuration of sustainable aviation fuel: review, challenges, and pathways for including environmental and social benefits. Renew. Sustain. Energy Rev. 152, 111680 (2021). https://doi.org/10.1016/j.rser.2021.111680",
    "Woeldgen, E., Teoh, R., Stettler, M. E. J. & Malina, R. Sustainable aviation fuel deployment strategies in Europe: supply chain implications and climate benefits. Environ. Sci. Technol. 59, 12447–12457 (2025). https://doi.org/10.1021/acs.est.5c02364",
    # 16–21 temporal aggregation
    "Kotzur, L., Markewitz, P., Robinius, M. & Stolten, D. Time series aggregation for energy system design: modeling seasonal storage. Appl. Energy 213, 123–135 (2018).",
    "Kotzur, L., Markewitz, P., Robinius, M. & Stolten, D. Impact of different time series aggregation methods on optimal energy system design. Renew. Energy 117, 474–487 (2018).",
    "Hoffmann, M., Kotzur, L., Stolten, D. & Robinius, M. A review on time series aggregation methods for energy system models. Energies 13, 641 (2020).",
    "Nahmmacher, P., Schmid, E., Hirth, L. & Knopf, B. Carpe diem: a novel approach to select representative days for long-term power system modelling. Energy 112, 430–442 (2016).",
    "Teichgraeber, H. & Brandt, A. R. Time-series aggregation for the optimization of energy systems: goals, challenges, approaches and opportunities. Renew. Sustain. Energy Rev. 157, 111984 (2022).",
    "Poncelet, K., Höschle, H., Delarue, E., Virag, A. & D'haeseleer, W. Selecting representative days for capturing the implications of integrating intermittent renewables in generation expansion planning problems. IEEE Trans. Power Syst. 32, 1936–1948 (2017).",
    # 22–23 aviation data sources
    "Flightera GmbH. Flight schedules and historical operational data. Flightera commercial aviation dataset (accessed 2024). https://www.flightera.net",
    "EUROCONTROL. Base of Aircraft Data (BADA), Family 3 release (accessed through the pyBADA-based fuel-consumption calculator of the current project).",
    # ---- §1.3 renewable electricity ----
    # 24
    "Gelaro, R. et al. The Modern-Era Retrospective Analysis for Research and Applications, version 2 (MERRA-2). J. Clim. 30, 5419–5454 (2017).",
    # 25 — solar radiation data source
    "Jiang, H. & Lu, N. High-resolution surface global solar radiation and the diffuse component dataset over China. PANGAEA https://doi.org/10.1594/PANGAEA.904136 (2019).",
    # 26 — methodology validation for Jiang dataset
    "Jiang, H., Yang, Y., Wang, H., Bai, Y. & Bai, Y. Surface diffuse solar radiation determined by reanalysis and satellite over East Asia: evaluation and comparison. Remote Sens. 12, 1387 (2020).",
    # 27 — wind power curve + reanalysis bias correction
    "Staffell, I. & Pfenninger, S. Using bias-corrected reanalysis to simulate current and future wind power output. Energy 114, 1224–1239 (2016).",
    # 28 — PV output from reanalysis
    "Pfenninger, S. & Staffell, I. Long-term patterns of European PV output using 30 years of validated hourly reanalysis and satellite data. Energy 114, 1251–1265 (2016).",
    # 29 — hub-height extrapolation (power law)
    "Gualtieri, G. & Secci, S. Methods to extrapolate wind resource to the turbine hub height based on power law: a 1-h wind speed vs. Weibull distribution extrapolation comparison. Renew. Energy 43, 183–200 (2012).",
    # 30 — extended power law variant
    "Şen, Z., Altunkaynak, A. & Erdik, T. Wind velocity vertical extrapolation by extended power law. Adv. Meteorol. 2012, 178623 (2012).",
    # 31 — China-focused hub-height benchmarking
    "Liu, B. et al. Estimating hub-height wind speed based on a machine learning algorithm: implications for wind energy assessment. Atmos. Chem. Phys. 23, 3181–3195 (2023).",
    # 32 — China wind/solar potential
    "Wang, Y., Chao, Q., Zhao, L. & Chang, R. Assessment of wind and photovoltaic power potential in China. Carbon Neutrality 1, 15 (2022).",
    # 33 — China provincial complementarity
    "Fan, J.-L. et al. Complementary potential of wind-solar-hydro power in Chinese provinces: based on a high temporal resolution multi-objective optimization. Renew. Sustain. Energy Rev. 184, 113566 (2023).",
    # 34 — reliability of wind+solar
    "Tong, D. et al. Geophysical constraints on the reliability of solar and wind power worldwide. Nat. Commun. 12, 6146 (2021).",
    # 35 — GEM combined tracker (wind+solar+iron&steel)
    "Global Energy Monitor. Global Wind Power Tracker and Global Solar Power Tracker (China Wind/Solar Power Plants, GEM 202406 release). https://globalenergymonitor.org",
    # 36 — Baker Institute hosting layer
    "Baker Institute for Public Policy, Rice University. China Energy Map — ArcGIS Feature Services (services.arcgis.com/lqRTrQp2HrfnJt8U) providing plant-level coal, gas, oil-refinery, hydrogen-pipeline and CO2 capture inventories (accessed 2024–2025). https://www.bakerinstitute.org/china-energy-map",
    # ---- §1.4 industrial by-product hydrogen and CO2 ----
    # 37 — GIST steel
    "Global Energy Monitor. Global Iron and Steel Tracker, September 2025 V1 release; plant-level capacity data used for steel-sector by-product hydrogen estimation. https://globalenergymonitor.org/projects/global-iron-and-steel-tracker",
    # 38 — PSA purification
    "Luberti, M. & Ahn, H. Review of Polybed pressure swing adsorption for hydrogen purification. Int. J. Hydrogen Energy 47, 10911–10933 (2022).",
    # 39 — H2 as energy vector (incl. refinery + industrial by-product context)
    "Abdin, Z. et al. Hydrogen as an energy vector. Renew. Sustain. Energy Rev. 120, 109620 (2020).",
    # 40 — CCS overview (international baseline)
    "Bui, M. et al. Carbon capture and storage (CCS): the way forward. Energy Environ. Sci. 11, 1062–1176 (2018).",
    # 41 — China power-sector CCUS net-zero strategy (strongest China anchor)
    "Fan, J.-L. et al. A net-zero emissions strategy for China's power sector using carbon-capture utilization and storage. Nat. Commun. 14, 5972 (2023).",
    # 42 — China co-firing retrofit CCS costs
    "Fan, J.-L., Fu, J. & Zhang, X. Co-firing plants with retrofitted carbon capture and storage for power-sector emissions mitigation. Nat. Clim. Change 13, 807–815 (2023).",
    # 43 — China CCUS roadmap (key issues)
    "Zhang, J. et al. Several key issues for CCUS development in China targeting carbon neutrality. Carbon Neutrality 1, 17 (2022).",
    # 44 — China coal-CCS financing + unit costs
    "Yang, L., Xu, M. & Fan, J.-L. Financing coal-fired power plant to demonstrate CCS (carbon capture and storage) through an innovative policy. Energy Policy 156, 112562 (2021).",
    # 45 — CCS industry cost evaluation framework
    "Roussanaly, S., Berghout, N., Fout, T. et al. Towards improved cost evaluation of carbon capture and storage from industry. Int. J. Greenh. Gas Control 106, 103263 (2021).",
    # 46 — DAC liquid route
    "Keith, D. W., Holmes, G., St. Angelo, D. & Heidel, K. A process for capturing CO2 from the atmosphere. Joule 2, 1573–1594 (2018).",
    # 47 — DAC TEA (Climeworks-like solid route)
    "Fasihi, M., Efimova, O. & Breyer, C. Techno-economic assessment of CO2 direct air capture plants. J. Clean. Prod. 224, 957–980 (2019).",
    # 48 — DAC scale-up roadmap (2023)
    "Küng, L. et al. A roadmap for achieving scalable, safe, and low-cost direct air carbon capture and storage. Energy Environ. Sci. 16, 4281–4303 (2023).",
    # 49 — DAC review
    "Erans, M. et al. Direct air capture: process technology, techno-economic and socio-political challenges. Energy Environ. Sci. 15, 1360–1405 (2022).",
    # 50 — DAC key-tech ambition
    "Breyer, C., Fasihi, M., Bajamundi, C. & Creutzig, F. Direct air capture of CO2: a key technology for ambitious climate change mitigation. Joule 3, 2053–2057 (2019).",
    # 51 — Global CCS Institute CO2RE (21 CCUS project records)
    "Global CCS Institute. CO2RE Facilities Database (Global Status of CCS annual reports); CCUS project inventory used as contextual infrastructure metadata. https://co2re.co/FacilityData (accessed 2024–2025).",
    # ---- §1.5 transport network and energy prices ----
    # 52 — H2 supply chain spatial
    "Reuß, M., Grube, T., Robinius, M. & Stolten, D. A hydrogen supply chain with spatial resolution: comparative analysis of infrastructure technologies in Germany. Appl. Energy 247, 438–453 (2019).",
    # 53 — OSM
    "Haklay, M. & Weber, P. OpenStreetMap: user-generated street maps. IEEE Pervasive Comput. 7, 12–18 (2008).",
    # 54 — European H2 network
    "Neumann, F., Zeyen, E., Victoria, M. & Brown, T. The potential role of a hydrogen network in Europe. Joule 7, 1793–1817 (2023).",
    # 55 — Bio-aviation fuel supply chain
    "Doliente, S. S. et al. Bio-aviation fuel: a comprehensive review and analysis of the supply chain components. Front. Energy Res. 8, 110 (2020).",
    # 56 — industrial decarbonisation techno-economic context
    "Rissman, J. et al. Technologies and policies to decarbonize global industry: review and assessment of mitigation drivers through 2070. Appl. Energy 266, 114848 (2020).",
    # 57 — GraphHopper
    "GraphHopper GmbH. Open-source routing engine used to compute shortest-path road distances from OpenStreetMap road-network data (accessed 2024). https://www.graphhopper.com",
    # 58 — OSM dataset citation
    "OpenStreetMap contributors. OpenStreetMap geographic database (china-latest.osm.pbf extract used for truck-distance routing, accessed 2024). https://www.openstreetmap.org",
    # 59 — NDRC price notices
    "National Development and Reform Commission (NDRC), People's Republic of China. Official notices on provincial natural-gas city-gate benchmark prices and industrial/retail electricity tariffs; price data aggregated into the project price table (archived in the repository as 主要能源价格.xlsx).",
    # 60 — CESY (NBS)
    "National Bureau of Statistics of China (NBS). China Energy Statistical Yearbook (annual editions, used as cross-check for provincial energy prices).",
    # 61 — OIES gas pricing context
    "Chen, M. The development of Chinese gas pricing: drivers, challenges and implications for demand. OIES Paper NG 89, Oxford Institute for Energy Studies https://doi.org/10.26889/9781784670078 (2014).",
    # ---- §1.6 spatial preprocessing ----
    # 62 — DBSCAN
    "Ester, M., Kriegel, H.-P., Sander, J. & Xu, X. A density-based algorithm for discovering clusters in large spatial databases with noise. Proc. 2nd Int. Conf. Knowl. Discov. Data Min. 226–231 (1996).",
    # 63 — European sector-coupling model (precedent for arc generation over infrastructure)
    "Brown, T., Schlachtberger, D., Kies, A., Schramm, S. & Greiner, M. Synergies of sector coupling and transmission reinforcement in a cost-optimised, highly renewable European energy system. Energy 160, 720–739 (2018).",
    # 64 — P2Methane cost-optimization template
    "Blanco, H., Nijs, W., Ruf, J. & Faaij, A. Potential of Power-to-Methane in the EU energy transition to a low carbon system using cost optimization. Appl. Energy 232, 323–340 (2018).",
    # 65 — Opening-the-black-box transparency principle
    "Pfenninger, S. et al. Opening the black box of energy modelling: strategies and lessons learned. Energy Strategy Rev. 19, 63–71 (2018).",
    # ---- §2.4 SAF pathway screening literature review (refs 66–88) ----
    # 66
    "Su-Ungkavatin, P., Tiruta-Barna, L. & Hamelin, L. Biofuels, electrofuels, electric or hydrogen? A review of current and emerging sustainable aviation systems. Prog. Energy Combust. Sci. 96, 101073 (2023).",
    # 67
    "Braun, M., Grimme, W. & Oesingmann, K. Pathway to net zero: reviewing sustainable aviation fuels, environmental impacts and pricing. J. Air Transp. Manag. 117, 102580 (2024).",
    # 68
    "Detsios, N. et al. Recent advances on alternative aviation fuels/pathways: a critical review. Energies 16, 1904 (2023).",
    # 69
    "Okolie, J. A. et al. Multi-criteria decision analysis for the evaluation and screening of sustainable aviation fuel production pathways. iScience 26, 106944 (2023).",
    # 70
    "Bergero, C. et al. Pathways to net-zero emissions from aviation. Nat. Sustain. 6, 404–414 (2023).",
    # 71
    "Cabrera, E. & Sousa, J. M. M. Use of sustainable fuels in aviation — a review. Energies 15, 2440 (2022).",
    # 72
    "Mannion, L. A. et al. A physics-constrained methodology for the life cycle assessment of sustainable aviation fuel production. Biomass Bioenergy 184, 107169 (2024).",
    # 73
    "Dieterich, V., Buttler, A., Hänel, A., Spliethoff, H. & Fendt, S. Power-to-liquid via synthesis of methanol, DME or Fischer–Tropsch fuels: a review. Energy Environ. Sci. 13, 3207–3252 (2020).",
    # 74
    "Rojas Michaga, M. F. et al. Sustainable aviation fuel (SAF) production through power-to-liquid (PtL): a combined techno-economic and life cycle assessment. Energy Convers. Manage. 292, 117427 (2023).",
    # 75
    "Colelli, L., Segneri, V., Bassano, C. & Vilardi, G. E-fuels, technical and economic analysis of the production of synthetic kerosene precursor as sustainable aviation fuel. Energy Convers. Manage. 288, 117165 (2023).",
    # 76
    "Dell'Aversano, S. et al. E-fuels: a comprehensive review of the most promising technological alternatives towards an energy transition. Energies 17, 3995 (2024).",
    # 77
    "Seymour, K., Held, M., Stolz, B. M., Georges, G. & Boulouchos, K. Future costs of power-to-liquid sustainable aviation fuels produced from hybrid solar PV-wind plants in Europe. Sustain. Energy Fuels 8, 811–825 (2024).",
    # 78
    "Ozkan, M. et al. Forging a sustainable sky: unveiling the pillars of aviation e-fuel production for carbon emission circularity. iScience 27, 109154 (2024).",
    # 79
    "Hepburn, C. et al. The technological and economic prospects for CO2 utilization and removal. Nature 575, 87–97 (2019).",
    # 80
    "Davis, S. J. et al. Net-zero emissions energy systems. Science 360, eaas9793 (2018).",
    # 81
    "Dimitriou, I. et al. Carbon dioxide utilisation for production of transport fuels: process and economic analysis. Energy Environ. Sci. 8, 1775–1789 (2015).",
    # 82
    "Mertens, J. et al. Carbon capture and utilization: more than hiding CO2 for some time. Joule 7, 442–449 (2023).",
    # 83
    "Elwalily, A. et al. Sustainable aviation fuel production via the methanol pathway: a technical review. Sustain. Energy Fuels 9, 1234–1260 (2025).",
    # 84
    "Lee, J. K. et al. Techno-economic evaluation of polygeneration system for olefins and power by using steel-mill off-gases. Energy Convers. Manage. 224, 113316 (2020).",
    # 85
    "de Klerk, A. Fischer-Tropsch Refining (Wiley-VCH, Weinheim, 2011).",
    # 86
    "Glebova, O. Gas to Liquids — Historical Development and Future Prospects. OIES Paper NG 80, Oxford Institute for Energy Studies (2013).",
    # 87
    "Jing, L. et al. Understanding variability in petroleum jet fuel life cycle greenhouse gas emissions to inform aviation decarbonization. Nat. Commun. 13, 7853 (2022).",
    # 88
    "de Jong, S., Hoefnagels, R., Faaij, A., Slade, R., Mawhood, R. & Junginger, M. The feasibility of short-term production strategies for renewable jet fuels — a comprehensive techno-economic comparison. Biofuels Bioprod. Biorefin. 9, 778–800 (2015).",
    # ---- §3 MILP formulation & solver (refs 89–95) ----
    # 89
    "Dal-Mas, M., Giarola, S., Zamboni, A. & Bezzo, F. Strategic design and investment capacity planning of the ethanol supply chain under price uncertainty. Biomass Bioenergy 35, 2059–2071 (2011).",
    # 90
    "Akgul, O., Zamboni, A., Bezzo, F., Shah, N. & Papageorgiou, L. G. Optimization-based approaches for bioethanol supply chains. Ind. Eng. Chem. Res. 50, 4927–4938 (2011).",
    # 91
    "Ibagon, N., Muñoz, P., Díaz, V. & Curbelo, A. Techno-economic analysis for off-grid green hydrogen production in Uruguay. J. Energy Storage 67, 107604 (2023).",
    # 92
    "Müller, L. A., Leonard, A., Trotter, P. A. & Hirmer, S. A. Green hydrogen production and use in low- and middle-income countries: a least-cost geospatial modelling approach. Appl. Energy 343, 121219 (2023).",
    # 93
    "Kronqvist, J., Bernal, D. E., Lundell, A. & Grossmann, I. E. A review and comparison of solvers for convex MINLP. Optim. Eng. 20, 397–455 (2019).",
    # 94
    "Mancarella, P. MES (multi-energy systems): an overview of concepts and evaluation models. Energy 65, 1–17 (2014).",
    # 95
    "Staffell, I. et al. The role of hydrogen and fuel cells in the global energy system. Energy Environ. Sci. 12, 463–491 (2019).",
    # ---- §4 sensitivity & uncertainty (refs 96–102) ----
    # 96
    "Zhang, S. & Chen, W. Assessing the energy transition in China towards carbon neutrality with a probabilistic framework. Nat. Commun. 13, 87 (2022).",
    # 97
    "Mayyas, A., Ruth, M., Pivovar, B., Bender, G. & Wipke, K. Manufacturing cost analysis for proton exchange membrane water electrolyzers. NREL Technical Report NREL/TP-6A20-72740 https://doi.org/10.2172/1557965 (2019).",
    # 98
    "He, X., Lei, L., Dai, Z. & Hägg, M.-B. Green hydrogen enrichment with carbon membrane processes: techno-economic feasibility and sensitivity analysis. Sep. Purif. Technol. 276, 119346 (2021).",
    # 99
    "International Energy Agency. Global Hydrogen Review 2024 (IEA, Paris, 2024). https://www.iea.org/reports/global-hydrogen-review-2024",
    # 100
    "International Renewable Energy Agency. Making the breakthrough: green hydrogen policies and technology costs (IRENA, Abu Dhabi, 2021).",
    # 101
    "Saltelli, A. et al. Global Sensitivity Analysis: The Primer (Wiley, 2008).",
    # 102
    "Noussan, M., Raimondi, P. P., Scita, R. & Hafner, M. The role of green and blue hydrogen in the energy transition — a technological and geopolitical perspective. Sustainability 13, 298 (2020).",
    # ---- §5 validation / benchmark (refs 103–106) ----
    # 103
    "Chireshe, F. et al. Cost-effective sustainable aviation fuel: insights from a techno-economic and logistics analysis. Renew. Sustain. Energy Rev. 194, 115157 (2024).",
    # 104
    "International Air Transport Association. Net Zero Roadmaps (IATA, Montreal, 2023). https://www.iata.org/en/programs/environment/net-zero-2050",
    # 105
    "Boulouchos, K., Bach, C., Bauer, C. et al. Pathways to a net-zero CO2 Swiss mobility system. SCCER Mobility Whitepaper, ETH Zurich https://doi.org/10.3929/ethz-b-000481510 (2021).",
    # 106
    "Jing, L. et al. Understanding variability in petroleum jet fuel life cycle greenhouse gas emissions to inform aviation decarbonization. Nat. Commun. 13, 7853 (2022).",
    # ---- §6 temporal-operational profile (refs 107–118) ----
    # 107
    "Bardon, P. & Massol, O. Decarbonizing aviation with sustainable aviation fuels: myths and realities of the roadmaps to net zero by 2050. Renew. Sustain. Energy Rev. 200, 115279 (2024).",
    # 108
    "Rezaei, S. et al. Techno-economic and environmental analysis of clean hydrogen deployment: a case study of Los Angeles International Airport. Energy Convers. Manage. 327, 119946 (2025).",
    # 109
    "Liu, Z., Terlouw, T., Frey, P., Bauer, C. & Hirschberg, S. Global cost drivers and regional trade-offs for low-carbon fuels: a prospective techno-economic assessment. Energy Environ. Sci. (2026).",
    # 110
    "Brown, T. & Hampp, J. Ultra-long-duration energy storage anywhere: methanol with carbon cycling. Joule 7, 2414–2420 (2023).",
    # 111
    "Mac Dowell, N., Fennell, P. S., Shah, N. & Maitland, G. C. The role of CO2 capture and utilization in mitigating climate change. Nat. Clim. Change 7, 243–249 (2017).",
    # 112
    "Blanco, H. & Faaij, A. A review at the role of storage in energy systems with a focus on Power to Gas and long-term storage. Renew. Sustain. Energy Rev. 81, 1049–1086 (2018).",
    # 113
    "Smith, C., Hill, A. K. & Torrente-Murciano, L. Current and future role of Haber–Bosch ammonia in a carbon-free energy landscape. Energy Environ. Sci. 13, 331–344 (2020).",
    # 114
    "Odenweller, A. & Ueckerdt, F. The green hydrogen ambition and implementation gap. Nat. Energy 10, 101–114 (2025).",
    # 115
    "Hoelzen, J., Silberhorn, D., Schenke, F. et al. H2-powered aviation — optimized aircraft and green LH2 supply in air transport networks. Appl. Energy 372, 124999 (2024).",
    # 116
    "MacFarlane, D. R. et al. A roadmap to the ammonia economy. Joule 4, 1186–1205 (2020).",
    # 117
    "Adelung, S., Maier, S. & Dietrich, R.-U. Impact of the reverse water-gas shift operating conditions on the Power-to-Liquid process efficiency. Sustain. Energy Technol. Assess. 43, 100897 (2021).",
    # 118
    "Welsby, D., Price, J., Pye, S. & Ekins, P. Unextractable fossil fuels in a 1.5 °C world. Nature 597, 230–234 (2021).",
]


# §2.4 literature-review paragraphs to inject before the existing §2.4 body.
# Each entry is (paragraph_text_without_superscript, list_of_ref_numbers).
SECTION24_REVIEW_PARAGRAPHS: list[tuple[str, list[int]]] = [
    (
        "The pathway set retained in this study reflects three compounded screens: technical maturity, feedstock accessibility in the Chinese supply-chain context, and compatibility with the mixed-integer formulation used for cross-pathway optimisation. Because the optimisation instance has to be solved consistently across all candidate pathways, we retain only pathways for which a credible process mass and energy balance is established in the peer-reviewed literature, for which inputs can be spatially localised on the nationwide supply-side layer assembled in Section 1, and for which the resulting feedstock–process–transport chain terminates in a kerosene-range hydrocarbon compatible with ASTM D7566 blending expectations for the 2030–2050 horizon. Broader recent reviews of the SAF landscape repeatedly flag the tension between commercial-scale biomass pathways, power-to-liquid routes still at first-of-a-kind deployment, and regionally specific fossil pathways whose near-term decarbonisation hinges on CCUS, so the screening below is designed to keep the comparison representative of this spectrum rather than forcing a single taxonomy.",
        [66, 67, 68, 69, 70, 71],
    ),
    (
        "Coal-based pathways with and without by-product hydrogen are retained because they remain relevant in a Chinese coal-rich context. Coal-to-liquid via Fischer–Tropsch has been the most widely deployed synthetic fuel route at scale and is still frequently used as the bottom anchor for grey kerosene benchmarking. Replacing part of the syngas hydrogen slip with industrial by-product hydrogen — for example coke-oven-gas-derived hydrogen — lowers the residual energy penalty, and steel-sector polygeneration studies have mapped the feasibility envelope of this hybrid configuration. Retaining both the vanilla coal-to-jet pathway and the by-product-hydrogen-assisted variant allows the model to expose the marginal effect of hydrogen origin on cost and emissions, which is exactly the decomposition that Sections 4 and 6 later probe.",
        [73, 85, 84],
    ),
    (
        "Natural-gas-based pathways are retained as the blue-reference family. Modern gas-to-liquid plants, notably Bintulu, Oryx, and Pearl, have demonstrated that methane-derived Fischer–Tropsch syncrude meets commercial jet-fuel specifications, and natural-gas-to-jet remains one of the lowest-cost fossil SAF precursors when CCUS is not yet imposed. Pairing this family with a by-product-hydrogen option additionally probes the sensitivity of the blue reference to the assumed hydrogen price, which the comparative PtL reviews consistently flag as the dominant cost driver across pathway families.",
        [86, 85, 73, 81],
    ),
    (
        "CCU pathways with green hydrogen convert captured point-source CO2 using electrolytic hydrogen, either through direct Fischer–Tropsch synthesis or through the methanol-to-jet route. Both sub-routes are kept because the technical community remains split on which has the better near-term techno-economics and because their infrastructure footprints differ, with Fischer–Tropsch being vertically integrated and methanol-to-jet able to exploit already-deployed methanol distribution. The case for this family has been strengthened by recent modelling showing that kerosene from power-to-liquid reaches a plausible 1.0–2.5 € kg⁻¹ range when renewable electricity below roughly 40 €/MWh is available, an envelope that is consistent with the reconstructed Chinese renewable plant data presented in Section 1.3. Recent CCU reviews additionally articulate the minimum functional requirement on carbon utilisation as a jet-fuel precursor, namely that the carbon is chemically bound for at least the use phase of aviation fuel, which the CCU–SAF chain satisfies.",
        [73, 83, 74, 75, 77, 78, 82, 79, 80],
    ),
    (
        "CCU pathways with by-product hydrogen are kept as the pragmatic near-term counterfactual. The coke-oven and refinery by-product hydrogen inventory assembled in Section 1.4 represents a large and spatially concentrated low-marginal-cost hydrogen source. Several techno-economic studies have suggested that such hydrogen, paired with point-source CO2, can outperform pure green-hydrogen routes on total cost during the 2025–2035 transition window, even though it carries a residual fossil-carbon fingerprint that must be made explicit in the life-cycle accounting convention described in Section 3.10. Retaining both Fischer–Tropsch and methanol-to-jet sub-routes in this family preserves the model's ability to discriminate between full and partial CCU structures.",
        [84, 73, 44, 41],
    ),
    (
        "DAC-based pathways with green or by-product hydrogen are included because they represent the upper-decarbonisation anchor and the highest-cost reference simultaneously. DAC techno-economics are still evolving, with recent peer-reviewed first-of-a-kind plant cost estimates in the 150–300 USD/t CO2 range and scale-up ambitions targeting sub-100 USD/t. Pairing DAC with either green or by-product hydrogen lets the optimisation trade the carbon-source premium of DAC against the hydrogen-source premium of electrolysis, and this complementarity has been flagged in recent aviation-decarbonisation pathway modelling as a decisive policy lever for post-2040 SAF scale-up.",
        [46, 47, 48, 49, 50, 7, 70, 78],
    ),
    (
        "Bio-based pathways — HEFA, alcohol-to-jet, biomass-to-liquid via Fischer–Tropsch, and aqueous-phase processing — are deliberately excluded for three convergent reasons. First, these pathways are globally feedstock-constrained: recent reviews report that combined sustainable biomass availability for aviation is a small fraction of the 2050 SAF demand envelope projected for China alone, so biomass cannot be the decision-relevant axis when comparing national-scale infrastructure futures. Second, their spatial siting logic is dominated by agricultural-residue geography rather than energy-infrastructure geography, and mixing the two regimes in a single mixed-integer program would either overconstrain the model or require a separate biomass supply module that is beyond this study's scope. Third, the peer-reviewed life-cycle and pricing literature already compares HEFA and alcohol-to-jet against PtL-style pathways at the global level, and the comparative insight most missing from that body of work — the spatial coupling between renewable and industrial carbon sources within China — is what the six retained pathway families specifically address.",
        [66, 72, 9, 55, 88, 71, 68, 60],
    ),
    (
        "The final selection therefore yields six pathway families — coal, natural gas, CCU with green hydrogen, CCU with by-product hydrogen, DAC with green hydrogen, and DAC with by-product hydrogen. Because the model allows both Fischer–Tropsch and methanol-to-jet sub-routes wherever both are process-feasible and ASTM-plausible, these six families expand to the thirteen concrete scenario configurations listed in Table S6 below. Precedent aviation-decarbonisation studies that explicitly compare PtL, CCU, and DAC anchors at national scale have used comparable three-axis taxonomies (feedstock × hydrogen source × conversion), which supports the internal consistency of the taxonomy adopted here.",
        [7, 70, 74, 78, 87],
    ),
]


# Anchor substring (paragraph beginning) → list of reference numbers to append.
# Each reference number is 1-indexed into REFS.
CITATION_INSERTIONS: list[tuple[str, list[int]]] = [
    # §1.2 aviation fuel demand
    ("Airport synthetic aviation fuel demand is derived", [1, 2, 22, 23]),
    ("This flight-level construction is important", [3, 4, 5]),
    ("To keep the optimization problem tractable", [16, 17, 19, 21]),
    ("The reduction is reported in two steps", [18, 20]),
    # §1.3 renewable electricity
    ("Renewable electricity is treated as a calculated input", [27, 28, 32, 35, 36]),
    # Note: the paragraph beginning with "For solar generation" also contains
    # the "For wind generation" sentence, so refs for both methodologies are
    # attached here.
    ("For solar generation", [24, 25, 26, 27, 28, 29, 30, 31]),
    ("In the representative four-week dataset", [33, 34]),
    ("Because green-hydrogen cost depends on resource quality", [32, 33]),
    # §1.4 by-product H2 and CO2
    ("Industrial by-product hydrogen is represented", [37, 38, 39]),
    ("These industrial hydrogen values are used as availability", [38, 39]),
    ("The refinery and steel layers are kept separate", [37, 36]),
    ("Industrial carbon data enter the framework", [36, 40]),
    ("The point-source CO₂ table is converted", [41, 42, 43, 44, 45]),
    # "For DAC scenarios" paragraph also contains the 21 CCUS records mention,
    # so CO2RE is attached here together with the DAC cost refs.
    ("For DAC scenarios", [46, 47, 48, 49, 50, 51]),
    # §1.5 transport network
    ("Transport-network construction combines existing infrastructure", [36, 52, 54]),
    ("The pipeline layer is not treated as a fully designed future network", [52, 54]),
    ("Downstream synthetic aviation fuel deliveries", [53, 57, 58, 55]),
    ("Economic routing assumptions are paired", [59, 60, 61]),
    # §1.6 preprocessing
    ("The raw spatial source tables are much denser", [63, 64, 65]),
    ("The repository applies DBSCAN clustering", [62]),
    ("Cluster-center selection is not based on geometric proximity", [63, 64]),
    ("After clustering, candidate logistics arcs are generated", [52, 54]),
    # §3 methodological supplement
    ("The 13 archived pathway models are written here as scenario-indexed instances", [89, 90, 91, 92, 94]),
    ("The core optimization objective minimizes annualized total system cost", [95, 17, 74]),
    ("The MILP uses a compact multi-commodity flow notation", [89, 92]),
    ("Binary siting variables are linked to continuous capacity", [89, 90]),
    ("Feedstock availability is enforced before conversion", [95, 102]),
    ("The active conversion block maps feedstocks to final synthetic aviation fuel output", [73, 75, 117]),
    ("Inventory states bridge variable production and scheduled logistics", [110, 112, 16]),
    ("Airport demand is enforced on a weekly basis", [92]),
    ("Emissions are calculated ex post on a well-to-wake basis", [9, 87, 106]),
    ("The archived models are implemented in Python and solved with Gurobi", [93]),
    # §4 sensitivity & uncertainty
    ("The three main results reported in the paper each rest on parameter values", [96, 101]),
    ("This analysis tests the main-text finding that GTL", [85, 86]),
    ("GTL-BH sources its hydrogen not from electrolysis or reforming", [84, 38, 98]),
    ("The main text identifies CCU-GH-FT as the deep-decarbonization archetype", [73, 74, 77]),
    ("The IEA 2030 scenario locates a specific point on this surface", [99, 100, 97]),
    ("The main text notes that DAC-linked pathways carry the highest LCO-SynAF values", [46, 47, 48, 49]),
    ("The core contribution of the main text is the identification of three distinct transition logics", [96, 70]),
    ("Three layers of uncertainty bear on the present model implementation", [101, 96, 18]),
    ("Despite these limitations, the findings of the sensitivity analysis are sufficient", [70, 96]),
    # §5 validation
    ("The methodological contribution of the current framework sits between plant-level techno-economic analysis", [103, 74, 13, 14, 15]),
    ("The archived temporal-robustness summary can also be read as an extreme-case operational test", [96, 107]),
    ("The archived cost–emissions frontier remains qualitatively stable across the supporting plots", [68, 70, 7]),
    ("Temporal penalties widen the economic separation between these pathways", [107, 114]),
    # §6 supplementary results
    ("Cost-optimal pathway selection does not guarantee smooth production", [107, 114, 115]),
    ("Facility utilization is computed at the 3-hour production-period level", [111, 112]),
    ("Green-hydrogen pathways (CCU-GH and DAC-GH families) show distinctly bimodal", [110, 113, 114]),
    ("Raw utilization rates do not directly quantify the implied cost premium", [107, 111, 112]),
    ("The three Pareto-representative pathways span the full range of temporal penalty behaviour", [70]),
    ("To attribute penalty differences among the three Pareto pathways", [111, 112, 107]),
    ("The following figures extend the main-text comparison by showing the full 13-pathway landscape", [67, 68]),
    ("Spatial outputs in the repository are summarised into clustered archetypes", [70, 92]),
    ("The appendix cost-decomposition panel focuses on the same three representative Pareto pathways", [73, 74, 75]),
    ("The emissions-decomposition panel uses the same representative pathways", [9, 87, 106]),
    ("The evidence assembled in Sections 6.1 through 6.7 establishes a triple discontinuity", [7, 70, 114, 118]),
    ("These three discontinuities jointly imply infrastructure non-inheritability", [118, 7, 80]),
]


def _clone_run_format(source_run, target_run) -> None:
    """Copy character formatting (font, size, color) from source run to target."""
    src_rpr = source_run._r.find(qn("w:rPr"))
    if src_rpr is None:
        return
    new_rpr = copy.deepcopy(src_rpr)
    existing = target_run._r.find(qn("w:rPr"))
    if existing is not None:
        target_run._r.remove(existing)
    target_run._r.insert(0, new_rpr)


def append_superscript(paragraph: Paragraph, ref_numbers: list[int]) -> None:
    """Append a Nature-style superscript citation to ``paragraph``."""
    if not paragraph.runs:
        raise RuntimeError(f"paragraph has no runs: {paragraph.text[:60]!r}")
    citation_text = ",".join(str(n) for n in ref_numbers)
    template_run = paragraph.runs[-1]
    sup_run = paragraph.add_run(citation_text)
    _clone_run_format(template_run, sup_run)
    sup_run.font.superscript = True


def find_paragraph(doc: Document, anchor: str) -> Paragraph | None:
    stripped = anchor.strip()
    for p in doc.paragraphs:
        if p.text.strip().startswith(stripped):
            return p
    return None


def insert_section24_review(doc: Document, anchor_substring: str = "The comparative analysis is restricted") -> int:
    """Insert literature-review paragraphs for §2.4 before the given anchor paragraph.

    Uses the anchor paragraph itself as a template for style and run formatting,
    then creates one new <w:p> element per review paragraph with the prose
    followed by a Nature-style superscript citation block.
    """
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(anchor_substring):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError(f"anchor not found: {anchor_substring!r}")
    if not anchor.runs:
        raise RuntimeError("anchor paragraph has no runs to clone formatting from")
    template_run = anchor.runs[0]
    anchor_style = anchor.style

    inserted = 0
    for text, refs in SECTION24_REVIEW_PARAGRAPHS:
        new_p_el = copy.deepcopy(anchor._element)
        # Strip all existing runs from the cloned element so we start from a
        # blank paragraph that still carries the paragraph-level properties.
        for r in list(new_p_el.findall(qn("w:r"))):
            new_p_el.remove(r)
        anchor._element.addprevious(new_p_el)
        new_para = Paragraph(new_p_el, anchor._parent)
        new_para.style = anchor_style
        prose_run = new_para.add_run(text)
        _clone_run_format(template_run, prose_run)
        citation_text = ",".join(str(n) for n in refs)
        sup_run = new_para.add_run(citation_text)
        _clone_run_format(template_run, sup_run)
        sup_run.font.superscript = True
        inserted += 1
    return inserted


def remove_existing_bibliography(doc: Document) -> int:
    """Delete every paragraph after (and not including) the Supplementary References heading."""
    heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Supplementary References":
            heading_idx = i
            break
    if heading_idx is None:
        raise RuntimeError("Supplementary References heading not found")
    for p in list(doc.paragraphs[heading_idx + 1 :]):
        p._element.getparent().remove(p._element)
    return heading_idx


def append_numbered_bibliography(doc: Document) -> None:
    style_names = {s.name for s in doc.styles}
    chosen_style = "Bibliography" if "Bibliography" in style_names else None
    for i, ref in enumerate(REFS, start=1):
        p = doc.add_paragraph(style=chosen_style)
        p.add_run(f"{i}. {ref}")


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    log.info("backup saved to %s", BACKUP_PATH)

    doc = Document(str(DOCX_PATH))

    missing: list[str] = []
    inserted = 0
    # To avoid double-annotating when unicode variants are both present (e.g. CO2 / CO₂),
    # track paragraphs we have already cited so an alternate anchor is a no-op.
    seen_paragraph_ids: set[int] = set()
    for anchor, refs in CITATION_INSERTIONS:
        p = find_paragraph(doc, anchor)
        if p is None:
            missing.append(anchor)
            continue
        pid = id(p._element)
        if pid in seen_paragraph_ids:
            log.info("skip duplicate anchor (same paragraph): %r", anchor[:60])
            continue
        append_superscript(p, refs)
        seen_paragraph_ids.add(pid)
        inserted += 1
        log.info("inserted %s after %r", refs, anchor[:60])
    if missing:
        log.warning("%d anchors not found:", len(missing))
        for a in missing:
            log.warning("  - %r", a[:80])

    review_inserted = insert_section24_review(doc)
    log.info("inserted %d §2.4 review paragraphs", review_inserted)

    remove_existing_bibliography(doc)
    append_numbered_bibliography(doc)
    log.info("rebuilt references block with %d entries", len(REFS))

    doc.save(str(DOCX_PATH))
    log.info("saved %s (inserted=%d, refs=%d, missing=%d)", DOCX_PATH, inserted, len(REFS), len(missing))


if __name__ == "__main__":
    main()
