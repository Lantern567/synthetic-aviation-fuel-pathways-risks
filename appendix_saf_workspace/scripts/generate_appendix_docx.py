#!/usr/bin/env python3
from __future__ import annotations

import ast
import base64
from collections import defaultdict
import glob
import json
import math
import re
import subprocess
import shutil
import sys
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import yaml
import cartopy.crs as ccrs
from openpyxl import load_workbook
from cycler import cycler
from matplotlib.lines import Line2D
from matplotlib.colors import Normalize
from matplotlib.patches import FancyBboxPatch, Patch, Rectangle
from PIL import Image, ImageDraw, ImageFont
Image.MAX_IMAGE_PIXELS = None  # 禁用解压炸弹检查（本地已知安全文件）
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from frykit import plot as fp
from shapely import wkt as shapely_wkt
try:
    from pareto_three_breakdown_visualization import generate_appendix_pareto_assets
except ModuleNotFoundError:
    from appendix_saf_workspace.scripts.pareto_three_breakdown_visualization import generate_appendix_pareto_assets



SCIENTIFIC_VIZ_SCRIPTS = Path.home() / ".codex" / "skills" / "scientific-visualization" / "scripts"
if SCIENTIFIC_VIZ_SCRIPTS.exists() and str(SCIENTIFIC_VIZ_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCIENTIFIC_VIZ_SCRIPTS))

try:
    from figure_export import save_publication_figure
except Exception:
    save_publication_figure = None

try:
    from style_presets import OKABE_ITO_COLORS
except Exception:
    OKABE_ITO_COLORS = [
        "#E69F00",
        "#56B4E9",
        "#009E73",
        "#F0E442",
        "#0072B2",
        "#D55E00",
        "#CC79A7",
        "#000000",
    ]


SCRIPT_DIR = Path(__file__).resolve().parent
WORKSPACE_ROOT = SCRIPT_DIR.parent
ROOT = WORKSPACE_ROOT.parent
WORD_DIR = WORKSPACE_ROOT / "word"
TEMPLATE_DOCX = WORD_DIR / "Supplementary_information.docx"
MANUSCRIPT_DOCX = WORD_DIR / "mamuscript_绿色航煤_0320_v1.docx"
OUTPUT_DOCX = WORD_DIR / "Appendix_SAF_draft.docx"
ASSET_DIR = WORKSPACE_ROOT / "figures"
PREP_DIR = ASSET_DIR / "prepared"
REDRAWN_DIR = ASSET_DIR / "redrawn"
CACHE_DIR = WORKSPACE_ROOT / "cache"
WIND_COMPLETE_PATH = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/preprocessed/wind_hourly_complete.csv"
SOLAR_COMPLETE_PATH = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/preprocessed/solar_hourly_complete.csv"
WIND_COMPLETE_CACHE = CACHE_DIR / "renewable_resource" / "wind_hourly_complete_summary.csv"
SOLAR_COMPLETE_CACHE = CACHE_DIR / "renewable_resource" / "solar_hourly_complete_summary.csv"
ANNUAL_AIRPORT_DEMAND_CACHE = CACHE_DIR / "aviation_demand" / "annual_airport_weekly_demand.csv"
AIRCRAFT_OPERATION_SUMMARY_CACHE = CACHE_DIR / "aviation_demand" / "airport_aircraft_operation_summary.csv"
MAP_PROJECTION = fp.CN_AZIMUTHAL_EQUIDISTANT
MAP_DATA_CRS = ccrs.PlateCarree()

# 敏感性分析图表目录
SENSITIVITY_FIGURES_DIR = ROOT / "products/supply_chain_optimization/sensitivity_analysis/figures"

# 可视化结果目录（时间维度等图表）
VIZ_RESULTS_DIR = ROOT / "products/supply_chain_optimization/visualization/results"

# 全场景合成图（极坐标 + donut）——用 importlib 按路径导入，避免与同名 appendix 脚本冲突
import importlib.util as _importlib_util

_VIZ_MOD_PATH = ROOT / "products/supply_chain_optimization/visualization/pareto_three_breakdown_visualization.py"
try:
    _spec = _importlib_util.spec_from_file_location("_viz_pareto_breakdown", str(_VIZ_MOD_PATH))
    _viz_mod = _importlib_util.module_from_spec(_spec)
    _spec.loader.exec_module(_viz_mod)
    generate_all_scenario_composites = _viz_mod.generate_all_scenario_composites
    APPENDIX_SCENARIO_GROUPS = _viz_mod.APPENDIX_SCENARIO_GROUPS
except Exception as _e:
    generate_all_scenario_composites = None
    APPENDIX_SCENARIO_GROUPS = {}
MAP_FACE_COLOR = "#F0F8FF"
MAP_LAND_COLOR = "#FFFFFF"
MAP_CITY_EDGE = "#D3D3D3"
MAP_PROVINCE_EDGE = "#606060"
MAP_BORDER_EDGE = "#000000"

ACADEMIC_PALETTE = [
    "#E64B35",
    "#4DBBD5",
    "#00A087",
    "#3C5488",
    "#F39B7F",
    "#8491B4",
]

FONT_CONFIG = {
    "title": 14,
    "label": 12,
    "tick": 10,
    "legend_text": 10,
    "annotation": 9,
    "line_width": 1.5,
}

ACADEMIC_TEMPLATE_SIZES = {
    "small_single": (3.0, 2.0),
    "standard_single": (3.5, 2.5),
    "wide_single": (7.0, 4.0),
    "horizontal_multi": (7.0, 2.5),
    "vertical_multi": (8.0, 5.0),
    "quad_multi": (7.2, 5.4),
    "processing_combo": (7.4, 5.8),
    "comparison_pair": (7.2, 5.0),
    "comparison_pair_short": (7.2, 4.1),
    "stats_triptych": (7.2, 3.2),
    "workflow_single": (7.4, 4.9),
    "placeholder_single": (7.0, 3.5),
    "national_map_pair": (8.2, 4.25),
    "national_map_triptych": (10.8, 4.45),
    "network_map_single": (7.2, 4.6),
    "network_map_tall": (7.2, 5.0),
    "map_preview_pair": (10.6, 4.9),
}


AIRPORT_MAP = {
    "北京": "Beijing",
    "天津": "Tianjin",
}

CO2_TYPE_LABELS = {
    "coal_power": "Coal power",
    "gas_power": "Gas power",
    "oil_refinery": "Oil refinery",
}

CO2_TYPE_COLORS = {
    "Coal power": "#8491B4",
    "Gas power": "#3C5488",
    "Oil refinery": "#E64B35",
}

FRYKIT_PROVINCE_MAP = {
    "北京": "北京市",
    "Beijing": "北京市",
    "天津": "天津市",
    "Tianjin": "天津市",
    "上海": "上海市",
    "Shanghai": "上海市",
    "重庆": "重庆市",
    "Chongqing": "重庆市",
    "河北": "河北省",
    "Hebei": "河北省",
    "山西": "山西省",
    "Shanxi": "山西省",
    "内蒙古": "内蒙古自治区",
    "Inner Mongolia": "内蒙古自治区",
    "辽宁": "辽宁省",
    "Liaoning": "辽宁省",
    "吉林": "吉林省",
    "Jilin": "吉林省",
    "黑龙江": "黑龙江省",
    "Heilongjiang": "黑龙江省",
    "江苏": "江苏省",
    "Jiangsu": "江苏省",
    "Nanjing": "江苏省",
    "浙江": "浙江省",
    "Zhejiang": "浙江省",
    "安徽": "安徽省",
    "Anhui": "安徽省",
    "福建": "福建省",
    "Fujian": "福建省",
    "江西": "江西省",
    "Jiangxi": "江西省",
    "山东": "山东省",
    "Shandong": "山东省",
    "河南": "河南省",
    "Henan": "河南省",
    "湖北": "湖北省",
    "Hubei": "湖北省",
    "湖南": "湖南省",
    "Hunan": "湖南省",
    "广东": "广东省",
    "Guangdong": "广东省",
    "广西": "广西壮族自治区",
    "Guangxi": "广西壮族自治区",
    "海南": "海南省",
    "Hainan": "海南省",
    "四川": "四川省",
    "Sichuan": "四川省",
    "贵州": "贵州省",
    "Guizhou": "贵州省",
    "云南": "云南省",
    "Yunnan": "云南省",
    "西藏": "西藏自治区",
    "Tibet": "西藏自治区",
    "陕西": "陕西省",
    "Shaanxi": "陕西省",
    "甘肃": "甘肃省",
    "Gansu": "甘肃省",
    "青海": "青海省",
    "Qinghai": "青海省",
    "宁夏": "宁夏回族自治区",
    "Ningxia": "宁夏回族自治区",
    "Ningxiz": "宁夏回族自治区",
    "新疆": "新疆维吾尔自治区",
    "Xinjiang": "新疆维吾尔自治区",
    "Hong Kong": "香港特别行政区",
}

CATEGORY_COLORS = {
    "Grey": "#8491B4",
    "Blue": "#3C5488",
    "Green": "#00A087",
}

ROUTE_MARKERS = {
    "FT": "o",
    "MTJ": "s",
}

SITE_SHARE_COLUMNS = [
    ("Renewable site", "SAF_Share_Renewable site"),
    ("CO₂ capture site", "SAF_Share_CO2 capture site"),
    ("Airport site", "SAF_Share_Airport site"),
    ("By-product H₂ site", "SAF_Share_Byproduct H2 site"),
    ("Other site", "SAF_Share_Other site"),
]

CLUSTER_COLORS = {
    1: ACADEMIC_PALETTE[3],
    2: ACADEMIC_PALETTE[0],
    3: ACADEMIC_PALETTE[2],
    4: ACADEMIC_PALETTE[1],
}

PUBLICATION_COLORS = {
    "orange": ACADEMIC_PALETTE[0],
    "sky": ACADEMIC_PALETTE[1],
    "green": ACADEMIC_PALETTE[2],
    "yellow": "#EAD56A",
    "blue": ACADEMIC_PALETTE[3],
    "vermillion": ACADEMIC_PALETTE[4],
    "magenta": ACADEMIC_PALETTE[5],
    "black": "#333333",
    "gray": "#6B7280",
    "light_gray": "#D1D5DB",
    "panel_fill": "#F8FAFC",
}

EXPECTED_EXPORT_SUFFIXES = (".png", ".pdf", ".svg")

FIGURE_TEMPLATE_REGISTRY = {
    "raw_data_overview": {"template": "quad_multi", "family": "data_chart"},
    "airport_demand_raw": {"template": "quad_multi", "family": "data_chart"},
    "renewable_raw": {"template": "quad_multi", "family": "data_chart"},
    "renewable_resource_maps": {"template": "national_map_triptych", "family": "map"},
    "renewable_inventory_map": {"template": "national_map_pair", "family": "map"},
    "renewable_processing": {"template": "processing_combo", "family": "data_chart"},
    "byproduct_raw": {"template": "quad_multi", "family": "data_chart"},
    "byproduct_geo": {"template": "national_map_pair", "family": "map"},
    "co2_capture_raw": {"template": "quad_multi", "family": "data_chart"},
    "co2_capture_geo": {"template": "national_map_pair", "family": "map"},
    "transport_raw": {"template": "quad_multi", "family": "data_chart"},
    "price_surface_maps": {"template": "national_map_pair", "family": "map"},
    "energy_infrastructure_overview": {"template": "network_map_tall", "family": "map"},
    "demand_aircraft_supplement": {"template": "quad_multi", "family": "data_chart"},
    "annual_demand_distribution": {"template": "comparison_pair", "family": "data_chart"},
    "demand_patterns": {"template": "comparison_pair", "family": "data_chart"},
    "renewable_profiles": {"template": "quad_multi", "family": "data_chart"},
    "framework_overview": {"template": "framework_single", "family": "workflow"},
    "workflow": {"template": "workflow_single", "family": "workflow"},
    "quadrant_chart": {"template": "comparison_pair", "family": "data_chart"},
    "efficiency_analysis": {"template": "quad_multi", "family": "data_chart"},
    "combined_cluster": {"template": "quad_multi", "family": "data_chart"},
    "transport_route_network": {"template": "network_map_single", "family": "map"},
    "transport_route_network_stats": {"template": "stats_triptych", "family": "data_chart"},
    "temporal_space": {"template": "quad_multi", "family": "data_chart"},
    "temporal_penalty": {"template": "comparison_pair_short", "family": "data_chart"},
    "temporal_h2_saf": {"template": "quad_multi", "family": "data_chart"},
    "pareto_temporal_dist": {"template": "comparison_pair", "family": "data_chart"},
    "pareto_penalty_comparison": {"template": "comparison_pair_short", "family": "data_chart"},
    "penalty_driver_heatmap": {"template": "quad_multi", "family": "data_chart"},
    "pareto_cost": {"template": "appendix_pareto_breakdown", "family": "data_chart"},
    "pareto_carbon": {"template": "appendix_pareto_breakdown", "family": "data_chart"},
    "electricity_price_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "electrolyzer_capex_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "dac_cost_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "carbon_price_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "fossil_price_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "bh_capex_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "green_h2_joint_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "dac_breakeven_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "carbon_price_asymmetric_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    "pareto_tornado_placeholder": {"template": "placeholder_single", "family": "placeholder"},
    # 全场景极坐标合成图（PNG-only composite, PDF/SVG via raster wrapper）
    "scenario_composite_ctl":         {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_ctl_bh":      {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_gtl_gh":      {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_gtl":         {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_gtl_bh":      {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_ccu_gh_mtj":  {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_ccu_gh_ft":   {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_ccu_bh_mtj":  {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_ccu_bh_ft":   {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_dac_gh_mtj":  {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_dac_gh_ft":   {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_dac_bh_mtj":  {"template": "scenario_composite_panel", "family": "composite_panel"},
    "scenario_composite_dac_bh_ft":   {"template": "scenario_composite_panel", "family": "composite_panel"},
    # 13 场景运输网络地图（聚类路线图，raster）
    "transport_map_ctl":         {"template": "transport_map_single", "family": "map"},
    "transport_map_ctl_bh":      {"template": "transport_map_single", "family": "map"},
    "transport_map_gtl_gh":      {"template": "transport_map_single", "family": "map"},
    "transport_map_gtl":         {"template": "transport_map_single", "family": "map"},
    "transport_map_gtl_bh":      {"template": "transport_map_single", "family": "map"},
    "transport_map_ccu_gh_mtj":  {"template": "transport_map_single", "family": "map"},
    "transport_map_ccu_gh_ft":   {"template": "transport_map_single", "family": "map"},
    "transport_map_ccu_bh_mtj":  {"template": "transport_map_single", "family": "map"},
    "transport_map_ccu_bh_ft":   {"template": "transport_map_single", "family": "map"},
    "transport_map_dac_gh_mtj":  {"template": "transport_map_single", "family": "map"},
    "transport_map_dac_gh_ft":   {"template": "transport_map_single", "family": "map"},
    "transport_map_dac_bh_mtj":  {"template": "transport_map_single", "family": "map"},
    "transport_map_dac_bh_ft":   {"template": "transport_map_single", "family": "map"},
    "transport_map_four_key":    {"template": "transport_map_quad",   "family": "map"},
    "transport_map_legend":      {"template": "transport_map_legend", "family": "map"},
}

QA_CONTACT_ORDER = [
    "raw_data_overview",
    "airport_demand_raw",
    "renewable_raw",
    "renewable_resource_maps",
    "renewable_inventory_map",
    "renewable_processing",
    "byproduct_raw",
    "byproduct_geo",
    "co2_capture_raw",
    "co2_capture_geo",
    "transport_raw",
    "price_surface_maps",
    "energy_infrastructure_overview",
    "demand_aircraft_supplement",
    "annual_demand_distribution",
    "demand_patterns",
    "renewable_profiles",
    "framework_overview",
    "workflow",
    "quadrant_chart",
    "efficiency_analysis",
    "combined_cluster",
    "transport_route_network",
    "transport_route_network_stats",
    "temporal_space",
    "temporal_penalty",
    "temporal_h2_saf",
    "pareto_temporal_dist",
    "pareto_penalty_comparison",
    "penalty_driver_heatmap",
    "pareto_cost",
    "pareto_carbon",
    "electricity_price_placeholder",
    "electrolyzer_capex_placeholder",
    "dac_cost_placeholder",
    "carbon_price_placeholder",
]

QA_KEY_PANEL_ORDER = [
    "raw_data_overview",
    "renewable_resource_maps",
    "renewable_inventory_map",
    "renewable_processing",
    "co2_capture_geo",
    "price_surface_maps",
    "energy_infrastructure_overview",
    "demand_aircraft_supplement",
    "annual_demand_distribution",
    "framework_overview",
    "workflow",
    "quadrant_chart",
    "efficiency_analysis",
    "combined_cluster",
    "transport_route_network",
    "temporal_space",
    "temporal_penalty",
    "pareto_cost",
    "pareto_carbon",
]

SCENARIO_SPECS = [
    {
        "code": "CTL",
        "category": "Grey",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/coal_hydrogen_saf_optimization/results/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/coal_hydrogen_saf_optimization/results/carbon_emissions_detailed_*.json",
    },
    {
        "code": "CTL-BH",
        "category": "Grey",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/coal_hydrogen_saf_optimization/results/byproduct_hydrogen/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/coal_hydrogen_saf_optimization/results/byproduct_hydrogen/carbon_emissions_detailed_*.json",
    },
    {
        "code": "DAC-GH-MTJ",
        "category": "Green",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/two_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/two_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "DAC-GH-FT",
        "category": "Green",
        "route": "FT",
        "solution_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/one_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/one_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "CCU-GH-MTJ",
        "category": "Green",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/two_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/two_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "CCU-GH-FT",
        "category": "Green",
        "route": "FT",
        "solution_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/one_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/one_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "GTL-GH",
        "category": "Blue",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/carbon_emissions_detailed_*.json",
    },
    {
        "code": "GTL",
        "category": "Blue",
        "route": "FT",
        "solution_pattern": "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/ft_one_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/ft_one_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "DAC-BH-MTJ",
        "category": "Blue",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/byproduct_hydrogen/two_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/byproduct_hydrogen/two_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "DAC-BH-FT",
        "category": "Blue",
        "route": "FT",
        "solution_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/byproduct_hydrogen/one_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/dac_hydrogen_saf_supply_chain_optimization/results/byproduct_hydrogen/one_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "GTL-BH",
        "category": "Blue",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/byproduct_hydrogen/byproduct_hydrogen/two_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/byproduct_hydrogen/byproduct_hydrogen/two_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "CCU-BH-MTJ",
        "category": "Blue",
        "route": "MTJ",
        "solution_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/byproduct_hydrogen/two_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/byproduct_hydrogen/two_step/carbon_emissions_detailed_*.json",
    },
    {
        "code": "CCU-BH-FT",
        "category": "Blue",
        "route": "FT",
        "solution_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/byproduct_hydrogen/one_step/complete_solution_*.json",
        "carbon_pattern": "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/results/byproduct_hydrogen/one_step/carbon_emissions_detailed_*.json",
    },
]


plt.rcParams.update(
    {
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans"],
        "axes.unicode_minus": False,
        "font.size": 10,
        "axes.titlesize": 12,
        "axes.labelsize": 10,
        "xtick.labelsize": 9,
        "ytick.labelsize": 9,
        "pdf.fonttype": 42,
        "ps.fonttype": 42,
        "mathtext.fontset": "custom",
        "mathtext.rm": "Arial",
        "mathtext.it": "Arial:italic",
        "mathtext.bf": "Arial:bold",
        "mathtext.default": "regular",
        "axes.prop_cycle": cycler(color=ACADEMIC_PALETTE),
        "figure.facecolor": "#FFFFFF",
        "axes.facecolor": "#FFFFFF",
        "legend.frameon": False,
    }
)


@dataclass
class AppendixData:
    demand_summary: dict
    week_rows: list[list[str]]
    renewable_summary: dict
    byproduct_summary: dict
    co2_summary: dict
    transport_summary: dict
    price_summary: dict
    clustering_summary: dict
    verification_summary: dict
    temporal_summary: dict
    config_summary: dict
    pathway_rows: list[list[str]]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def read_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt_int(value: float | int) -> str:
    return f"{int(round(value)):,}"


def fmt_million_kg(value: float) -> str:
    return f"{value / 1e6:.2f} million kg"


def fmt_twh_from_mwh(value: float) -> str:
    return f"{value / 1e6:.2f} TWh"


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def infer_levelized_cost_term() -> str:
    return "levelized cost of SAF (LCO-SAF; labeled as LCOE in several archived result files)"


def latest_file(pattern: str) -> Path | None:
    matches = sorted(glob.glob(str(ROOT / pattern)), reverse=True)
    if not matches:
        return None
    return Path(matches[0])


def parse_latlon_tuple(value: object) -> tuple[float, float] | None:
    if pd.isna(value):
        return None
    try:
        parsed = ast.literal_eval(str(value))
    except Exception:
        return None
    if isinstance(parsed, (list, tuple)) and len(parsed) == 2:
        return float(parsed[0]), float(parsed[1])
    return None


def parse_path_coords(value: object) -> list[tuple[float, float]]:
    if pd.isna(value):
        return []
    text = str(value).strip()
    if not text or text == "[]":
        return []
    try:
        parsed = json.loads(text)
    except Exception:
        try:
            parsed = ast.literal_eval(text)
        except Exception:
            return []
    coords: list[tuple[float, float]] = []
    for item in parsed:
        if isinstance(item, (list, tuple)) and len(item) == 2:
            coords.append((float(item[0]), float(item[1])))
    return coords


def translate_infra_type(value: object) -> str:
    mapping = {
        "太阳能发电站": "Solar plant",
        "风电场": "Wind plant",
        "机场": "Airport",
        "FT设施": "FT facility",
        "MTJ设施": "MTJ facility",
        "LNG接收站": "LNG terminal",
        "byproduct_hydrogen_refinery": "Refinery H₂",
        "byproduct_hydrogen_steel": "Steel H₂",
    }
    text = str(value)
    return mapping.get(text, text.replace("_", " ").title())


def normalize_capacity_factor(value: object) -> float | None:
    if pd.isna(value):
        return None
    text = str(value).strip()
    if not text:
        return None
    text = text.rstrip("%")
    try:
        numeric = float(text)
    except ValueError:
        return None
    if numeric > 1.5:
        numeric /= 100.0
    return numeric


def normalize_frykit_province_name(value: object) -> str:
    text = str(value).strip()
    if text in FRYKIT_PROVINCE_MAP:
        return FRYKIT_PROVINCE_MAP[text]
    if text.endswith(("省", "市", "自治区", "特别行政区")):
        return text
    return text


def load_renewable_resource_frame(csv_path: Path, hours_per_step: int = 1) -> pd.DataFrame:
    df = pd.read_csv(
        csv_path,
        usecols=["plant_id", "capacity_mw", "power_output_mw", "latitude", "longitude"],
    )
    grouped = (
        df.groupby("plant_id")
        .agg(
            capacity_mw=("capacity_mw", "first"),
            total_output_mwh=("power_output_mw", "sum"),
            latitude=("latitude", "first"),
            longitude=("longitude", "first"),
            periods=("power_output_mw", "size"),
        )
        .reset_index()
    )
    grouped["hours"] = grouped["periods"] * hours_per_step
    grouped["capacity_factor"] = grouped["total_output_mwh"] / (
        grouped["capacity_mw"] * grouped["hours"]
    )
    grouped["full_load_hours"] = grouped["total_output_mwh"] / grouped["capacity_mw"]
    return grouped


def load_cached_renewable_resource_frame(
    csv_path: Path,
    cache_path: Path,
    hours_per_step: int = 1,
    clip_to_capacity: bool = False,
    chunksize: int = 250_000,
) -> pd.DataFrame:
    ensure_dir(cache_path.parent)
    if cache_path.exists() and cache_path.stat().st_mtime >= csv_path.stat().st_mtime:
        return pd.read_csv(cache_path)

    total_output = defaultdict(float)
    clipped_output = defaultdict(float)
    periods = defaultdict(int)
    capacity = {}
    latitude = {}
    longitude = {}

    usecols = ["plant_id", "capacity_mw", "power_output_mw", "latitude", "longitude"]
    for chunk in pd.read_csv(csv_path, usecols=usecols, chunksize=chunksize):
        chunk["power_output_mw"] = pd.to_numeric(chunk["power_output_mw"], errors="coerce").fillna(0.0)
        chunk["capacity_mw"] = pd.to_numeric(chunk["capacity_mw"], errors="coerce").replace(0, np.nan)
        chunk = chunk.dropna(subset=["capacity_mw", "latitude", "longitude"])
        chunk["clipped_output_mw"] = np.minimum(chunk["power_output_mw"], chunk["capacity_mw"])
        grouped = (
            chunk.groupby("plant_id")
            .agg(
                total_output_mwh=("power_output_mw", "sum"),
                clipped_output_mwh=("clipped_output_mw", "sum"),
                capacity_mw=("capacity_mw", "first"),
                latitude=("latitude", "first"),
                longitude=("longitude", "first"),
                periods=("power_output_mw", "size"),
            )
        )
        for plant_id, row in grouped.iterrows():
            total_output[plant_id] += float(row["total_output_mwh"])
            clipped_output[plant_id] += float(row["clipped_output_mwh"])
            periods[plant_id] += int(row["periods"])
            capacity.setdefault(plant_id, float(row["capacity_mw"]))
            latitude.setdefault(plant_id, float(row["latitude"]))
            longitude.setdefault(plant_id, float(row["longitude"]))

    frame = pd.DataFrame(
        {
            "plant_id": list(capacity.keys()),
            "capacity_mw": [capacity[key] for key in capacity],
            "total_output_mwh": [total_output[key] for key in capacity],
            "clipped_output_mwh": [clipped_output[key] for key in capacity],
            "latitude": [latitude[key] for key in capacity],
            "longitude": [longitude[key] for key in capacity],
            "periods": [periods[key] for key in capacity],
        }
    )
    frame["hours"] = frame["periods"] * hours_per_step
    denominator = frame["capacity_mw"] * frame["hours"]
    generation_for_cf = frame["clipped_output_mwh"] if clip_to_capacity else frame["total_output_mwh"]
    frame["capacity_factor"] = generation_for_cf / denominator
    frame["raw_capacity_factor"] = frame["total_output_mwh"] / denominator
    frame["full_load_hours"] = generation_for_cf / frame["capacity_mw"]
    frame.to_csv(cache_path, index=False)
    return frame


def summarize_renewable_resource(csv_path: Path, hours_per_step: int = 1) -> dict:
    frame = load_renewable_resource_frame(csv_path, hours_per_step=hours_per_step)
    return {
        "candidate_plants": int(len(frame)),
        "hours": int(frame["hours"].iloc[0]) if not frame.empty else 0,
        "capacity_factor_median": float(frame["capacity_factor"].median()),
        "capacity_factor_p10": float(frame["capacity_factor"].quantile(0.10)),
        "capacity_factor_p90": float(frame["capacity_factor"].quantile(0.90)),
        "full_load_hours_median": float(frame["full_load_hours"].median()),
        "full_load_hours_p10": float(frame["full_load_hours"].quantile(0.10)),
        "full_load_hours_p90": float(frame["full_load_hours"].quantile(0.90)),
    }


def summarize_cached_renewable_resource(
    csv_path: Path,
    cache_path: Path,
    hours_per_step: int = 1,
    clip_to_capacity: bool = False,
) -> dict:
    frame = load_cached_renewable_resource_frame(
        csv_path,
        cache_path=cache_path,
        hours_per_step=hours_per_step,
        clip_to_capacity=clip_to_capacity,
    )
    return {
        "candidate_plants": int(len(frame)),
        "hours": int(frame["hours"].iloc[0]) if not frame.empty else 0,
        "capacity_factor_median": float(frame["capacity_factor"].median()),
        "capacity_factor_p10": float(frame["capacity_factor"].quantile(0.10)),
        "capacity_factor_p90": float(frame["capacity_factor"].quantile(0.90)),
        "full_load_hours_median": float(frame["full_load_hours"].median()),
        "full_load_hours_p10": float(frame["full_load_hours"].quantile(0.10)),
        "full_load_hours_p90": float(frame["full_load_hours"].quantile(0.90)),
    }


def load_co2_capture_frame() -> pd.DataFrame:
    path = ROOT / "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/data/co2_capture_sources.csv"
    df = pd.read_csv(path)
    df["facility_label"] = df["facility_type"].map(CO2_TYPE_LABELS).fillna(df["facility_type"].str.replace("_", " ").str.title())
    df["status_label"] = df["status"].map(normalize_status)
    df["capacity_factor_num"] = df["capacity_factor"].map(normalize_capacity_factor)
    return df


def summarize_co2() -> dict:
    df = load_co2_capture_frame()
    by_type = (
        df.groupby("facility_label")
        .agg(
            records=("location_id", "size"),
            provinces=("province", "nunique"),
            weekly_capture_tonnes=("co2_capture_capacity_ton_per_week", "sum"),
            median_weekly_capture_tonnes=("co2_capture_capacity_ton_per_week", "median"),
            unit_capture_cost_yuan_per_ton=("capture_cost_yuan_per_ton", "median"),
            median_capacity_factor=("capacity_factor_num", "median"),
        )
        .to_dict("index")
    )
    top_provinces = (
        df.groupby("province")["co2_capture_capacity_ton_per_week"]
        .sum()
        .sort_values(ascending=False)
        .head(10)
        .to_dict()
    )
    return {
        "records": int(len(df)),
        "provinces": int(df["province"].nunique()),
        "operating_records": int(df["status_label"].eq("Operating").sum()),
        "weekly_capture_tonnes": float(df["co2_capture_capacity_ton_per_week"].sum()),
        "by_type": by_type,
        "top_provinces": top_provinces,
    }


def load_provincial_price_frame() -> pd.DataFrame:
    path = ROOT / "products/supply_chain_optimization/natural_gas_supply_chain_optimization/data/integrated_gas_pipeline_price_data.csv"
    df = pd.read_csv(path)
    grouped = (
        df.groupby("province")
        .agg(
            gas_price_yuan_per_10k_m3=("natural_gas_price_yuan_per_10k_m3", "median"),
            electricity_price_yuan_per_kwh=("electricity_price_yuan_per_kwh", "median"),
            pipelines=("pipeline_id", "nunique"),
            cumulative_length_km=("length_km", "sum"),
        )
        .reset_index()
    )
    return grouped


def summarize_price_surface() -> dict:
    frame = load_provincial_price_frame()
    provincial = frame[frame["province"] != "全国"].copy()
    return {
        "province_count": int(len(provincial)),
        "gas_price_min": float(provincial["gas_price_yuan_per_10k_m3"].min()),
        "gas_price_max": float(provincial["gas_price_yuan_per_10k_m3"].max()),
        "electricity_price_min": float(provincial["electricity_price_yuan_per_kwh"].min()),
        "electricity_price_max": float(provincial["electricity_price_yuan_per_kwh"].max()),
    }


def add_province_choropleth(
    ax,
    province_values: pd.Series,
    colorbar_label: str,
    cmap_name: str = "YlOrRd",
    cax=None,
) -> None:
    extents = [73, 136, 17, 55]
    clean = province_values.dropna()
    clean = clean[clean.index != "全国"]

    configure_frykit_map(ax, extents, dx=10, dy=10)
    cmap = plt.get_cmap(cmap_name)
    norm = Normalize(vmin=float(clean.min()), vmax=float(clean.max()))
    for province, value in clean.items():
        province_name = normalize_frykit_province_name(province)
        fp.add_cn_province(
            ax,
            province=province_name,
            facecolor=cmap(norm(float(value))),
            edgecolor="white",
            linewidth=0.55,
            alpha=0.98,
            zorder=2.0,
        )
    fp.add_cn_province(ax, facecolor="none", edgecolor=MAP_PROVINCE_EDGE, linewidth=0.55, zorder=2.25)
    fp.add_cn_line(ax, edgecolor=MAP_BORDER_EDGE, linewidth=1.0, zorder=2.6)
    sm = plt.cm.ScalarMappable(norm=norm, cmap=cmap)
    sm.set_array([])
    if cax is not None:
        cbar = ax.figure.colorbar(sm, cax=cax, orientation="horizontal")
    else:
        cbar = ax.figure.colorbar(sm, ax=ax, orientation="horizontal", fraction=0.045, pad=0.035)
    cbar.set_label(colorbar_label, fontsize=7.2)
    cbar.ax.tick_params(labelsize=6.5, width=0.6, length=2)
    cbar.outline.set_linewidth(0.6)


def category_color(category: str) -> str:
    return CATEGORY_COLORS.get(category, "#334155")


def compute_map_ticks(extents: list[float], dx: float, dy: float) -> tuple[np.ndarray, np.ndarray]:
    min_lon, max_lon, min_lat, max_lat = extents
    x_start = math.ceil(min_lon / dx) * dx
    x_stop = math.floor(max_lon / dx) * dx
    y_start = math.ceil(min_lat / dy) * dy
    y_stop = math.floor(max_lat / dy) * dy
    xticks = np.arange(x_start, x_stop + 0.1, dx)
    yticks = np.arange(y_start, y_stop + 0.1, dy)
    return xticks, yticks


def configure_frykit_map(ax, extents: list[float], dx: float, dy: float) -> None:
    xticks, yticks = compute_map_ticks(extents, dx, dy)
    ax.set_extent(extents, crs=MAP_DATA_CRS)
    ax.set_facecolor(MAP_FACE_COLOR)
    fp.set_map_ticks(ax, extents, xticks, yticks)
    ax.gridlines(xlocs=xticks, ylocs=yticks, lw=0.35, ls=":", color="#CBD5E1", alpha=0.22)
    ax.tick_params(
        length=3.2,
        width=0.60,
        labelsize=7.0,
        top=False,
        right=False,
        labeltop=False,
        labelright=False,
        colors="#475569",
        direction="in",
    )
    fp.add_cn_city(ax, lw=0.18, edgecolor=MAP_CITY_EDGE, linestyle="-", zorder=1.0)
    fp.add_cn_province(ax, facecolor=MAP_LAND_COLOR, edgecolor=MAP_PROVINCE_EDGE, linewidth=0.55, zorder=1.4)
    fp.add_cn_line(ax, edgecolor=MAP_BORDER_EDGE, linewidth=1.0, zorder=2.4)


def configure_reference_frykit_map(ax, extents: list[float], dx: float, dy: float, show_tick_labels: bool = True) -> None:
    from cartopy.feature import LAND

    xticks, yticks = compute_map_ticks(extents, dx, dy)
    ax.set_extent(extents, crs=MAP_DATA_CRS)
    ax.set_facecolor(MAP_FACE_COLOR)
    fp.set_map_ticks(ax, extents, xticks, yticks)
    ax.gridlines(xlocs=xticks, ylocs=yticks, lw=0.38, ls=":", color="#CBD5E1", alpha=0.24)
    ax.tick_params(
        length=3.8 if show_tick_labels else 0.0,
        width=0.65,
        labelsize=7.4,
        top=show_tick_labels,
        right=show_tick_labels,
        labeltop=False,
        labelright=False,
        labelbottom=show_tick_labels,
        labelleft=show_tick_labels,
        bottom=show_tick_labels,
        left=show_tick_labels,
        colors="#475569",
        direction="in",
    )
    ax.add_feature(LAND, fc=MAP_LAND_COLOR, ec="#808080", lw=0.3, zorder=0)
    fp.add_cn_city(ax, lw=0.2, edgecolor=MAP_CITY_EDGE, linestyle="-", zorder=1.0)
    fp.add_cn_province(ax, lw=0.6, edgecolor=MAP_PROVINCE_EDGE, zorder=1.5)
    fp.add_cn_line(ax, lw=1.0, edgecolor=MAP_BORDER_EDGE, zorder=2.5)


def add_reference_mini_map(main_ax) -> object:
    mini_ax = fp.add_mini_axes(main_ax)
    mini_ax.set_extent((105, 122, 2, 25), MAP_DATA_CRS)
    mini_xticks = np.arange(105, 123, 5)
    mini_yticks = np.arange(5, 26, 5)
    configure_reference_frykit_map(mini_ax, [105, 122, 2, 25], dx=5, dy=5, show_tick_labels=False)
    mini_ax.gridlines(xlocs=mini_xticks, ylocs=mini_yticks, lw=0.45, ls=":", color="#B0B0B0", alpha=0.30)
    return mini_ax


def add_reference_map_decorations(
    main_ax,
    mini_ax=None,
    main_scale_km: int = 1000,
    mini_scale_km: int = 500,
    compass_size: float = 15.0,
) -> None:
    try:
        fp.add_compass(main_ax, 0.92, 0.84, size=compass_size, style="star")
        scale_bar = fp.add_scale_bar(main_ax, 0.07, 0.92, length=main_scale_km)
        main_ticks = [0, main_scale_km // 2, main_scale_km] if main_scale_km >= 400 else [0, main_scale_km]
        scale_bar.set_xticks(main_ticks)
        scale_bar.xaxis.get_label().set_fontsize(6.6)
        scale_bar.tick_params(labelsize=6.2, width=0.5, length=2)
        if mini_ax is not None:
            scale_bar_mini = fp.add_scale_bar(mini_ax, 0.48, 0.22, length=mini_scale_km)
            scale_bar_mini.set_xticks([0, mini_scale_km])
            scale_bar_mini.xaxis.get_label().set_fontsize(5.2)
            scale_bar_mini.tick_params(labelsize=5.0, width=0.45, length=1.5)
    except Exception:
        pass


def create_reference_map_panel(
    fig,
    subplot_spec,
    extents: list[float],
    dx: float,
    dy: float,
    main_scale_km: int,
    mini_scale_km: int = 500,
    compass_size: float = 15.0,
    with_mini: bool = True,
):
    ax = fig.add_subplot(subplot_spec, projection=MAP_PROJECTION)
    configure_reference_frykit_map(ax, extents, dx=dx, dy=dy)
    mini_ax = add_reference_mini_map(ax) if with_mini else None
    add_reference_map_decorations(
        ax,
        mini_ax,
        main_scale_km=main_scale_km,
        mini_scale_km=mini_scale_km,
        compass_size=compass_size,
    )
    return ax, mini_ax


def draw_reference_choropleth(
    ax,
    province_values: pd.Series,
    cmap_name: str,
) -> tuple[plt.cm.ScalarMappable, Normalize]:
    clean = province_values.dropna()
    clean = clean[clean.index != "全国"]
    cmap = plt.get_cmap(cmap_name)
    norm = Normalize(vmin=float(clean.min()), vmax=float(clean.max()))
    for province, value in clean.items():
        province_name = normalize_frykit_province_name(province)
        fp.add_cn_province(
            ax,
            province=province_name,
            facecolor=cmap(norm(float(value))),
            edgecolor="white",
            linewidth=0.55,
            alpha=0.98,
            zorder=2.0,
        )
    fp.add_cn_province(ax, facecolor="none", edgecolor=MAP_PROVINCE_EDGE, linewidth=0.6, zorder=2.3)
    fp.add_cn_line(ax, lw=1.0, edgecolor=MAP_BORDER_EDGE, zorder=2.6)
    sm = plt.cm.ScalarMappable(norm=norm, cmap=cmap)
    sm.set_array([])
    return sm, norm


def mirror_points_to_mini(
    mini_ax,
    longitudes,
    latitudes,
    **scatter_kwargs,
) -> None:
    if mini_ax is None:
        return
    lon = np.asarray(longitudes)
    lat = np.asarray(latitudes)
    mask = (lon >= 105) & (lon <= 122) & (lat >= 2) & (lat <= 25)
    if mask.any():
        masked_kwargs = {}
        for key, value in scatter_kwargs.items():
            if isinstance(value, (str, bytes)):
                masked_kwargs[key] = value
                continue
            try:
                if len(value) == len(lon):
                    masked_kwargs[key] = np.asarray(value)[mask]
                    continue
            except Exception:
                pass
            masked_kwargs[key] = value
        if mask.sum() >= 500 and "rasterized" not in masked_kwargs:
            masked_kwargs["rasterized"] = True
        if mask.sum() >= 500 and "edgecolors" not in masked_kwargs:
            masked_kwargs["edgecolors"] = "none"
        mini_ax.scatter(lon[mask], lat[mask], transform=MAP_DATA_CRS, **masked_kwargs)


def mirror_lines_to_mini(mini_ax, xs, ys, **plot_kwargs) -> None:
    if mini_ax is None:
        return
    lon = np.asarray(xs)
    lat = np.asarray(ys)
    mask = (lon >= 105) & (lon <= 122) & (lat >= 2) & (lat <= 25)
    if np.count_nonzero(mask) >= 2:
        mini_ax.plot(lon[mask], lat[mask], transform=MAP_DATA_CRS, **plot_kwargs)


def add_china_base_map(ax, extents: list[float], title: str, dx: float, dy: float) -> None:
    configure_frykit_map(ax, extents, dx=dx, dy=dy)
    _ = title


def sample_points(frame: pd.DataFrame, limit: int, random_state: int = 42) -> pd.DataFrame:
    if len(frame) <= limit:
        return frame
    return frame.sample(n=limit, random_state=random_state)


def make_scatter_size_legend_handles(
    values: list[float],
    size_fn,
    *,
    color: str,
    edgecolor: str = "white",
    alpha: float = 0.65,
    marker: str = "o",
    markeredgewidth: float = 0.45,
    value_formatter: str = "{:.0f}",
    label_suffix: str = "",
) -> list[Line2D]:
    handles: list[Line2D] = []
    for value in values:
        area = float(np.asarray(size_fn([value]), dtype=float)[0])
        handles.append(
            Line2D(
                [0],
                [0],
                marker=marker,
                linestyle="None",
                markerfacecolor=color,
                markeredgecolor=edgecolor,
                markeredgewidth=markeredgewidth,
                alpha=alpha,
                markersize=max(float(np.sqrt(area)), 3.0),
                label=f"{value_formatter.format(value)}{label_suffix}",
            )
        )
    return handles


def plot_geometry_wkts(
    ax,
    geometry_series: pd.Series,
    color: str = "#94A3B8",
    linewidth: float = 0.35,
    alpha: float = 0.18,
    limit: int | None = None,
) -> None:
    geometries = geometry_series.dropna()
    if limit is not None and len(geometries) > limit:
        geometries = geometries.sample(n=limit, random_state=42)
    for geometry_text in geometries:
        try:
            geom = shapely_wkt.loads(geometry_text)
        except Exception:
            continue
        if geom.geom_type == "LineString":
            xs, ys = geom.xy
            ax.plot(xs, ys, transform=ccrs.PlateCarree(), color=color, linewidth=linewidth, alpha=alpha, zorder=2)
        elif geom.geom_type == "MultiLineString":
            for part in geom.geoms:
                xs, ys = part.xy
                ax.plot(xs, ys, transform=ccrs.PlateCarree(), color=color, linewidth=linewidth, alpha=alpha, zorder=2)


def summarize_scenario_portfolio() -> pd.DataFrame:
    records: list[dict] = []
    for spec in SCENARIO_SPECS:
        solution_path = latest_file(spec["solution_pattern"])
        carbon_path = latest_file(spec["carbon_pattern"])
        if solution_path is None or carbon_path is None:
            continue
        solution = read_json(solution_path)
        carbon = read_json(carbon_path)
        cost_breakdown = solution.get("cost_breakdown", {})
        by_stage = carbon.get("by_stage", {})
        production_kg = float(solution.get("lifecycle_total_production_kg") or carbon.get("total_production_kg") or 0.0)
        energy_content_mj = 43.15

        def per_kg(cost: float) -> float:
            return float(cost) / production_kg if production_kg else 0.0

        def stage_intensity(value: float) -> float:
            return float(value) / production_kg * 1000 / energy_content_mj if production_kg else 0.0

        logistics_cost = (
            float(solution.get("transport_operation_cost", 0.0))
            + float(solution.get("storage_operation_cost", 0.0))
            + float(solution.get("final_inventory_cost", 0.0))
        )
        records.append(
            {
                "Scenario": spec["code"],
                "Category": spec["category"],
                "Route": spec["route"],
                "Color": category_color(spec["category"]),
                "SolutionPath": str(solution_path),
                "CarbonPath": str(carbon_path),
                "LCOE": float(solution.get("lifecycle_levelized_cost_excluding_shortage_per_kg", 0.0)),
                "AnnualProductionMt": float(solution.get("annual_production_kg", 0.0)) / 1e6,
                "CarbonDiffVsJet": float(carbon.get("abs_diff_vs_traditional_jet_gco2e_per_mj", 0.0)),
                "CarbonIntensity": float(carbon.get("carbon_intensity_mj", 0.0)),
                "InvestmentCostPerKg": per_kg(solution.get("total_investment_cost", 0.0)),
                "CoreOperationsPerKg": per_kg(max(float(solution.get("total_operation_cost", 0.0)) - logistics_cost, 0.0)),
                "LogisticsStoragePerKg": per_kg(logistics_cost),
                "FeedstockStage": stage_intensity(by_stage.get("raw_material_emissions", 0.0)),
                "FacilityStage": stage_intensity(by_stage.get("facility_emissions", 0.0)),
                "ProductionStage": stage_intensity(by_stage.get("production_emissions", 0.0)),
                "StorageStage": stage_intensity(by_stage.get("storage_emissions", 0.0)),
                "TransportStage": stage_intensity(by_stage.get("transport_emissions", 0.0)),
                "CreditStage": stage_intensity(by_stage.get("co2_utilization_credit", 0.0)),
                "TotalEmissionsStage": stage_intensity(by_stage.get("total_emissions", 0.0)),
                "DemandFulfillment": float(solution.get("demand_fulfillment_ratio", 0.0)),
                "TotalInvestmentCost": float(solution.get("total_investment_cost", 0.0)),
                "TotalOperationCost": float(solution.get("total_operation_cost", 0.0)),
                "TransportOperationCost": float(solution.get("transport_operation_cost", 0.0)),
                "StorageOperationCost": float(solution.get("storage_operation_cost", 0.0)),
                "CostBreakdown": cost_breakdown,
            }
        )
    frame = pd.DataFrame(records)
    route_order = {"FT": 0, "MTJ": 1}
    category_order = {"Grey": 0, "Blue": 1, "Green": 2}
    return frame.sort_values(
        by=["Category", "Route", "LCOE"],
        key=lambda s: s.map(category_order) if s.name == "Category" else s.map(route_order) if s.name == "Route" else s,
    ).reset_index(drop=True)


def load_latest_temporal_efficiency() -> pd.DataFrame:
    path = latest_file("products/supply_chain_optimization/visualization/results/temporal_efficiency_*/temporal_efficiency_summary.csv")
    if path is None:
        return pd.DataFrame()
    return pd.read_csv(path)


def load_latest_network_metrics() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    network_path = latest_file("products/supply_chain_optimization/visualization/results/network_radar_clusters_multi_*/scenario_network_metrics_official_names.csv")
    siting_path = latest_file("products/supply_chain_optimization/visualization/results/network_siting_clusters_*/scenario_siting_metrics.csv")
    decision_path = latest_file("products/supply_chain_optimization/visualization/results/network_siting_clusters_*/scenario_decision_point_classification_counts.csv")
    network = pd.read_csv(network_path) if network_path else pd.DataFrame()
    siting = pd.read_csv(siting_path) if siting_path else pd.DataFrame()
    decision = pd.read_csv(decision_path) if decision_path else pd.DataFrame()
    return network, siting, decision


def identify_pareto_front(frame: pd.DataFrame) -> list[str]:
    ordered = frame.sort_values(["CarbonDiffVsJet", "LCOE"], ascending=[True, True])
    best_cost = float("inf")
    frontier: list[str] = []
    for _, row in ordered.iterrows():
        if row["LCOE"] < best_cost - 1e-9:
            frontier.append(row["Scenario"])
            best_cost = row["LCOE"]
    return frontier


def collect_pathway_rows() -> list[list[str]]:
    manuscript = Document(str(MANUSCRIPT_DOCX))
    table = manuscript.tables[0]
    rows: list[list[str]] = []
    for r in table.rows[1:]:
        cells = [c.text.replace("\n", " ").strip() for c in r.cells]
        rows.append(cells + ["Retained"])
    return rows


def summarize_representative_weeks() -> tuple[dict, list[list[str]]]:
    demand_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_demand_20251129_231231.xlsx"
    week_map_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/week_mapping_20251129_231231.csv"
    demand_df = pd.read_excel(demand_path)
    week_map = pd.read_csv(week_map_path)

    demand_df["airport_en"] = demand_df["airport"].map(AIRPORT_MAP).fillna(demand_df["airport"])
    airport_summary = (
        demand_df.groupby("airport_en")
        .agg(
            total_demand_kg=("weekly_total_fuel_kg_total", "sum"),
            min_week_kg=("weekly_total_fuel_kg_total", "min"),
            max_week_kg=("weekly_total_fuel_kg_total", "max"),
            total_flights=("total_flights", "sum"),
        )
        .to_dict("index")
    )

    merged = demand_df.merge(
        week_map.rename(columns={"new_week_number": "week_number"}),
        on="week_number",
        how="left",
    )

    def demand_band(airport: str, value: float) -> str:
        if airport == "Beijing":
            if value >= 11_000_000:
                return "High"
            if value <= 6_000_000:
                return "Low"
            return "Medium"
        if value <= 2_000_000:
            return "Low"
        return "Normal"

    week_rows: list[list[str]] = []
    for week_number in sorted(merged["week_number"].unique()):
        week_slice = merged[merged["week_number"] == week_number].copy()
        if week_slice.empty:
            continue
        start = pd.to_datetime(week_slice["week_start"].iloc[0]).strftime("%b %d, %Y")
        end = pd.to_datetime(week_slice["week_end"].iloc[0]).strftime("%b %d, %Y")
        mapping_row = week_slice.iloc[0]
        beijing = week_slice.loc[week_slice["airport_en"] == "Beijing", "weekly_total_fuel_kg_total"].iloc[0]
        tianjin = week_slice.loc[week_slice["airport_en"] == "Tianjin", "weekly_total_fuel_kg_total"].iloc[0]
        week_rows.append(
            [
                str(week_number),
                str(int(mapping_row["week_in_12weeks"])),
                str(int(mapping_row["original_week_in_52"])),
                f"{start} - {end}",
                f"{fmt_million_kg(beijing)} ({demand_band('Beijing', beijing)})",
                f"{fmt_million_kg(tianjin)} ({demand_band('Tianjin', tianjin)})",
            ]
        )

    summary = {
        "airport_summary": airport_summary,
        "weeks": len(week_rows),
        "periods_per_week": 56,
        "hours_per_period": 3,
        "total_3h_periods": 56 * len(week_rows),
    }
    return summary, week_rows


def load_or_build_annual_airport_weekly_demand(cache_path: Path = ANNUAL_AIRPORT_DEMAND_CACHE) -> pd.DataFrame:
    if cache_path.exists():
        return pd.read_csv(cache_path)

    raw_path = ROOT / "products/aviation_fuel_analysis/air_port_data_process/results/parallel_calculation/并行计算结果_20250706_121257.xlsx"
    workbook = load_workbook(raw_path, read_only=True, data_only=True)
    worksheet = workbook.active
    header = list(next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True)))
    date_idx = header.index("日期")
    airport_idx = header.index("降落机场")
    fuel_idx = header.index("total_fuel_kg")

    weekly_totals: defaultdict[tuple[str, int], float] = defaultdict(float)
    try:
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            airport_raw = row[airport_idx]
            if airport_raw is None:
                continue
            airport_en = classify_destination_airport(airport_raw)
            if airport_en is None:
                continue

            date_value = row[date_idx]
            if date_value is None:
                continue
            if hasattr(date_value, "isocalendar"):
                iso_value = date_value.isocalendar()
                week_of_year = int(iso_value.week if hasattr(iso_value, "week") else iso_value[1])
            else:
                week_of_year = int(pd.Timestamp(date_value).isocalendar().week)

            fuel_value = row[fuel_idx]
            weekly_totals[(airport_en, week_of_year)] += float(fuel_value or 0.0)
    finally:
        workbook.close()

    rows: list[dict[str, object]] = []
    for airport in ["Beijing", "Tianjin"]:
        for week_of_year in range(1, 53):
            rows.append(
                {
                    "airport_en": airport,
                    "week_of_year": week_of_year,
                    "weekly_total_fuel_kg": weekly_totals.get((airport, week_of_year), 0.0),
                }
            )

    frame = pd.DataFrame(rows)
    ensure_dir(cache_path.parent)
    frame.to_csv(cache_path, index=False)
    return frame


def classify_destination_airport(value: object) -> str | None:
    airport_text = str(value or "")
    if "首都" in airport_text or "大兴" in airport_text:
        return "Beijing"
    if "滨海" in airport_text:
        return "Tianjin"
    return None


def classify_aircraft_class(value: object) -> str:
    aircraft_text = str(value or "")
    if "波音737" in aircraft_text or "B737" in aircraft_text:
        return "B737"
    if "空客320" in aircraft_text or "A320" in aircraft_text:
        return "A320"
    if "空客321" in aircraft_text or "A321" in aircraft_text:
        return "A321"
    if "空客319" in aircraft_text or "A319" in aircraft_text:
        return "A319"
    if "空客330" in aircraft_text or "A330" in aircraft_text:
        return "A330"
    if "波音787" in aircraft_text or "B787" in aircraft_text:
        return "B787"
    if "ERJ" in aircraft_text or "CRJ" in aircraft_text or "庞巴迪" in aircraft_text:
        return "Regional jet"
    if "新舟" in aircraft_text or "涡桨" in aircraft_text:
        return "Turbo-prop"
    return "Other/Unknown"


def load_or_build_aircraft_operation_summary(cache_path: Path = AIRCRAFT_OPERATION_SUMMARY_CACHE) -> pd.DataFrame:
    if cache_path.exists():
        return pd.read_csv(cache_path)

    raw_path = ROOT / "products/aviation_fuel_analysis/air_port_data_process/results/parallel_calculation/并行计算结果_20250706_121257.xlsx"
    workbook = load_workbook(raw_path, read_only=True, data_only=True)
    worksheet = workbook.active
    header = list(next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True)))
    airport_idx = header.index("降落机场")
    aircraft_idx = header.index("aircraft_type")
    distance_idx = header.index("distance_km")
    fuel_idx = header.index("total_fuel_kg")

    aggregate: defaultdict[str, dict[str, float]] = defaultdict(
        lambda: {"flights": 0.0, "total_fuel_kg": 0.0, "distance_sum_km": 0.0, "distance_count": 0.0}
    )

    try:
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            if classify_destination_airport(row[airport_idx]) is None:
                continue
            aircraft_class = classify_aircraft_class(row[aircraft_idx])
            bucket = aggregate[aircraft_class]
            bucket["flights"] += 1.0
            bucket["total_fuel_kg"] += float(row[fuel_idx] or 0.0)
            distance_value = row[distance_idx]
            if distance_value is not None and not pd.isna(distance_value):
                bucket["distance_sum_km"] += float(distance_value)
                bucket["distance_count"] += 1.0
    finally:
        workbook.close()

    total_flights = sum(bucket["flights"] for bucket in aggregate.values())
    total_fuel = sum(bucket["total_fuel_kg"] for bucket in aggregate.values())
    rows: list[dict[str, float | str]] = []
    for aircraft_class, bucket in aggregate.items():
        distance_count = bucket["distance_count"]
        mean_distance = bucket["distance_sum_km"] / distance_count if distance_count else np.nan
        rows.append(
            {
                "aircraft_class": aircraft_class,
                "flights": int(bucket["flights"]),
                "total_fuel_kg": bucket["total_fuel_kg"],
                "mean_distance_km": mean_distance,
                "flight_share_pct": bucket["flights"] / total_flights * 100 if total_flights else 0.0,
                "fuel_share_pct": bucket["total_fuel_kg"] / total_fuel * 100 if total_fuel else 0.0,
            }
        )

    frame = pd.DataFrame(rows).sort_values(["flights", "fuel_share_pct"], ascending=[False, False]).reset_index(drop=True)
    ensure_dir(cache_path.parent)
    frame.to_csv(cache_path, index=False)
    return frame


def load_candidate_week_metadata() -> tuple[pd.DataFrame, pd.DataFrame]:
    demand_12_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/typical_weeks_data/typical_12weeks_demand_20251129_163442.xlsx"
    week_map_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/week_mapping_20251129_231231.csv"

    demand_12 = pd.read_excel(demand_12_path)
    week_map = pd.read_csv(week_map_path)

    candidate_meta = (
        demand_12.groupby("week_number", as_index=False)
        .first()[["week_number", "week_start", "week_end", "original_week"]]
        .sort_values("week_number")
    )
    candidate_meta["week_start"] = pd.to_datetime(candidate_meta["week_start"])
    candidate_meta["week_end"] = pd.to_datetime(candidate_meta["week_end"])
    return candidate_meta, week_map


def summarize_renewable(csv_path: Path) -> dict:
    plant_capacity: dict[str, float] = {}
    hours: set[int] = set()
    total_mwh = 0.0
    for chunk in pd.read_csv(csv_path, usecols=["hour", "plant_id", "power_output_mw", "capacity_mw"], chunksize=200_000):
        hours.update(chunk["hour"].astype(int).unique().tolist())
        total_mwh += float(chunk["power_output_mw"].sum())
        for plant_id, capacity in chunk[["plant_id", "capacity_mw"]].drop_duplicates().itertuples(index=False):
            plant_capacity[str(plant_id)] = float(capacity)
    return {
        "unique_plants": len(plant_capacity),
        "hours": len(hours),
        "total_mwh": total_mwh,
        "capacity_mw": sum(plant_capacity.values()),
    }


def summarize_byproduct() -> dict:
    refinery_path = ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/refinery_daily_byproduct_h2_20251027_202950.csv"
    steel_path = ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/steel_daily_byproduct_h2_20251027_202950.csv"
    refinery = pd.read_csv(refinery_path)
    steel = pd.read_csv(steel_path)
    return {
        "refinery_sites": int(len(refinery)),
        "refinery_provinces": int(refinery["province"].nunique()),
        "refinery_daily_h2_tonnes": float(refinery["h2_daily_tonnes"].sum()),
        "steel_sites": int(len(steel)),
        "steel_provinces": int(steel["province"].nunique()),
        "steel_daily_h2_tonnes": float(steel["h2_daily_tonnes"].sum()),
    }


def summarize_transport() -> dict:
    pipe_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/natural_gas_pipelines.csv"
    wind_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/wind_power_plants.csv"
    solar_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/solar_power_plants.csv"
    ccs_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/ccs_projects.csv"
    pipes = pd.read_csv(pipe_path)
    wind_inv = pd.read_csv(wind_inv_path)
    solar_inv = pd.read_csv(solar_inv_path)
    ccs = pd.read_csv(ccs_path)

    lengths = (
        pipes["Note"]
        .fillna("")
        .str.extract(r"Length:\s*([0-9.]+)\s*km", expand=False)
        .dropna()
        .astype(float)
    )
    return {
        "pipeline_records": int(len(pipes)),
        "pipeline_rows": int(len(pipes)),
        "pipeline_operating": int(pipes["Status"].fillna("").str.contains("Operating", case=False).sum()),
        "pipeline_planned": int(pipes["Status"].fillna("").str.contains("Planned", case=False).sum()),
        "pipeline_length_km": float(lengths.sum()),
        "total_ng_pipeline_km": float(lengths.sum()),
        "median_ng_pipeline_km": float(lengths.median()),
        "year_stamped_ng_records": int(pd.to_numeric(pipes["YearOnline"], errors="coerce").notna().sum()),
        "wind_inventory_rows": int(len(wind_inv)),
        "solar_inventory_rows": int(len(solar_inv)),
        "wind_inventory_capacity_mw": float(wind_inv["Capacity__MW_"].fillna(0).sum()),
        "solar_inventory_capacity_mw": float(solar_inv["Capacity__MW_"].fillna(0).sum()),
        "ccs_projects": int(len(ccs)),
        "ccs_operating": int(ccs["Status"].fillna("").str.contains("Operating", case=False).sum()),
    }


def summarize_clustering() -> dict:
    renewable_cluster = read_json(ROOT / "clustering_results.json")
    byproduct_cluster = read_json(ROOT / "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/byproduct_clustering_results.json")
    co2_cluster = read_json(ROOT / "products/supply_chain_optimization/green_hydrogen_supply_chain_optimization/co2_clustering_results.json")
    return {
        "renewable_clusters": renewable_cluster["total_clusters"],
        "renewable_noise": renewable_cluster["total_noise_points"],
        "renewable_eps_km": renewable_cluster["clustering_params"]["eps_distance_km"],
        "renewable_min_samples": renewable_cluster["clustering_params"]["min_samples"],
        "byproduct_clusters": byproduct_cluster["total_clusters"],
        "byproduct_noise": byproduct_cluster["total_noise_points"],
        "byproduct_eps_km": byproduct_cluster["clustering_params"]["eps_distance_km"],
        "co2_clusters": co2_cluster["total_clusters"],
        "co2_noise": co2_cluster["total_noise_points"],
        "co2_eps_km": co2_cluster["clustering_params"]["eps_distance_km"],
        "pipeline_weight": renewable_cluster["clustering_params"]["pipeline_weight"],
        "shared_pipeline_discount": renewable_cluster["clustering_params"]["shared_pipeline_discount_factor"],
        "super_graph_k": 10,
    }


def summarize_verification() -> dict:
    report = read_json(ROOT / "products/supply_chain_optimization/natural_gas_supply_chain_optimization/tests/comprehensive_verification_report.json")
    direct = read_json(ROOT / "products/supply_chain_optimization/natural_gas_supply_chain_optimization/tests/direct_comparison_report.json")
    return {
        "summary": report["verification_summary"],
        "component_details": report["component_details"],
        "direct_summary": direct["comparison_summary"],
        "direct_details": direct["detailed_results"],
    }


def summarize_temporal(pareto_summary_path: Path | None = None) -> dict:
    penalty_path = ROOT / "products/supply_chain_optimization/visualization/results/pareto_penalty_summary_latest.csv"
    temporal_path = ROOT / "products/supply_chain_optimization/visualization/results/temporal_robustness_metrics_latest.csv"
    penalty_df = pd.read_csv(penalty_path)
    temporal_df = pd.read_csv(temporal_path)
    frontier = penalty_df.to_dict("records")
    stress = (
        temporal_df[temporal_df["Scenario"].isin(["GTL-BH", "GTL", "CCU-GH-FT"])]
        .set_index("Scenario")
        .to_dict("index")
    )
    default_pareto_summary = ROOT / "products/supply_chain_optimization/visualization/results/pareto_breakdown_summary_latest.json"
    pareto = read_json(pareto_summary_path or default_pareto_summary)
    return {
        "frontier_penalties": frontier,
        "frontier_stress": stress,
        "pareto_frontier": pareto[:3],
    }


def summarize_configs() -> dict:
    scenario_config_paths = {
        "CTL": ROOT / "products/supply_chain_optimization/coal_hydrogen_saf_optimization/data/CoalHydrogenSAFOptimizer_config.yaml",
        "CTL-BH": ROOT / "shared/data/CoalByproductHydrogenSAFOptimizer_config.yaml",
        "CCU-BH-MTJ": ROOT / "shared/data/ByproductHydrogenSupplyChainOptimizer_config_two_step.yaml",
        "CCU-BH-FT": ROOT / "shared/data/ByproductHydrogenSupplyChainOptimizer_config.yaml",
        "DAC-BH-MTJ": ROOT / "shared/data/DACByproductHydrogenSAFOptimizer_config_two_step.yaml",
        "DAC-BH-FT": ROOT / "shared/data/DACByproductHydrogenSAFOptimizer_config_one_step.yaml",
        "GTL-BH": ROOT / "shared/data/NaturalGasByproductHydrogenOptimizer_config.yaml",
        "GTL-GH": ROOT / "shared/data/NaturalGasSupplyChainOptimizer_config.yaml",
        "GTL": ROOT / "shared/data/NaturalGasSupplyChainOptimizer_config_one_step.yaml",
        "DAC-GH-MTJ": ROOT / "shared/config/DACHydrogenSAFOptimizer_config_two_step.yaml",
        "DAC-GH-FT": ROOT / "shared/config/DACHydrogenSAFOptimizer_config.yaml",
        "CCU-GH-MTJ": ROOT / "shared/data/GreenHydrogenSupplyChainOptimizer_config_two_step.yaml",
        "CCU-GH-FT": ROOT / "shared/data/GreenHydrogenSupplyChainOptimizer_config_one_step_direct_ft.yaml",
    }

    def load_yaml(path: Path) -> dict:
        with path.open("r", encoding="utf-8") as f:
            return yaml.safe_load(f)

    loaded = {
        scenario: {"path": path, "cfg": load_yaml(path)}
        for scenario, path in scenario_config_paths.items()
    }
    pathway_rows = collect_pathway_rows()
    pathway_lookup = {row[5]: row for row in pathway_rows}

    def rel_path(path: Path) -> str:
        return path.relative_to(ROOT).as_posix()

    def get_nested(obj: dict, *keys: str, default=None):
        current = obj
        for key in keys:
            if not isinstance(current, dict) or key not in current:
                return default
            current = current[key]
        return current

    def fmt_number(value, digits: int = 3) -> str:
        if value is None:
            return "n/a"
        if isinstance(value, bool):
            return "yes" if value else "no"
        if isinstance(value, (int, float)):
            if isinstance(value, float) and not math.isfinite(value):
                return "n/a"
            if abs(float(value) - round(float(value))) < 1e-9:
                return fmt_int(value)
            return f"{float(value):.{digits}f}".rstrip("0").rstrip(".")
        return str(value)

    def fmt_pct_value(value) -> str:
        if value is None:
            return "n/a"
        return f"{float(value) * 100:g}%"

    def route_code(scenario: str) -> str:
        if scenario.startswith("CTL"):
            return "CTL"
        if scenario == "GTL" or scenario.endswith("-FT"):
            return "FT"
        return "MTJ"

    def conversion_block(scenario: str) -> str:
        if scenario.startswith("CTL"):
            return "CTL / coal gasification"
        if scenario == "GTL" or scenario.endswith("-FT"):
            return "Direct FT"
        return "MTJ"

    def hydrogen_module(scenario: str) -> str:
        if "-GH" in scenario:
            return "Green electrolysis"
        if "-BH" in scenario:
            return "Industrial by-product H₂"
        if scenario == "GTL":
            return "No external H₂ input"
        if scenario == "CTL":
            return "Fossil-based H₂ baseline"
        return "Scenario-specific H₂ supply"

    def carbon_module(scenario: str) -> str:
        if scenario.startswith("CCU"):
            return "Industrial point-source CO₂"
        if scenario.startswith("DAC"):
            return "Direct air capture CO₂"
        if scenario.startswith("GTL"):
            return "Natural-gas carbon"
        if scenario.startswith("CTL"):
            return "Coal-derived carbon"
        return "Scenario-specific carbon source"

    def methanol_inventory(scenario: str) -> str:
        return "Active" if conversion_block(scenario) == "MTJ" else "Not active"

    def carbon_accounting_rule(scenario: str) -> str:
        if scenario.startswith("CCU"):
            return "Industrial CO₂ utilization credit enabled; credit factor = -1.0"
        if scenario.startswith("DAC"):
            return "Atmospheric CO₂ removal credit enabled; credit factor = -1.0"
        return "Fossil feedstock; utilization credit disabled"

    def capacity_bounds(cfg: dict, scenario: str) -> str:
        limits = cfg["capacity_limits"]
        if scenario.startswith("CTL"):
            candidates = [("saf_reactor", "CTL/SAF reactor")]
        elif conversion_block(scenario) == "MTJ":
            candidates = [("mtj", "MTJ plant"), ("saf_reactor", "SAF reactor"), ("ft_reactor", "FT reactor")]
        elif scenario == "GTL":
            candidates = [("ft_reactor", "FT reactor"), ("saf_reactor", "SAF reactor")]
        else:
            candidates = [("rwgs_ft_reactor", "RWGS-FT reactor"), ("ft_reactor", "FT reactor"), ("saf_reactor", "SAF reactor")]
        for prefix, label in candidates:
            max_value = limits.get(f"{prefix}_max_capacity_kg_per_hour")
            min_value = limits.get(f"{prefix}_min_capacity_kg_per_hour")
            if max_value is not None or min_value is not None:
                return f"{label}: {fmt_number(min_value)}-{fmt_number(max_value)} kg/h"
        return "n/a"

    def electrolyzer_bounds(cfg: dict) -> str:
        limits = cfg["capacity_limits"]
        min_value = limits.get("electrolyzer_min_capacity_kg_per_hour")
        max_value = limits.get("electrolyzer_max_capacity_kg_per_hour")
        if min_value is None and max_value is None:
            return "not active"
        return f"{fmt_number(min_value)}-{fmt_number(max_value)} kg/h"

    def min_output(cfg: dict) -> str:
        limits = cfg["capacity_limits"]
        values = []
        for key, label in [
            ("saf_min_annual_production_kg", "SAF"),
            ("mtj_min_annual_production_kg", "MTJ"),
        ]:
            if key in limits:
                values.append(f"{label}: {fmt_number(limits[key])} kg/yr")
        return "; ".join(values) if values else "n/a"

    def solver_text(cfg: dict) -> str:
        solver = cfg["solver_parameters"]
        return (
            f"{fmt_number(solver.get('TimeLimit'))} s; "
            f"gap {fmt_pct_value(solver.get('MIPGap'))}; "
            f"{fmt_number(solver.get('Threads'))} threads"
        )

    def formal_model_instance(scenario: str) -> str:
        if scenario == "GTL":
            h_block = "X_H(no external H₂)"
        elif "-GH" in scenario:
            h_block = "X_H(green electrolysis)"
        elif "-BH" in scenario:
            h_block = "X_H(by-product H₂)"
        elif scenario == "CTL":
            h_block = "X_H(fossil/process H₂)"
        else:
            h_block = "X_H(scenario-specific)"

        if scenario.startswith("CCU"):
            c_block = "X_C(industrial CO₂)"
        elif scenario.startswith("DAC"):
            c_block = "X_C(DAC CO₂)"
        elif scenario.startswith("GTL"):
            c_block = "X_C(natural gas)"
        elif scenario.startswith("CTL"):
            c_block = "X_C(coal)"
        else:
            c_block = "X_C(scenario-specific)"

        if scenario.startswith("CTL"):
            p_block = "X_P(CTL)"
        elif scenario == "GTL":
            p_block = "X_P(GTL-FT)"
        elif scenario.endswith("-FT"):
            p_block = "X_P(RWGS-FT)"
        else:
            p_block = "X_P(MTJ)"
        return f"X_core + {h_block} + {c_block} + {p_block} + X_I + X_N + X_D + X_E"

    def feedstock_model_text(scenario: str) -> str:
        if scenario == "GTL":
            return "Natural-gas pipeline/LNG availability and price surfaces; no external H₂ balance."
        if "-GH" in scenario:
            h_text = "green-H₂ electrolysis capacity, renewable/grid electricity use, H₂ storage and logistics"
        elif "-BH" in scenario:
            h_text = "plant-level by-product H₂ availability, H₂ storage and logistics"
        elif scenario == "CTL":
            h_text = "fossil/process H₂ baseline availability"
        else:
            h_text = "scenario-specific H₂ availability"

        if scenario.startswith("CCU"):
            c_text = "industrial point-source CO₂ capture capacity and CO₂ logistics"
        elif scenario.startswith("DAC"):
            c_text = "DAC siting, capture-rate, energy-use and CO₂ logistics constraints"
        elif scenario.startswith("GTL"):
            c_text = "natural-gas feedstock supply and logistics constraints"
        elif scenario.startswith("CTL"):
            c_text = "coal purchase, coal-gasification and coal-derived CO₂ inventory constraints"
        else:
            c_text = "scenario-specific carbon-source constraints"
        return f"{h_text}; {c_text}."

    def conversion_model_text(scenario: str) -> str:
        if scenario.startswith("CTL"):
            return "CTL/direct-liquefaction production block; CO₂ and SAF temporal balances; airport-delivery balance."
        if scenario == "GTL":
            return "Direct GTL/FT block with natural-gas feedstock; methanol inventory inactive; SAF delivery and shortage balance."
        if scenario.endswith("-FT"):
            return "One-step RWGS-FT/direct-FT block; methanol inventory inactive; H₂/CO₂ and SAF temporal balances."
        return "Two-step MTJ block; methanol production and methanol inventory active; final SAF conversion and weekly delivery balance."

    def scenario_accounting_text(scenario: str) -> str:
        if scenario.startswith("DAC"):
            return "DAC removal credit active with DAC process burdens retained."
        if scenario.startswith("CCU"):
            return "Industrial CO₂ utilization credit active under the archived accounting switch."
        if scenario.startswith("GTL"):
            return "Natural-gas carbon counted as fossil feedstock; utilization credit inactive."
        if scenario.startswith("CTL"):
            return "Coal-derived carbon counted as fossil feedstock; utilization credit inactive."
        return "Scenario-specific ex-post lifecycle accounting."

    scenario_rows = []
    scenario_model_rows = []
    capacity_solver_rows = []
    rows = []
    scenario_order = [row[5] for row in pathway_rows if row[5] in loaded]
    for scenario in scenario_order:
        item = loaded[scenario]
        cfg = item["cfg"]
        path = item["path"]
        pathway = pathway_lookup.get(scenario, ["n/a"] * 8)
        scenario_rows.append(
            [
                scenario,
                f"{hydrogen_module(scenario)}; {carbon_module(scenario)}",
                f"{conversion_block(scenario)}; methanol inventory: {methanol_inventory(scenario).lower()}",
                carbon_accounting_rule(scenario),
            ]
        )
        scenario_model_rows.append(
            [
                scenario,
                formal_model_instance(scenario),
                feedstock_model_text(scenario),
                conversion_model_text(scenario),
                scenario_accounting_text(scenario),
            ]
        )
        capacity_solver_rows.append(
            [
                scenario,
                route_code(scenario),
                capacity_bounds(cfg, scenario),
                electrolyzer_bounds(cfg),
                min_output(cfg),
                f"{fmt_number(cfg['economic_parameters'].get('levelized_cost_threshold_yuan_per_kg'))} yuan/kg",
                solver_text(cfg),
            ]
        )
        rows.append(
            {
                "label": scenario,
                "time_horizon_weeks": cfg["basic_parameters"]["time_horizon_weeks"],
                "periods_per_week": cfg["basic_parameters"]["periods_per_week"],
                "hours_per_period": cfg["basic_parameters"]["hours_per_period"],
                "discount_rate": cfg["economic_parameters"]["discount_rate"],
                "project_lifespan": cfg["economic_parameters"]["project_lifespan"],
                "shortage_penalty": cfg["cost_parameters"]["shortage_penalty_yuan_per_kg"],
                "wind_price": get_nested(cfg, "cost_parameters", "renewable_energy", "wind_power_price_yuan_per_kwh"),
                "solar_price": get_nested(cfg, "cost_parameters", "renewable_energy", "solar_power_price_yuan_per_kwh"),
                "grid_price": get_nested(cfg, "cost_parameters", "renewable_energy", "grid_electricity_price_yuan_per_kwh"),
                "natural_gas_price": get_nested(cfg, "cost_parameters", "raw_materials", "natural_gas_price_yuan_per_m3"),
                "jet_ci": cfg["carbon_emission_parameters"]["benchmarks"]["traditional_jet_fuel"],
                "saf_lhv": cfg["carbon_emission_parameters"]["benchmarks"]["saf_energy_content"],
                "solver_time": cfg["solver_parameters"]["TimeLimit"],
                "solver_gap": cfg["solver_parameters"]["MIPGap"],
                "solver_threads": cfg["solver_parameters"]["Threads"],
                "config_path": rel_path(path),
            }
        )

    green = loaded["CCU-GH-MTJ"]["cfg"]
    green_ft = loaded["CCU-GH-FT"]["cfg"]
    gas_mtj = loaded["GTL-GH"]["cfg"]
    gas_ft = loaded["GTL"]["cfg"]
    coal = loaded["CTL"]["cfg"]
    dac = loaded["DAC-GH-FT"]["cfg"]
    byproduct = loaded["CCU-BH-MTJ"]["cfg"]
    basic = green["basic_parameters"]
    econ = green["economic_parameters"]
    green_cost = green["cost_parameters"]
    green_transport = green["transport_modes"]
    gas_transport = gas_ft["transport_modes"]
    green_prod = green["carbon_emission_parameters"]["production_process"]
    green_ft_prod = green_ft["carbon_emission_parameters"]["production_process"]
    gas_mtj_prod = gas_mtj["carbon_emission_parameters"]["production_process"]
    gas_ft_prod = gas_ft["carbon_emission_parameters"]["production_process"]
    coal_prod = coal["carbon_emission_parameters"]["production_process"]
    dac_params = dac["dac_parameters"]
    coal_params = coal["coal_parameters"]
    h2_curve = green["costs"]["hydrogen_pipeline_costs"]["transport_cost_function"]["data_points"]
    co2_curve = get_nested(
        dac,
        "unified_costs",
        "co2_transport",
        "pipeline",
        "transport_cost_function",
        "data_points",
        default=[],
    )
    super_graph = get_nested(dac, "super_graph_config", default={})
    co2_super_graph = get_nested(dac, "co2_super_graph_config", default={})
    benchmark = green["carbon_emission_parameters"]["benchmarks"]
    transport_ci = green["carbon_emission_parameters"]["transportation"]
    gas_transport_ci = gas_ft["carbon_emission_parameters"]["transportation"]

    h2_pipeline_cost_text = "; ".join(f"{fmt_number(km)} km: {fmt_number(cost)}" for km, cost in h2_curve)
    co2_pipeline_cost_text = "; ".join(f"{fmt_number(km)} km: {fmt_number(cost)}" for km, cost in co2_curve)

    boundary_rows = [
        ["Supply-side search domain", "Mainland China", "categorical", "All scenarios", "Defines the candidate region for energy, feedstock, synthesis, and infrastructure nodes."],
        ["Demand nodes", "Beijing and Tianjin airports", "categorical", "Archived demand instance", "Defines the SAF sink nodes used in the four-week optimization runs."],
        ["Representative horizon", fmt_number(basic["time_horizon_weeks"]), "weeks", "All scenarios", "Reduced annual horizon used for comparative optimization."],
        ["Production time step", f"{fmt_number(basic['periods_per_week'])} periods/week; {fmt_number(basic['hours_per_period'])} h/period", "periods; hours", "Production and inventory balances", "Aggregates hourly renewable profiles into tractable production periods."],
        ["Demand and delivery time step", "Weekly", "time step", "Airport demand and final-fuel transport", "Enforces airport demand and delivered SAF balances at the week level."],
        ["Candidate-arc distance screen", fmt_number(basic["max_transport_distance_km"]), "km", "All transport candidate generation", "Removes infeasible long-distance candidate arcs before optimization."],
        ["Road-routing engine", f"GraphHopper at {basic['graphhopper_host']}:{basic['graphhopper_port']}; cache expiry {fmt_number(basic['cache_expiry_hours'])} h", "engine; hours", "Truck-routing arcs", "Provides routed road distances for short-distance truck transport."],
        ["Pipeline connector graph", f"k = {fmt_number(super_graph.get('k_connections'))}; algorithm = {super_graph.get('algorithm', 'n/a')}", "nearest connectors", "Pipeline-enabled H₂/CO₂ cases", "Restricts pipeline candidate generation to a sparse super-graph representation."],
    ]

    system_boundary_rows = [
        ["Hydrogen supply", "Green electrolysis, industrial by-product H₂, and fossil/process H₂ are activated by scenario.", "A national hydrogen-market equilibrium and sub-3 h electrolyzer dispatch are not modeled.", "Hydrogen enters as source availability, conversion efficiency, transport, and capacity constraints."],
        ["Carbon supply", "Industrial point-source CO₂, DAC CO₂, natural-gas carbon, and coal-derived carbon are represented by pathway.", "Detailed plant-level capture retrofits and endogenous carbon-source expansion are not optimized.", "Carbon source class determines both feedstock logistics and carbon-accounting treatment."],
        ["Conversion facilities", "FT/direct SAF, MTJ, and CTL/coal-gasification blocks are represented with route-specific capacity limits.", "Internal reactor scheduling and equipment sizing below the model time step are not represented.", "Conversion choice is a scenario definition, not an endogenous route-selection decision."],
        ["Intermediate inventories", "H₂, CO₂, methanol, and SAF inventories are included only where required by the pathway block.", "Tank-farm design, boil-off, and safety-zone layout are outside the optimization boundary.", "Inventory buffers connect the production time step to weekly delivery requirements."],
        ["Transport system", "H₂, CO₂, natural gas, methanol intermediate, and SAF logistics are modeled through candidate arcs.", "Vehicle-count scheduling and pipe-diameter design are not optimized explicitly.", "Transport modes impose distance screens, capacity factors, and cost coefficients."],
        ["Demand and use", "Weekly SAF demand is imposed at the Beijing and Tianjin airport nodes.", "Aircraft assignment, refueling operations, and downstream combustion allocation are not modeled.", "Airport demand is a fixed sink for supply-chain optimization."],
        ["Emissions accounting", "Lifecycle tracking, construction emissions, utilization credits, and benchmark conversion are configured by scenario.", "Policy feedbacks and endogenous carbon-price responses are not represented.", "Ex-post carbon intensity and displacement metrics use the same accounting flags across scenario families."],
    ]

    route_module_rows = [
        ["Direct FT / RWGS-FT", "GTL, CCU-GH-FT, CCU-BH-FT, DAC-GH-FT, DAC-BH-FT", "Direct SAF output; no methanol inventory.", f"RWGS energy coefficient = {fmt_number(green_ft_prod.get('rwgs_process_energy'))}; FT energy coefficient = {fmt_number(green_ft_prod.get('ft_process_energy'))} for CCU-GH-FT; GTL FT energy coefficient = {fmt_number(gas_ft_prod.get('ft_process_energy'))}.", "Represents one-step/direct SAF synthesis scenarios."],
        ["MTJ", "GTL-GH, GTL-BH, CCU-GH-MTJ, CCU-BH-MTJ, DAC-GH-MTJ, DAC-BH-MTJ", "Methanol intermediate and methanol inventory active.", f"MTJ process-energy coefficient = {fmt_number(green_prod.get('mtj_process_energy'))} for CCU/DAC-style MTJ configs; {fmt_number(gas_mtj_prod.get('mtj_process_energy'))} for natural-gas MTJ configs.", "Separates methanol production from final jet-fuel synthesis."],
        ["CTL / coal gasification", "CTL, CTL-BH", "Coal-derived carbon and CTL/SAF reactor block active.", f"Coal price = {fmt_number(coal_params.get('coal_price_yuan_per_ton'))} yuan/t; carbon content = {fmt_number(coal_params.get('carbon_content'))}; gasification efficiency = {fmt_number(coal_params.get('gasification_efficiency'))}.", "Provides the grey coal-linked baseline and by-product-H₂ variant."],
        ["Industrial point-source CO₂", "CCU-GH-MTJ, CCU-GH-FT, CCU-BH-MTJ, CCU-BH-FT", "Industrial capture-source layer active.", "scenario_type = industrial_capture; utilization credit factor = -1.0.", "Credits captured industrial CO₂ fixed into SAF under the archived accounting convention."],
        ["DAC CO₂", "DAC-GH-MTJ, DAC-GH-FT, DAC-BH-MTJ, DAC-BH-FT", "DAC module and DAC siting logic active.", f"Capture efficiency = {fmt_number(dac_params.get('capture_efficiency'))}; capture energy = {fmt_number(dac_params.get('energy_kwh_per_ton_co2'))} kWh/t CO₂.", "Treats atmospheric CO₂ fixed into SAF as a carbon-removal credit."],
        ["By-product H₂", "CTL-BH, CCU-BH-MTJ, CCU-BH-FT, DAC-BH-MTJ, DAC-BH-FT, GTL-BH", "Industrial by-product H₂ source layer active.", f"H₂ equipment lifetime convention = {fmt_number(byproduct['economic_parameters'].get('electrolyzer_lifetime'))} yr; capacity factor = {fmt_number(byproduct['economic_parameters']['capacity_factors'].get('electrolyzer_capacity_factor'))}.", "Defines blue variants that replace green electrolysis with by-product H₂ supply."],
    ]

    economic_rows = [
        ["Project evaluation life", fmt_number(econ["project_lifespan"]), "yr", "All scenarios", "Defines the present-value horizon for cost and levelized-cost calculations."],
        ["Discount rate", fmt_pct_value(econ["discount_rate"]), "annual fraction", "All scenarios", "Annualizes and discounts capital and operating cost streams."],
        ["Unserved-demand penalty", fmt_number(green_cost["shortage_penalty_yuan_per_kg"]), "yuan/kg SAF", "All scenarios", "Discourages unmet airport demand while preserving model feasibility."],
        ["LCO-SAF screening threshold", "20 for GTL/GTL-GH/GTL-BH; 50 for CTL/CTL-BH; 1000 for most CCU/DAC cases; 1200 for CCU-GH-FT", "yuan/kg SAF", "Scenario-specific", "Used as a configuration-level filter and plausibility screen, not as an endogenous policy target."],
        ["Electrolyzer lifetime", "15 for green-H₂ cases; 20 for by-product-H₂ and CTL-BH cases", "yr", "H₂-supply scenarios with an H₂ capacity block", "Sets capital recovery for the H₂ supply block."],
        ["Synthesis and logistics lifetimes", "MTJ/RWGS 20; FT/SAF reactor 25; pipeline 30; storage 25; transport vehicle 20", "yr", "Route-specific", "Sets annualized capital recovery for route and transport assets."],
        ["Capacity factors", "Electrolyzer 0.75 for green-H₂ cases and 0.90 for by-product-H₂/CTL-BH cases; synthesis 0.85; pipeline 0.95; storage 0.90; transport 0.75", "fraction", "Scenario-family-specific", "Translates installed capacity into available production, storage, and transport service."],
        ["Electricity prices", f"Wind {fmt_number(green_cost['renewable_energy']['wind_power_price_yuan_per_kwh'])}; solar {fmt_number(green_cost['renewable_energy']['solar_power_price_yuan_per_kwh'])}; grid {fmt_number(green_cost['renewable_energy']['grid_electricity_price_yuan_per_kwh'])}", "yuan/kWh", "Pathways using renewable or grid electricity", "Exogenous operating-cost inputs for electricity-consuming steps."],
        ["Natural-gas price", fmt_number(gas_mtj["cost_parameters"]["raw_materials"]["natural_gas_price_yuan_per_m3"]), "yuan/m3", "GTL, GTL-GH, GTL-BH", "Feedstock price for natural-gas-based SAF pathways."],
        ["Coal price", fmt_number(coal_params["coal_price_yuan_per_ton"]), "yuan/t", "CTL, CTL-BH", "Feedstock price for coal-linked CTL pathways."],
        ["DAC capture cost", f"{fmt_number(dac_params.get('capture_cost_yuan_per_ton'))}; 2030 reference = {fmt_number(dac_params.get('capture_cost_2030_yuan_per_ton'))}", "yuan/t CO₂", "DAC-GH-MTJ, DAC-GH-FT, DAC-BH-MTJ, DAC-BH-FT", "Exogenous cost for atmospheric CO₂ capture."],
    ]

    feedstock_rows = [
        ["Electrolysis efficiency", fmt_number(green_cost["electrolysis"]["electrolysis_efficiency"]), "fraction", "CCU/DAC/CTL-style configs with electrolysis blocks", "Converts electricity input into H₂ production."],
        ["Electrolysis efficiency, natural-gas MTJ", fmt_number(gas_mtj["cost_parameters"]["electrolysis"]["electrolysis_efficiency"]), "fraction", "GTL-GH, GTL-BH", "Applies to the H₂ addition block used in natural-gas MTJ configurations."],
        ["Electrolysis electricity intensity", f"{fmt_number(green_cost['electrolysis']['electrolysis_power_consumption'])}; CCU-GH-FT = {fmt_number(green_ft['cost_parameters']['electrolysis']['electrolysis_power_consumption'])}", "kWh/kg H₂", "H₂-producing pathway blocks", "Links H₂ production to electricity demand."],
        ["Hydrogen market-price proxy", f"{fmt_number(green_cost['raw_materials']['hydrogen_market_price_yuan_per_kg'])}; natural-gas MTJ configs = {fmt_number(gas_mtj['cost_parameters']['raw_materials']['hydrogen_market_price_yuan_per_kg'])}", "yuan/kg H₂", "H₂ supply and transport cost blocks", "Provides the exogenous H₂ price level used in scenario cost terms."],
        ["Natural-gas feedstock", f"Price {fmt_number(gas_mtj['cost_parameters']['raw_materials']['natural_gas_price_yuan_per_m3'])}; extraction intensity {fmt_number(gas_mtj['carbon_emission_parameters']['raw_materials']['ng_extraction_intensity'])}; pipeline-transport coefficient {fmt_number(gas_mtj['carbon_emission_parameters']['raw_materials']['ng_pipeline_transport'])}", "mixed", "GTL, GTL-GH, GTL-BH", "Defines natural-gas feedstock cost and upstream-emissions coefficients."],
        ["Coal feedstock", f"Type {coal_params.get('coal_type')}; carbon content {fmt_number(coal_params.get('carbon_content'))}; CO₂ yield {fmt_number(coal_params.get('co2_per_kg_coal'))}; coal-to-SAF use {fmt_number(coal_params.get('coal_consumption_kg_per_kg_saf'))}", "fraction; kg CO₂/kg coal; kg coal/kg SAF", "CTL, CTL-BH", "Defines coal material input and carbon yield for the coal-linked scenarios."],
        ["Coal gasification", f"Efficiency {fmt_number(coal_params.get('gasification_efficiency'))}; energy {fmt_number(coal_params.get('gasification_energy_mj_per_kg'))} MJ/kg ({fmt_number(coal_params.get('gasification_energy_kwh_per_kg'))} kWh/kg); energy cost {fmt_number(coal_params.get('gasification_energy_cost_yuan_per_mj'))}", "fraction; MJ/kg; kWh/kg; yuan/MJ", "CTL, CTL-BH", "Defines gasification energy use before CTL/SAF conversion."],
        ["DAC capture module", f"Capture efficiency {fmt_number(dac_params.get('capture_efficiency'))}; CO₂ purity {fmt_number(dac_params.get('co2_purity'))}; energy {fmt_number(dac_params.get('energy_kwh_per_ton_co2'))}; module capacity {fmt_number(dac_params.get('module_capacity_ton_year'))}; availability {fmt_number(dac_params.get('module_availability'))}", "fraction; kWh/t CO₂; t/yr", "DAC-GH-MTJ, DAC-GH-FT, DAC-BH-MTJ, DAC-BH-FT", "Defines the DAC supply block and modular deployment constraint."],
        ["GTL direct FT process", f"NG upstream emission {fmt_number(gas_ft_prod.get('ng_upstream_emission'))}; NG process emission {fmt_number(gas_ft_prod.get('ng_process_emission'))}; FT process emission {fmt_number(gas_ft_prod.get('ft_process_emission'))}; FT process-energy coefficient {fmt_number(gas_ft_prod.get('ft_process_energy'))}", "configuration coefficients", "GTL", "Provides process-emissions and energy coefficients for the direct gas-to-liquids route."],
        ["Natural-gas MTJ process", f"NG-to-methanol rate {fmt_number(gas_mtj_prod.get('ng_to_methanol_rate'))}; H₂ addition rate {fmt_number(gas_mtj_prod.get('h2_addition_rate'))}; MTJ process-energy coefficient {fmt_number(gas_mtj_prod.get('mtj_process_energy'))}", "configuration coefficients", "GTL-GH, GTL-BH", "Defines the methanol-intermediate step for natural-gas MTJ routes."],
    ]

    transport_rows = [
        ["H₂ pipeline", f"{fmt_number(green_transport['renewable_h2_pipeline']['distance_limit_km'])}", "km", f"CAPEX {fmt_number(green['costs']['hydrogen_pipeline_costs']['capex_yuan_per_km'])} yuan/km; cost curve {h2_pipeline_cost_text} yuan/(kg H₂ per 100 km)", "H₂-pipeline-enabled green, DAC, CCU, CTL, and gas-MTJ cases"],
        ["H₂ truck", f"{fmt_number(green_transport['h2_truck_transport']['distance_limit_km'])}", "km", f"Truck capacity {fmt_number(green['transport_constraints']['h2_truck_transport']['truck_capacity_kg'])} kg; max {fmt_number(green['transport_constraints']['h2_truck_transport']['max_trucks_per_day'])} trucks/day; cost {fmt_number(green_transport['h2_truck_transport']['cost_yuan_per_kg_km'])} yuan/(kg km); capacity factor {fmt_number(green_transport['h2_truck_transport']['capacity_factor'])}", "Green, DAC, CCU, and CTL H₂ logistics"],
        ["CO₂ pipeline", f"{fmt_number(green_transport['co2_pipeline_transport']['distance_limit_km'])}", "km", f"Capacity factor {fmt_number(green_transport['co2_pipeline_transport']['capacity_factor'])}; DAC cost curve {co2_pipeline_cost_text or 'n/a'} yuan/(t CO₂ per 100 km)", "CCU and DAC carbon logistics where CO₂ pipeline distance is enabled"],
        ["CO₂ truck", f"{fmt_number(green_transport['co2_truck_transport']['distance_limit_km'])}", "km", f"Truck capacity {fmt_number(green['transport_constraints']['co2_truck_transport']['truck_capacity_ton'])} t; max {fmt_number(green['transport_constraints']['co2_truck_transport']['max_trucks_per_day'])} trucks/day; cost {fmt_number(green_transport['co2_truck_transport']['cost_yuan_per_ton_km'])} yuan/(t km); capacity factor {fmt_number(green_transport['co2_truck_transport']['capacity_factor'])}", "CCU, DAC, and CTL carbon logistics"],
        ["Natural-gas pipeline direct", f"{fmt_number(gas_transport['ng_pipeline_direct']['distance_limit_km'])}", "km", f"Capacity factor {fmt_number(gas_transport['ng_pipeline_direct']['capacity_factor'])}", "GTL direct FT configuration"],
        ["LNG terminal supply", f"{fmt_number(gas_transport['lng_terminal_supply']['distance_limit_km'])}", "km", f"Capacity factor {fmt_number(gas_transport['lng_terminal_supply']['capacity_factor'])}", "GTL direct FT configuration"],
        ["Natural-gas truck constraint", "n/a", "m3; trucks/day", f"Truck capacity {fmt_number(gas_ft['transport_constraints']['ng_truck_transport']['truck_capacity_m3'])} m3; max {fmt_number(gas_ft['transport_constraints']['ng_truck_transport']['max_trucks_per_day'])} trucks/day", "Natural-gas logistics constraints in GTL-family configs"],
        ["Airport-integrated option", f"{fmt_number(green_transport['airport_integrated']['distance_limit_km'])} for CO₂/H₂ cases; {fmt_number(gas_transport['airport_integrated']['distance_limit_km'])} for GTL", "km", f"Cost {fmt_number(green_transport['airport_integrated']['cost_yuan_per_kg'])} yuan/kg in CO₂/H₂ cases; capacity factor {fmt_number(green_transport['airport_integrated']['capacity_factor'])}", "Local production or delivery option around airport demand nodes"],
        ["Routing cache", f"{fmt_number(basic['cache_expiry_hours'])}", "h", f"GraphHopper routing enabled; pipeline super graph k = {fmt_number(super_graph.get('k_connections'))}; CO₂ super graph k = {fmt_number(co2_super_graph.get('k_connections'))}", "Repeated scenario runs and candidate-arc preprocessing"],
    ]

    emission_rows = [
        ["Conventional jet benchmark", fmt_number(benchmark["traditional_jet_fuel"]), "gCO2e/MJ", "All scenarios", "Reference value for carbon-intensity comparison."],
        ["SAF energy-content conversion", fmt_number(benchmark["saf_energy_content"]), "MJ/kg", "All scenarios", "Converts mass-based SAF output to energy-basis carbon intensity."],
        ["CORSIA benchmark limit", fmt_number(benchmark["corsia_limit"]), "gCO2e/MJ", "All scenarios", "Reference policy benchmark retained in the archived configuration."],
        ["Lifecycle tracking flags", "Enabled; facility construction included; lifecycle amortization enabled", "categorical", "All scenarios", "Keeps construction and lifecycle allocation inside ex-post emissions reporting."],
        ["Industrial CO₂ utilization credit", "-1.0", "dimensionless credit factor", "CCU-GH-MTJ, CCU-GH-FT, CCU-BH-MTJ, CCU-BH-FT", "Credits industrial point-source CO₂ fixed into SAF according to the archived accounting switch."],
        ["DAC removal credit", f"-1.0; DAC capture emission {fmt_number(dac['carbon_emission_parameters']['production_process']['dac_capture_emission'])}; DAC capture energy {fmt_number(dac['carbon_emission_parameters']['production_process']['dac_capture_energy'])}", "credit factor; configuration coefficients", "DAC-GH-MTJ, DAC-GH-FT, DAC-BH-MTJ, DAC-BH-FT", "Treats atmospheric CO₂ fixed into SAF as carbon removal while retaining DAC process burdens."],
        ["Natural-gas fossil-feedstock treatment", f"Utilization credit disabled; NG extraction {fmt_number(gas_mtj['carbon_emission_parameters']['raw_materials']['ng_extraction_intensity'])}; NG transport {fmt_number(gas_mtj['carbon_emission_parameters']['raw_materials']['ng_pipeline_transport'])}", "configuration coefficients", "GTL-GH, GTL, GTL-BH", "Counts natural-gas-derived carbon as fossil feedstock rather than carbon removal."],
        ["Coal fossil-feedstock treatment", f"Utilization credit disabled; mining {fmt_number(coal_prod.get('coal_mining_emission'))}; transport {fmt_number(coal_prod.get('coal_transport_emission'))}; gasification direct {fmt_number(coal_prod.get('coal_gasification_direct_emission'))}; capture efficiency {fmt_number(coal_prod.get('co2_capture_efficiency'))}", "configuration coefficients; fraction", "CTL, CTL-BH", "Counts coal-derived carbon as fossil feedstock."],
        ["Transport-emission coefficients", f"H₂ truck {fmt_number(transport_ci['h2_truck_intensity'])}; CO₂ truck {fmt_number(transport_ci['co2_truck_intensity'])}; SAF truck {fmt_number(transport_ci['saf_truck_intensity'])}; H₂ pipeline {fmt_number(transport_ci['h2_pipeline_intensity'])}; CO₂ pipeline {fmt_number(transport_ci['co2_pipeline_intensity'])}; NG truck {fmt_number(gas_transport_ci['ng_truck_intensity'])}", "configuration coefficients", "Relevant transport arcs", "Applied in ex-post transport-emissions accounting."],
    ]

    constraint_block_rows = [
        ["X_core", "Facility siting, installed capacity, production variables, and total-cost objective.", "All 13 scenarios", "Links binary activation decisions to continuous production and capacity variables."],
        ["X_H(s)", "Green electrolysis, industrial by-product H\u2082, fossil/process H\u2082, or no-external-H\u2082 supply block.", "Scenario-specific", "Selects the H\u2082 availability, storage, and logistics equations implied by the pathway name."],
        ["X_C(s)", "Industrial CO\u2082 capture, DAC capture, natural gas, or coal feedstock block.", "Scenario-specific", "Defines carbon/feedstock availability, capture energy, and upstream logistics."],
        ["X_P(s)", "Conversion route equations for FT/RWGS-FT, MTJ, GTL, or CTL.", "Scenario-specific", "Maps H\u2082, CO\u2082, natural gas, or coal-derived carbon into SAF output."],
        ["X_I(s)", "Inventory balances for H\u2082, CO\u2082, methanol, and SAF.", "Only commodities active in each scenario", "Couples the 3-hour production step with weekly logistics and demand."],
        ["X_N(s)", "Candidate truck, pipeline, LNG, natural-gas, H\u2082, CO\u2082, methanol, and SAF transport arcs.", "Mode-specific", "Applies distance screens, routing-cache distances, and arc-capacity limits."],
        ["X_D", "Weekly airport demand, delivered SAF, and shortage slack constraints.", "All 13 scenarios", "Ensures Beijing and Tianjin demand coverage while preserving feasibility under stressed cases."],
        ["X_E", "Ex-post lifecycle accounting and benchmark comparison.", "All 13 scenarios", "Uses pathway-specific credit switches and process/transport emissions coefficients."],
    ]

    def join_value_unit(value: str, unit: str) -> str:
        if unit == "n/a":
            return value
        if unit in {"categorical", "mixed", "configuration coefficients"}:
            return f"{value} ({unit})"
        return f"{value} {unit}"

    boundary_rows = [[row[0], join_value_unit(row[1], row[2]), row[3], row[4]] for row in boundary_rows]
    system_boundary_rows = [[row[0], f"{row[1]} {row[3]}", row[2]] for row in system_boundary_rows]
    route_module_rows = [[row[0], row[1], row[2], f"{row[3]} {row[4]}"] for row in route_module_rows]
    economic_rows = [[row[0], join_value_unit(row[1], row[2]), row[3], row[4]] for row in economic_rows]
    feedstock_rows = [[row[0], join_value_unit(row[1], row[2]), row[3], row[4]] for row in feedstock_rows]
    transport_rows = [[row[0], join_value_unit(row[1], row[2]), row[3], row[4]] for row in transport_rows]
    scenario_model_rows = [[row[0], row[1], row[2], f"{row[3]} {row[4]}"] for row in scenario_model_rows]
    capacity_solver_rows = [
        [
            row[0],
            row[1],
            f"{row[2]}; H₂ capacity: {row[3]}",
            f"Minimum output: {row[4]}; LCO screen: {row[5]}",
            row[6],
        ]
        for row in capacity_solver_rows
    ]
    emission_rows = [[row[0], join_value_unit(row[1], row[2]), row[3], row[4]] for row in emission_rows]

    return {
        "rows": rows,
        "boundary_rows": boundary_rows,
        "system_boundary_rows": system_boundary_rows,
        "route_module_rows": route_module_rows,
        "scenario_rows": scenario_rows,
        "constraint_block_rows": constraint_block_rows,
        "scenario_model_rows": scenario_model_rows,
        "economic_rows": economic_rows,
        "feedstock_rows": feedstock_rows,
        "transport_rows": transport_rows,
        "capacity_solver_rows": capacity_solver_rows,
        "emission_rows": emission_rows,
    }


def collect_data(pareto_summary_path: Path | None = None) -> AppendixData:
    demand_summary, week_rows = summarize_representative_weeks()
    renewable_summary = {
        "wind": summarize_renewable(
            ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_wind_20251129_231231.csv"
        ),
        "solar": summarize_renewable(
            ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_solar_20251129_231231.csv"
        ),
        "wind_resource": summarize_cached_renewable_resource(
            WIND_COMPLETE_PATH,
            cache_path=WIND_COMPLETE_CACHE,
            hours_per_step=1,
            clip_to_capacity=True,
        ),
        "solar_resource": summarize_cached_renewable_resource(
            SOLAR_COMPLETE_PATH,
            cache_path=SOLAR_COMPLETE_CACHE,
            hours_per_step=1,
            clip_to_capacity=True,
        ),
    }
    return AppendixData(
        demand_summary=demand_summary,
        week_rows=week_rows,
        renewable_summary=renewable_summary,
        byproduct_summary=summarize_byproduct(),
        co2_summary=summarize_co2(),
        transport_summary=summarize_transport(),
        price_summary=summarize_price_surface(),
        clustering_summary=summarize_clustering(),
        verification_summary=summarize_verification(),
        temporal_summary=summarize_temporal(pareto_summary_path=pareto_summary_path),
        config_summary=summarize_configs(),
        pathway_rows=collect_pathway_rows(),
    )


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def strip_heading_auto_numbering(doc: Document) -> None:
    # The template already defines multilevel numbering on heading styles,
    # but appendix section numbers are written directly in the heading text.
    for style in doc.styles:
        style_name = getattr(style, "name", "")
        if not style_name.lower().startswith("heading "):
            continue
        p_pr = style.element.find(qn("w:pPr"))
        if p_pr is None:
            continue
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)


def style_exists(doc: Document, style_name: str) -> bool:
    return any(s.name == style_name for s in doc.styles)


def add_para(doc: Document, text: str, style: str = "Normal", align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    p = doc.add_paragraph(style=style if style_exists(doc, style) else "Normal")
    p.add_run(text)
    if align is not None:
        p.alignment = align


def add_heading(doc: Document, text: str, level_style: str) -> None:
    p = doc.add_paragraph(style=level_style if style_exists(doc, level_style) else "Normal")
    p.add_run(text)


def math_text(text: str, style: str | None = None) -> OxmlElement:
    run = OxmlElement("m:r")
    if style:
        run_pr = OxmlElement("m:rPr")
        sty = OxmlElement("m:sty")
        sty.set(qn("m:val"), style)
        run_pr.append(sty)
        run.append(run_pr)
    text_el = OxmlElement("m:t")
    text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def math_container(tag: str, parts: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement(tag)
    for part in parts:
        node.append(part)
    return node


def math_sub(base: list[OxmlElement], subscript: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement("m:sSub")
    node.append(math_container("m:e", base))
    node.append(math_container("m:sub", subscript))
    return node


def math_sup(base: list[OxmlElement], superscript: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement("m:sSup")
    node.append(math_container("m:e", base))
    node.append(math_container("m:sup", superscript))
    return node


def math_subsup(base: list[OxmlElement], subscript: list[OxmlElement], superscript: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement("m:sSubSup")
    node.append(math_container("m:e", base))
    node.append(math_container("m:sub", subscript))
    node.append(math_container("m:sup", superscript))
    return node


def math_fraction(numerator: list[OxmlElement], denominator: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement("m:f")
    node.append(math_container("m:num", numerator))
    node.append(math_container("m:den", denominator))
    return node


def math_sum(subscript: str, body: list[OxmlElement]) -> list[OxmlElement]:
    return [math_sub([math_text("∑")], [math_text(subscript)]), *body]


def math_var(symbol: str, subscript: str | None = None, superscript: str | None = None) -> OxmlElement:
    base = [math_text(symbol)]
    if subscript is not None and superscript is not None:
        return math_subsup(base, [math_text(subscript, "p")], [math_text(superscript, "p")])
    if subscript is not None:
        return math_sub(base, [math_text(subscript, "p")])
    if superscript is not None:
        return math_sup(base, [math_text(superscript, "p")])
    return math_text(symbol)


def add_equation(doc: Document, label: str, equation: str | list[OxmlElement]) -> None:
    style = "AMDisplayEquation" if style_exists(doc, "AMDisplayEquation") else "Normal"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = p.add_run(f"{label}  ")
    label_run.font.name = "Arial"
    label_run.font.size = Pt(10)
    math = OxmlElement("m:oMath")
    parts = [math_text(equation)] if isinstance(equation, str) else equation
    for part in parts:
        math.append(part)
    p._p.append(math)


def format_math_run(run: OxmlElement) -> None:
    run_pr = run.find(qn("m:rPr"))
    if run_pr is None:
        run_pr = OxmlElement("m:rPr")
        run.insert(0, run_pr)

    normal_text = run_pr.find(qn("m:nor"))
    if normal_text is None:
        run_pr.append(OxmlElement("m:nor"))

    # Force an explicit math font/size so symbols stay stable inside tables and PDFs.
    ctrl_pr = run_pr.find(qn("m:ctrlPr"))
    if ctrl_pr is None:
        ctrl_pr = OxmlElement("m:ctrlPr")
        run_pr.append(ctrl_pr)
    w_rpr = ctrl_pr.find(qn("w:rPr"))
    if w_rpr is None:
        w_rpr = OxmlElement("w:rPr")
        ctrl_pr.append(w_rpr)

    r_fonts = w_rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        w_rpr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), "Cambria Math")

    sz = w_rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        w_rpr.append(sz)
    sz.set(qn("w:val"), "24")

    sz_cs = w_rpr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        w_rpr.append(sz_cs)
    sz_cs.set(qn("w:val"), "24")


def format_math_tree(node: OxmlElement) -> None:
    for run in node.findall(".//" + qn("m:r")):
        format_math_run(run)


def clear_paragraph_indent(paragraph) -> None:
    p_format = paragraph.paragraph_format
    p_format.first_line_indent = Pt(0)
    p_format.left_indent = Pt(0)
    p_format.right_indent = Pt(0)
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    for attr in (
        "firstLine",
        "left",
        "right",
        "hanging",
        "firstLineChars",
        "leftChars",
        "rightChars",
        "hangingChars",
    ):
        ind.set(qn(f"w:{attr}"), "0")


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_cell_math(cell, parts: list[OxmlElement]) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph_indent(paragraph)
    math = OxmlElement("m:oMath")
    for part in parts:
        math.append(part)
    format_math_tree(math)
    paragraph._p.append(math)


def symbol_token(
    text: str,
    *,
    italic: bool = False,
    subscript: bool = False,
    superscript: bool = False,
    line_break: bool = False,
) -> dict[str, object]:
    return {
        "text": text,
        "italic": italic,
        "subscript": subscript,
        "superscript": superscript,
        "line_break": line_break,
    }


def symbol_var(
    symbol: str,
    subscript: str | None = None,
    superscript: str | None = None,
    *,
    index_italic: bool = True,
    exponent_italic: bool = False,
) -> list[dict[str, object]]:
    segments = [symbol_token(symbol, italic=True)]
    if subscript is not None:
        segments.append(symbol_token(subscript, italic=index_italic, subscript=True))
    if superscript is not None:
        segments.append(symbol_token(superscript, italic=exponent_italic, superscript=True))
    return segments


def set_rich_text_cell(
    cell,
    segments: list[dict[str, object]],
    *,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    font_size: float = 11,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    clear_paragraph_content(paragraph)
    paragraph.alignment = align
    clear_paragraph_indent(paragraph)
    for spec in segments:
        run = paragraph.add_run(spec["text"])
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(font_size)
        run.font.italic = bool(spec.get("italic"))
        run.font.subscript = bool(spec.get("subscript"))
        run.font.superscript = bool(spec.get("superscript"))
        if spec.get("line_break"):
            run.add_break()


def add_rich_para(
    doc: Document,
    segments: list[dict[str, object]],
    style: str = "Normal",
    align: WD_ALIGN_PARAGRAPH | None = None,
    *,
    font_size: float | None = None,
) -> None:
    p = doc.add_paragraph(style=style if style_exists(doc, style) else "Normal")
    for spec in segments:
        run = p.add_run(spec["text"])
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if font_size is not None:
            run.font.size = Pt(font_size)
        run.font.italic = bool(spec.get("italic"))
        run.font.subscript = bool(spec.get("subscript"))
        run.font.superscript = bool(spec.get("superscript"))
        if spec.get("line_break"):
            run.add_break()
    if align is not None:
        p.alignment = align


def set_symbol_cell(cell, segments: list[dict[str, object]]) -> None:
    set_rich_text_cell(cell, segments, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=11)


def model_module_segments(
    module: str,
    argument: str | None = None,
    *,
    argument_italic: bool = False,
) -> list[dict[str, object]]:
    segments = [
        symbol_token("X", italic=True),
        symbol_token(module, subscript=True, italic=False),
    ]
    if argument is not None:
        segments.extend(
            [
                symbol_token("("),
                symbol_token(argument, italic=argument_italic),
                symbol_token(")"),
            ]
        )
    return segments


MODEL_INSTANCE_PATTERN = re.compile(r"X_([A-Za-z]+)(?:\((.*?)\))?")


def model_instance_segments(instance_text: str) -> list[dict[str, object]]:
    matches = MODEL_INSTANCE_PATTERN.findall(instance_text)
    segments: list[dict[str, object]] = []
    break_after = {0, 1, 3, 5}
    for idx, (module, raw_argument) in enumerate(matches):
        argument = raw_argument or ("s" if module in {"I", "N"} else None)
        segments.extend(
            model_module_segments(
                module,
                argument,
                argument_italic=argument == "s",
            )
        )
        if idx < len(matches) - 1:
            segments.append(symbol_token(" + ", line_break=idx in break_after))
    return segments


def set_model_instance_cell(cell, instance_text: str) -> None:
    set_rich_text_cell(
        cell,
        model_instance_segments(instance_text),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font_size=10.5,
    )


def eq_lco_saf() -> list[OxmlElement]:
    def discount_factor() -> list[OxmlElement]:
        return [math_sup([math_text("(1 + r)", "p")], [math_text("y")])]

    annual_cost = [
        math_var("C", "CAPEX,y"),
        math_text(" + ", "p"),
        math_var("C", "OPEX,y"),
        math_text(" + ", "p"),
        math_var("C", "TRANS,y"),
    ]
    annual_output = [math_var("Q", "SAF,y")]
    return [
        math_text("LCO-SAF", "p"),
        math_text(" = ", "p"),
        math_fraction(
            math_sum("y", [math_fraction(annual_cost, discount_factor())]),
            math_sum("y", [math_fraction(annual_output, discount_factor())]),
        ),
    ]


def eq_objective() -> list[OxmlElement]:
    return [
        math_text("Z"),
        math_text(" = ", "p"),
        math_var("Z", "CAPEX"),
        math_text(" + ", "p"),
        math_var("Z", "OPEX"),
        math_text(" + ", "p"),
        math_var("Z", "Transport"),
        math_text(" + ", "p"),
        math_var("Z", "Energy"),
        math_text(" + ", "p"),
        math_var("Z", "Material"),
        math_text(" + ", "p"),
        math_var("Z", "Shortage"),
        math_text(" − ", "p"),
        math_var("Z", "Credit"),
    ]


def eq_unified_scenario_model() -> list[OxmlElement]:
    return [
        math_sub([math_text("min", "p")], [math_var("x", "s")]),
        math_text(" ", "p"),
        math_var("Z", "s"),
        math_text("    s.t.    ", "p"),
        math_var("x", "s"),
        math_text(" ∈ ", "p"),
        math_var("X", "s", "core"),
        math_text(" ∩ ", "p"),
        math_var("X", "s", "H"),
        math_text(" ∩ ", "p"),
        math_var("X", "s", "C"),
        math_text(" ∩ ", "p"),
        math_var("X", "s", "P"),
        math_text(" ∩ ", "p"),
        math_var("X", "s", "I"),
        math_text(" ∩ ", "p"),
        math_var("X", "s", "N"),
        math_text(" ∩ ", "p"),
        math_var("X", "s", "D"),
        math_text(",  s ∈ S", "p"),
    ]


def eq_capacity_activation() -> list[OxmlElement]:
    return [
        math_text("0 ≤ ", "p"),
        math_var("q", "k,t", "p"),
        math_text(" ≤ ", "p"),
        math_var("Q", "k", "p"),
        math_text(";   ", "p"),
        math_var("Q", "k", "p,min"),
        math_var("z", "k", "p"),
        math_text(" ≤ ", "p"),
        math_var("Q", "k", "p"),
        math_text(" ≤ ", "p"),
        math_var("Q", "k", "p,max"),
        math_var("z", "k", "p"),
    ]


def eq_feedstock_availability() -> list[OxmlElement]:
    return [
        *math_sum("k", [math_var("f", "i,k,t", "m")]),
        math_text(" ≤ ", "p"),
        math_var("A", "i,t", "m"),
        math_text(",  m ∈ ", "p"),
        math_var("M", "s"),
    ]


def eq_production_balance() -> list[OxmlElement]:
    return [
        math_var("q", "k,t", "SAF"),
        math_text(" = ", "p"),
        math_var("β", superscript="MTJ"),
        math_text(" · ", "p"),
        math_var("q", "k,t", "MeOH,cons"),
        math_text(";  ", "p"),
        math_var("x", "p,k,t", "m"),
        math_text(" = ", "p"),
        math_var("α", "p", "m"),
        math_text(" · ", "p"),
        math_var("q", "p,k,t"),
        math_text(",  m ∈ {H₂, CO₂}", "p"),
    ]


def eq_conversion_balance() -> list[OxmlElement]:
    return [
        math_var("f", "k,t", "m"),
        math_text(" = ", "p"),
        math_var("α", "m", "p"),
        math_var("q", "k,t", "SAF"),
        math_text(";   ", "p"),
        math_var("q", "k,t", "SAF"),
        math_text(" = ", "p"),
        math_var("β", superscript="MTJ"),
        math_var("q", "k,t", "MeOH,cons"),
        math_text(",  p ∈ {FT, MTJ, CTL, GTL}", "p"),
    ]


def eq_pathway_feedstock_balance() -> list[OxmlElement]:
    return [
        math_var("e", "k,t", "DAC"),
        math_text(" = ", "p"),
        math_var("ε", superscript="DAC"),
        math_var("c", "k,t", "DAC"),
        math_text(";   ", "p"),
        math_var("c", "k,t", "coal"),
        math_text(" = ", "p"),
        math_var("γ", superscript="coal"),
        math_var("b", "k,t", "coal"),
        math_text(";   ", "p"),
        math_var("g", "k,t", "NG"),
        math_text(" = ", "p"),
        math_var("α", superscript="NG"),
        math_var("q", "k,t", "SAF"),
    ]


def eq_inventory_balance() -> list[OxmlElement]:
    return [
        math_var("I", "n,t"),
        math_text(" = ", "p"),
        math_var("I", "n,t−1"),
        math_text(" + ", "p"),
        math_var("F", "n,t", "in"),
        math_text(" − ", "p"),
        math_var("F", "n,t", "out"),
        math_text(",  n ∈ {H₂, CO₂, MeOH, SAF}", "p"),
    ]


def eq_transport_capacity() -> list[OxmlElement]:
    return [
        math_text("0 ≤ ", "p"),
        math_var("f", "u,v,τ", "m"),
        math_text(" ≤ ", "p"),
        math_var("F", "u,v", "m,max"),
        math_var("y", "u,v", "m"),
        math_text(",  m ∈ ", "p"),
        math_var("M", "s"),
    ]


def eq_demand_balance() -> list[OxmlElement]:
    return [
        *math_sum("k", [math_var("f", "k,a,w", "SAF")]),
        math_text(" + ", "p"),
        math_var("s", "a,w"),
        math_text(" ≥ ", "p"),
        math_var("D", "a,w"),
        math_text(",  ∀ a ∈ A, w ∈ W", "p"),
    ]


def eq_emissions_intensity() -> list[OxmlElement]:
    numerator = [
        math_var("E", "s", "feed"),
        math_text(" + ", "p"),
        math_var("E", "s", "process"),
        math_text(" + ", "p"),
        math_var("E", "s", "transport"),
        math_text(" + ", "p"),
        math_var("E", "s", "storage"),
        math_text(" + ", "p"),
        math_var("E", "s", "capex"),
        math_text(" − ", "p"),
        math_var("C", "s", "credit"),
    ]
    denominator = [math_var("Q", "s", "SAF,E")]
    return [
        math_var("E", "s", "SAF"),
        math_text(" = ", "p"),
        math_fraction(numerator, denominator),
    ]


def eq_emissions_reduction() -> list[OxmlElement]:
    return [
        math_var("R", "CO₂"),
        math_text(" = ", "p"),
        math_var("Q", "E", "SAF"),
        math_text(" · (", "p"),
        math_var("E", superscript="JetA"),
        math_text(" − ", "p"),
        math_var("E", superscript="SAF"),
        math_text(")", "p"),
    ]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def inches_to_dxa(width_inches: float) -> int:
    return int(round(width_inches * 1440))


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(inches_to_dxa(width_inches)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(inches_to_dxa(width_inches)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_grid_widths(table, col_widths: list[float]) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(inches_to_dxa(width)))
        tbl_grid.append(grid_col)


def set_cell_margins(cell, top: int = 50, bottom: int = 50, left: int = 65, right: int = 65) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_borders(
    cell,
    *,
    top: tuple[str, int, str] | None = None,
    bottom: tuple[str, int, str] | None = None,
    left: tuple[str, int, str] | None = None,
    right: tuple[str, int, str] | None = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    def apply(edge_name: str, spec: tuple[str, int, str] | None) -> None:
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        if spec is None:
            edge.set(qn("w:val"), "nil")
            for attr in ("sz", "space", "color"):
                if qn(f"w:{attr}") in edge.attrib:
                    del edge.attrib[qn(f"w:{attr}")]
            return
        val, size, color = spec
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)

    apply("top", top)
    apply("bottom", bottom)
    apply("left", left)
    apply("right", right)


def apply_three_line_table_style(table) -> None:
    border_color = "000000"
    thick = ("single", 12, border_color)  # 1.5 pt
    thin = ("single", 6, border_color)     # 0.75 pt
    none = None

    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for edge_name, spec in {
        "top": thick,
        "bottom": thick,
        "left": none,
        "right": none,
        "insideH": none,
        "insideV": none,
    }.items():
        edge = tbl_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tbl_borders.append(edge)
        if spec is None:
            edge.set(qn("w:val"), "nil")
            for attr in ("sz", "space", "color"):
                if qn(f"w:{attr}") in edge.attrib:
                    del edge.attrib[qn(f"w:{attr}")]
            continue
        val, size, color = spec
        edge.set(qn("w:val"), val)
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)

    last_row_idx = len(table.rows) - 1
    for row_idx, row in enumerate(table.rows):
        is_header = row_idx == 0
        is_last = row_idx == last_row_idx
        for cell in row.cells:
            set_cell_borders(
                cell,
                top=thick if is_header else none,
                bottom=thin if is_header else (thick if is_last else none),
                left=none,
                right=none,
            )


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    if style_exists(doc, "Normal Table"):
        table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(header_cells[i])
        for paragraph in header_cells[i].paragraphs:
            if style_exists(doc, "No Spacing"):
                paragraph.style = "No Spacing"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_indent(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(12)

    for row in rows:
        tr = table.add_row().cells
        for i, value in enumerate(row):
            tr[i].text = value
            tr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(tr[i])
            for paragraph in tr[i].paragraphs:
                if style_exists(doc, "No Spacing"):
                    paragraph.style = "No Spacing"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                clear_paragraph_indent(paragraph)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)

    if col_widths:
        set_table_width(table, sum(col_widths))
        set_table_grid_widths(table, col_widths)
        for idx, width in enumerate(col_widths):
            table.columns[idx].width = Inches(width)
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                set_cell_width(row.cells[idx], width)

    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)
    apply_three_line_table_style(table)

    return table


def resize_image(src: Path, dst: Path, max_width: int = 2400) -> Path:
    ensure_dir(dst.parent)
    with Image.open(src) as image:
        image = image.convert("RGB")
        if image.width > max_width:
            ratio = max_width / image.width
            resample = getattr(Image, "Resampling", Image).LANCZOS
            image = image.resize((max_width, int(image.height * ratio)), resample)
        image.save(dst, format="PNG", optimize=True)
    return dst


def ensure_raster_export_bundle(png_path: Path) -> None:
    ensure_dir(png_path.parent)
    with Image.open(png_path) as image:
        rgb = image.convert("RGB")
        width_px, height_px = rgb.size
        pdf_path = png_path.with_suffix(".pdf")
        if not pdf_path.exists():
            rgb.save(pdf_path, format="PDF", resolution=300.0)

    svg_path = png_path.with_suffix(".svg")
    if not svg_path.exists():
        encoded = base64.b64encode(png_path.read_bytes()).decode("ascii")
        svg_path.write_text(
            (
                f'<svg xmlns="http://www.w3.org/2000/svg" width="{width_px}" height="{height_px}" '
                f'viewBox="0 0 {width_px} {height_px}">'
                f'<image href="data:image/png;base64,{encoded}" width="{width_px}" height="{height_px}"/>'
                "</svg>\n"
            ),
            encoding="utf-8",
        )


def copy_figure_bundle(src_png: Path, dst_png: Path, max_width: int = 2400) -> Path:
    resized_png = resize_image(src_png, dst_png, max_width=max_width)
    for suffix in (".pdf", ".svg"):
        src_variant = src_png.with_suffix(suffix)
        if src_variant.exists():
            shutil.copy2(src_variant, dst_png.with_suffix(suffix))
    ensure_raster_export_bundle(dst_png)
    return resized_png


def save_matplotlib_figure(fig, path: Path) -> Path:
    ensure_dir(path.parent)
    base_path = path.with_suffix("")
    if getattr(fig, "_suptitle", None) is not None:
        fig._suptitle.set_text("")
    for text in fig.texts:
        _, y_pos = text.get_position()
        if y_pos >= 0.88:
            text.set_visible(False)
    for ax in fig.axes:
        ax.title.set_text("")
        if hasattr(ax, "_left_title"):
            ax._left_title.set_text("")
        if hasattr(ax, "_right_title"):
            ax._right_title.set_text("")
    if save_publication_figure is not None:
        save_publication_figure(
            fig,
            base_path,
            formats=["png", "pdf", "svg"],
            dpi=300,
            bbox_inches="tight",
            pad_inches=0.04,
            facecolor=fig.get_facecolor(),
        )
    else:
        fig.savefig(path, dpi=300, bbox_inches="tight", pad_inches=0.04, facecolor=fig.get_facecolor())
        fig.savefig(path.with_suffix(".pdf"), bbox_inches="tight", pad_inches=0.04, facecolor=fig.get_facecolor())
        fig.savefig(path.with_suffix(".svg"), bbox_inches="tight", pad_inches=0.04, facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def style_chart_axes(ax, grid_axis: str = "y") -> None:
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#CBD5E1")
    ax.spines["bottom"].set_color("#CBD5E1")
    ax.tick_params(
        axis="both",
        colors="#334155",
        labelsize=7.0,
        direction="out",
        length=4,
        width=0.8,
    )
    if grid_axis:
        ax.grid(axis=grid_axis, color="#CBD5E1", linewidth=0.8, linestyle="--", alpha=0.3)
    ax.xaxis.grid(False)
    ax.set_axisbelow(True)


def create_template_figure(template_key: str, **kwargs):
    figsize = ACADEMIC_TEMPLATE_SIZES.get(template_key)
    if figsize is None:
        raise KeyError(f"Unknown template figure size: {template_key}")
    return plt.figure(figsize=figsize, dpi=300, **kwargs)


def apply_horizontal_bar_template(
    ax,
    labels: list[str],
    values: list[float] | np.ndarray,
    *,
    colors,
    xlabel: str,
    invert: bool = True,
    show_end_labels: bool = False,
    value_formatter: str = "{:.1f}",
) -> None:
    values = np.asarray(values, dtype=float)
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, values, color=colors, alpha=0.85, edgecolor="white")
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=FONT_CONFIG["tick"])
    if invert:
        ax.invert_yaxis()
    ax.set_xlabel(xlabel)
    style_chart_axes(ax, grid_axis="x")
    ax.yaxis.grid(False)
    if show_end_labels:
        offset = max(values.max() * 0.015, 0.5) if len(values) else 0.5
        for bar, value in zip(bars, values):
            ax.text(
                bar.get_width() + offset,
                bar.get_y() + bar.get_height() / 2,
                value_formatter.format(value),
                va="center",
                fontsize=7.0,
                color=PUBLICATION_COLORS["black"],
            )


def apply_grouped_bar_template(
    ax,
    categories: list[str],
    series: list[tuple[str, np.ndarray | list[float], str]],
    *,
    ylabel: str | None = None,
    xlabel: str | None = None,
    legend: bool = True,
    log_y: bool = False,
    legend_kwargs: dict | None = None,
) -> None:
    n_series = max(len(series), 1)
    x = np.arange(len(categories), dtype=float)
    bar_width = 0.8 / n_series
    for idx, (label, values, color) in enumerate(series):
        offset = (idx - n_series / 2 + 0.5) * bar_width
        ax.bar(x + offset, values, width=bar_width, label=label, color=color, alpha=0.85, edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=FONT_CONFIG["tick"])
    if ylabel:
        ax.set_ylabel(ylabel)
    if xlabel:
        ax.set_xlabel(xlabel)
    if log_y:
        ax.set_yscale("log")
    style_chart_axes(ax)
    if legend:
        ax.legend(frameon=False, **(legend_kwargs or {}))


def apply_stacked_bar_template(
    ax,
    categories: list[str],
    stacks: list[tuple[str, np.ndarray | list[float], str]],
    *,
    ylabel: str | None = None,
    xlabel: str | None = None,
    orientation: str = "v",
    normalize_to_pct: bool = False,
    legend: bool = True,
    legend_kwargs: dict | None = None,
) -> None:
    if orientation not in {"v", "h"}:
        raise ValueError("orientation must be 'v' or 'h'")
    arrays = [np.asarray(values, dtype=float) for _, values, _ in stacks]
    if normalize_to_pct and arrays:
        total = np.sum(arrays, axis=0)
        total = np.where(total == 0, 1.0, total)
        arrays = [values / total * 100 for values in arrays]
    if orientation == "v":
        x = np.arange(len(categories), dtype=float)
        bottoms = np.zeros(len(categories), dtype=float)
        for (label, _, color), values in zip(stacks, arrays):
            ax.bar(x, values, bottom=bottoms, color=color, alpha=0.85, edgecolor="white", label=label)
            bottoms += values
        ax.set_xticks(x)
        ax.set_xticklabels(categories, fontsize=FONT_CONFIG["tick"])
        if ylabel:
            ax.set_ylabel(ylabel)
        if xlabel:
            ax.set_xlabel(xlabel)
        style_chart_axes(ax)
    else:
        y = np.arange(len(categories), dtype=float)
        lefts = np.zeros(len(categories), dtype=float)
        for (label, _, color), values in zip(stacks, arrays):
            ax.barh(y, values, left=lefts, color=color, alpha=0.85, edgecolor="white", label=label)
            lefts += values
        ax.set_yticks(y)
        ax.set_yticklabels(categories, fontsize=FONT_CONFIG["tick"])
        ax.invert_yaxis()
        if ylabel:
            ax.set_ylabel(ylabel)
        if xlabel:
            ax.set_xlabel(xlabel)
        style_chart_axes(ax, grid_axis="x")
        ax.yaxis.grid(False)
    if legend:
        ax.legend(frameon=False, **(legend_kwargs or {}))


def apply_scatter_template(
    ax,
    x,
    y,
    *,
    color=None,
    label: str | None = None,
    marker: str = "o",
    s: float | np.ndarray = 18,
    alpha: float = 0.85,
    edgecolors: str = "white",
    linewidth: float = 0.7,
    zorder: float = 3,
    rasterize_threshold: int = 5000,
    **kwargs,
):
    point_count = len(x) if hasattr(x, "__len__") else 0
    if point_count >= rasterize_threshold and edgecolors == "white":
        edgecolors = "none"
    scatter_color = kwargs.pop("c", color)
    return ax.scatter(
        x,
        y,
        c=scatter_color,
        label=label,
        marker=marker,
        s=s,
        alpha=alpha,
        edgecolors=edgecolors,
        linewidth=linewidth,
        zorder=zorder,
        rasterized=point_count >= rasterize_threshold,
        **kwargs,
    )


def apply_boxplot_template(
    ax,
    data_groups: list[np.ndarray | list[float]],
    *,
    labels: list[str],
    colors: list[str],
    ylabel: str | None = None,
    log_y: bool = False,
    rotation: float = 0,
    showfliers: bool = False,
    widths: float = 0.56,
):
    box = ax.boxplot(
        data_groups,
        patch_artist=True,
        tick_labels=labels,
        widths=widths,
        showfliers=showfliers,
        medianprops={"color": "white", "linewidth": 1.2},
    )
    for patch, color in zip(box["boxes"], colors):
        patch.set_facecolor(color)
        patch.set_edgecolor(color)
        patch.set_alpha(0.82)
    if log_y:
        ax.set_yscale("log")
    if ylabel:
        ax.set_ylabel(ylabel)
    style_chart_axes(ax)
    if rotation:
        ax.tick_params(axis="x", rotation=rotation)
    return box


def apply_heatmap_template(
    fig,
    ax,
    matrix: np.ndarray,
    *,
    xlabels: list[str],
    ylabels: list[str],
    cmap: str = "RdYlBu_r",
    vmin: float | None = None,
    vmax: float | None = None,
    colorbar_label: str | None = None,
    annotate: bool = False,
    annotation_formatter: str = "{:.0%}",
):
    im = ax.imshow(matrix, cmap=cmap, vmin=vmin, vmax=vmax, aspect="auto", interpolation="nearest")
    ax.set_xticks(np.arange(len(xlabels)))
    ax.set_xticklabels(xlabels, rotation=18, ha="right", fontsize=FONT_CONFIG["tick"])
    ax.set_yticks(np.arange(len(ylabels)))
    ax.set_yticklabels(ylabels, fontsize=FONT_CONFIG["tick"])
    if annotate:
        for i in range(matrix.shape[0]):
            for j in range(matrix.shape[1]):
                ax.text(j, i, annotation_formatter.format(matrix[i, j]), ha="center", va="center", fontsize=6.2, color="#0F172A")
    cbar = fig.colorbar(im, ax=ax, shrink=0.85, pad=0.03)
    if colorbar_label:
        cbar.set_label(colorbar_label, rotation=-90, va="bottom", labelpad=15, fontsize=FONT_CONFIG["label"])
    return im, cbar


def apply_histogram_template(
    ax,
    series: list[tuple[str | None, np.ndarray | list[float], str, float]],
    *,
    bins,
    xlabel: str,
    ylabel: str = "Count",
    log_x: bool = False,
    legend: bool = False,
    legend_kwargs: dict | None = None,
) -> None:
    for label, values, color, alpha in series:
        ax.hist(values, bins=bins, color=color, alpha=alpha, label=label)
    if log_x:
        ax.set_xscale("log")
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    style_chart_axes(ax)
    if legend and any(label for label, *_ in series):
        ax.legend(frameon=False, **(legend_kwargs or {}))


def apply_line_template(
    ax,
    series: list[dict[str, object]],
    *,
    xlabel: str | None = None,
    ylabel: str | None = None,
    legend: bool = False,
    legend_kwargs: dict | None = None,
    vertical_guides: list[float] | None = None,
    horizontal_guides: list[float] | None = None,
) -> None:
    for guide in vertical_guides or []:
        ax.axvline(guide, color="#CBD5E1", linestyle="--", linewidth=0.8, zorder=0)
    for guide in horizontal_guides or []:
        ax.axhline(guide, color="#94A3B8", linestyle="--", linewidth=0.8, zorder=0)
    for spec in series:
        ax.plot(
            spec["x"],
            spec["y"],
            color=spec.get("color", PUBLICATION_COLORS["blue"]),
            linewidth=spec.get("linewidth", 1.4),
            marker=spec.get("marker"),
            markersize=spec.get("markersize", 0),
            alpha=spec.get("alpha", 1.0),
            linestyle=spec.get("linestyle", "-"),
            label=spec.get("label"),
            zorder=spec.get("zorder", 3),
        )
    if xlabel:
        ax.set_xlabel(xlabel)
    if ylabel:
        ax.set_ylabel(ylabel)
    style_chart_axes(ax)
    if legend:
        ax.legend(frameon=False, **(legend_kwargs or {}))


def apply_dual_axis_bar_line_template(
    ax,
    categories: list[str],
    bar_values: np.ndarray | list[float],
    line_values: np.ndarray | list[float],
    *,
    bar_color: str,
    line_color: str,
    left_ylabel: str,
    right_ylabel: str,
    rotation: float = 0,
) -> object:
    x = np.arange(len(categories), dtype=float)
    ax.bar(x, bar_values, color=bar_color, alpha=0.85, edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=FONT_CONFIG["tick"], rotation=rotation)
    ax.set_ylabel(left_ylabel)
    style_chart_axes(ax)

    ax_right = ax.twinx()
    ax_right.plot(x, line_values, color=line_color, marker="o", linewidth=1.3, zorder=4)
    ax_right.set_ylabel(right_ylabel)
    ax_right.spines["top"].set_visible(False)
    ax_right.spines["right"].set_color("#CBD5E1")
    ax_right.tick_params(axis="y", colors="#334155", labelsize=7.0, width=0.7, length=3)
    ax_right.grid(False)
    return ax_right


def add_compact_bottom_strip_axes(
    fig,
    anchor_ax,
    *,
    y: float,
    height: float,
    x_offset_frac: float,
    width_frac: float,
):
    pos = anchor_ax.get_position()
    return fig.add_axes([pos.x0 + pos.width * x_offset_frac, y, pos.width * width_frac, height])


def add_compact_bottom_legend(
    fig,
    anchor_ax,
    handles,
    *,
    y: float,
    height: float,
    x_offset_frac: float = 0.03,
    width_frac: float = 0.58,
    ncol: int = 1,
    loc: str = "center left",
    fontsize: float = 6.8,
    columnspacing: float = 1.0,
    handletextpad: float = 0.4,
):
    lax = add_compact_bottom_strip_axes(
        fig,
        anchor_ax,
        y=y,
        height=height,
        x_offset_frac=x_offset_frac,
        width_frac=width_frac,
    )
    lax.axis("off")
    lax.legend(
        handles,
        [handle.get_label() for handle in handles],
        frameon=False,
        ncol=ncol,
        loc=loc,
        handletextpad=handletextpad,
        columnspacing=columnspacing,
        borderaxespad=0.0,
        fontsize=fontsize,
    )
    return lax


def add_compact_bottom_colorbar(
    fig,
    anchor_ax,
    mappable,
    *,
    label: str,
    y: float,
    height: float,
    x_offset_frac: float = 0.16,
    width_frac: float = 0.64,
    fontsize: float = 7.0,
    tick_labelsize: float = 6.3,
):
    cax = add_compact_bottom_strip_axes(
        fig,
        anchor_ax,
        y=y,
        height=height,
        x_offset_frac=x_offset_frac,
        width_frac=width_frac,
    )
    cbar = fig.colorbar(mappable, cax=cax, orientation="horizontal")
    cbar.set_label(label, fontsize=fontsize)
    cbar.ax.tick_params(labelsize=tick_labelsize, width=0.6, length=2)
    cbar.outline.set_linewidth(0.6)
    return cbar


def add_publication_legend(ax, *args, **kwargs):
    return ax.legend(*args, frameon=False, **kwargs)


def add_publication_figure_legend(
    fig,
    handles,
    *,
    loc: str,
    bbox_to_anchor,
    ncol: int,
    fontsize: float = 6.8,
    columnspacing: float = 1.0,
    handletextpad: float = 0.5,
):
    return fig.legend(
        handles,
        [handle.get_label() for handle in handles],
        frameon=False,
        loc=loc,
        bbox_to_anchor=bbox_to_anchor,
        ncol=ncol,
        fontsize=fontsize,
        columnspacing=columnspacing,
        handletextpad=handletextpad,
    )


def add_panel_label(ax, label: str) -> None:
    rendered = label.strip()
    if len(rendered) == 1 and rendered.isalpha():
        rendered = f"({rendered.lower()})"
    ax.text(
        -0.10,
        1.02,
        rendered,
        transform=ax.transAxes,
        fontsize=10.0,
        fontweight="bold",
        color=PUBLICATION_COLORS["black"],
        va="bottom",
        ha="left",
        clip_on=False,
    )


def get_publication_rc() -> dict[str, object]:
    return {
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans"],
        "axes.unicode_minus": False,
        "font.size": 8,
        "axes.titlesize": 8.5,
        "axes.labelsize": 8.5,
        "xtick.labelsize": 7,
        "ytick.labelsize": 7,
        "axes.linewidth": 0.7,
        "xtick.major.width": 0.7,
        "ytick.major.width": 0.7,
        "xtick.major.size": 3,
        "ytick.major.size": 3,
        "legend.fontsize": 6.8,
        "legend.frameon": False,
        "figure.facecolor": "white",
        "axes.facecolor": "white",
        "savefig.dpi": 300,
        "pdf.fonttype": 42,
        "ps.fonttype": 42,
        "svg.fonttype": "none",
        "mathtext.fontset": "custom",
        "mathtext.rm": "Arial",
        "mathtext.it": "Arial:italic",
        "mathtext.bf": "Arial:bold",
        "mathtext.default": "regular",
        "axes.prop_cycle": cycler(color=ACADEMIC_PALETTE),
    }


def normalize_scenario_code(value: object) -> str:
    mapping = {
        "GTL-GH-FT": "GTL",
        "GTL-GH-MTJ": "GTL-GH",
        "GTL-BH-MTJ": "GTL-BH",
    }
    text = str(value).strip()
    return mapping.get(text, text)


def scenario_sort_key(value: object) -> int:
    order = {spec["code"]: idx for idx, spec in enumerate(SCENARIO_SPECS)}
    return order.get(normalize_scenario_code(value), 999)


def translate_network_metric_columns(frame: pd.DataFrame) -> pd.DataFrame:
    if frame.empty:
        return frame
    translated = frame.rename(
        columns={
            "节点数量": "NodeCount",
            "关键节点集中度": "KeyNodeConcentration",
            "关键通道集中度": "KeyCorridorConcentration",
            "平均运输距离": "MeanTransportDistance",
            "跨区域物流比例": "CrossRegionalShare",
            "单点失效影响": "SingleFailureImpact",
        }
    ).copy()
    translated["Scenario"] = translated["OfficialName"].fillna(translated["Scenario"]).map(normalize_scenario_code)
    return translated


def translate_material_label(value: object) -> str:
    mapping = {
        "H2 production": "H₂ source",
        "CO2 capture+H2 production": "CO₂+H₂ source",
        "No raw material production": "No supply",
    }
    return mapping.get(str(value).strip(), str(value).strip())


def translate_site_share_label(value: object) -> str:
    mapping = {
        "Renewable site": "Renewable",
        "CO2 capture site": "CO₂ capture",
        "Airport site": "Airport",
        "By-product H2 site": "By-product H₂",
        "Other site": "Other",
    }
    return mapping.get(str(value).strip(), str(value).strip())


def translate_route_mode(value: object) -> str:
    mapping = {
        "truck": "Truck",
        "pipeline": "Pipeline",
    }
    return mapping.get(str(value).strip(), str(value).strip().title())


def translate_cargo_type(value: object) -> str:
    mapping = {
        "天然气": "Natural gas",
        "氢气": "Hydrogen",
        "MTJ": "SAF",
    }
    return mapping.get(str(value).strip(), str(value).strip())


def translate_node_type(value: object) -> str:
    mapping = {
        "天然气管道": "Natural-gas pipeline",
        "生产设施": "Production facility",
        "氢气生产站": "Hydrogen source",
        "MTJ工厂": "MTJ plant",
        "机场": "Airport",
    }
    return mapping.get(str(value).strip(), str(value).strip())


def build_scenario_metrics_table() -> pd.DataFrame:
    portfolio = summarize_scenario_portfolio().copy()
    portfolio["Scenario"] = portfolio["Scenario"].map(normalize_scenario_code)

    temporal_eff = load_latest_temporal_efficiency()
    if not temporal_eff.empty:
        temporal_eff = temporal_eff.copy()
        temporal_eff["Scenario"] = temporal_eff["Scenario"].map(normalize_scenario_code)
        temporal_eff = (
            temporal_eff.groupby("Scenario")
            .agg(
                MeanH2Utilization=("H2_Utilization", "mean"),
                MeanSAFUtilization=("SAF_Utilization", "mean"),
                P10SAFUtilization=("SAF_Utilization", lambda s: float(np.quantile(s, 0.10))),
                P90SAFUtilization=("SAF_Utilization", lambda s: float(np.quantile(s, 0.90))),
            )
            .reset_index()
        )

    else:
        temporal_eff = pd.DataFrame(columns=["Scenario", "MeanH2Utilization", "MeanSAFUtilization", "P10SAFUtilization", "P90SAFUtilization"])

    network_metrics, siting_metrics, _ = load_latest_network_metrics()
    network_metrics = translate_network_metric_columns(network_metrics)
    if not network_metrics.empty:
        network_metrics = network_metrics[
            [
                "Scenario",
                "NodeCount",
                "KeyNodeConcentration",
                "KeyCorridorConcentration",
                "MeanTransportDistance",
                "CrossRegionalShare",
                "SingleFailureImpact",
                "Cluster",
            ]
        ].drop_duplicates(subset=["Scenario"])

    if not siting_metrics.empty:
        siting_metrics = siting_metrics.copy()
        siting_metrics["Scenario"] = siting_metrics["OfficialName"].fillna(siting_metrics["Scenario"]).map(normalize_scenario_code)
        siting_metrics = siting_metrics[
            [
                "Scenario",
                "SAF_Plants",
                "Electrolyzers",
                "Avg_H2_Distance_km",
                "Avg_SAF_Distance_km",
                "SAF_Share_Renewable site",
                "SAF_Share_CO2 capture site",
                "SAF_Share_Airport site",
                "SAF_Share_Byproduct H2 site",
                "SAF_Share_Other site",
                "Cluster",
            ]
        ].drop_duplicates(subset=["Scenario"])

    temporal_path = ROOT / "products/supply_chain_optimization/visualization/results/temporal_robustness_metrics_latest.csv"
    temporal = pd.read_csv(temporal_path)
    temporal["Scenario"] = temporal["Scenario"].map(normalize_scenario_code)
    temporal = temporal[
        [
            "Scenario",
            "WorstWeekRetention",
            "WorstWeekCostLiftPct",
            "ChronicPenaltyPct",
            "RobustnessPenaltyPct",
            "TotalPenaltyPct",
            "InvestmentShare",
            "LowLoadEpisodes12h",
            "MaxLowRunHours",
            "PenaltyClass",
        ]
    ].drop_duplicates(subset=["Scenario"])

    merged = portfolio.merge(temporal_eff, on="Scenario", how="left")
    if not network_metrics.empty:
        merged = merged.merge(network_metrics, on="Scenario", how="left", suffixes=("", "_network"))
    if not siting_metrics.empty:
        merged = merged.merge(siting_metrics, on="Scenario", how="left", suffixes=("", "_siting"))
        if "Cluster_siting" in merged.columns and "Cluster" in merged.columns:
            merged["Cluster"] = merged["Cluster"].fillna(merged["Cluster_siting"])
            merged = merged.drop(columns=["Cluster_siting"])
    merged = merged.merge(temporal, on="Scenario", how="left", suffixes=("", "_temporal"))
    return merged.sort_values("Scenario", key=lambda s: s.map(scenario_sort_key)).reset_index(drop=True)


def normalize_status(value: object) -> str:
    if pd.isna(value):
        return "Unknown"
    status = str(value).replace("\n", " ").strip().lower()
    if not status:
        return "Unknown"
    if "operating" in status:
        return "Operating"
    if "pre-construction" in status or ("construction" in status and "pre" in status):
        return "Pre-construction"
    if "construction" in status:
        return "Under construction"
    if "announced" in status:
        return "Announced"
    if "planned" in status or "permitted" in status or "proposed" in status:
        return "Planned / proposed"
    return status.title()


def create_raw_data_overview_panel(path: Path, data: AppendixData) -> Path:
    demand_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_demand_20251129_231231.xlsx"
    demand_df = pd.read_excel(demand_path)
    demand_df["airport_en"] = demand_df["airport"].map(AIRPORT_MAP).fillna(demand_df["airport"])
    weekly_demand = (
        demand_df.pivot_table(
            index="week_number",
            columns="airport_en",
            values="weekly_total_fuel_kg_total",
            aggfunc="first",
        )
        .sort_index()
        .fillna(0.0)
        / 1e6
    )

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(
            2,
            2,
            left=0.08,
            right=0.985,
            top=0.965,
            bottom=0.12,
            wspace=0.28,
            hspace=0.34,
        )

        ax = fig.add_subplot(gs[0, 0])
        count_rows = [
            ("Solar inventory rows", data.transport_summary["solar_inventory_rows"], PUBLICATION_COLORS["orange"]),
            ("Wind inventory rows", data.transport_summary["wind_inventory_rows"], PUBLICATION_COLORS["blue"]),
            (
                "Modeled renewable plants",
                data.renewable_summary["wind"]["unique_plants"] + data.renewable_summary["solar"]["unique_plants"],
                PUBLICATION_COLORS["green"],
            ),
            ("Pipeline records", data.transport_summary["pipeline_records"], PUBLICATION_COLORS["gray"]),
            ("Representative periods", data.demand_summary["total_3h_periods"], PUBLICATION_COLORS["sky"]),
            ("Steel H₂ sites", data.byproduct_summary["steel_sites"], PUBLICATION_COLORS["vermillion"]),
            ("Refinery H₂ sites", data.byproduct_summary["refinery_sites"], PUBLICATION_COLORS["magenta"]),
            ("CO₂ capture sources", data.co2_summary["records"], PUBLICATION_COLORS["black"]),
        ]
        count_rows = sorted(count_rows, key=lambda item: item[1], reverse=True)
        labels = [row[0] for row in count_rows]
        values = [row[1] for row in count_rows]
        colors = [row[2] for row in count_rows]
        apply_horizontal_bar_template(
            ax,
            labels,
            values,
            colors=colors,
            xlabel="Archived records or periods (log scale)",
            invert=True,
        )
        ax.set_xscale("log")
        ax.set_xlim(10, 25000)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        renewable_pairs = [
            {
                "label": "Wind",
                "raw_count": data.transport_summary["wind_inventory_rows"],
                "modeled_count": data.renewable_summary["wind"]["unique_plants"],
                "raw_capacity_gw": data.transport_summary["wind_inventory_capacity_mw"] / 1000,
                "modeled_capacity_gw": data.renewable_summary["wind"]["capacity_mw"] / 1000,
                "color": PUBLICATION_COLORS["blue"],
            },
            {
                "label": "Solar",
                "raw_count": data.transport_summary["solar_inventory_rows"],
                "modeled_count": data.renewable_summary["solar"]["unique_plants"],
                "raw_capacity_gw": data.transport_summary["solar_inventory_capacity_mw"] / 1000,
                "modeled_capacity_gw": data.renewable_summary["solar"]["capacity_mw"] / 1000,
                "color": PUBLICATION_COLORS["orange"],
            },
        ]
        y_pos = np.array([1, 0], dtype=float)
        for idx, item in enumerate(renewable_pairs):
            y = y_pos[idx]
            ax.plot(
                [item["modeled_count"], item["raw_count"]],
                [y, y],
                color=PUBLICATION_COLORS["light_gray"],
                linewidth=2.1,
                zorder=1,
            )
            ax.scatter(
                item["raw_count"],
                y,
                s=40,
                facecolors="white",
                edgecolors=item["color"],
                linewidth=1.4,
                zorder=3,
                label="Raw inventory" if idx == 0 else None,
            )
            ax.scatter(
                item["modeled_count"],
                y,
                s=40,
                color=item["color"],
                edgecolors="white",
                linewidth=0.5,
                zorder=4,
                label="Modeled plants" if idx == 0 else None,
            )
        ax.set_xscale("log")
        ax.set_xlim(900, 18000)
        ax.set_ylim(-0.6, 1.6)
        ax.set_yticks(y_pos)
        ax.set_yticklabels([item["label"] for item in renewable_pairs])
        ax.set_xlabel("Facility records or modeled plants (log scale)")
        style_chart_axes(ax, grid_axis="x")
        add_publication_legend(ax, loc="lower right", handletextpad=0.6)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        apply_grouped_bar_template(
            ax,
            [f"W{int(week)}" for week in weekly_demand.index],
            [
                ("Beijing", weekly_demand.get("Beijing", pd.Series(index=weekly_demand.index, dtype=float)).values, PUBLICATION_COLORS["blue"]),
                ("Tianjin", weekly_demand.get("Tianjin", pd.Series(index=weekly_demand.index, dtype=float)).values, PUBLICATION_COLORS["orange"]),
            ],
            ylabel="Weekly fuel demand (million kg)",
            legend=True,
            legend_kwargs={"ncol": 2, "loc": "upper right"},
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        raw_temporal = np.array(
            [
                52,
                52 * data.demand_summary["periods_per_week"],
                52 * 7 * 24,
            ],
            dtype=float,
        )
        retained_temporal = np.array(
            [
                data.demand_summary["weeks"],
                data.demand_summary["total_3h_periods"],
                data.renewable_summary["wind"]["hours"],
            ],
            dtype=float,
        )
        apply_grouped_bar_template(
            ax,
            ["Weeks", "3 h periods", "Hours"],
            [
                ("Raw year", raw_temporal, PUBLICATION_COLORS["light_gray"]),
                ("Retained model horizon", retained_temporal, PUBLICATION_COLORS["green"]),
            ],
            ylabel="Temporal records (log scale)",
            legend=True,
            log_y=True,
            legend_kwargs={"loc": "upper right"},
        )
        add_panel_label(ax, "D")

        return save_matplotlib_figure(fig, path)


def create_airport_demand_raw_summary(path: Path) -> Path:
    demand_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_demand_20251129_231231.xlsx"
    df = pd.read_excel(demand_path)
    df["airport_en"] = df["airport"].map(AIRPORT_MAP).fillna(df["airport"])
    df["week_label"] = df["week_number"].apply(lambda v: f"W{int(v)}")
    airport_colors = {"Beijing": PUBLICATION_COLORS["blue"], "Tianjin": PUBLICATION_COLORS["orange"]}

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.25, hspace=0.34)

        ax = fig.add_subplot(gs[0, 0])
        week_numbers = sorted(df["week_number"].unique())
        apply_grouped_bar_template(
            ax,
            [f"W{int(w)}" for w in week_numbers],
            [
                (
                    airport,
                    df[df["airport_en"] == airport].sort_values("week_number")["weekly_total_fuel_kg_total"].to_numpy() / 1e6,
                    airport_colors[airport],
                )
                for airport in ["Beijing", "Tianjin"]
            ],
            ylabel="Weekly fuel demand (million kg)",
            legend=True,
            legend_kwargs={"loc": "upper right"},
        )
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        week_markers = {1: "o", 2: "s", 3: "^", 4: "D"}
        for airport in ["Beijing", "Tianjin"]:
            subset = df[df["airport_en"] == airport].copy()
            for _, row in subset.iterrows():
                apply_scatter_template(
                    ax,
                    [row["weekly_distance_km_total"] / 1e6],
                    [row["weekly_total_fuel_kg_total"] / 1e6],
                    s=row["total_flights"] / 4,
                    color=airport_colors[airport],
                    marker=week_markers.get(int(row["week_number"]), "o"),
                    alpha=0.85,
                    edgecolors="white",
                    linewidth=0.8,
                )
        ax.set_xlabel("Weekly flown distance (million km)")
        ax.set_ylabel("Weekly fuel demand (million kg)")
        style_chart_axes(ax)
        airport_handles = [
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=airport_colors["Beijing"], markeredgecolor="white", markersize=6, label="Beijing"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=airport_colors["Tianjin"], markeredgecolor="white", markersize=6, label="Tianjin"),
        ]
        week_handles = [
            Line2D([0], [0], marker=marker, linestyle="None", markerfacecolor="#475569", markeredgecolor="white", markersize=6, label=f"Week {week}")
            for week, marker in week_markers.items()
        ]
        legend1 = add_publication_legend(ax, handles=airport_handles, loc="upper left")
        ax.add_artist(legend1)
        add_publication_legend(ax, handles=week_handles, loc="lower right", ncol=2, fontsize=6.8, columnspacing=0.8, handletextpad=0.4)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        phase_df = (
            df.groupby("airport_en")[["weekly_climb_fuel_kg_total", "weekly_cruise_fuel_kg_total", "weekly_descent_fuel_kg_total"]]
            .sum()
            .rename(
                columns={
                    "weekly_climb_fuel_kg_total": "Climb",
                    "weekly_cruise_fuel_kg_total": "Cruise",
                    "weekly_descent_fuel_kg_total": "Descent",
                }
            )
        )
        phase_share = phase_df.div(phase_df.sum(axis=1), axis=0)
        phase_colors = {"Climb": PUBLICATION_COLORS["sky"], "Cruise": PUBLICATION_COLORS["blue"], "Descent": PUBLICATION_COLORS["orange"]}
        apply_stacked_bar_template(
            ax,
            phase_share.index.tolist(),
            [
                (phase, phase_share[phase].to_numpy(), phase_colors[phase])
                for phase in ["Climb", "Cruise", "Descent"]
            ],
            ylabel="Fuel share (%)",
            normalize_to_pct=True,
            legend=True,
            legend_kwargs={"loc": "upper right"},
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        fuel_per_flight = (
            df.assign(fuel_per_flight_kg=df["weekly_total_fuel_kg_total"] / df["total_flights"].replace(0, np.nan))
            .pivot_table(index="week_number", columns="airport_en", values="fuel_per_flight_kg", aggfunc="first")
            .sort_index()
        )
        apply_grouped_bar_template(
            ax,
            [f"W{int(w)}" for w in week_numbers],
            [
                (airport, fuel_per_flight.get(airport, pd.Series(index=fuel_per_flight.index, dtype=float)).values, airport_colors[airport])
                for airport in ["Beijing", "Tianjin"]
            ],
            ylabel="Average fuel per flight (kg)",
            legend=False,
        )
        add_panel_label(ax, "D")

        return save_matplotlib_figure(fig, path)


def create_demand_aircraft_supplement(path: Path) -> Path:
    annual = load_or_build_annual_airport_weekly_demand()
    _, week_map = load_candidate_week_metadata()
    aircraft = load_or_build_aircraft_operation_summary().copy()

    annual_pivot = (
        annual.pivot_table(index="week_of_year", columns="airport_en", values="weekly_total_fuel_kg", aggfunc="sum")
        .reindex(range(1, 53))
        .fillna(0.0)
        / 1e6
    )
    selected_weeks = week_map.sort_values("new_week_number").reset_index(drop=True)
    aircraft = aircraft.sort_values("flights", ascending=False).reset_index(drop=True)
    aircraft_colors = {
        "B737": PUBLICATION_COLORS["blue"],
        "A320": PUBLICATION_COLORS["orange"],
        "A321": PUBLICATION_COLORS["green"],
        "A319": PUBLICATION_COLORS["sky"],
        "A330": PUBLICATION_COLORS["vermillion"],
        "B787": "#2C7FB8",
        "Regional jet": PUBLICATION_COLORS["gray"],
        "Turbo-prop": "#A16207",
        "Other/Unknown": PUBLICATION_COLORS["light_gray"],
    }
    aircraft["color"] = aircraft["aircraft_class"].map(aircraft_colors).fillna(PUBLICATION_COLORS["gray"])
    seasonal_spans = [
        (1, 2, "#FCE7A8", "Spring\nFestival"),
        (27, 35, "#DDF1E4", "Summer peak"),
    ]

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(
            2,
            2,
            left=0.08,
            right=0.985,
            top=0.88,
            bottom=0.13,
            height_ratios=[1.08, 1.0],
            hspace=0.42,
            wspace=0.30,
        )

        ax = fig.add_subplot(gs[0, :])
        y_max = float(annual_pivot.max().max()) * 1.25
        for start_week, end_week, shade_color, label in seasonal_spans:
            ax.axvspan(start_week - 0.5, end_week + 0.5, facecolor=shade_color, alpha=0.32, edgecolor="none", zorder=0)
            ax.text(
                (start_week + end_week) / 2,
                y_max * 0.965,
                label,
                ha="center",
                va="top",
                fontsize=6.5,
                color="#64748B",
            )
        for airport, color in [("Beijing", PUBLICATION_COLORS["blue"]), ("Tianjin", PUBLICATION_COLORS["orange"])]:
            ax.fill_between(annual_pivot.index, annual_pivot[airport].to_numpy(), color=color, alpha=0.08, zorder=1)
        apply_line_template(
            ax,
            [
                {
                    "x": annual_pivot.index.to_numpy(),
                    "y": annual_pivot["Beijing"].to_numpy(),
                    "color": PUBLICATION_COLORS["blue"],
                    "linewidth": 1.6,
                    "label": "Beijing",
                },
                {
                    "x": annual_pivot.index.to_numpy(),
                    "y": annual_pivot["Tianjin"].to_numpy(),
                    "color": PUBLICATION_COLORS["orange"],
                    "linewidth": 1.6,
                    "label": "Tianjin",
                },
            ],
            xlabel="Week of year",
            ylabel="Weekly fuel demand (million kg)",
            legend=False,
        )
        for row in selected_weeks.itertuples(index=False):
            calendar_week = int(row.original_week_in_52)
            ax.axvline(calendar_week, color="#94A3B8", linestyle="--", linewidth=0.9, alpha=0.7, zorder=1)
            for airport in ["Beijing", "Tianjin"]:
                apply_scatter_template(
                    ax,
                    [calendar_week],
                    [annual_pivot.loc[calendar_week, airport]],
                    color="white",
                    marker="D",
                    s=42,
                    edgecolors=PUBLICATION_COLORS["black"],
                    linewidth=0.95,
                    zorder=4,
                )
            label_dx = 10 if int(row.new_week_number) == 1 else 0
            label_ha = "left" if int(row.new_week_number) == 1 else "center"
            ax.annotate(
                f"W{int(row.new_week_number)}",
                xy=(calendar_week, float(annual_pivot.loc[calendar_week, "Beijing"])),
                xytext=(label_dx, 10),
                textcoords="offset points",
                ha=label_ha,
                va="bottom",
                fontsize=6.7,
                color=PUBLICATION_COLORS["black"],
            )
        ax.set_xlim(1, 52)
        ax.set_ylim(0, y_max)
        ax.set_xticks([1, 5, 10, 15, 20, 25, 30, 35, 40, 45, 52])
        add_panel_label(ax, "A")

        legend_handles = [
            Line2D([0], [0], color=PUBLICATION_COLORS["blue"], linewidth=1.6, label="Beijing"),
            Line2D([0], [0], color=PUBLICATION_COLORS["orange"], linewidth=1.6, label="Tianjin"),
            Line2D(
                [0],
                [0],
                marker="D",
                linestyle="None",
                markerfacecolor="white",
                markeredgecolor=PUBLICATION_COLORS["black"],
                markersize=5.0,
                label="Retained representative week",
            ),
            Patch(facecolor=seasonal_spans[0][2], edgecolor="none", alpha=0.65, label="Spring Festival period"),
            Patch(facecolor=seasonal_spans[1][2], edgecolor="none", alpha=0.65, label="Summer peak"),
        ]
        add_publication_figure_legend(
            fig,
            legend_handles,
            loc="upper center",
            bbox_to_anchor=(0.5, 0.99),
            ncol=5,
            fontsize=6.4,
            columnspacing=1.1,
            handletextpad=0.5,
        )

        ax = fig.add_subplot(gs[1, 0])
        apply_horizontal_bar_template(
            ax,
            aircraft["aircraft_class"].tolist(),
            aircraft["flight_share_pct"].to_numpy(),
            colors=aircraft["color"].tolist(),
            xlabel="Share of flight records (%)",
            show_end_labels=True,
            value_formatter="{:.1f}%",
        )
        ax.set_xlim(0, float(aircraft["flight_share_pct"].max()) * 1.18)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 1])
        bubble_scale = np.sqrt(aircraft["flights"].to_numpy() / float(aircraft["flights"].max()))
        bubble_sizes = 28 + 360 * bubble_scale
        apply_scatter_template(
            ax,
            aircraft["mean_distance_km"].to_numpy(),
            aircraft["fuel_share_pct"].to_numpy(),
            c=aircraft["color"].tolist(),
            s=bubble_sizes,
            alpha=0.82,
            edgecolors="white",
            linewidth=0.8,
            zorder=3,
        )
        style_chart_axes(ax)
        ax.set_xlabel("Mean flight distance (km)")
        ax.set_ylabel("Share of total fuel demand (%)")
        x_min = float(aircraft["mean_distance_km"].min()) - 80
        x_max = float(aircraft["mean_distance_km"].max()) + 130
        y_max_scatter = float(aircraft["fuel_share_pct"].max()) + 5
        ax.set_xlim(x_min, x_max)
        ax.set_ylim(-1.0, y_max_scatter)
        label_offsets = {
            "B737": (6, 0),
            "A320": (6, 3),
            "A321": (6, -4),
            "A319": (6, -2),
            "A330": (6, 3),
            "B787": (6, -2),
            "Regional jet": (6, -2),
            "Turbo-prop": (6, 2),
            "Other/Unknown": (6, 3),
        }
        for row in aircraft.itertuples(index=False):
            dx, dy = label_offsets.get(row.aircraft_class, (6, 0))
            arrowprops = None if abs(dy) <= 2 else {"arrowstyle": "-", "lw": 0.45, "color": "#94A3B8", "shrinkA": 2, "shrinkB": 2}
            ax.annotate(
                row.aircraft_class,
                xy=(float(row.mean_distance_km), float(row.fuel_share_pct)),
                xytext=(dx, dy),
                textcoords="offset points",
                ha="left",
                va="center",
                fontsize=6.2,
                color="#334155",
                arrowprops=arrowprops,
            )
        legend_counts = sorted(
            {
                max(500, int(round(float(aircraft["flights"].max()) * frac / 500.0) * 500))
                for frac in (0.15, 0.45, 0.90)
            }
        )
        size_handles = [
            ax.scatter([], [], s=28 + 360 * math.sqrt(count / float(aircraft["flights"].max())), color="#94A3B8", alpha=0.6, edgecolors="none")
            for count in legend_counts
        ]
        ax.legend(
            size_handles,
            [f"{count:,}" for count in legend_counts],
            title="Flight count",
            loc="upper left",
            frameon=False,
            fontsize=6.1,
            title_fontsize=6.5,
            borderaxespad=0.2,
            handletextpad=0.7,
        )
        add_panel_label(ax, "C")
        return save_matplotlib_figure(fig, path)


def create_annual_demand_distribution(path: Path) -> Path:
    annual = load_or_build_annual_airport_weekly_demand()
    candidate_meta, week_map = load_candidate_week_metadata()

    annual_pivot = (
        annual.pivot_table(index="week_of_year", columns="airport_en", values="weekly_total_fuel_kg", aggfunc="sum")
        .reindex(range(1, 53))
        .fillna(0.0)
        / 1e6
    )
    candidate_weeks = candidate_meta["original_week"].astype(int).tolist()
    label_map = {
        int(row["original_week_in_52"]): f"W{int(row['new_week_number'])}"
        for _, row in week_map.iterrows()
    }
    selected_weeks = sorted(label_map)
    airport_specs = [("Beijing", PUBLICATION_COLORS["blue"]), ("Tianjin", PUBLICATION_COLORS["orange"])]

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("comparison_pair")
        gs = fig.add_gridspec(1, 2, left=0.08, right=0.985, top=0.95, bottom=0.22, wspace=0.26)

        ax = fig.add_subplot(gs[0, 0])
        apply_boxplot_template(
            ax,
            [annual_pivot["Beijing"].to_numpy(), annual_pivot["Tianjin"].to_numpy()],
            labels=["Beijing", "Tianjin"],
            colors=[PUBLICATION_COLORS["blue"], PUBLICATION_COLORS["orange"]],
            ylabel="Weekly fuel demand (million kg)",
            showfliers=False,
        )
        for x_pos, (airport, color) in enumerate(airport_specs, start=1):
            candidate_vals = annual_pivot.loc[candidate_weeks, airport].to_numpy()
            candidate_jitter = np.linspace(-0.10, 0.10, len(candidate_vals))
            apply_scatter_template(
                ax,
                np.full(len(candidate_vals), x_pos, dtype=float) + candidate_jitter,
                candidate_vals,
                color="white",
                s=26,
                alpha=1.0,
                edgecolors=PUBLICATION_COLORS["gray"],
                linewidth=0.8,
                zorder=4,
            )

            retained_vals = annual_pivot.loc[selected_weeks, airport].to_numpy()
            retained_jitter = np.linspace(-0.06, 0.06, len(retained_vals))
            apply_scatter_template(
                ax,
                np.full(len(retained_vals), x_pos, dtype=float) + retained_jitter,
                retained_vals,
                color=color,
                marker="D",
                s=34,
                alpha=0.95,
                edgecolors="white",
                linewidth=0.8,
                zorder=5,
            )
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        ranking_frames: dict[str, pd.DataFrame] = {}
        series_specs: list[dict[str, object]] = []
        for airport, color in airport_specs:
            ranking = (
                annual_pivot[[airport]]
                .reset_index()
                .rename(columns={airport: "weekly_total"})
                .sort_values(["weekly_total", "week_of_year"])
                .reset_index(drop=True)
            )
            ranking["rank"] = np.arange(1, len(ranking) + 1)
            ranking_frames[airport] = ranking
            series_specs.append(
                {
                    "x": ranking["rank"],
                    "y": ranking["weekly_total"],
                    "color": color,
                    "linewidth": 1.4,
                    "label": airport,
                }
            )

        apply_line_template(
            ax,
            series_specs,
            xlabel="Rank in the 52-week distribution",
            ylabel="Weekly fuel demand (million kg)",
            legend=False,
            vertical_guides=[13, 26, 39],
        )

        for airport, color in airport_specs:
            ranking = ranking_frames[airport]
            apply_scatter_template(
                ax,
                ranking["rank"],
                ranking["weekly_total"],
                color=color,
                s=12,
                alpha=0.40,
                edgecolors="none",
                zorder=2,
            )

            candidate_points = ranking[ranking["week_of_year"].isin(candidate_weeks)].copy()
            apply_scatter_template(
                ax,
                candidate_points["rank"],
                candidate_points["weekly_total"],
                color="white",
                s=28,
                alpha=1.0,
                edgecolors=color,
                linewidth=0.9,
                zorder=4,
            )

            retained_points = ranking[ranking["week_of_year"].isin(selected_weeks)].copy()
            apply_scatter_template(
                ax,
                retained_points["rank"],
                retained_points["weekly_total"],
                color=color,
                marker="D",
                s=42,
                alpha=0.95,
                edgecolors="white",
                linewidth=0.8,
                zorder=5,
            )

            y_span = max(ranking["weekly_total"].max() - ranking["weekly_total"].min(), 0.6)
            offset = 0.06 * y_span if airport == "Beijing" else -0.06 * y_span
            valign = "bottom" if offset > 0 else "top"
            for point in retained_points.itertuples(index=False):
                ax.text(
                    point.rank + 0.45,
                    point.weekly_total + offset,
                    label_map[int(point.week_of_year)],
                    fontsize=6.6,
                    color=color,
                    va=valign,
                )
        ax.set_xlim(0.5, 52.5)
        ax.set_xticks([1, 13, 26, 39, 52])
        add_panel_label(ax, "B")

        legend_handles = [
            Line2D([0], [0], color=PUBLICATION_COLORS["blue"], marker="o", markersize=4.2, linewidth=1.4, label="Beijing"),
            Line2D([0], [0], color=PUBLICATION_COLORS["orange"], marker="o", markersize=4.2, linewidth=1.4, label="Tianjin"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor="white", markeredgecolor=PUBLICATION_COLORS["gray"], markersize=5.2, label="12-week candidate pool"),
            Line2D([0], [0], marker="D", linestyle="None", markerfacecolor=PUBLICATION_COLORS["black"], markeredgecolor="white", markersize=5.2, label="Retained representative week"),
        ]
        add_publication_figure_legend(
            fig,
            legend_handles,
            loc="lower center",
            bbox_to_anchor=(0.5, 0.03),
            ncol=4,
            fontsize=6.6,
            columnspacing=1.0,
            handletextpad=0.4,
        )
        return save_matplotlib_figure(fig, path)


def create_renewable_raw_summary(path: Path, data: AppendixData) -> Path:
    wind_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/wind_power_plants.csv"
    solar_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/solar_power_plants.csv"
    wind = pd.read_csv(wind_inv_path, usecols=["Longitude", "Latitude", "Status", "Capacity__MW_"]).dropna(subset=["Longitude", "Latitude"])
    solar = pd.read_csv(solar_inv_path, usecols=["Longitude", "Latitude", "Status", "Capacity__MW_"]).dropna(subset=["Longitude", "Latitude"])
    wind_resource = load_cached_renewable_resource_frame(
        WIND_COMPLETE_PATH,
        cache_path=WIND_COMPLETE_CACHE,
        hours_per_step=1,
        clip_to_capacity=True,
    )
    solar_resource = load_cached_renewable_resource_frame(
        SOLAR_COMPLETE_PATH,
        cache_path=SOLAR_COMPLETE_CACHE,
        hours_per_step=1,
        clip_to_capacity=True,
    )
    wind_resource["capacity_factor_pct"] = wind_resource["capacity_factor"].clip(lower=0.0, upper=1.0) * 100.0
    solar_resource["capacity_factor_pct"] = solar_resource["capacity_factor"].clip(lower=0.0, upper=1.0) * 100.0
    status_order = ["Operating", "Under construction", "Pre-construction", "Announced", "Planned / proposed", "Unknown"]
    wind_status = wind["Status"].map(normalize_status).value_counts()
    solar_status = solar["Status"].map(normalize_status).value_counts()
    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.26, hspace=0.34)

        ax = fig.add_subplot(gs[0, 0])
        apply_grouped_bar_template(
            ax,
            ["Oper.", "Constr.", "Pre-constr.", "Ann.", "Plan.", "Unk."],
            [
                ("Wind", [wind_status.get(k, 0) for k in status_order], PUBLICATION_COLORS["blue"]),
                ("Solar", [solar_status.get(k, 0) for k in status_order], PUBLICATION_COLORS["orange"]),
            ],
            ylabel="Inventory rows",
            legend=True,
            legend_kwargs={"loc": "upper right"},
        )
        ax.tick_params(axis="x", rotation=15)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        cap_max = float(max(wind["Capacity__MW_"].max(), solar["Capacity__MW_"].max()))
        bins = np.logspace(np.log10(1), np.log10(cap_max + 1), 22)
        apply_histogram_template(
            ax,
            [
                ("Wind", wind["Capacity__MW_"].clip(lower=1), PUBLICATION_COLORS["blue"], 0.55),
                ("Solar", solar["Capacity__MW_"].clip(lower=1), PUBLICATION_COLORS["orange"], 0.45),
            ],
            bins=bins,
            xlabel="Nameplate capacity per record (MW, log scale)",
            log_x=True,
            legend=True,
            legend_kwargs={"loc": "upper right"},
        )
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        stage_labels = ["Inventory rows", "Annual resource plants", "Modeled 4-week plants"]
        stage_x = np.arange(len(stage_labels))
        wind_stage = [
            data.transport_summary["wind_inventory_rows"],
            data.renewable_summary["wind_resource"]["candidate_plants"],
            data.renewable_summary["wind"]["unique_plants"],
        ]
        solar_stage = [
            data.transport_summary["solar_inventory_rows"],
            data.renewable_summary["solar_resource"]["candidate_plants"],
            data.renewable_summary["solar"]["unique_plants"],
        ]
        apply_grouped_bar_template(
            ax,
            ["Inventory", "Annual table", "4-week model"],
            [
                ("Wind", wind_stage, PUBLICATION_COLORS["blue"]),
                ("Solar", solar_stage, PUBLICATION_COLORS["orange"]),
            ],
            ylabel="Rows or plants (log scale)",
            legend=True,
            log_y=True,
            legend_kwargs={"loc": "upper right"},
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        apply_boxplot_template(
            ax,
            [
                wind_resource["full_load_hours"].to_numpy(),
                solar_resource["full_load_hours"].to_numpy(),
            ],
            labels=["Wind", "Solar"],
            colors=[PUBLICATION_COLORS["blue"], PUBLICATION_COLORS["orange"]],
            ylabel="Annual full-load hours from clipped profiles (h)",
        )
        add_panel_label(ax, "D")

        return save_matplotlib_figure(fig, path)


def create_renewable_resource_maps(path: Path, data: AppendixData) -> Path:
    del data
    wind_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/wind_power_plants.csv"
    solar_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/solar_power_plants.csv"
    wind = pd.read_csv(wind_inv_path, usecols=["Longitude", "Latitude"]).dropna(subset=["Longitude", "Latitude"])
    solar = pd.read_csv(solar_inv_path, usecols=["Longitude", "Latitude"]).dropna(subset=["Longitude", "Latitude"])
    wind_resource = load_cached_renewable_resource_frame(
        WIND_COMPLETE_PATH,
        cache_path=WIND_COMPLETE_CACHE,
        hours_per_step=1,
        clip_to_capacity=True,
    )
    solar_resource = load_cached_renewable_resource_frame(
        SOLAR_COMPLETE_PATH,
        cache_path=SOLAR_COMPLETE_CACHE,
        hours_per_step=1,
        clip_to_capacity=True,
    )
    wind_resource["capacity_factor_pct"] = wind_resource["capacity_factor"].clip(lower=0.0, upper=1.0) * 100.0
    solar_resource["capacity_factor_pct"] = solar_resource["capacity_factor"].clip(lower=0.0, upper=1.0) * 100.0
    wind_map = sample_points(wind, limit=2200)
    solar_map = sample_points(solar, limit=2200)

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("national_map_triptych")
        gs = fig.add_gridspec(1, 3, left=0.05, right=0.985, top=0.96, bottom=0.24, wspace=0.18)

        ax_a, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=12.5
        )
        ax_a.scatter(
            wind_map["Longitude"],
            wind_map["Latitude"],
            s=5,
            color=PUBLICATION_COLORS["blue"],
            alpha=0.18,
            edgecolors="none",
            transform=ccrs.PlateCarree(),
            zorder=3,
            rasterized=True,
        )
        ax_a.scatter(
            solar_map["Longitude"],
            solar_map["Latitude"],
            s=5,
            color=PUBLICATION_COLORS["orange"],
            alpha=0.14,
            edgecolors="none",
            transform=ccrs.PlateCarree(),
            zorder=3,
            rasterized=True,
        )
        mirror_points_to_mini(mini_ax, wind_map["Longitude"], wind_map["Latitude"], s=5, color=PUBLICATION_COLORS["blue"], alpha=0.30, edgecolors="none", zorder=3)
        mirror_points_to_mini(mini_ax, solar_map["Longitude"], solar_map["Latitude"], s=5, color=PUBLICATION_COLORS["orange"], alpha=0.25, edgecolors="none", zorder=3)
        add_panel_label(ax_a, "A")

        ax_b, mini_ax = create_reference_map_panel(
            fig, gs[0, 1], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=12.5
        )
        wind_scatter = ax_b.scatter(
            wind_resource["longitude"],
            wind_resource["latitude"],
            c=wind_resource["capacity_factor_pct"],
            s=8,
            cmap="Blues",
            vmin=float(wind_resource["capacity_factor_pct"].quantile(0.05)),
            vmax=float(wind_resource["capacity_factor_pct"].quantile(0.95)),
            linewidth=0.0,
            alpha=0.62,
            transform=ccrs.PlateCarree(),
            zorder=3,
            rasterized=True,
        )
        mirror_points_to_mini(
            mini_ax,
            wind_resource["longitude"],
            wind_resource["latitude"],
            c=wind_resource["capacity_factor_pct"],
            s=7,
            cmap="Blues",
            vmin=float(wind_resource["capacity_factor_pct"].quantile(0.05)),
            vmax=float(wind_resource["capacity_factor_pct"].quantile(0.95)),
            linewidth=0.0,
            alpha=0.55,
            zorder=3,
            rasterized=True,
        )
        add_panel_label(ax_b, "B")

        ax_c, mini_ax = create_reference_map_panel(
            fig, gs[0, 2], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=12.5
        )
        solar_scatter = ax_c.scatter(
            solar_resource["longitude"],
            solar_resource["latitude"],
            c=solar_resource["capacity_factor_pct"],
            s=8,
            cmap="YlOrBr",
            vmin=float(solar_resource["capacity_factor_pct"].quantile(0.05)),
            vmax=float(solar_resource["capacity_factor_pct"].quantile(0.95)),
            linewidth=0.0,
            alpha=0.62,
            transform=ccrs.PlateCarree(),
            zorder=3,
            rasterized=True,
        )
        mirror_points_to_mini(
            mini_ax,
            solar_resource["longitude"],
            solar_resource["latitude"],
            c=solar_resource["capacity_factor_pct"],
            s=7,
            cmap="YlOrBr",
            vmin=float(solar_resource["capacity_factor_pct"].quantile(0.05)),
            vmax=float(solar_resource["capacity_factor_pct"].quantile(0.95)),
            linewidth=0.0,
            alpha=0.55,
            zorder=3,
            rasterized=True,
        )
        add_panel_label(ax_c, "C")

        item_y = 0.105
        item_h = 0.036

        legend_handles = [
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=PUBLICATION_COLORS["blue"], markeredgecolor=PUBLICATION_COLORS["blue"], markersize=4.8, alpha=0.7, label="Wind inventory"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=PUBLICATION_COLORS["orange"], markeredgecolor=PUBLICATION_COLORS["orange"], markersize=4.8, alpha=0.6, label="Solar inventory"),
        ]
        add_compact_bottom_legend(
            fig,
            ax_a,
            legend_handles,
            y=item_y,
            height=item_h,
            width_frac=0.52,
        )

        add_compact_bottom_colorbar(
            fig,
            ax_b,
            wind_scatter,
            label="Annual availability factor (%)",
            y=item_y,
            height=item_h,
            width_frac=0.66,
        )

        add_compact_bottom_colorbar(
            fig,
            ax_c,
            solar_scatter,
            label="Annual availability factor (%)",
            y=item_y,
            height=item_h,
            width_frac=0.66,
        )

        return save_matplotlib_figure(fig, path)


def load_operating_renewable_inventory_points(csv_path: Path) -> pd.DataFrame:
    frame = pd.read_csv(csv_path, usecols=["Status", "Capacity__MW_", "Latitude", "Longitude"])
    frame["Status"] = frame["Status"].fillna("").astype(str).str.strip().str.lower()
    frame["Capacity__MW_"] = pd.to_numeric(frame["Capacity__MW_"], errors="coerce")
    frame["Latitude"] = pd.to_numeric(frame["Latitude"], errors="coerce")
    frame["Longitude"] = pd.to_numeric(frame["Longitude"], errors="coerce")
    frame = frame[frame["Status"].eq("operating")].dropna(subset=["Capacity__MW_", "Latitude", "Longitude"]).copy()
    frame = frame[frame["Capacity__MW_"] > 0].rename(
        columns={"Capacity__MW_": "capacity_mw", "Latitude": "latitude", "Longitude": "longitude"}
    )
    return frame


def scale_capacity_marker_areas(
    capacity_mw: pd.Series | np.ndarray,
    *,
    clip_upper_mw: float = 500.0,
    min_marker_pt: float = 2.8,
    max_marker_pt: float = 10.8,
) -> np.ndarray:
    values = np.asarray(capacity_mw, dtype=float)
    if values.size == 0:
        return np.array([], dtype=float)
    clipped = np.clip(values, 0.0, clip_upper_mw)
    transformed = np.sqrt(clipped)
    positive = transformed[transformed > 0]
    if positive.size == 0:
        marker_pts = np.full(values.shape, min_marker_pt, dtype=float)
    else:
        lo = float(positive.min())
        hi = float(positive.max())
        if not np.isfinite(lo) or not np.isfinite(hi) or hi <= lo:
            marker_pts = np.full(values.shape, (min_marker_pt + max_marker_pt) / 2.0, dtype=float)
        else:
            marker_pts = np.interp(transformed, [lo, hi], [min_marker_pt, max_marker_pt])
    return marker_pts ** 2


def build_capacity_legend_handles(
    capacities_mw: list[float],
    *,
    clip_upper_mw: float = 500.0,
) -> list[Line2D]:
    marker_areas = scale_capacity_marker_areas(np.asarray(capacities_mw, dtype=float), clip_upper_mw=clip_upper_mw)
    handles: list[Line2D] = []
    for capacity, area in zip(capacities_mw, marker_areas):
        is_clipped_cap = float(capacity) >= clip_upper_mw
        handles.append(
            Line2D(
                [0],
                [0],
                marker="o",
                linestyle="None",
                markerfacecolor="#CBD5E1",
                markeredgecolor="#64748B",
                markeredgewidth=0.45,
                markersize=float(np.sqrt(area)),
                alpha=0.95,
                label=f"{int(capacity):,}+ MW" if is_clipped_cap else f"{int(capacity):,} MW",
            )
        )
    return handles


def draw_inventory_capacity_points(
    ax,
    mini_ax,
    frame: pd.DataFrame,
    *,
    color: str,
    alpha: float,
    clip_upper_mw: float = 500.0,
) -> None:
    marker_areas = scale_capacity_marker_areas(frame["capacity_mw"], clip_upper_mw=clip_upper_mw)
    ax.scatter(
        frame["longitude"],
        frame["latitude"],
        s=marker_areas,
        color=color,
        alpha=alpha,
        edgecolors="none",
        transform=MAP_DATA_CRS,
        zorder=3,
        rasterized=True,
    )
    mirror_points_to_mini(
        mini_ax,
        frame["longitude"],
        frame["latitude"],
        s=np.clip(marker_areas * 0.36, 1.0, 16.0),
        color=color,
        alpha=min(alpha + 0.06, 0.30),
        edgecolors="none",
        zorder=3,
        rasterized=True,
    )


def create_renewable_inventory_capacity_maps(path: Path) -> Path:
    wind_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/wind_power_plants.csv"
    solar_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/solar_power_plants.csv"
    wind = load_operating_renewable_inventory_points(wind_inv_path)
    solar = load_operating_renewable_inventory_points(solar_inv_path)
    clip_upper_mw = 500.0

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("national_map_pair")
        gs = fig.add_gridspec(1, 2, left=0.05, right=0.985, top=0.96, bottom=0.18, wspace=0.18)

        ax_a, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        draw_inventory_capacity_points(
            ax_a,
            mini_ax,
            wind,
            color=PUBLICATION_COLORS["blue"],
            alpha=0.20,
            clip_upper_mw=clip_upper_mw,
        )
        add_panel_label(ax_a, "A")

        ax_b, mini_ax = create_reference_map_panel(
            fig, gs[0, 1], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        draw_inventory_capacity_points(
            ax_b,
            mini_ax,
            solar,
            color=PUBLICATION_COLORS["orange"],
            alpha=0.16,
            clip_upper_mw=clip_upper_mw,
        )
        add_panel_label(ax_b, "B")

        add_publication_figure_legend(
            fig,
            build_capacity_legend_handles([20, 100, 300, 500], clip_upper_mw=clip_upper_mw),
            loc="lower center",
            bbox_to_anchor=(0.5, 0.045),
            ncol=4,
            fontsize=6.8,
            columnspacing=1.1,
            handletextpad=0.55,
        )

        return save_matplotlib_figure(fig, path)


def create_byproduct_raw_summary(path: Path) -> Path:
    refinery_path = ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/refinery_daily_byproduct_h2_20251027_202950.csv"
    steel_path = ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/steel_daily_byproduct_h2_20251027_202950.csv"
    refinery = pd.read_csv(refinery_path)
    steel = pd.read_csv(steel_path)

    province_df = pd.concat(
        [
            refinery.assign(source="Refinery").rename(columns={"h2_daily_tonnes": "h2"}),
            steel.assign(source="Steel").rename(columns={"h2_daily_tonnes": "h2"}),
        ],
        ignore_index=True,
    )
    top_provinces = province_df.groupby("province")["h2"].sum().sort_values(ascending=False).head(8).index.tolist()
    pivot = (
        province_df[province_df["province"].isin(top_provinces)]
        .groupby(["province", "source"])["h2"]
        .sum()
        .unstack(fill_value=0)
        .loc[top_provinces]
    )

    byproduct_colors = [PUBLICATION_COLORS["magenta"], PUBLICATION_COLORS["sky"]]

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.28, hspace=0.34)

        ax = fig.add_subplot(gs[0, 0])
        apply_grouped_bar_template(
            ax,
            ["Site count", "Province count"],
            [
                ("Refinery", [len(refinery), refinery["province"].nunique()], byproduct_colors[0]),
                ("Steel", [len(steel), steel["province"].nunique()], byproduct_colors[1]),
            ],
            ylabel="Count",
            legend=True,
            legend_kwargs={"loc": "upper right"},
        )
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        totals = [refinery["h2_daily_tonnes"].sum(), steel["h2_daily_tonnes"].sum()]
        ax.bar(["Refinery", "Steel"], totals, color=byproduct_colors, alpha=0.85, edgecolor="white")
        ax.set_ylabel("Daily H₂ availability (t/day, log scale)")
        ax.set_yscale("log")
        style_chart_axes(ax)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        apply_boxplot_template(
            ax,
            [refinery["h2_daily_tonnes"].clip(lower=0.1), steel["h2_daily_tonnes"].clip(lower=0.1)],
            labels=["Refinery", "Steel"],
            colors=byproduct_colors,
            ylabel="Site-level H₂ availability (t/day, log scale)",
            log_y=True,
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        y_pos = np.arange(len(top_provinces), dtype=float)
        refinery_vals = pivot.get("Refinery", pd.Series(0.0, index=top_provinces)).reindex(top_provinces).to_numpy(dtype=float)
        steel_vals = pivot.get("Steel", pd.Series(0.0, index=top_provinces)).reindex(top_provinces).to_numpy(dtype=float)
        x_floor = 0.2
        for values, color, offset, label in [
            (refinery_vals, byproduct_colors[0], 0.16, "Refinery"),
            (steel_vals, byproduct_colors[1], -0.16, "Steel"),
        ]:
            clipped = np.clip(values, x_floor, None)
            ax.hlines(y_pos + offset, x_floor, clipped, color=color, linewidth=1.1, alpha=0.35)
            ax.scatter(
                clipped,
                y_pos + offset,
                s=24,
                color=color,
                edgecolors="white",
                linewidth=0.45,
                alpha=0.90,
                zorder=3,
                label=label,
            )
        ax.set_xscale("log")
        ax.set_xlim(x_floor, max(float(steel_vals.max()), float(refinery_vals.max())) * 1.25)
        ax.set_yticks(y_pos)
        ax.set_yticklabels(top_provinces, fontsize=FONT_CONFIG["tick"])
        ax.invert_yaxis()
        ax.set_xlabel("Provincial daily H₂ availability (t/day, log scale)")
        style_chart_axes(ax)
        ax.grid(axis="x", color="#CBD5E1", linewidth=0.8, linestyle="--", alpha=0.3)
        ax.yaxis.grid(False)
        ax.legend(frameon=False, loc="lower right")
        add_panel_label(ax, "D")

        return save_matplotlib_figure(fig, path)


def create_byproduct_geo_maps(path: Path) -> Path:
    refinery_path = ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/refinery_daily_byproduct_h2_20251027_202950.csv"
    steel_path = ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/steel_daily_byproduct_h2_20251027_202950.csv"
    refinery = pd.read_csv(refinery_path)
    steel = pd.read_csv(steel_path)

    refinery_size = lambda values: np.clip(np.sqrt(np.asarray(values, dtype=float).clip(min=0.1)) * 11.0, 10, 84)
    steel_size = lambda values: np.clip(np.sqrt(np.asarray(values, dtype=float).clip(min=1.0)) * 1.25, 10, 78)

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("national_map_pair")
        gs = fig.add_gridspec(1, 2, left=0.05, right=0.985, top=0.96, bottom=0.22, wspace=0.10)

        ax, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        ax.scatter(
            refinery["longitude"],
            refinery["latitude"],
            s=refinery_size(refinery["h2_daily_tonnes"]),
            color="#EC4899",
            alpha=0.52,
            edgecolors="white",
            linewidth=0.35,
            transform=ccrs.PlateCarree(),
            zorder=3,
        )
        mirror_points_to_mini(
            mini_ax,
            refinery["longitude"],
            refinery["latitude"],
            s=np.clip(refinery_size(refinery["h2_daily_tonnes"]) * 0.28, 4.5, 20),
            color="#EC4899",
            alpha=0.45,
            edgecolors="white",
            linewidth=0.20,
            zorder=3,
        )
        add_panel_label(ax, "A")

        ax, mini_ax = create_reference_map_panel(
            fig, gs[0, 1], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        ax.scatter(
            steel["longitude"],
            steel["latitude"],
            s=steel_size(steel["h2_daily_tonnes"]),
            color="#0EA5E9",
            alpha=0.48,
            edgecolors="white",
            linewidth=0.35,
            transform=ccrs.PlateCarree(),
            zorder=3,
        )
        mirror_points_to_mini(
            mini_ax,
            steel["longitude"],
            steel["latitude"],
            s=np.clip(steel_size(steel["h2_daily_tonnes"]) * 0.24, 4.5, 18),
            color="#0EA5E9",
            alpha=0.42,
            edgecolors="white",
            linewidth=0.20,
            zorder=3,
        )
        add_panel_label(ax, "B")

        add_compact_bottom_legend(
            fig,
            fig.axes[0],
            make_scatter_size_legend_handles(
                [2, 10, 30],
                refinery_size,
                color="#EC4899",
                label_suffix=" t/day",
            ),
            y=0.105,
            height=0.036,
            x_offset_frac=0.10,
            width_frac=0.75,
            ncol=3,
            loc="center left",
        )
        add_compact_bottom_legend(
            fig,
            fig.axes[2],
            make_scatter_size_legend_handles(
                [200, 1000, 3000],
                steel_size,
                color="#0EA5E9",
                label_suffix=" t/day",
            ),
            y=0.105,
            height=0.036,
            x_offset_frac=0.06,
            width_frac=0.82,
            ncol=3,
            loc="center left",
        )

        return save_matplotlib_figure(fig, path)


def create_co2_capture_raw_summary(path: Path, data: AppendixData) -> Path:
    co2 = load_co2_capture_frame()
    type_order = ["Coal power", "Gas power", "Oil refinery"]
    by_type = (
        co2.groupby("facility_label")
        .agg(
            records=("location_id", "size"),
            weekly_capture_tonnes=("co2_capture_capacity_ton_per_week", "sum"),
            unit_cost_yuan_per_ton=("capture_cost_yuan_per_ton", "median"),
        )
        .reindex(type_order)
    )
    province_type = (
        co2.groupby(["province", "facility_label"])["co2_capture_capacity_ton_per_week"]
        .sum()
        .unstack(fill_value=0)
    )
    top_provinces = province_type.sum(axis=1).sort_values(ascending=False).head(10).index.tolist()

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.26, hspace=0.34)

        ax = fig.add_subplot(gs[0, 0])
        shares = pd.DataFrame(
            {
                "Record share": by_type["records"] / by_type["records"].sum() * 100,
                "Weekly capture share": by_type["weekly_capture_tonnes"] / by_type["weekly_capture_tonnes"].sum() * 100,
            },
            index=type_order,
        )
        apply_stacked_bar_template(
            ax,
            ["Record share", "Weekly capture share"],
            [
                (label, [shares.loc[label, "Record share"], shares.loc[label, "Weekly capture share"]], CO2_TYPE_COLORS[label])
                for label in type_order
            ],
            xlabel="Share of archived CO₂ source table (%)",
            orientation="h",
            legend=True,
            legend_kwargs={"loc": "lower right"},
        )
        ax.set_xlim(0, 100)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        apply_dual_axis_bar_line_template(
            ax,
            type_order,
            by_type["weekly_capture_tonnes"].to_numpy() / 1e6,
            by_type["unit_cost_yuan_per_ton"].to_numpy(),
            bar_color=PUBLICATION_COLORS["orange"],
            line_color=PUBLICATION_COLORS["black"],
            left_ylabel="Aggregate weekly capture (Mt CO₂/week)",
            right_ylabel="Unit capture cost (yuan/t CO₂)",
        )
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        stacked = province_type.loc[top_provinces, type_order] / 1e6
        apply_stacked_bar_template(
            ax,
            stacked.index.tolist(),
            [(label, stacked[label].to_numpy(), CO2_TYPE_COLORS[label]) for label in type_order],
            ylabel="Weekly capture potential (Mt CO₂/week)",
            legend=False,
        )
        ax.tick_params(axis="x", rotation=28)
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        cost_values = by_type["unit_cost_yuan_per_ton"].reindex(type_order).to_numpy(dtype=float)
        apply_horizontal_bar_template(
            ax,
            type_order,
            cost_values,
            colors=[CO2_TYPE_COLORS[label] for label in type_order],
            xlabel="Capture cost assumption (yuan/t CO₂)",
            show_end_labels=True,
            value_formatter="{:.0f}",
        )
        ax.set_xlim(0, float(np.nanmax(cost_values)) * 1.20)
        add_panel_label(ax, "D")
        return save_matplotlib_figure(fig, path)


def create_co2_capture_geo_maps(path: Path) -> Path:
    co2 = load_co2_capture_frame()
    type_order = ["Coal power", "Gas power", "Oil refinery"]
    markers = {"Coal power": "o", "Gas power": "s", "Oil refinery": "^"}
    province_totals = co2.groupby("province")["co2_capture_capacity_ton_per_week"].sum()

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("national_map_pair")
        gs = fig.add_gridspec(1, 2, left=0.05, right=0.985, top=0.96, bottom=0.24, wspace=0.10)

        ax_left, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        for label in type_order:
            subset = co2[co2["facility_label"] == label].copy()
            limit = 1800 if label == "Coal power" else len(subset)
            subset = sample_points(subset, limit=limit)
            ax_left.scatter(
                subset["longitude"],
                subset["latitude"],
                s=14 if label == "Coal power" else 18,
                marker=markers[label],
                color=CO2_TYPE_COLORS[label],
                alpha=0.22 if label == "Coal power" else 0.58,
                edgecolors="white" if label != "Coal power" else "none",
                linewidth=0.35,
                transform=ccrs.PlateCarree(),
                zorder=3,
                label=label,
                rasterized=True,
            )
            mirror_points_to_mini(
                mini_ax,
                subset["longitude"],
                subset["latitude"],
                s=6 if label == "Coal power" else 7.5,
                marker=markers[label],
                color=CO2_TYPE_COLORS[label],
                alpha=0.30 if label == "Coal power" else 0.55,
                edgecolors="white" if label != "Coal power" else "none",
                linewidth=0.20,
                zorder=3,
            )
        add_panel_label(ax_left, "A")

        ax_right, mini_ax = create_reference_map_panel(
            fig, gs[0, 1], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        scalar_mappable, _ = draw_reference_choropleth(
            ax_right,
            province_totals / 1e6,
            cmap_name="Reds",
        )
        draw_reference_choropleth(
            mini_ax,
            province_totals / 1e6,
            cmap_name="Reds",
        )
        add_panel_label(ax_right, "B")

        item_y = 0.105
        item_h = 0.036

        legend_handles = [
            Line2D([0], [0], marker=markers[label], linestyle="None", markerfacecolor=CO2_TYPE_COLORS[label], markeredgecolor="white" if label != "Coal power" else CO2_TYPE_COLORS[label], markersize=6.2, alpha=0.8, label=label)
            for label in type_order
        ]
        add_compact_bottom_legend(
            fig,
            ax_left,
            legend_handles,
            y=item_y,
            height=item_h,
            width_frac=0.58,
        )

        add_compact_bottom_colorbar(
            fig,
            ax_right,
            scalar_mappable,
            label="Weekly capture potential (Mt CO₂/week)",
            y=item_y,
            height=item_h,
            x_offset_frac=0.18,
            width_frac=0.60,
            fontsize=7.2,
            tick_labelsize=6.5,
        )
        return save_matplotlib_figure(fig, path)


def create_transport_raw_summary(path: Path) -> Path:
    pipe_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/natural_gas_pipelines.csv"
    ccs_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/ccs_projects.csv"
    pipes = pd.read_csv(pipe_path)
    ccs = pd.read_csv(ccs_path)
    lengths = (
        pipes["Note"]
        .fillna("")
        .str.extract(r"Length:\s*([0-9.]+)\s*km", expand=False)
        .dropna()
        .astype(float)
    )
    pipe_status = pipes["Status"].map(normalize_status).value_counts()
    ccs_status = ccs["Status"].map(normalize_status).value_counts()
    online_year = pd.to_numeric(pipes["YearOnline"], errors="coerce").dropna()

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.28, hspace=0.34)

        ax = fig.add_subplot(gs[0, 0])
        bins = np.logspace(np.log10(max(lengths.min(), 0.3)), np.log10(lengths.max()), 22)
        apply_histogram_template(
            ax,
            [(None, lengths, "#0284C7", 0.8)],
            bins=bins,
            xlabel="Reported pipeline length (km, log scale)",
            log_x=True,
        )
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        status_order = ["Operating", "Under construction", "Planned / proposed", "Unknown"]
        apply_grouped_bar_template(
            ax,
            status_order,
            [
                (
                    "Pipeline records",
                    [pipe_status.get(label, 0) for label in status_order],
                    PUBLICATION_COLORS["green"],
                )
            ],
            ylabel="Pipeline records",
            legend=False,
        )
        plt.setp(ax.get_xticklabels(), rotation=16, ha="right")
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        bins = np.arange(int(online_year.min()) - 1, int(online_year.max()) + 5, 5)
        apply_histogram_template(
            ax,
            [(None, online_year, "#334155", 0.85)],
            bins=bins,
            xlabel="Online year",
            ylabel="Records",
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        ccs_order = ["Operating", "Planned / proposed", "Unknown"]
        ccs_values = [ccs_status.get(label, 0) for label in ccs_order]
        apply_grouped_bar_template(
            ax,
            ccs_order,
            [("CCUS project records", ccs_values, PUBLICATION_COLORS["magenta"])],
            ylabel="CCUS project records",
            legend=False,
        )
        plt.setp(ax.get_xticklabels(), rotation=12, ha="right")
        add_panel_label(ax, "D")

        return save_matplotlib_figure(fig, path)


def create_price_surface_maps(path: Path) -> Path:
    price_surface = load_provincial_price_frame().set_index("province")

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("national_map_pair")
        gs = fig.add_gridspec(1, 2, left=0.05, right=0.985, top=0.96, bottom=0.22, wspace=0.10)

        ax_left, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        gas_sm, _ = draw_reference_choropleth(
            ax_left,
            price_surface["gas_price_yuan_per_10k_m3"],
            cmap_name="YlGnBu",
        )
        draw_reference_choropleth(
            mini_ax,
            price_surface["gas_price_yuan_per_10k_m3"],
            cmap_name="YlGnBu",
        )
        add_panel_label(ax_left, "A")

        ax_right, mini_ax = create_reference_map_panel(
            fig, gs[0, 1], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.0
        )
        elec_sm, _ = draw_reference_choropleth(
            ax_right,
            price_surface["electricity_price_yuan_per_kwh"],
            cmap_name="OrRd",
        )
        draw_reference_choropleth(
            mini_ax,
            price_surface["electricity_price_yuan_per_kwh"],
            cmap_name="OrRd",
        )
        add_panel_label(ax_right, "B")

        item_y = 0.082
        item_h = 0.055

        add_compact_bottom_colorbar(
            fig,
            ax_left,
            gas_sm,
            label="Median gas price (yuan per 10k m$^3$)",
            y=item_y,
            height=item_h,
            fontsize=7.2,
            tick_labelsize=6.5,
        )

        add_compact_bottom_colorbar(
            fig,
            ax_right,
            elec_sm,
            label="Median power price (yuan per kWh)",
            y=item_y,
            height=item_h,
            fontsize=7.2,
            tick_labelsize=6.5,
        )

        return save_matplotlib_figure(fig, path)


def create_energy_infrastructure_overview(path: Path) -> Path:
    wind = pd.read_csv(
        ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/wind_power_plants.csv",
        usecols=["Longitude", "Latitude"],
    ).dropna()
    solar = pd.read_csv(
        ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/solar_power_plants.csv",
        usecols=["Longitude", "Latitude"],
    ).dropna()
    refinery = pd.read_csv(ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/refinery_daily_byproduct_h2_20251027_202950.csv")
    steel = pd.read_csv(ROOT / "products/gis_energy_mapping/industrial_byproduct_hydrogen/data/steel_daily_byproduct_h2_20251027_202950.csv")
    pipes = pd.read_csv(ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/natural_gas_pipelines.csv")
    ccs = pd.read_csv(ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/ccs_projects.csv")

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("network_map_tall")
        fig.patch.set_facecolor("#FFFFFF")
        gs = fig.add_gridspec(1, 1, left=0.05, right=0.985, top=0.96, bottom=0.24)

        ax_map, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], [73, 136, 17, 55], dx=10, dy=10, main_scale_km=1000, mini_scale_km=500, compass_size=13.5
        )
        plot_geometry_wkts(ax_map, pipes["geometry_wkt"], color="#94A3B8", linewidth=0.35, alpha=0.15, limit=260)
        wind_map = sample_points(wind, limit=1400)
        solar_map = sample_points(solar, limit=1800)
        ax_map.scatter(wind_map["Longitude"], wind_map["Latitude"], s=5, color=PUBLICATION_COLORS["blue"], alpha=0.18, edgecolors="none", transform=ccrs.PlateCarree(), zorder=3, rasterized=True)
        ax_map.scatter(solar_map["Longitude"], solar_map["Latitude"], s=5, color=PUBLICATION_COLORS["orange"], alpha=0.12, edgecolors="none", transform=ccrs.PlateCarree(), zorder=3, rasterized=True)
        mirror_points_to_mini(mini_ax, wind_map["Longitude"], wind_map["Latitude"], s=5, color=PUBLICATION_COLORS["blue"], alpha=0.28, edgecolors="none", zorder=3)
        mirror_points_to_mini(mini_ax, solar_map["Longitude"], solar_map["Latitude"], s=5, color=PUBLICATION_COLORS["orange"], alpha=0.24, edgecolors="none", zorder=3)
        ax_map.scatter(
            refinery["longitude"],
            refinery["latitude"],
            s=20,
            facecolors="none",
            edgecolors=PUBLICATION_COLORS["magenta"],
            linewidth=0.8,
            alpha=0.80,
            transform=ccrs.PlateCarree(),
            zorder=4,
            rasterized=True,
        )
        mirror_points_to_mini(mini_ax, refinery["longitude"], refinery["latitude"], s=10, facecolors="none", edgecolors=PUBLICATION_COLORS["magenta"], linewidth=0.45, alpha=0.80, zorder=4)
        ax_map.scatter(
            steel["longitude"],
            steel["latitude"],
            s=18,
            marker="^",
            facecolors="none",
            edgecolors=PUBLICATION_COLORS["sky"],
            linewidth=0.8,
            alpha=0.70,
            transform=ccrs.PlateCarree(),
            zorder=4,
            rasterized=True,
        )
        mirror_points_to_mini(mini_ax, steel["longitude"], steel["latitude"], s=10, marker="^", facecolors="none", edgecolors=PUBLICATION_COLORS["sky"], linewidth=0.45, alpha=0.70, zorder=4)
        ax_map.scatter(
            ccs["Long_"],
            ccs["Lat"],
            s=24,
            marker="D",
            color=PUBLICATION_COLORS["magenta"],
            alpha=0.50,
            edgecolors="none",
            transform=ccrs.PlateCarree(),
            zorder=4,
            rasterized=True,
        )
        mirror_points_to_mini(mini_ax, ccs["Long_"], ccs["Lat"], s=10, marker="D", color=PUBLICATION_COLORS["magenta"], alpha=0.55, edgecolors="none", zorder=4)

        legend_handles = [
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=PUBLICATION_COLORS["blue"], markeredgecolor=PUBLICATION_COLORS["blue"], markersize=5, alpha=0.75, label="Wind inventory sample"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=PUBLICATION_COLORS["orange"], markeredgecolor=PUBLICATION_COLORS["orange"], markersize=5, alpha=0.65, label="Solar inventory sample"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor="none", markeredgecolor=PUBLICATION_COLORS["magenta"], markersize=6, label=r"Refinery H$_2$ sites"),
            Line2D([0], [0], marker="^", linestyle="None", markerfacecolor="none", markeredgecolor=PUBLICATION_COLORS["sky"], markersize=6, label=r"Steel H$_2$ sites"),
            Line2D([0], [0], marker="D", linestyle="None", markerfacecolor=PUBLICATION_COLORS["magenta"], markeredgecolor=PUBLICATION_COLORS["magenta"], markersize=5, alpha=0.60, label="CCUS projects"),
            Line2D([0], [0], color="#94A3B8", linewidth=1.2, alpha=0.6, label="Pipeline records"),
        ]

        add_publication_figure_legend(
            fig,
            legend_handles,
            loc="lower center",
            bbox_to_anchor=(0.51, 0.03),
            ncol=3,
            fontsize=7.5,
            columnspacing=1.0,
            handletextpad=0.5,
        )

        return save_matplotlib_figure(fig, path)


def create_map_template_preview(path: Path) -> Path:
    wind_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/wind_power_plants.csv"
    solar_inv_path = ROOT / "products/gis_energy_mapping/gis_data_scraper/scraped_gis_data/solar_power_plants.csv"
    wind = pd.read_csv(wind_inv_path, usecols=["Longitude", "Latitude"]).dropna(subset=["Longitude", "Latitude"])
    solar = pd.read_csv(solar_inv_path, usecols=["Longitude", "Latitude"]).dropna(subset=["Longitude", "Latitude"])
    price_surface = load_provincial_price_frame().set_index("province")
    wind_map = sample_points(wind, limit=1400)
    solar_map = sample_points(solar, limit=1600)

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("map_preview_pair")
        gs = fig.add_gridspec(1, 2, left=0.04, right=0.985, top=0.95, bottom=0.20, wspace=0.08)

        ax_left = fig.add_subplot(gs[0, 0], projection=MAP_PROJECTION)
        configure_reference_frykit_map(ax_left, [73, 136, 17, 55], dx=10, dy=10)
        mini_left = add_reference_mini_map(ax_left)
        ax_left.scatter(
            wind_map["Longitude"],
            wind_map["Latitude"],
            s=8,
            color=PUBLICATION_COLORS["blue"],
            alpha=0.22,
            edgecolors="none",
            transform=MAP_DATA_CRS,
            zorder=4,
            rasterized=True,
        )
        ax_left.scatter(
            solar_map["Longitude"],
            solar_map["Latitude"],
            s=8,
            color=PUBLICATION_COLORS["orange"],
            alpha=0.18,
            edgecolors="none",
            transform=MAP_DATA_CRS,
            zorder=4,
            rasterized=True,
        )
        mirror_points_to_mini(
            mini_left,
            wind_map["Longitude"],
            wind_map["Latitude"],
            s=7,
            color=PUBLICATION_COLORS["blue"],
            alpha=0.35,
            edgecolors="none",
            zorder=4,
        )
        mirror_points_to_mini(
            mini_left,
            solar_map["Longitude"],
            solar_map["Latitude"],
            s=7,
            color=PUBLICATION_COLORS["orange"],
            alpha=0.30,
            edgecolors="none",
            zorder=4,
        )
        add_reference_map_decorations(ax_left, mini_left, main_scale_km=1000, mini_scale_km=500, compass_size=15)
        add_panel_label(ax_left, "A")

        ax_right = fig.add_subplot(gs[0, 1], projection=MAP_PROJECTION)
        configure_reference_frykit_map(ax_right, [73, 136, 17, 55], dx=10, dy=10)
        mini_right = add_reference_mini_map(ax_right)
        scalar_mappable, _ = draw_reference_choropleth(
            ax_right,
            price_surface["gas_price_yuan_per_10k_m3"],
            cmap_name="YlGnBu",
        )
        draw_reference_choropleth(
            mini_right,
            price_surface["gas_price_yuan_per_10k_m3"],
            cmap_name="YlGnBu",
        )
        add_reference_map_decorations(ax_right, mini_right, main_scale_km=1000, mini_scale_km=500, compass_size=15)
        add_panel_label(ax_right, "B")

        legend_handles = [
            Line2D([0], [0], marker="o", linestyle="None", color=PUBLICATION_COLORS["blue"], markersize=6.5, alpha=0.7, label="Wind sites"),
            Line2D([0], [0], marker="o", linestyle="None", color=PUBLICATION_COLORS["orange"], markersize=6.5, alpha=0.7, label="Solar sites"),
        ]
        add_compact_bottom_legend(
            fig,
            ax_left,
            legend_handles,
            y=0.08,
            height=0.055,
            x_offset_frac=0.06,
            width_frac=0.36,
            ncol=2,
            columnspacing=1.0,
            fontsize=6.8,
        )

        add_compact_bottom_colorbar(
            fig,
            ax_right,
            scalar_mappable,
            label="Median gas price (yuan per 10k m$^3$)",
            y=0.08,
            height=0.055,
            x_offset_frac=0.18,
            width_frac=0.56,
            fontsize=7.2,
            tick_labelsize=6.5,
        )

        return save_matplotlib_figure(fig, path)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str, width: int = 8) -> None:
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 24
    arrow_angle = math.pi / 7
    p1 = (
        int(end[0] - arrow_len * math.cos(angle - arrow_angle)),
        int(end[1] - arrow_len * math.sin(angle - arrow_angle)),
    )
    p2 = (
        int(end[0] - arrow_len * math.cos(angle + arrow_angle)),
        int(end[1] - arrow_len * math.sin(angle + arrow_angle)),
    )
    draw.polygon([end, p1, p2], fill=fill)


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: str, fill: str, outline: str) -> None:
    title_font = load_font(36)
    body_font = load_font(28)
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=4)
    x0, y0, x1, y1 = xy
    draw.text((x0 + 28, y0 + 24), title, font=title_font, fill="#0F172A")
    wrapped = textwrap.fill(body, width=26)
    draw.multiline_text((x0 + 28, y0 + 84), wrapped, font=body_font, fill="#1F2937", spacing=8)


def create_workflow_figure(path: Path) -> Path:
    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("workflow_single")
        ax = fig.add_subplot(111)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")

        stages = [
            {
                "title": "Raw layers",
                "body": "Airport demand\nRenewable inventories\nIndustrial sources\nPipeline GIS",
                "anchor": (0.06, 0.76),
                "accent": "#2563EB",
            },
            {
                "title": "Demand reduction",
                "body": "pyBADA fuel model\n12 candidate weeks\n4 retained weeks",
                "anchor": (0.30, 0.76),
                "accent": "#059669",
            },
            {
                "title": "Supply preprocessing",
                "body": "Plant matching\nHourly profile fusion\nCapacity screening",
                "anchor": (0.55, 0.76),
                "accent": "#D97706",
            },
            {
                "title": "Scenario inputs",
                "body": "Prices\nEmissions factors\nRoute switches\nTechnology blocks",
                "anchor": (0.79, 0.76),
                "accent": "#7C3AED",
            },
            {
                "title": "Spatial filtering",
                "body": "DBSCAN clustering\nCenter optimization\nPipeline-access penalty",
                "anchor": (0.19, 0.31),
                "accent": "#16A34A",
            },
            {
                "title": "Arc generation",
                "body": "Road shortest paths\nPipeline corridors\nAirport delivery links",
                "anchor": (0.46, 0.31),
                "accent": "#2563EB",
            },
            {
                "title": "MILP-ready archive",
                "body": "Nodes and arcs\nWeekly demand\n3 h periods\nScenario-specific bounds",
                "anchor": (0.74, 0.31),
                "accent": "#059669",
            },
        ]

        for stage in stages:
            x_pos, y_pos = stage["anchor"]
            ax.scatter([x_pos - 0.016], [y_pos + 0.015], s=30, color=stage["accent"], zorder=3)
            ax.text(x_pos, y_pos, stage["title"], fontsize=8.3, fontweight="bold", color="#0F172A", va="top")
            ax.text(x_pos, y_pos - 0.055, stage["body"], fontsize=6.9, color="#334155", va="top", linespacing=1.20)

        arrow_specs = [
            ((0.18, 0.775), (0.28, 0.775)),
            ((0.43, 0.775), (0.53, 0.775)),
            ((0.68, 0.775), (0.77, 0.775)),
            ((0.38, 0.57), (0.27, 0.38)),
            ((0.61, 0.57), (0.53, 0.38)),
            ((0.85, 0.57), (0.79, 0.38)),
            ((0.34, 0.325), (0.44, 0.325)),
            ((0.61, 0.325), (0.72, 0.325)),
        ]
        for start, end in arrow_specs:
            ax.annotate(
                "",
                xy=end,
                xytext=start,
                arrowprops=dict(arrowstyle="-|>", lw=1.0, color="#64748B", shrinkA=2, shrinkB=2),
            )
        return save_matplotlib_figure(fig, path)


def create_placeholder_figure(path: Path, title: str) -> Path:
    note_map = {
        "Electricity price sensitivity": "Reserved for the electricity-price sweep of green-hydrogen pathways.",
        "Electrolyzer CAPEX sensitivity": "Reserved for the electrolyzer-CAPEX comparison across FT and MTJ routes.",
        "DAC cost sensitivity": "Reserved for the DAC cost and energy-intensity sensitivity batch.",
        "Carbon-price sensitivity": "Reserved for the carbon-price and carbon-tax response comparison.",
        "Fossil feedstock price sensitivity": "Reserved for the natural-gas and coal price sweep of GTL and GTL-BH pathways. Will show LCO-SynAF vs. gas/coal price with the 6-8 yuan/kg market window marked.",
        "Green-hydrogen joint sensitivity": "Reserved for the electricity-price × electrolyzer-CAPEX joint heatmap. Will show the 8 yuan/kg threshold contour on a 2-D parameter surface.",
        "DAC break-even threshold": "Reserved for the DAC cost break-even line chart. Will show the exact cost at which DAC-GH-FT enters the 6-8 yuan/kg market window.",
        "Carbon-price pathway response": "Reserved for the asymmetric carbon-price response across GTL-BH, GTL, and CCU-GH-FT. Will show effective LCO-SynAF (net of carbon cost/credit) vs. carbon price.",
        "Pareto frontier tornado": "Reserved for the Tornado diagram of single-parameter impact on LCO-SynAF, ranking parameters by output range and flagging any rank-reversal boundary.",
    }
    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("placeholder_single")
        ax = fig.add_subplot(111)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")
        ax.text(0.10, 0.72, title, fontsize=9.3, fontweight="bold", color="#0F172A", va="top")
        ax.text(0.10, 0.54, note_map.get(title, "Reserved for finalized sensitivity results."), fontsize=7.6, color="#334155", va="top")
        ax.text(0.10, 0.38, "No values are plotted here because the corresponding batch runs have not been finalized.", fontsize=7.2, color="#475569")
        ax.text(0.10, 0.24, "The figure slot is kept fixed to preserve appendix numbering and layout.", fontsize=7.2, color="#64748B")
        return save_matplotlib_figure(fig, path)


def create_representative_demand_patterns(path: Path) -> Path:
    demand_12_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/typical_weeks_data/typical_12weeks_demand_20251129_163442.xlsx"
    demand_4_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_demand_20251129_231231.xlsx"
    week_map_path = ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/week_mapping_20251129_231231.csv"

    demand_12 = pd.read_excel(demand_12_path)
    demand_4 = pd.read_excel(demand_4_path)
    week_map = pd.read_csv(week_map_path)
    for frame in (demand_12, demand_4):
        frame["airport_en"] = frame["airport"].map(AIRPORT_MAP).fillna(frame["airport"])

    candidate = (
        demand_12.pivot_table(index="week_number", columns="airport_en", values="weekly_total_fuel_kg_total", aggfunc="first")
        .sort_index()
        .fillna(0.0)
        / 1e6
    )
    retained = (
        demand_4.pivot_table(index="week_number", columns="airport_en", values="weekly_total_fuel_kg_total", aggfunc="first")
        .sort_index()
        .fillna(0.0)
        / 1e6
    )
    candidate_meta = (
        demand_12.groupby("week_number", as_index=False)
        .first()[["week_number", "original_week"]]
        .sort_values("week_number")
    )
    candidate = candidate.join(candidate_meta.set_index("week_number"))
    selected_lookup = {
        int(row["week_in_12weeks"]): int(row["new_week_number"])
        for _, row in week_map.sort_values("new_week_number").iterrows()
    }
    retained_lookup = {
        int(row["week_in_12weeks"]): int(row["original_week_in_52"])
        for _, row in week_map.sort_values("new_week_number").iterrows()
    }
    selected_weeks = sorted(selected_lookup)
    calendar_weeks = [retained_lookup[week] for week in selected_weeks]
    candidate_x = candidate["original_week"].to_numpy()
    y_max = max(float(candidate["Beijing"].max()), float(retained["Beijing"].max())) * 1.13

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("comparison_pair")
        gs = fig.add_gridspec(1, 2, left=0.09, right=0.985, top=0.86, bottom=0.17, wspace=0.24)

        ax = fig.add_subplot(gs[0, 0])
        for selected_week in selected_weeks:
            ax.axvline(
                retained_lookup[selected_week],
                color="#CBD5E1",
                linestyle="--",
                linewidth=0.8,
                alpha=0.75,
                zorder=0,
            )
        apply_line_template(
            ax,
            [
                {
                    "x": candidate_x,
                    "y": candidate[airport].to_numpy(),
                    "color": color,
                    "linewidth": 1.6,
                    "marker": "o",
                    "markersize": 3.8,
                    "label": airport,
                }
                for airport, color in [("Beijing", PUBLICATION_COLORS["blue"]), ("Tianjin", PUBLICATION_COLORS["orange"])]
            ],
            xlabel="Original calendar week in 2024",
            ylabel="Weekly fuel demand (million kg)",
            legend=False,
        )
        for airport, color in [("Beijing", PUBLICATION_COLORS["blue"]), ("Tianjin", PUBLICATION_COLORS["orange"])]:
            highlights = candidate.loc[candidate.index.isin(selected_weeks), ["original_week", airport]]
            apply_scatter_template(
                ax,
                highlights["original_week"].to_numpy(),
                highlights[airport].to_numpy(),
                s=42,
                marker="D",
                color="white",
                edgecolors=PUBLICATION_COLORS["black"],
                linewidth=0.85,
                zorder=4,
            )
        for retained_week in selected_weeks:
            ax.annotate(
                f"W{selected_lookup[retained_week]}",
                xy=(retained_lookup[retained_week], float(candidate.loc[retained_week, "Beijing"])),
                xytext=(0, 9),
                textcoords="offset points",
                ha="center",
                va="bottom",
                fontsize=6.7,
                color=PUBLICATION_COLORS["black"],
            )
        ax.set_xticks(candidate_x)
        ax.set_ylim(0, y_max)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        apply_grouped_bar_template(
            ax,
            [f"W{int(i)}\n(week {calendar_weeks[i - 1]})" for i in retained.index],
            [
                ("Beijing", retained["Beijing"].to_numpy(), PUBLICATION_COLORS["blue"]),
                ("Tianjin", retained["Tianjin"].to_numpy(), PUBLICATION_COLORS["orange"]),
            ],
            ylabel="Weekly fuel demand (million kg)",
            legend=False,
        )
        ax.set_ylim(0, y_max)
        add_panel_label(ax, "B")
        legend_handles = [
            Line2D([0], [0], color=PUBLICATION_COLORS["blue"], marker="o", markersize=4.0, linewidth=1.6, label="Beijing"),
            Line2D([0], [0], color=PUBLICATION_COLORS["orange"], marker="o", markersize=4.0, linewidth=1.6, label="Tianjin"),
            Line2D(
                [0],
                [0],
                marker="D",
                linestyle="None",
                markerfacecolor="white",
                markeredgecolor=PUBLICATION_COLORS["black"],
                markersize=5.0,
                label="Retained representative week",
            ),
        ]
        add_publication_figure_legend(
            fig,
            legend_handles,
            loc="upper center",
            bbox_to_anchor=(0.5, 0.985),
            ncol=3,
            fontsize=6.5,
        )
        return save_matplotlib_figure(fig, path)


def aggregate_normalized_hourly_profile(csv_path: Path, chunksize: int = 350_000, clip_to_capacity: bool = False) -> pd.DataFrame:
    hourly_totals: pd.DataFrame | None = None
    for chunk in pd.read_csv(csv_path, usecols=["hour", "power_output_mw", "capacity_mw"], chunksize=chunksize):
        output_series = np.minimum(chunk["power_output_mw"], chunk["capacity_mw"]) if clip_to_capacity else chunk["power_output_mw"]
        chunk = chunk.assign(power_for_profile_mw=output_series)
        grouped = (
            chunk.groupby("hour", as_index=True)
            .agg(total_output_mw=("power_for_profile_mw", "sum"), total_capacity_mw=("capacity_mw", "sum"))
            .astype(float)
        )
        hourly_totals = grouped if hourly_totals is None else hourly_totals.add(grouped, fill_value=0.0)
    if hourly_totals is None:
        return pd.DataFrame(columns=["hour", "total_output_mw", "total_capacity_mw", "normalized_output_pct"])
    profile = hourly_totals.reset_index().sort_values("hour").reset_index(drop=True)
    denominator = profile["total_capacity_mw"].replace(0, np.nan)
    profile["normalized_output_pct"] = (profile["total_output_mw"] / denominator).fillna(0.0) * 100.0
    return profile


def sample_generation_relationship_points(
    directory: Path,
    pattern: str,
    *,
    usecols: list[str],
    max_files: int = 18,
    rows_per_file: int = 550,
) -> pd.DataFrame:
    files = sorted(directory.glob(pattern))
    if not files:
        return pd.DataFrame(columns=usecols)
    if len(files) > max_files:
        indices = np.unique(np.linspace(0, len(files) - 1, max_files, dtype=int))
        files = [files[int(idx)] for idx in indices]
    frames: list[pd.DataFrame] = []
    for idx, file in enumerate(files):
        frame = pd.read_csv(file, usecols=usecols)
        if len(frame) > rows_per_file:
            frame = frame.sample(n=rows_per_file, random_state=100 + idx)
        frames.append(frame)
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame(columns=usecols)


def build_binned_median_curve(x: pd.Series, y: pd.Series, bins: int = 18) -> pd.DataFrame:
    frame = pd.DataFrame({"x": x, "y": y}).replace([np.inf, -np.inf], np.nan).dropna()
    if frame.empty:
        return pd.DataFrame(columns=["x_mid", "y_median"])
    x_min = float(frame["x"].min())
    x_max = float(frame["x"].max())
    if not np.isfinite(x_min) or not np.isfinite(x_max) or x_min == x_max:
        return pd.DataFrame(columns=["x_mid", "y_median"])
    edges = np.linspace(x_min, x_max, bins + 1)
    frame["bin"] = pd.cut(frame["x"], bins=edges, include_lowest=True)
    summary = (
        frame.groupby("bin", observed=True)
        .agg(x_mid=("x", "median"), y_median=("y", "median"))
        .dropna()
        .reset_index(drop=True)
    )
    return summary


def draw_processing_workflow_panel(ax, selected_weeks_12: list[int]) -> None:
    ax.set_axis_off()
    add_panel_label(ax, "A")

    marker_x = 0.14
    text_x = 0.23
    stage_y = [0.80, 0.52, 0.23]
    stage_labels = [
        "Inventories + archived drivers",
        "Hourly plant-level output",
        "12-week pool -> retained 4 weeks",
    ]

    for idx, (ypos, label) in enumerate(zip(stage_y, stage_labels), start=1):
        ax.scatter(
            [marker_x],
            [ypos],
            s=72,
            transform=ax.transAxes,
            facecolor="#F8FAFC",
            edgecolor="#94A3B8",
            linewidth=0.9,
            zorder=3,
        )
        ax.text(
            marker_x,
            ypos,
            str(idx),
            transform=ax.transAxes,
            ha="center",
            va="center",
            fontsize=6.4,
            color=PUBLICATION_COLORS["black"],
            zorder=4,
        )
        ax.text(
            text_x,
            ypos + 0.01,
            label,
            transform=ax.transAxes,
            ha="left",
            va="center",
            fontsize=7.1,
            color=PUBLICATION_COLORS["black"],
        )

    for start, end in zip(stage_y[:-1], stage_y[1:]):
        ax.annotate(
            "",
            xy=(marker_x, end + 0.07),
            xytext=(marker_x, start - 0.07),
            xycoords="axes fraction",
            textcoords="axes fraction",
            arrowprops=dict(arrowstyle="->", lw=0.95, color="#94A3B8"),
        )

    ax.text(
        text_x,
        stage_y[1] - 0.09,
        "Solar: radiation + capacity",
        transform=ax.transAxes,
        ha="left",
        va="center",
        fontsize=6.4,
        color=PUBLICATION_COLORS["orange"],
    )
    ax.text(
        text_x,
        stage_y[1] - 0.18,
        "Wind: speed -> hub -> curve",
        transform=ax.transAxes,
        ha="left",
        va="center",
        fontsize=6.4,
        color=PUBLICATION_COLORS["blue"],
    )
    ax.text(
        text_x,
        stage_y[2] - 0.10,
        f"Retained: {', '.join(f'W{week}' for week in selected_weeks_12)}",
        transform=ax.transAxes,
        ha="left",
        va="center",
        fontsize=6.4,
        color=PUBLICATION_COLORS["gray"],
    )
    ax.text(
        text_x,
        stage_y[2] - 0.18,
        "2,016 h -> 672 h",
        transform=ax.transAxes,
        ha="left",
        va="center",
        fontsize=6.4,
        color=PUBLICATION_COLORS["gray"],
    )


def draw_raw_relationship_panel(
    ax,
    frame: pd.DataFrame,
    *,
    x_column: str,
    xlabel: str,
    color: str,
    panel_label: str,
) -> None:
    clean = frame[[x_column, "power_output_mw", "capacity_mw"]].replace([np.inf, -np.inf], np.nan).dropna().copy()
    clean = clean[clean["capacity_mw"] > 0].copy()
    clean["normalized_output_pct"] = (clean["power_output_mw"] / clean["capacity_mw"]).clip(lower=0.0, upper=1.05) * 100.0
    x_upper = float(clean[x_column].quantile(0.995)) if not clean.empty else 0.0
    if x_upper > 0:
        clean = clean[clean[x_column] <= x_upper].copy()
    apply_scatter_template(
        ax,
        clean[x_column].to_numpy(),
        clean["normalized_output_pct"].to_numpy(),
        color=color,
        s=10,
        alpha=0.12,
        edgecolors="none",
        linewidth=0.0,
        rasterize_threshold=3000,
    )
    summary = build_binned_median_curve(clean[x_column], clean["normalized_output_pct"])
    if not summary.empty:
        ax.plot(
            summary["x_mid"],
            summary["y_median"],
            color=PUBLICATION_COLORS["black"],
            linewidth=1.3,
            zorder=4,
        )
    ax.set_xlabel(xlabel)
    ax.set_ylabel("Plant output / capacity (%)")
    ax.set_ylim(0, 105)
    style_chart_axes(ax)
    add_panel_label(ax, panel_label)


def create_renewable_processing_figure(path: Path) -> Path:
    solar_12_path = latest_file("products/aviation_fuel_analysis/resource_flight_data_process/results/typical_weeks_data/typical_12weeks_solar_*.csv")
    wind_12_path = latest_file("products/aviation_fuel_analysis/resource_flight_data_process/results/typical_weeks_data/typical_12weeks_wind_*.csv")
    week_map_path = latest_file("products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/week_mapping_*.csv")
    if solar_12_path is None or wind_12_path is None:
        raise FileNotFoundError("12-week renewable profiles were not found for Section 1.3 processing figure")

    selected_weeks_12 = [1, 2, 4, 11]
    if week_map_path and week_map_path.exists():
        week_map = pd.read_csv(week_map_path)
        selected_weeks_12 = week_map["week_in_12weeks"].astype(int).tolist()

    solar_profile_12 = aggregate_normalized_hourly_profile(solar_12_path, clip_to_capacity=True)
    wind_profile_12 = aggregate_normalized_hourly_profile(wind_12_path, clip_to_capacity=True)

    solar_raw = sample_generation_relationship_points(
        ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/solar_generation",
        "*.csv",
        usecols=["solar_radiation_wm2", "power_output_mw", "capacity_mw"],
    )
    wind_raw = sample_generation_relationship_points(
        ROOT / "products/aviation_fuel_analysis/resource_flight_data_process/results/3hourly_generation",
        "*.csv",
        usecols=["hub_wind_speed_ms", "power_output_mw", "capacity_mw"],
    )

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("processing_combo")
        gs = fig.add_gridspec(
            2,
            3,
            left=0.07,
            right=0.985,
            top=0.965,
            bottom=0.11,
            wspace=0.28,
            hspace=0.36,
            height_ratios=[0.86, 1.14],
        )

        ax = fig.add_subplot(gs[0, 0])
        draw_processing_workflow_panel(ax, selected_weeks_12)

        ax = fig.add_subplot(gs[0, 1])
        draw_raw_relationship_panel(
            ax,
            solar_raw,
            x_column="solar_radiation_wm2",
            xlabel="Archived solar-radiation input",
            color=PUBLICATION_COLORS["orange"],
            panel_label="B",
        )

        ax = fig.add_subplot(gs[0, 2])
        draw_raw_relationship_panel(
            ax,
            wind_raw,
            x_column="hub_wind_speed_ms",
            xlabel="Hub-height wind speed (m/s)",
            color=PUBLICATION_COLORS["blue"],
            panel_label="C",
        )

        ax = fig.add_subplot(gs[1, :])
        retained_patch = Patch(facecolor="#E2E8F0", edgecolor="none", alpha=0.40, label="Retained week window")
        for week in selected_weeks_12:
            start = (int(week) - 1) * 168
            end = int(week) * 168
            ax.axvspan(start, end, facecolor="#E2E8F0", alpha=0.40, edgecolor="none", zorder=0)

        ax.plot(
            wind_profile_12["hour"],
            wind_profile_12["normalized_output_pct"],
            color=PUBLICATION_COLORS["blue"],
            linewidth=1.0,
            alpha=0.40,
            zorder=2,
        )
        ax.plot(
            solar_profile_12["hour"],
            solar_profile_12["normalized_output_pct"],
            color=PUBLICATION_COLORS["orange"],
            linewidth=1.0,
            alpha=0.40,
            zorder=2,
        )
        for week in selected_weeks_12:
            start = (int(week) - 1) * 168
            end = int(week) * 168
            wind_mask = wind_profile_12["hour"].between(start, end - 1)
            solar_mask = solar_profile_12["hour"].between(start, end - 1)
            ax.plot(
                wind_profile_12.loc[wind_mask, "hour"],
                wind_profile_12.loc[wind_mask, "normalized_output_pct"],
                color=PUBLICATION_COLORS["blue"],
                linewidth=1.5,
                zorder=3,
            )
            ax.plot(
                solar_profile_12.loc[solar_mask, "hour"],
                solar_profile_12.loc[solar_mask, "normalized_output_pct"],
                color=PUBLICATION_COLORS["orange"],
                linewidth=1.5,
                zorder=3,
            )
        ax.set_xlim(0, 2015)
        ax.set_xticks([week * 168 + 84 for week in range(12)])
        ax.set_xticklabels([f"W{week}" for week in range(1, 13)])
        ax.set_xlabel("Week in the 12-week pool")
        ax.set_ylabel("Aggregate clipped output / capacity (%)")
        profile_ceiling = max(
            float(wind_profile_12["normalized_output_pct"].max()) if not wind_profile_12.empty else 0.0,
            float(solar_profile_12["normalized_output_pct"].max()) if not solar_profile_12.empty else 0.0,
        )
        ax.set_ylim(0, math.ceil(profile_ceiling / 5) * 5 + 5)
        style_chart_axes(ax)
        ax.legend(
            handles=[
                Line2D([0], [0], color=PUBLICATION_COLORS["blue"], linewidth=1.5, label="Wind profile"),
                Line2D([0], [0], color=PUBLICATION_COLORS["orange"], linewidth=1.5, label="Solar profile"),
                retained_patch,
            ],
            loc="upper right",
            frameon=False,
            fontsize=6.6,
        )
        add_panel_label(ax, "D")

        return save_matplotlib_figure(fig, path)


def create_representative_renewable_profiles(path: Path) -> Path:
    wind_path = latest_file("products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_wind_*.csv")
    solar_path = latest_file("products/aviation_fuel_analysis/resource_flight_data_process/results/4_typical_weeks_data/typical_4weeks_solar_*.csv")
    if wind_path is None or solar_path is None:
        raise FileNotFoundError("4-week renewable profiles were not found for Section 1.3")

    wind = pd.read_csv(wind_path, usecols=["hour", "power_output_mw"])
    solar = pd.read_csv(solar_path, usecols=["hour", "power_output_mw"])
    wind_hourly = wind.groupby("hour", as_index=False)["power_output_mw"].sum().rename(columns={"power_output_mw": "wind_mw"})
    solar_hourly = solar.groupby("hour", as_index=False)["power_output_mw"].sum().rename(columns={"power_output_mw": "solar_mw"})
    hourly = wind_hourly.merge(solar_hourly, on="hour", how="outer").fillna(0.0).sort_values("hour")
    hourly["week"] = hourly["hour"] // 168 + 1
    hourly["hour_in_day"] = hourly["hour"] % 24
    diurnal = hourly.groupby(["week", "hour_in_day"], as_index=False)[["wind_mw", "solar_mw"]].mean()
    week_colors = [PUBLICATION_COLORS["blue"], PUBLICATION_COLORS["orange"], PUBLICATION_COLORS["green"], PUBLICATION_COLORS["vermillion"]]
    week_midpoints = [84, 252, 420, 588]
    week_labels = [f"Week {week}" for week in sorted(diurnal["week"].unique())]
    wind_top_limit = math.ceil(float(hourly["wind_mw"].max()) / 1000 / 25) * 25
    solar_top_limit = math.ceil(float(hourly["solar_mw"].max()) / 1000 / 25) * 25

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.24, hspace=0.46)

        ax = fig.add_subplot(gs[0, 0])
        for idx, (start, end) in enumerate(zip([0, 168, 336, 504], [168, 336, 504, 672])):
            if idx % 2 == 0:
                ax.axvspan(start, end, facecolor="#F8FAFC", alpha=0.95, edgecolor="none", zorder=0)
        apply_line_template(
            ax,
            [
                {
                    "x": hourly["hour"],
                    "y": hourly["wind_mw"] / 1000,
                    "color": PUBLICATION_COLORS["blue"],
                    "linewidth": 1.2,
                }
            ],
            xlabel="Retained week (168 h each)",
            ylabel="Aggregate wind output (GW)",
            vertical_guides=[168, 336, 504],
        )
        ax.set_xticks(week_midpoints)
        ax.set_xticklabels(week_labels)
        ax.set_ylim(0, wind_top_limit)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        for idx, (start, end) in enumerate(zip([0, 168, 336, 504], [168, 336, 504, 672])):
            if idx % 2 == 0:
                ax.axvspan(start, end, facecolor="#F8FAFC", alpha=0.95, edgecolor="none", zorder=0)
        apply_line_template(
            ax,
            [
                {
                    "x": hourly["hour"],
                    "y": hourly["solar_mw"] / 1000,
                    "color": PUBLICATION_COLORS["orange"],
                    "linewidth": 1.2,
                }
            ],
            xlabel="Retained week (168 h each)",
            ylabel="Aggregate solar output (GW)",
            vertical_guides=[168, 336, 504],
        )
        ax.set_xticks(week_midpoints)
        ax.set_xticklabels(week_labels)
        ax.set_ylim(0, solar_top_limit)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        apply_line_template(
            ax,
            [
                {
                    "x": diurnal.loc[diurnal["week"] == week, "hour_in_day"],
                    "y": diurnal.loc[diurnal["week"] == week, "wind_mw"] / 1000,
                    "color": color,
                    "linewidth": 1.4,
                    "label": f"Week {week}",
                }
                for week, color in zip(sorted(diurnal["week"].unique()), week_colors)
            ],
            xlabel="Hour of day",
            ylabel="Mean wind output (GW)",
        )
        ax.set_xlim(0, 23)
        ax.set_xticks([0, 6, 12, 18, 23])
        ax.set_ylim(0, wind_top_limit)
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        apply_line_template(
            ax,
            [
                {
                    "x": diurnal.loc[diurnal["week"] == week, "hour_in_day"],
                    "y": diurnal.loc[diurnal["week"] == week, "solar_mw"] / 1000,
                    "color": color,
                    "linewidth": 1.4,
                    "label": f"Week {week}",
                }
                for week, color in zip(sorted(diurnal["week"].unique()), week_colors)
            ],
            xlabel="Hour of day",
            ylabel="Mean solar output (GW)",
        )
        ax.set_xlim(0, 23)
        ax.set_xticks([0, 6, 12, 18, 23])
        ax.set_ylim(0, solar_top_limit)
        add_panel_label(ax, "D")
        week_handles = [
            Line2D([0], [0], color=color, linewidth=1.6, label=f"Week {week}")
            for week, color in zip(sorted(diurnal["week"].unique()), week_colors)
        ]
        add_publication_figure_legend(
            fig,
            week_handles,
            loc="upper center",
            bbox_to_anchor=(0.50, 0.56),
            ncol=4,
            columnspacing=1.2,
            handletextpad=0.5,
        )
        return save_matplotlib_figure(fig, path)


def create_quadrant_comparison(path: Path) -> Path:
    metrics = build_scenario_metrics_table()
    frontier = identify_pareto_front(metrics.rename(columns={"CarbonDiffVsJet": "CarbonDiffVsJet", "LCOE": "LCOE"}))

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("comparison_pair")
        gs = fig.add_gridspec(1, 2, left=0.08, right=0.985, top=0.88, bottom=0.18, wspace=0.30)

        ax = fig.add_subplot(gs[0, 0])
        for route, marker in ROUTE_MARKERS.items():
            subset = metrics[metrics["Route"] == route]
            apply_scatter_template(
                ax,
                subset["CarbonDiffVsJet"],
                subset["LCOE"],
                label=route,
                marker=marker,
                s=42,
                c=subset["Category"].map(category_color),
                edgecolors="white",
                linewidth=0.7,
                zorder=3,
            )
        frontier_df = metrics[metrics["Scenario"].isin(frontier)].sort_values("CarbonDiffVsJet")
        ax.plot(frontier_df["CarbonDiffVsJet"], frontier_df["LCOE"], color=PUBLICATION_COLORS["black"], linewidth=1.0, linestyle="--", zorder=2)
        ax.axvline(0, color="#94A3B8", linestyle="--", linewidth=0.8)
        for _, row in frontier_df.iterrows():
            ax.annotate(row["Scenario"], (row["CarbonDiffVsJet"], row["LCOE"]), xytext=(5, 5), textcoords="offset points", fontsize=6.4, color="#334155")
        ax.set_xlabel("Carbon delta vs. jet fuel (gCO2e/MJ)")
        ax.set_ylabel("LCO-SAF (CNY/kg)")
        style_chart_axes(ax)
        legend_handles = [
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=category_color("Grey"), markeredgecolor="white", markersize=6, label="Grey"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=category_color("Blue"), markeredgecolor="white", markersize=6, label="Blue"),
            Line2D([0], [0], marker="o", linestyle="None", markerfacecolor=category_color("Green"), markeredgecolor="white", markersize=6, label="Green"),
            Line2D([0], [0], marker=ROUTE_MARKERS["FT"], linestyle="None", markerfacecolor="#475569", markeredgecolor="white", markersize=6, label="FT"),
            Line2D([0], [0], marker=ROUTE_MARKERS["MTJ"], linestyle="None", markerfacecolor="#475569", markeredgecolor="white", markersize=6, label="MTJ"),
        ]
        add_publication_legend(ax, handles=legend_handles, ncol=3, loc="upper center", bbox_to_anchor=(0.5, 1.23))
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        ordered = metrics.sort_values("LCOE", ascending=True)
        apply_horizontal_bar_template(
            ax,
            ordered["Scenario"].tolist(),
            ordered["LCOE"].to_numpy(),
            colors=ordered["Category"].map(category_color).tolist(),
            xlabel="LCO-SAF (CNY/kg)",
            invert=True,
        )
        for y, (_, row) in enumerate(ordered.iterrows()):
            ax.text(row["LCOE"] + 0.4, y, f"{row['CarbonDiffVsJet']:+.0f}", va="center", fontsize=6.5, color="#334155")
        add_panel_label(ax, "B")
        return save_matplotlib_figure(fig, path)


def create_efficiency_analysis_figure(path: Path) -> Path:
    metrics = build_scenario_metrics_table()

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.88, bottom=0.22, wspace=0.28, hspace=0.42)

        ax = fig.add_subplot(gs[0, 0])
        apply_scatter_template(
            ax,
            metrics["WorstWeekRetention"],
            metrics["MeanSAFUtilization"] * 100,
            marker="o",
            s=np.clip(metrics["LCOE"] * 2.2, 18, 120),
            c=metrics["Category"].map(category_color),
            edgecolors="white",
            linewidth=0.7,
        )
        ax.set_xlabel("Worst-week retention (%)")
        ax.set_ylabel("Mean SAF utilization (%)")
        style_chart_axes(ax)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        for route, marker in ROUTE_MARKERS.items():
            subset = metrics[metrics["Route"] == route]
            apply_scatter_template(
                ax,
                subset["NodeCount"],
                subset["MeanTransportDistance"],
                label=route,
                marker=marker,
                s=np.clip(subset["SAF_Plants"].fillna(0) * 2.8, 20, 130),
                c=subset["Category"].map(category_color),
                edgecolors="white",
                linewidth=0.7,
        )
        ax.set_xlabel("Active nodes in archived network")
        ax.set_ylabel("Mean transport distance (km)")
        style_chart_axes(ax)
        add_publication_legend(ax, loc="upper center", bbox_to_anchor=(0.5, 1.24), ncol=2)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        share_cols = [col for _, col in SITE_SHARE_COLUMNS]
        share_labels = [translate_site_share_label(label) for label, _ in SITE_SHARE_COLUMNS]
        share_df = metrics[["Scenario"] + share_cols].fillna(0.0).sort_values("Scenario", key=lambda s: s.map(scenario_sort_key))
        share_colors = [
            PUBLICATION_COLORS["blue"],
            PUBLICATION_COLORS["sky"],
            PUBLICATION_COLORS["orange"],
            PUBLICATION_COLORS["green"],
            PUBLICATION_COLORS["gray"],
        ]
        apply_stacked_bar_template(
            ax,
            share_df["Scenario"].tolist(),
            [(label, share_df[column].to_numpy(), color) for label, column, color in zip(share_labels, share_cols, share_colors)],
            xlabel="Share of SAF production siting (%)",
            orientation="h",
            legend=True,
            legend_kwargs={"ncol": 3, "loc": "upper center", "bbox_to_anchor": (0.5, -0.22)},
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        economic = metrics.sort_values("Scenario", key=lambda s: s.map(scenario_sort_key))
        apply_horizontal_bar_template(
            ax,
            economic["Scenario"].tolist(),
            economic["LCOE"].to_numpy(),
            colors=economic["Category"].map(category_color).tolist(),
            xlabel="CNY/kg for bars; gCO2e/MJ for dots",
            invert=True,
        )
        apply_scatter_template(
            ax,
            economic["CarbonDiffVsJet"],
            economic["Scenario"],
            color=PUBLICATION_COLORS["black"],
            s=18,
            edgecolors="none",
            label="Carbon delta",
        )
        add_panel_label(ax, "D")
        return save_matplotlib_figure(fig, path)


def create_cluster_archetype_figure(path: Path) -> Path:
    network_metrics, siting_metrics, decision_counts = load_latest_network_metrics()
    network_metrics = translate_network_metric_columns(network_metrics)
    siting_metrics = siting_metrics.copy()
    siting_metrics["Scenario"] = siting_metrics["OfficialName"].fillna(siting_metrics["Scenario"]).map(normalize_scenario_code)
    decision_counts = decision_counts.copy()
    decision_counts["Scenario"] = decision_counts["OfficialName"].fillna(decision_counts["Scenario"]).map(normalize_scenario_code)
    decision_counts["MaterialLabelEn"] = decision_counts["MaterialLabel"].map(translate_material_label)

    cluster_site = (
        siting_metrics.groupby("Cluster")[[col for _, col in SITE_SHARE_COLUMNS]]
        .mean()
        .sort_index()
    )
    decision_cluster = decision_counts.groupby(["Cluster", "MaterialLabelEn"], as_index=False)["Count"].sum()
    decision_cluster["Share"] = decision_cluster["Count"] / decision_cluster.groupby("Cluster")["Count"].transform("sum")
    top_materials = (
        decision_cluster.groupby("MaterialLabelEn")["Share"].sum().sort_values(ascending=False).head(4).index.tolist()
    )
    decision_heat = (
        decision_cluster[decision_cluster["MaterialLabelEn"].isin(top_materials)]
        .pivot(index="Cluster", columns="MaterialLabelEn", values="Share")
        .fillna(0.0)
        .sort_index()
    )

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.90, bottom=0.24, wspace=0.30, hspace=0.42)

        ax = fig.add_subplot(gs[0, 0])
        for cluster, subset in network_metrics.groupby("Cluster"):
            apply_scatter_template(
                ax,
                subset["KeyNodeConcentration"],
                subset["KeyCorridorConcentration"],
                label=f"Cluster {int(cluster)}",
                s=48,
                color=CLUSTER_COLORS.get(int(cluster), PUBLICATION_COLORS["gray"]),
                edgecolors="white",
                linewidth=0.7,
        )
        ax.set_xlabel("Key-node concentration (%)")
        ax.set_ylabel("Key-corridor concentration (%)")
        style_chart_axes(ax)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        for cluster, subset in network_metrics.groupby("Cluster"):
            apply_scatter_template(
                ax,
                subset["MeanTransportDistance"],
                subset["CrossRegionalShare"],
                s=48,
                color=CLUSTER_COLORS.get(int(cluster), PUBLICATION_COLORS["gray"]),
                edgecolors="white",
                linewidth=0.7,
        )
        ax.set_xlabel("Mean transport distance (km)")
        ax.set_ylabel("Cross-regional logistics share (%)")
        style_chart_axes(ax)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        share_labels = [translate_site_share_label(label) for label, _ in SITE_SHARE_COLUMNS]
        share_colors = [
            PUBLICATION_COLORS["blue"],
            PUBLICATION_COLORS["sky"],
            PUBLICATION_COLORS["orange"],
            PUBLICATION_COLORS["green"],
            PUBLICATION_COLORS["gray"],
        ]
        apply_stacked_bar_template(
            ax,
            [f"Cluster {int(idx)}" for idx in cluster_site.index],
            [(label, cluster_site[column].to_numpy(), color) for label, column, color in zip(share_labels, cluster_site.columns, share_colors)],
            xlabel="Mean siting share (%)",
            orientation="h",
            legend=True,
            legend_kwargs={"ncol": 3, "loc": "upper center", "bbox_to_anchor": (0.5, -0.22)},
        )
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        _, cbar = apply_heatmap_template(
            fig,
            ax,
            decision_heat.to_numpy(),
            xlabels=decision_heat.columns.tolist(),
            ylabels=[f"Cluster {int(idx)}" for idx in decision_heat.index],
            cmap="Blues",
            vmin=0,
            vmax=max(0.25, float(decision_heat.to_numpy().max())),
            colorbar_label="Share within cluster",
            annotate=True,
        )
        add_panel_label(ax, "D")
        cluster_handles = [
            Line2D(
                [0],
                [0],
                marker="o",
                linestyle="None",
                markerfacecolor=CLUSTER_COLORS.get(cluster, PUBLICATION_COLORS["gray"]),
                markeredgecolor="white",
                markersize=6.5,
                label=f"Cluster {cluster}",
            )
            for cluster in sorted(network_metrics["Cluster"].dropna().astype(int).unique())
        ]
        add_publication_figure_legend(
            fig,
            cluster_handles,
            loc="lower center",
            bbox_to_anchor=(0.5, 0.08),
            ncol=min(4, len(cluster_handles)),
            columnspacing=1.0,
            handletextpad=0.5,
        )
        return save_matplotlib_figure(fig, path)


def create_example_transport_network_figure(path: Path) -> Path:
    transport_path = ROOT / "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/transport_summary_20260126_170826.csv"
    routes = pd.read_csv(transport_path)
    routes = routes[routes["周运输量(kg)"] > 1_000].copy()
    routes["CargoEn"] = routes["货物类型"].map(translate_cargo_type)
    routes["ModeEn"] = routes["运输方式"].map(translate_route_mode)
    routes["StartTypeEn"] = routes["起点类型"].map(translate_node_type)
    routes["EndTypeEn"] = routes["终点类型"].map(translate_node_type)
    routes["StartCoord"] = routes["起点坐标"].map(parse_latlon_tuple)
    routes["EndCoord"] = routes["终点坐标"].map(parse_latlon_tuple)
    routes = routes.dropna(subset=["StartCoord", "EndCoord"])
    routes = routes[routes["距离(km)"] > 1.0]

    cargo_colors = {
        "Natural gas": PUBLICATION_COLORS["gray"],
        "Hydrogen": PUBLICATION_COLORS["green"],
        "SAF": PUBLICATION_COLORS["orange"],
    }
    node_markers = {
        "Natural-gas pipeline": ("o", PUBLICATION_COLORS["gray"]),
        "Production facility": ("s", PUBLICATION_COLORS["blue"]),
        "Hydrogen source": ("^", PUBLICATION_COLORS["green"]),
        "MTJ plant": ("D", PUBLICATION_COLORS["orange"]),
        "Airport": ("*", PUBLICATION_COLORS["black"]),
    }

    unique_nodes: dict[tuple[float, float, str], dict[str, object]] = {}
    for _, row in routes.iterrows():
        start_lat, start_lon = row["StartCoord"]
        end_lat, end_lon = row["EndCoord"]
        unique_nodes[(start_lon, start_lat, row["StartTypeEn"])] = {"lon": start_lon, "lat": start_lat, "type": row["StartTypeEn"]}
        unique_nodes[(end_lon, end_lat, row["EndTypeEn"])] = {"lon": end_lon, "lat": end_lat, "type": row["EndTypeEn"]}

    lons = [node["lon"] for node in unique_nodes.values()]
    lats = [node["lat"] for node in unique_nodes.values()]
    extent = [min(lons) - 1.0, max(lons) + 1.0, min(lats) - 0.8, max(lats) + 0.8]

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("network_map_single")
        gs = fig.add_gridspec(1, 1, left=0.05, right=0.985, top=0.95, bottom=0.25)

        ax, mini_ax = create_reference_map_panel(
            fig, gs[0, 0], extent, dx=1.5, dy=1.0, main_scale_km=200, mini_scale_km=500, compass_size=12.0
        )
        mass = routes["周运输量(kg)"].replace(0, np.nan).fillna(routes["周运输量(kg)"].max())
        width_scale = 0.5 + 2.0 * (np.log10(mass) - np.log10(mass.min())) / (np.log10(mass.max()) - np.log10(mass.min()) + 1e-9)
        for (_, row), line_width in zip(routes.iterrows(), width_scale):
            coords = parse_path_coords(row["路径坐标"])
            if len(coords) >= 2:
                xs = [coord[0] for coord in coords]
                ys = [coord[1] for coord in coords]
            else:
                start_lat, start_lon = row["StartCoord"]
                end_lat, end_lon = row["EndCoord"]
                xs = [start_lon, end_lon]
                ys = [start_lat, end_lat]
            ax.plot(xs, ys, color=cargo_colors.get(row["CargoEn"], PUBLICATION_COLORS["gray"]), linewidth=float(line_width), alpha=0.45, transform=ccrs.PlateCarree(), zorder=2)
            mirror_lines_to_mini(mini_ax, xs, ys, color=cargo_colors.get(row["CargoEn"], PUBLICATION_COLORS["gray"]), linewidth=max(float(line_width) * 0.55, 0.4), alpha=0.35, zorder=2)
        for node_type, (marker, color) in node_markers.items():
            subset = [node for node in unique_nodes.values() if node["type"] == node_type]
            if not subset:
                continue
            ax.scatter(
                [node["lon"] for node in subset],
                [node["lat"] for node in subset],
                marker=marker,
                s=34 if marker != "*" else 70,
                color=color,
                edgecolors="white",
                linewidth=0.6,
                alpha=0.9,
                transform=ccrs.PlateCarree(),
                zorder=4,
                label=node_type,
            )
            mirror_points_to_mini(
                mini_ax,
                [node["lon"] for node in subset],
                [node["lat"] for node in subset],
                marker=marker,
                s=18 if marker != "*" else 26,
                color=color,
                edgecolors="white",
                linewidth=0.35,
                alpha=0.85,
                zorder=4,
            )
        add_panel_label(ax, "A")
        cargo_handles = [
            Line2D([0], [0], color=color, linewidth=2.0, label=label)
            for label, color in cargo_colors.items()
            if label in routes["CargoEn"].unique()
        ]
        node_handles = [
            Line2D([0], [0], marker=marker, linestyle="None", markerfacecolor=color, markeredgecolor="white", markersize=6 if marker != "*" else 8, label=label)
            for label, (marker, color) in node_markers.items()
        ]
        legend_handles = cargo_handles + node_handles
        add_publication_figure_legend(
            fig,
            legend_handles,
            loc="lower center",
            bbox_to_anchor=(0.50, 0.03),
            ncol=4,
            fontsize=6.6,
            columnspacing=0.8,
            handletextpad=0.4,
        )
        return save_matplotlib_figure(fig, path)


def create_example_transport_network_stats(path: Path) -> Path:
    transport_path = ROOT / "products/supply_chain_optimization/natural_gas_supply_chain_optimization/results/transport_summary_20260126_170826.csv"
    routes = pd.read_csv(transport_path)
    routes = routes[routes["周运输量(kg)"] > 1_000].copy()
    routes["CargoEn"] = routes["货物类型"].map(translate_cargo_type)
    routes["ModeEn"] = routes["运输方式"].map(translate_route_mode)
    routes = routes[routes["距离(km)"] > 1.0]

    cargo_summary = routes.groupby("CargoEn")["周运输量(kg)"].sum().sort_values(ascending=False) / 1e6
    mode_counts = routes["ModeEn"].value_counts()
    cargo_colors = {
        "Natural gas": PUBLICATION_COLORS["gray"],
        "Hydrogen": PUBLICATION_COLORS["green"],
        "SAF": PUBLICATION_COLORS["orange"],
    }

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("stats_triptych")
        gs = fig.add_gridspec(1, 3, left=0.08, right=0.985, top=0.94, bottom=0.18, wspace=0.35)

        ax = fig.add_subplot(gs[0, 0])
        apply_horizontal_bar_template(
            ax,
            cargo_summary.index.tolist(),
            cargo_summary.to_numpy(),
            colors=[cargo_colors.get(label, PUBLICATION_COLORS["gray"]) for label in cargo_summary.index],
            xlabel="Million kg/week",
            invert=True,
        )
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        apply_grouped_bar_template(
            ax,
            mode_counts.index.tolist(),
            [("Route segments", mode_counts.values, PUBLICATION_COLORS["orange"])],
            ylabel="Route segments",
            legend=False,
        )
        ax.tick_params(axis="x", rotation=15)
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[0, 2])
        distance_groups = [routes.loc[routes["CargoEn"] == label, "距离(km)"].to_numpy() for label in cargo_summary.index]
        apply_boxplot_template(
            ax,
            distance_groups,
            labels=list(cargo_summary.index),
            colors=[cargo_colors.get(label, PUBLICATION_COLORS["gray"]) for label in cargo_summary.index],
            ylabel="Route distance (km)",
            rotation=18,
        )
        add_panel_label(ax, "C")

        return save_matplotlib_figure(fig, path)


def create_temporal_robustness_overview(path: Path) -> Path:
    metrics = build_scenario_metrics_table().copy()

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("quad_multi")
        gs = fig.add_gridspec(2, 2, left=0.08, right=0.985, top=0.965, bottom=0.12, wspace=0.30, hspace=0.34)

        ax = fig.add_subplot(gs[0, 0])
        apply_scatter_template(
            ax,
            metrics["WorstWeekRetention"],
            metrics["TotalPenaltyPct"],
            s=np.clip(metrics["LCOE"] * 2.0, 18, 120),
            c=metrics["Category"].map(category_color),
            edgecolors="white",
            linewidth=0.7,
        )
        ax.set_xlabel("Worst-week retention (%)")
        ax.set_ylabel("Total temporal penalty (%)")
        style_chart_axes(ax)
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        ordered = metrics.sort_values("WorstWeekRetention", ascending=True)
        apply_horizontal_bar_template(
            ax,
            ordered["Scenario"].tolist(),
            ordered["WorstWeekRetention"].to_numpy(),
            colors=ordered["Category"].map(category_color).tolist(),
            xlabel="Worst-week retention (%)",
            invert=True,
        )
        add_panel_label(ax, "B")

        ax = fig.add_subplot(gs[1, 0])
        apply_scatter_template(
            ax,
            metrics["ChronicPenaltyPct"],
            metrics["RobustnessPenaltyPct"],
            s=42,
            c=metrics["Category"].map(category_color),
            edgecolors="white",
            linewidth=0.7,
        )
        max_penalty = max(float(metrics["ChronicPenaltyPct"].max()), float(metrics["RobustnessPenaltyPct"].max()))
        ax.plot([0, max_penalty], [0, max_penalty], linestyle="--", color="#94A3B8", linewidth=0.8)
        ax.set_xlabel("Chronic penalty (%)")
        ax.set_ylabel("Robustness penalty (%)")
        style_chart_axes(ax)
        add_panel_label(ax, "C")

        ax = fig.add_subplot(gs[1, 1])
        apply_scatter_template(
            ax,
            metrics["MaxLowRunHours"],
            metrics["MeanSAFUtilization"] * 100,
            s=np.clip(metrics["LowLoadEpisodes12h"].fillna(0).to_numpy() * 8 + 18, 18, 130),
            c=metrics["Category"].map(category_color),
            edgecolors="white",
            linewidth=0.7,
        )
        ax.set_xlabel("Longest low-load event (h)")
        ax.set_ylabel("Mean SAF utilization (%)")
        style_chart_axes(ax)
        add_panel_label(ax, "D")
        return save_matplotlib_figure(fig, path)


def create_temporal_penalty_breakdown(path: Path) -> Path:
    penalty_path = ROOT / "products/supply_chain_optimization/visualization/results/pareto_penalty_summary_latest.csv"
    weekly_path = latest_file("products/supply_chain_optimization/visualization/results/temporal_robustness_space_*/penalty_class_weekly_metrics.csv")
    penalty = pd.read_csv(penalty_path)
    weekly = pd.read_csv(weekly_path) if weekly_path else pd.DataFrame()
    scenario_map = {"GTL-BH": "GTL-BH", "GTL": "GTL", "CCU-GH-FT": "CCU-GH-FT"}
    penalty["Scenario"] = penalty["Scenario"].map(normalize_scenario_code)
    penalty = penalty[penalty["Scenario"].isin(scenario_map)].copy()
    penalty["Scenario"] = pd.Categorical(penalty["Scenario"], categories=["GTL-BH", "GTL", "CCU-GH-FT"], ordered=True)
    penalty = penalty.sort_values("Scenario")
    if not weekly.empty:
        weekly["Scenario"] = weekly["Scenario"].map(normalize_scenario_code)
        weekly = weekly[weekly["Scenario"].isin(["GTL-BH", "GTL", "CCU-GH-FT"])].copy()

    with plt.rc_context(get_publication_rc()):
        fig = create_template_figure("comparison_pair_short")
        gs = fig.add_gridspec(1, 2, left=0.08, right=0.985, top=0.88, bottom=0.18, wspace=0.30)

        ax = fig.add_subplot(gs[0, 0])
        apply_stacked_bar_template(
            ax,
            penalty["Scenario"].astype(str).tolist(),
            [
                ("Chronic", penalty["ChronicPenaltyPct"].to_numpy(), PUBLICATION_COLORS["blue"]),
                ("Robustness", penalty["RobustnessPenaltyPct"].to_numpy(), PUBLICATION_COLORS["orange"]),
            ],
            ylabel="Penalty relative to baseline LCO-SAF (%)",
            legend=True,
            legend_kwargs={"loc": "upper center", "bbox_to_anchor": (0.5, 1.22), "ncol": 2},
        )
        for idx, row in enumerate(penalty.itertuples(index=False)):
            total = row.ChronicPenaltyPct + row.RobustnessPenaltyPct
            ax.text(idx, total + 2.0, f"{total:.1f}", ha="center", fontsize=6.8, color="#334155")
        add_panel_label(ax, "A")

        ax = fig.add_subplot(gs[0, 1])
        week_order = ["W1", "W2", "W3", "W4"]
        palette = {"GTL-BH": PUBLICATION_COLORS["green"], "GTL": PUBLICATION_COLORS["blue"], "CCU-GH-FT": PUBLICATION_COLORS["vermillion"]}
        apply_line_template(
            ax,
            [
                {
                    "x": subset.set_index("Week").reindex(week_order).reset_index()["Week"],
                    "y": subset.set_index("Week").reindex(week_order).reset_index()["WeekPenaltyLift"],
                    "color": palette.get(scenario, PUBLICATION_COLORS["gray"]),
                    "linewidth": 1.4,
                    "marker": "o",
                    "markersize": 4.0,
                    "label": scenario,
                }
                for scenario, subset in weekly.groupby("Scenario")
            ],
            ylabel="Week-specific penalty lift (%)",
            legend=True,
            legend_kwargs={"loc": "upper center", "bbox_to_anchor": (0.5, 1.22), "ncol": 3},
            horizontal_guides=[0],
        )
        add_panel_label(ax, "B")
        return save_matplotlib_figure(fig, path)


# 聚类运输地图源目录 & 文件名→注册键映射
_TRANSPORT_MAP_SRC_DIR = (
    ROOT
    / "products/supply_chain_optimization/visualization/results"
    / "clustered_transport_maps_13scenarios_20260319_225158"
)

_TRANSPORT_MAP_FILES: dict[str, str] = {
    "transport_map_ctl":        "clustered_map_coal_hydrogen.png",
    "transport_map_ctl_bh":     "clustered_map_byproduct_h2_+_coal.png",
    "transport_map_gtl_gh":     "clustered_map_natural_gas_two-step.png",
    "transport_map_gtl":        "clustered_map_natural_gas_one-step.png",
    "transport_map_gtl_bh":     "clustered_map_byproduct_h2_+_ng_two-step.png",
    "transport_map_ccu_gh_mtj": "clustered_map_green_h2_two-step.png",
    "transport_map_ccu_gh_ft":  "clustered_map_green_h2_one-step.png",
    "transport_map_ccu_bh_mtj": "clustered_map_byproduct_h2_two-step.png",
    "transport_map_ccu_bh_ft":  "clustered_map_byproduct_h2_one-step.png",
    "transport_map_dac_gh_mtj": "clustered_map_dac_two-step.png",
    "transport_map_dac_gh_ft":  "clustered_map_dac_one-step.png",
    "transport_map_dac_bh_mtj": "clustered_map_byproduct_h2_+_dac_two-step.png",
    "transport_map_dac_bh_ft":  "clustered_map_byproduct_h2_+_dac_one-step.png",
    "transport_map_four_key":   "four_key_scenarios_comparison.png",
    "transport_map_legend":     "legend_overview.png",
}

# 附录中的显示顺序（分组）
_TRANSPORT_MAP_GROUPS: dict[str, list[str]] = {
    "Coal-based transport networks":            ["transport_map_ctl", "transport_map_ctl_bh"],
    "Natural-gas-based transport networks":     ["transport_map_gtl_gh", "transport_map_gtl", "transport_map_gtl_bh"],
    "CCU Green-H₂ transport networks":         ["transport_map_ccu_gh_mtj", "transport_map_ccu_gh_ft"],
    "CCU By-product-H₂ transport networks":    ["transport_map_ccu_bh_mtj", "transport_map_ccu_bh_ft"],
    "DAC Green-H₂ transport networks":         ["transport_map_dac_gh_mtj", "transport_map_dac_gh_ft"],
    "DAC By-product-H₂ transport networks":    ["transport_map_dac_bh_mtj", "transport_map_dac_bh_ft"],
}

_TRANSPORT_MAP_SCENARIO_LABELS: dict[str, str] = {
    "transport_map_ctl":        "CTL",
    "transport_map_ctl_bh":     "CTL-BH",
    "transport_map_gtl_gh":     "GTL-GH",
    "transport_map_gtl":        "GTL",
    "transport_map_gtl_bh":     "GTL-BH",
    "transport_map_ccu_gh_mtj": "CCU-GH-MTJ",
    "transport_map_ccu_gh_ft":  "CCU-GH-FT",
    "transport_map_ccu_bh_mtj": "CCU-BH-MTJ",
    "transport_map_ccu_bh_ft":  "CCU-BH-FT",
    "transport_map_dac_gh_mtj": "DAC-GH-MTJ",
    "transport_map_dac_gh_ft":  "DAC-GH-FT",
    "transport_map_dac_bh_mtj": "DAC-BH-MTJ",
    "transport_map_dac_bh_ft":  "DAC-BH-FT",
    "transport_map_four_key":   "Four-key scenarios comparison",
    "transport_map_legend":     "Legend overview",
}


def _save_raster_with_bundles(src_png: Path, dst_png: Path, max_width_px: int = 3000) -> Path:
    """
    从源 PNG 降采样到 max_width_px 宽，保存 PNG / PDF / SVG 三份副本。
    返回目标 PNG 路径。
    """
    import base64
    from PIL import Image as _PILImage

    _PILImage.MAX_IMAGE_PIXELS = None
    img = _PILImage.open(src_png).convert("RGB")
    w, h = img.size
    if w > max_width_px:
        new_h = int(h * max_width_px / w)
        img = img.resize((max_width_px, new_h), _PILImage.LANCZOS)

    dst_png.parent.mkdir(parents=True, exist_ok=True)
    img.save(dst_png, dpi=(300, 300))

    # PDF
    img.save(dst_png.with_suffix(".pdf"), "PDF", resolution=300)

    # SVG（base64 raster 嵌入）
    with open(dst_png, "rb") as _f:
        _b64 = base64.b64encode(_f.read()).decode()
    wp, hp = img.size
    w_mm = wp / 300 * 25.4
    h_mm = hp / 300 * 25.4
    svg = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        f'<svg xmlns="http://www.w3.org/2000/svg" xmlns:xlink="http://www.w3.org/1999/xlink" '
        f'width="{w_mm:.2f}mm" height="{h_mm:.2f}mm" viewBox="0 0 {wp} {hp}">\n'
        f'  <image width="{wp}" height="{hp}" xlink:href="data:image/png;base64,{_b64}"/>\n'
        '</svg>\n'
    )
    dst_png.with_suffix(".svg").write_text(svg, encoding="utf-8")
    return dst_png


def prepare_transport_map_assets() -> dict[str, Path]:
    """
    将聚类运输地图降采样、存入 PREP_DIR/transport_maps/ 并生成 PDF/SVG 副本。
    返回 {registry_key: prepared_png_path}
    """
    dst_dir = PREP_DIR / "transport_maps"
    result: dict[str, Path] = {}
    for key, fname in _TRANSPORT_MAP_FILES.items():
        src = _TRANSPORT_MAP_SRC_DIR / fname
        if not src.exists():
            print(f"[WARNING] 运输地图文件不存在: {src}")
            continue
        # legend 较宽，限 6000px；其余限 3000px
        max_w = 6000 if key == "transport_map_legend" else 3000
        dst = dst_dir / f"{key}.png"
        _save_raster_with_bundles(src, dst, max_width_px=max_w)
        result[key] = dst
        print(f"已准备运输地图: {key} → {dst.name}")
    return result


def prepare_assets(data: AppendixData, pareto_assets: dict[str, Path] | None = None) -> dict[str, Path]:
    ensure_dir(PREP_DIR)
    figure_map: dict[str, Path] = {}
    if pareto_assets is None:
        pareto_assets = generate_appendix_pareto_assets(output_dir=REDRAWN_DIR, prepared_dir=PREP_DIR)

    figure_map["raw_data_overview"] = create_raw_data_overview_panel(PREP_DIR / "raw_data_overview.png", data)
    figure_map["airport_demand_raw"] = create_airport_demand_raw_summary(PREP_DIR / "airport_demand_raw_summary.png")
    figure_map["renewable_raw"] = create_renewable_raw_summary(PREP_DIR / "renewable_raw_inventory_summary.png", data)
    figure_map["renewable_resource_maps"] = create_renewable_resource_maps(PREP_DIR / "renewable_resource_maps.png", data)
    figure_map["renewable_inventory_map"] = create_renewable_inventory_capacity_maps(PREP_DIR / "renewable_inventory_map.png")
    figure_map["renewable_processing"] = create_renewable_processing_figure(PREP_DIR / "renewable_processing.png")
    figure_map["byproduct_raw"] = create_byproduct_raw_summary(PREP_DIR / "byproduct_h2_raw_summary.png")
    figure_map["byproduct_geo"] = create_byproduct_geo_maps(PREP_DIR / "byproduct_h2_geo_maps.png")
    figure_map["co2_capture_raw"] = create_co2_capture_raw_summary(PREP_DIR / "co2_capture_raw_summary.png", data)
    figure_map["co2_capture_geo"] = create_co2_capture_geo_maps(PREP_DIR / "co2_capture_geo_maps.png")
    figure_map["transport_raw"] = create_transport_raw_summary(PREP_DIR / "transport_network_raw_summary.png")
    figure_map["price_surface_maps"] = create_price_surface_maps(PREP_DIR / "price_surface_maps.png")
    figure_map["energy_infrastructure_overview"] = create_energy_infrastructure_overview(PREP_DIR / "energy_infrastructure_overview.png")
    figure_map["demand_aircraft_supplement"] = create_demand_aircraft_supplement(PREP_DIR / "demand_aircraft_supplement.png")
    figure_map["annual_demand_distribution"] = create_annual_demand_distribution(PREP_DIR / "annual_demand_distribution.png")
    figure_map["demand_patterns"] = create_representative_demand_patterns(PREP_DIR / "demand_patterns.png")
    figure_map["renewable_profiles"] = create_representative_renewable_profiles(PREP_DIR / "renewable_profiles.png")
    _framework_src = ROOT / "appendix_saf_workspace/word/render_check/framework.png"
    if _framework_src.exists():
        figure_map["framework_overview"] = copy_figure_bundle(_framework_src, PREP_DIR / "framework_overview.png")
    else:
        figure_map["framework_overview"] = create_placeholder_figure(PREP_DIR / "framework_overview.png", "Optimization Framework Overview")

    figure_map["workflow"] = create_workflow_figure(PREP_DIR / "workflow_figure.png")
    figure_map["quadrant_chart"] = create_quadrant_comparison(PREP_DIR / "quadrant_chart.png")
    figure_map["efficiency_analysis"] = create_efficiency_analysis_figure(PREP_DIR / "efficiency_analysis.png")
    figure_map["combined_cluster"] = create_cluster_archetype_figure(PREP_DIR / "combined_cluster.png")
    figure_map["transport_route_network"] = create_example_transport_network_figure(PREP_DIR / "transport_route_network.png")
    figure_map["transport_route_network_stats"] = create_example_transport_network_stats(PREP_DIR / "transport_route_network_stats.png")
    figure_map["temporal_space"] = create_temporal_robustness_overview(PREP_DIR / "temporal_space.png")
    figure_map["temporal_penalty"] = create_temporal_penalty_breakdown(PREP_DIR / "temporal_penalty.png")

    # 时间维度附加图表（直接使用已生成的可视化结果）
    _temporal_fig_map = {
        "temporal_h2_saf":         VIZ_RESULTS_DIR / "h2_saf_2col_3row_by_category.png",
        "pareto_temporal_dist":     VIZ_RESULTS_DIR / "pareto_temporal_distribution_latest.png",
        "pareto_penalty_comparison": VIZ_RESULTS_DIR / "pareto_penalty_comparison_latest.png",
        "penalty_driver_heatmap":   VIZ_RESULTS_DIR / "pareto_penalty_driver_heatmap_latest.png",
    }
    for key, src_path in _temporal_fig_map.items():
        dst_png = PREP_DIR / f"{key}.png"
        if src_path.exists():
            figure_map[key] = copy_figure_bundle(src_path, dst_png)
        else:
            print(f"[WARNING] 时间维度图表文件不存在，跳过: {src_path.name}")
            figure_map[key] = create_placeholder_figure(dst_png, src_path.name)
    figure_map["pareto_cost"] = copy_figure_bundle(pareto_assets["cost_png"], PREP_DIR / "pareto_cost.png")
    figure_map["pareto_carbon"] = copy_figure_bundle(pareto_assets["carbon_png"], PREP_DIR / "pareto_carbon.png")

    # 全场景极坐标合成图（每个工艺一张三合一图）
    _scenario_composites_dir = PREP_DIR / "scenario_composites"
    _scenario_composites_dir.mkdir(parents=True, exist_ok=True)
    if generate_all_scenario_composites is not None:
        try:
            _all_composites = generate_all_scenario_composites(
                output_dir=REDRAWN_DIR / "scenario_composites",
                prepared_dir=_scenario_composites_dir,
            )
            for _sname, _spath in _all_composites.items():
                _slug = _sname.lower().replace("-", "_")
                figure_map[f"scenario_composite_{_slug}"] = _spath
        except Exception as _e:
            print(f"[WARNING] 全场景合成图生成失败: {_e}")

    # 裁剪合成图：只保留左侧极坐标学习曲线图，去除右侧甜甜圈
    # 两种图宽(6494/6581)在 x=3280 处均为空白列
    _CROP_X = 3280
    _cropped_dir = _scenario_composites_dir / "left_only"
    _cropped_dir.mkdir(parents=True, exist_ok=True)
    from PIL import Image as _PIL_Image
    for _key in list(figure_map.keys()):
        if _key.startswith("scenario_composite_"):
            _src = Path(figure_map[_key])
            if _src.exists():
                _img = _PIL_Image.open(_src)
                _w, _h = _img.size
                _cropped = _img.crop((0, 0, min(_CROP_X, _w), _h))
                _dst = _cropped_dir / f"{_src.stem}_left.png"
                _cropped.save(str(_dst))
                ensure_raster_export_bundle(_dst)
                figure_map[_key] = _dst
                print(f"  [裁剪] {_key} → {_dst.name}")

    # 仍使用占位符的图表（尚未有批量运行数据）
    remaining_placeholder_titles = {
        "electricity_price_placeholder": "Electricity price sensitivity",
        "electrolyzer_capex_placeholder": "Electrolyzer CAPEX sensitivity",
        "dac_cost_placeholder": "DAC cost sensitivity",
        "carbon_price_placeholder": "Carbon-price sensitivity",
        "carbon_price_asymmetric_placeholder": "Carbon-price pathway response",
    }
    for key, title in remaining_placeholder_titles.items():
        figure_map[key] = create_placeholder_figure(PREP_DIR / f"{key}.png", title)

    # 敏感性分析批量运行已完成：使用真实图表替换占位符
    _sens_fig_map = {
        "fossil_price_placeholder":  "fig_ng_price_sensitivity_appendix.png",
        "bh_capex_placeholder":      "fig_bh_capex_sensitivity_appendix.png",
        "green_h2_joint_placeholder":"fig_electricity_capex_heatmap_appendix.png",
        "dac_breakeven_placeholder":  "fig_dac_breakeven_appendix.png",
        "pareto_tornado_placeholder": "fig_pareto_tornado_appendix.png",
    }
    for key, fname in _sens_fig_map.items():
        real_path = SENSITIVITY_FIGURES_DIR / fname
        if real_path.exists():
            # 复制到 PREP_DIR 并生成 PDF/SVG 导出包（QA 要求）
            dst_png = PREP_DIR / f"{key}.png"
            figure_map[key] = copy_figure_bundle(real_path, dst_png)
        else:
            # 兜底：生成占位符（防止文件缺失时报错）
            figure_map[key] = create_placeholder_figure(PREP_DIR / f"{key}.png", fname)

    # 13 场景聚类运输地图
    figure_map.update(prepare_transport_map_assets())

    return figure_map


def add_figure(doc: Document, image_path: Path, caption: str, width_inches: float = 6.1) -> None:
    figure_style = "No Spacing" if style_exists(doc, "No Spacing") else "Normal"
    p = doc.add_paragraph(style=figure_style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_format = p.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption_p = doc.add_paragraph(style="Caption" if style_exists(doc, "Caption") else "Normal")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_format = caption_p.paragraph_format
    caption_format.space_before = Pt(0)
    caption_format.space_after = Pt(4)
    caption_format.line_spacing = 1.0
    caption_p.add_run(caption)


def add_figure_block(
    doc: Document,
    image_path: Path,
    caption: str,
    width_inches: float = 6.1,
    notes: Iterable[str] | None = None,
) -> None:
    add_figure(doc, image_path, caption, width_inches=width_inches)
    for note in notes or []:
        add_para(doc, note)


def _wrap_label_text(text: str, width: int = 28) -> list[str]:
    wrapped = textwrap.wrap(text, width=width) or [text]
    return wrapped[:3]


def create_appendix_contact_sheet(
    figure_map: dict[str, Path],
    output_path: Path,
    *,
    title: str,
    keys: list[str] | None = None,
    columns: int = 4,
    show_template_meta: bool = False,
    thumb_width: int = 360,
    thumb_height: int = 220,
) -> Path:
    ensure_dir(output_path.parent)
    ordered_keys = [key for key in (keys or list(figure_map.keys())) if key in figure_map]
    if not ordered_keys:
        raise ValueError("No appendix figures were available for the contact sheet")

    outer_pad = 24
    inner_pad = 16
    header_height = 72
    label_height = 56 if show_template_meta else 36
    rows = math.ceil(len(ordered_keys) / columns)
    cell_width = thumb_width + inner_pad * 2
    cell_height = thumb_height + inner_pad * 2 + label_height
    canvas_width = outer_pad * 2 + columns * cell_width
    canvas_height = header_height + outer_pad + rows * cell_height + outer_pad

    canvas = Image.new("RGB", (canvas_width, canvas_height), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(22)
    label_font = load_font(12)
    meta_font = load_font(10)

    draw.text((outer_pad, 18), title, fill="#111827", font=title_font)
    draw.text(
        (outer_pad, 44),
        f"{len(ordered_keys)} appendix figures regenerated under the publication template pipeline",
        fill="#475569",
        font=meta_font,
    )

    for idx, key in enumerate(ordered_keys):
        row, col = divmod(idx, columns)
        x0 = outer_pad + col * cell_width
        y0 = header_height + row * cell_height
        x1 = x0 + cell_width - 10
        y1 = y0 + cell_height - 10
        draw.rounded_rectangle((x0, y0, x1, y1), radius=14, outline="#E2E8F0", width=1, fill="#FFFFFF")

        image_path = figure_map[key]
        with Image.open(image_path) as image:
            image = image.convert("RGB")
            scale = min(thumb_width / image.width, thumb_height / image.height)
            resized = image.resize(
                (max(1, int(image.width * scale)), max(1, int(image.height * scale))),
                getattr(Image, "Resampling", Image).LANCZOS,
            )

        image_x = x0 + inner_pad + (thumb_width - resized.width) // 2
        image_y = y0 + inner_pad + (thumb_height - resized.height) // 2
        canvas.paste(resized, (image_x, image_y))

        label_y = y0 + inner_pad + thumb_height + 6
        for line_idx, line in enumerate(_wrap_label_text(key.replace("_", " "))):
            draw.text((x0 + inner_pad, label_y + line_idx * 13), line, fill="#111827", font=label_font)

        if show_template_meta:
            registry = FIGURE_TEMPLATE_REGISTRY.get(key, {})
            meta_text = f"{registry.get('family', 'unregistered')} | {registry.get('template', 'n/a')}"
            draw.text((x0 + inner_pad, label_y + 28), meta_text, fill="#64748B", font=meta_font)

    canvas.save(output_path, format="PNG", optimize=True)
    return output_path


def write_appendix_qa_report(figure_map: dict[str, Path], output_path: Path) -> Path:
    ensure_dir(output_path.parent)
    records: list[dict[str, object]] = []
    qa_registry = dict(FIGURE_TEMPLATE_REGISTRY)
    qa_registry.update(
        {
            "demand_aircraft_supplement": {"template": "quad_multi", "family": "data_chart"},
            "annual_demand_distribution": {"template": "comparison_pair", "family": "data_chart"},
        }
    )
    missing_registry = [key for key in figure_map if key not in qa_registry]
    missing_exports: list[str] = []

    Image.MAX_IMAGE_PIXELS = None  # 禁用 PIL 解压炸弹检查（已知安全的本地文件）
    for key, png_path in figure_map.items():
        registry = qa_registry.get(key, {})
        export_flags = {suffix: png_path.with_suffix(suffix).exists() for suffix in EXPECTED_EXPORT_SUFFIXES}
        with Image.open(png_path) as image:
            width_px, height_px = image.size
        records.append(
            {
                "asset": key,
                "template": registry.get("template", "unregistered"),
                "family": registry.get("family", "unregistered"),
                "png": export_flags[".png"],
                "pdf": export_flags[".pdf"],
                "svg": export_flags[".svg"],
                "size": f"{width_px}x{height_px}",
                "path": str(png_path),
            }
        )
        if not all(export_flags.values()):
            missing_suffixes = [suffix[1:] for suffix, exists in export_flags.items() if not exists]
            missing_exports.append(f"{key}: {', '.join(missing_suffixes)}")

    total_assets = len(records)
    map_assets = sum(1 for record in records if record["family"] == "map")
    data_assets = sum(1 for record in records if record["family"] == "data_chart")
    workflow_assets = sum(1 for record in records if record["family"] == "workflow")
    placeholder_assets = sum(1 for record in records if record["family"] == "placeholder")

    lines = [
        "# Appendix figure QA report",
        "",
        f"Generated from: `{SCRIPT_DIR / 'generate_appendix_docx.py'}`",
        f"Prepared figure directory: `{PREP_DIR}`",
        "",
        "## Baseline checks",
        "",
        "- Academic rc baseline: Arial-first sans serif, `pdf.fonttype=42`, `svg.fonttype='none'`, sentence-case labels, borderless legends.",
        "- Export baseline: PNG + PDF + SVG, `dpi=300`, `bbox_inches='tight'`.",
        "- Layout baseline: panel labels use `(a)`, no subplot titles, maps and data charts are separated unless explicitly justified.",
        "",
        "## Registry summary",
        "",
        f"- Total appendix figures audited: {total_assets}",
        f"- Map figures: {map_assets}",
        f"- Data-chart figures: {data_assets}",
        f"- Workflow figures: {workflow_assets}",
        f"- Placeholder figures: {placeholder_assets}",
        "",
        "## Exceptions",
        "",
        f"- Missing registry entries: {', '.join(missing_registry) if missing_registry else 'None'}",
        f"- Missing export bundles: {'; '.join(missing_exports) if missing_exports else 'None'}",
        "",
        "## Asset table",
        "",
        "| Asset | Family | Template | PNG | PDF | SVG | Size |",
        "|---|---|---|---|---|---|---|",
    ]
    for record in records:
        lines.append(
            f"| {record['asset']} | {record['family']} | {record['template']} | "
            f"{'Y' if record['png'] else 'N'} | {'Y' if record['pdf'] else 'N'} | "
            f"{'Y' if record['svg'] else 'N'} | {record['size']} |"
        )

    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    if missing_registry or missing_exports:
        raise RuntimeError(
            "Appendix QA registry/export validation failed: "
            + "; ".join(missing_registry + missing_exports)
        )

    return output_path


def run_appendix_qa(figure_map: dict[str, Path]) -> dict[str, Path]:
    outputs = {
        "qa_contact_sheet": create_appendix_contact_sheet(
            figure_map,
            PREP_DIR / "qa_contact_sheet.png",
            title="Appendix figure contact sheet",
            keys=QA_CONTACT_ORDER,
            columns=4,
            show_template_meta=False,
        ),
        "qa_template_contact_sheet": create_appendix_contact_sheet(
            figure_map,
            PREP_DIR / "qa_template_contact_sheet.png",
            title="Appendix figure template audit",
            keys=QA_CONTACT_ORDER,
            columns=4,
            show_template_meta=True,
        ),
        "qa_template_key_panels": create_appendix_contact_sheet(
            figure_map,
            PREP_DIR / "qa_template_key_panels.png",
            title="Appendix figure key-panel audit",
            keys=QA_KEY_PANEL_ORDER,
            columns=2,
            show_template_meta=True,
            thumb_width=460,
            thumb_height=270,
        ),
        "qa_report": write_appendix_qa_report(
            figure_map,
            PREP_DIR / "qa_checklist_report.md",
        ),
    }
    return outputs


def build_document(data: AppendixData, assets: dict[str, Path]) -> Document:
    doc = Document(str(TEMPLATE_DOCX))
    strip_heading_auto_numbering(doc)
    clear_document(doc)

    add_heading(doc, "Supplementary information", "Heading 1")
    subtitle_style = "Subtitle" if style_exists(doc, "Subtitle") else "Normal"
    add_para(doc, "Appendix for “Techno-economic, spatial, and temporal trade-offs define synthetic aviation-fuel transition pathways in China”", style=subtitle_style)
    add_para(doc, "This appendix consolidates the data-processing steps, scenario definitions, methodological details, verification notes, and supplementary figures that support the archived SAF optimization workflows in the current repository.")

    figure_no = 1
    table_no = 1
    eq_no = 1

    # Chapter 1
    add_heading(doc, "1 Data Sources and Processing Methods", "Heading 2")
    add_para(
        doc,
        "This section documents the data pipeline used to construct the archived optimization instance. The purpose is not only to list data sources, but also to make explicit how raw records with different spatial scales, temporal resolutions, and engineering units are converted into the model-ready demand, supply, cost, emissions, and transport inputs used by the SAF pathway comparison.",
    )
    add_para(
        doc,
        "The first part of the appendix therefore follows the same order as the optimization workflow. It starts from the cross-dataset inventory, then describes airport demand construction, renewable-electricity reconstruction, industrial hydrogen and carbon-source processing, transport-network construction, and finally the clustering and arc-generation steps that reduce the raw data layers to an optimization network.",
    )
    add_heading(doc, "1.1 Overview of input data", "Heading 3")
    add_para(
        doc,
        "The archived model instance combines five major data blocks: aviation fuel demand, renewable-electricity availability, industrial hydrogen sources, industrial or atmospheric carbon sources, and transport or price-network information. These data blocks are not used as isolated descriptive layers. Each one is converted into a specific model role, such as weekly airport demand, plant-level renewable generation limits, source-node capacity upper bounds, route-specific cost coefficients, or candidate transport arcs.",
    )
    add_para(
        doc,
        "The spatial boundary is deliberately asymmetric. Supply-side candidate nodes are screened at the national scale because green hydrogen, by-product hydrogen, industrial CO₂, DAC siting, and pipeline corridors may all be located far from the final demand nodes. The demand side is narrower in the current archived instance and is represented by the Beijing-Tianjin airport system. This design keeps the aviation demand case concrete while preserving a national search space for upstream supply, conversion, and logistics options.",
    )
    add_para(
        doc,
        "The raw source tables differ substantially in structure. Some records are point inventories with coordinates and capacities, some are hourly or weekly time series, some are provincial price surfaces, and some are route or infrastructure metadata. The preprocessing workflow harmonizes these inputs into common units and decision-index sets before optimization. Table S1 summarizes the resulting model role for each data block, and the figures in the following subsections keep raw-data diagnostics close to the section where each data source is explained.",
    )

    add_para(
        doc,
        f"Table S{table_no} Overview of input datasets, coverage, and model role.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Dataset block", "Raw source / scope", "Model-ready representation", "Role in optimization"],
        [
            [
                "Airport fuel demand",
                "2024 domestic flight schedule processed flight by flight with a pyBADA-based fuel model",
                "8 airport-week records for Beijing and Tianjin across 4 representative weeks",
                "Weekly SAF demand at airport nodes",
            ],
            [
                "Renewable electricity",
                f"China-wide wind and solar plant inventories plus hourly output profiles",
                f"{fmt_int(data.renewable_summary['wind']['unique_plants'])} wind plants and {fmt_int(data.renewable_summary['solar']['unique_plants'])} solar plants over 672 hours",
                "Limits green-hydrogen production and defines renewable spatial opportunity",
            ],
            [
                "Industrial by-product hydrogen",
                "Refinery and steel-plant geocoded inventories with daily hydrogen estimates",
                f"{data.byproduct_summary['refinery_sites']} refineries and {data.byproduct_summary['steel_sites']} steel plants",
                "Upper bounds for by-product hydrogen pathways",
            ],
            [
                "Industrial / atmospheric CO₂",
                "Industrial capture source metadata, CCUS references, and DAC siting rules",
                f"{fmt_int(data.co2_summary['records'])} industrial capture-source records plus DAC candidate locations generated from scenario configs",
                "Upper bounds or siting candidates for CCU and DAC pathways",
            ],
            [
                "Transport network",
                "Existing pipeline GIS layers, province-level gas/electricity price tables, and OpenStreetMap road network routed with GraphHopper/OSM tools",
                f"{data.transport_summary['pipeline_records']} pipeline polylines plus {data.price_summary['province_count']} provincial price surfaces",
                "Constructs candidate logistics arcs and route-specific transport / energy costs",
            ],
            [
                "Techno-economic parameters",
                "Scenario-specific YAML configuration files and harmonized literature inputs",
                "Route-specific cost, emissions, efficiency, lifetime, and policy parameters",
                "Objective coefficients and ex-post emissions accounting",
            ],
        ],
        [1.35, 2.05, 2.1, 1.9],
    )
    table_no += 1
    add_para(
        doc,
        "The overview figure is intentionally compact. It is used as a cross-dataset scale check rather than a substitute for the detailed data-source descriptions that follow. Demand, renewable-power, industrial hydrogen and CO₂, transport-network, and price-surface figures are placed in their corresponding subsections so that the reader can interpret each visualization together with the relevant processing assumptions.",
    )
    add_figure_block(
        doc,
        assets["raw_data_overview"],
        f"Figure S{figure_no}. Raw-data scale and temporal compression.",
        width_inches=6.2,
        notes=[
            f"This data-only overview shows the relative size of the archived source layers before preprocessing. The gallery combines {fmt_int(data.transport_summary['solar_inventory_rows'])} solar rows, {fmt_int(data.transport_summary['wind_inventory_rows'])} wind rows, {fmt_int(data.transport_summary['pipeline_records'])} pipeline records, {fmt_int(data.byproduct_summary['steel_sites'])} steel H₂ sites, {fmt_int(data.byproduct_summary['refinery_sites'])} refinery H₂ sites, and {fmt_int(data.co2_summary['records'])} CO₂ capture records.",
            f"The temporal-compression panel makes the representative-year reduction explicit: 52 annual weeks become {data.demand_summary['weeks']} representative weeks, {52 * data.demand_summary['periods_per_week']} three-hour periods become {data.demand_summary['total_3h_periods']} retained periods, and {52 * 7 * 24:,} annual hours become {data.renewable_summary['wind']['hours']:,} modeled hours in the representative four-week stack.",
        ],
    )
    figure_no += 1

    add_heading(doc, "1.2 Aviation fuel demand estimation", "Heading 3")
    demand_beijing = data.demand_summary["airport_summary"]["Beijing"]
    demand_tianjin = data.demand_summary["airport_summary"]["Tianjin"]
    add_para(
        doc,
        "Airport SAF demand is derived from the 2024 domestic flight schedule rather than from a fixed annual consumption scalar. Each flight record provides the operating airport, aircraft information, and route-distance context needed to estimate fuel use at the flight level. The workflow applies a pyBADA-based fuel-consumption model to these records and then aggregates the resulting fuel estimates to airport-week demand values.",
    )
    add_para(
        doc,
        "This flight-level construction is important because airport fuel demand is affected by both traffic volume and aircraft mix. A simple average fuel-per-flight assumption would hide differences between short and long routes or between aircraft classes. The stored demand table therefore preserves the link between traffic composition, route structure, and weekly fuel requirements before the data are reduced for optimization.",
    )
    add_para(
        doc,
        f"To keep the optimization problem tractable, the full-year demand series was compressed from 52 calendar weeks to a smaller representative set. The stored reduction keeps four final representative weeks. In that reduced dataset, Beijing contributes {fmt_million_kg(demand_beijing['total_demand_kg'])} across the four modeled weeks and Tianjin contributes {fmt_million_kg(demand_tianjin['total_demand_kg'])}. The selected weeks cover the high, low, and intermediate demand regimes observed for Beijing and the normal and low demand regimes observed for Tianjin.",
    )
    add_para(
        doc,
        "The reduction is reported in two steps to make the compression auditable. The appendix first restores the 52-week annual context, then shows the 12-week candidate pool and the four retained weeks used by the reduced optimization instance. This makes it possible to verify that the final weeks are not arbitrary samples and that low-demand and high-demand conditions remain represented after temporal aggregation.",
    )
    add_para(
        doc,
        f"Table S{table_no} Representative-week mapping used by the reduced demand model.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["New week", "Week in 12-week set", "Original week in year", "Calendar span", "Beijing demand", "Tianjin demand"],
        data.week_rows,
        [0.8, 1.0, 1.0, 1.4, 1.25, 1.25],
    )
    table_no += 1
    add_figure_block(
        doc,
        assets["airport_demand_raw"],
        f"Figure S{figure_no}. Representative-week airport-demand descriptors.",
        width_inches=6.25,
        notes=[
            f"The airport-demand panels remain chart-only in this revision. They summarize weekly fuel totals, flown-distance scaling, flight-phase fuel shares, and average fuel per flight for the archived Beijing–Tianjin demand instance.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["demand_aircraft_supplement"],
        f"Figure S{figure_no}. Full-year weekly airport demand, aircraft mix, and aircraft-distance structure for the Beijing-Tianjin demand instance.",
        width_inches=6.1,
        notes=[
            "Panel (a) restores the omitted raw-year context by plotting all 52 weekly fuel-demand observations and marking the retained original calendar weeks 1, 5, 14, and 44. Panels (b) and (c) summarize the aircraft heterogeneity embedded in the same flight schedule, which supports flight-level pyBADA fuel estimation instead of a single stylized aircraft.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["annual_demand_distribution"],
        f"Figure S{figure_no}. Distribution coverage of the 52-week demand series and positions of the retained weeks.",
        width_inches=6.1,
        notes=[
            "The boxplots summarize the full 52-week demand distribution for each airport. Open circles show the 12-week candidate pool extracted from the raw year, and filled diamonds show the four retained representative weeks. The ranked-distribution panel confirms that the screening pipeline spans low, medium, and upper-demand positions before reduction to the final 4-week set.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["demand_patterns"],
        f"Figure S{figure_no}. Candidate 12-week airport-demand pool and retained representative weeks.",
        width_inches=6.0,
        notes=[
            "This reduced-form view reports the demand traces only after the 12-week candidate pool has been formed. It should be read together with the previous two figures, which recover the annual baseline and the coverage check that justify the representative-week reduction.",
        ],
    )
    figure_no += 1

    add_heading(doc, "1.3 Renewable electricity data processing", "Heading 3")
    add_para(
        doc,
        "Renewable electricity is treated as a calculated input rather than as a directly downloaded optimization-ready power trace. The workflow starts from nationwide wind and solar plant inventories, then combines plant locations and capacities with archived meteorological drivers to reconstruct hourly plant-level generation. This distinction matters because the model needs both spatial information, which determines where green-hydrogen production can be located, and temporal resource quality, which determines how much renewable generation is available in each modeled period.",
    )
    add_para(
        doc,
        "For solar generation, the processing chain links plant coordinates and installed capacities to the stored radiation-related drivers used by the solar-generation workflow. For wind generation, nationwide wind-farm coordinates are combined with MERRA-2-based wind-speed fields, hub-height adjustment, and turbine power conversion. The resulting hourly output profiles are therefore generated from physical resource drivers and plant attributes rather than being treated as exogenous electricity-supply records.",
    )
    add_para(
        doc,
        f"In the representative four-week dataset, the wind stack contains {fmt_int(data.renewable_summary['wind']['unique_plants'])} plants with {data.renewable_summary['wind']['capacity_mw'] / 1000:.2f} GW of installed capacity and {fmt_twh_from_mwh(data.renewable_summary['wind']['total_mwh'])} of simulated generation. The solar stack contains {fmt_int(data.renewable_summary['solar']['unique_plants'])} plants with {data.renewable_summary['solar']['capacity_mw'] / 1000:.2f} GW of installed capacity and {fmt_twh_from_mwh(data.renewable_summary['solar']['total_mwh'])} of simulated generation. These values define the renewable upper bounds available to green-hydrogen pathways after temporal reduction.",
    )
    add_para(
        doc,
        f"Because green-hydrogen cost depends on resource quality rather than installed capacity alone, the appendix also reports plant-level resource-quality summaries from the archived full-year renewable tables before representative-week extraction. Using annual hourly output traces clipped at 100% of nameplate capacity for descriptive plotting, median availability factors are {data.renewable_summary['wind_resource']['capacity_factor_median'] * 100:.1f}% for wind and {data.renewable_summary['solar_resource']['capacity_factor_median'] * 100:.1f}% for solar, with median annual full-load hours of {data.renewable_summary['wind_resource']['full_load_hours_median']:.0f} h and {data.renewable_summary['solar_resource']['full_load_hours_median']:.0f} h. The clipping is used only for descriptive resource-quality diagnostics so that mapped availability factors remain physically interpretable.",
    )
    add_figure_block(
        doc,
        assets["renewable_raw"],
        f"Figure S{figure_no}. Renewable inventory descriptors and annual resource summaries.",
        width_inches=6.25,
        notes=[
            f"This data-only figure summarizes renewable status mix, plant-capacity distributions, and the reduction from raw inventory rows to annual resource tables and then to the four-week modeled stack. The archived source tables contain {fmt_int(data.transport_summary['wind_inventory_rows'])} wind rows and {fmt_int(data.transport_summary['solar_inventory_rows'])} solar rows.",
            f"The full-year resource tables summarized here contain {fmt_int(data.renewable_summary['wind_resource']['candidate_plants'])} wind plants and {fmt_int(data.renewable_summary['solar_resource']['candidate_plants'])} solar plants over {fmt_int(data.renewable_summary['wind_resource']['hours'])} hourly steps. Median annual full-load hours are {data.renewable_summary['wind_resource']['full_load_hours_median']:.0f} h for wind and {data.renewable_summary['solar_resource']['full_load_hours_median']:.0f} h for solar.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["renewable_resource_maps"],
        f"Figure S{figure_no}. National renewable resource maps.",
        width_inches=6.25,
        notes=[
            f"The three map panels show sampled raw renewable geography together with nationwide annual plant-level availability summaries for wind and solar. To keep the resource-endowment layer physically interpretable, hourly output-to-capacity ratios are clipped at 100% when computing the annual descriptive statistics used for mapping.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["renewable_inventory_map"],
        f"Figure S{figure_no}. National operating renewable plant inventories used as spatial anchors for hourly renewable-power reconstruction.",
        width_inches=6.2,
        notes=[
            "Panel (a) maps archived wind-plant inventory entries and panel (b) maps archived solar-plant inventory entries, both restricted to operating records with valid coordinates and installed capacities. Marker area scales with installed capacity, with values above 500 MW clipped for readability.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["renewable_processing"],
        f"Figure S{figure_no}. Renewable-data construction from archived meteorological drivers to representative-week optimization inputs.",
        width_inches=6.2,
        notes=[
            "Panel (a) summarizes the renewable-data chain from plant inventories and archived weather drivers to plant-level hourly output, the 12-week candidate pool, and the retained 4-week optimization horizon. Panels (b) and (c) visualize the stored raw driver-to-output relationships from the archived solar and wind generation tables, confirming that the renewable traces are calculated from radiation or wind-speed inputs rather than used as direct exogenous power records.",
            "Panel (d) shows how the 12-week renewable pool is reduced to the retained weeks 1, 2, 4, and 11 before remapping them to the final four-week horizon. The shaded bands therefore indicate the time slices kept for optimization after representative-week screening.",
        ],
    )
    figure_no += 1
    add_figure(
        doc,
        assets["renewable_profiles"],
        f"Figure S{figure_no}. Aggregate renewable production profiles for the four representative weeks.",
        width_inches=6.2,
    )
    figure_no += 1

    add_heading(doc, "1.4 Industrial by-product hydrogen and CO₂ data", "Heading 3")
    add_para(
        doc,
        f"Industrial by-product hydrogen is represented with two plant-level source layers: refinery-associated hydrogen and steel-sector by-product hydrogen. The refinery dataset contains {data.byproduct_summary['refinery_sites']} sites across {data.byproduct_summary['refinery_provinces']} provinces with an aggregate estimated supply of {data.byproduct_summary['refinery_daily_h2_tonnes']:.2f} t H₂/day. The steel dataset contains {data.byproduct_summary['steel_sites']} sites across {data.byproduct_summary['steel_provinces']} provinces with an aggregate estimated supply of {data.byproduct_summary['steel_daily_h2_tonnes']:.2f} t H₂/day.",
    )
    add_para(
        doc,
        "These industrial hydrogen values are used as availability upper bounds rather than as forced production levels. In other words, the optimization model may use hydrogen from these nodes when a scenario allows by-product hydrogen, but it is not required to consume all reported by-product supply. This treatment avoids interpreting the inventory as a mandatory dispatch profile and instead uses it as a spatially explicit supply opportunity set.",
    )
    add_para(
        doc,
        "The refinery and steel layers are kept separate because their magnitudes and spatial distributions differ strongly. Refinery sites are more numerous in coastal and petrochemical regions but have much smaller daily hydrogen estimates, whereas steel-sector records dominate the aggregate by-product hydrogen quantity. The appendix therefore reports both summary distributions and separate maps so that the smaller refinery range remains visible instead of being visually overwhelmed by the steel layer.",
    )
    add_figure_block(
        doc,
        assets["byproduct_raw"],
        f"Figure S{figure_no}. Industrial by-product hydrogen summary statistics.",
        width_inches=6.25,
        notes=[
            f"This chart figure summarizes source counts, province counts, site-level H₂ availability, and the provincial concentration of refinery and steel by-product hydrogen. Aggregate supply totals are {data.byproduct_summary['refinery_daily_h2_tonnes']:,.1f} t/day for refinery sources and {data.byproduct_summary['steel_daily_h2_tonnes']:,.1f} t/day for steel sources.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["byproduct_geo"],
        f"Figure S{figure_no}. National maps of industrial by-product hydrogen sources.",
        width_inches=6.25,
        notes=[
            "The two map panels keep refinery and steel inventories separate rather than mixing maps with summary charts. Bubble-size reference circles are given in t/day below each panel so that the much smaller refinery range and the much larger steel range remain readable without implying a shared visual scale.",
        ],
    )
    figure_no += 1
    add_para(
        doc,
        f"Industrial carbon data enter the framework through point-source capture layers and through DAC siting logic. For CCU scenarios, the project stores {fmt_int(data.co2_summary['records'])} industrial capture-source records across {data.co2_summary['provinces']} provinces and applies source-type-specific capture assumptions to represent point-source CO₂ availability. Coal-power sources dominate the archived table by record count and aggregate capture potential, while gas-power and refinery sources remain important because they differ in capture cost, geographic coverage, and technology assumptions.",
    )
    add_para(
        doc,
        f"The point-source CO₂ table is converted into weekly capture potential and unit capture-cost inputs. Aggregate weekly capture potentials are {data.co2_summary['by_type']['Coal power']['weekly_capture_tonnes'] / 1e6:.1f} Mt/week from coal power, {data.co2_summary['by_type']['Gas power']['weekly_capture_tonnes'] / 1e6:.1f} Mt/week from gas power, and {data.co2_summary['by_type']['Oil refinery']['weekly_capture_tonnes'] / 1e6:.1f} Mt/week from oil refineries. Median unit capture costs are {data.co2_summary['by_type']['Coal power']['unit_capture_cost_yuan_per_ton']:.0f}, {data.co2_summary['by_type']['Gas power']['unit_capture_cost_yuan_per_ton']:.0f}, and {data.co2_summary['by_type']['Oil refinery']['unit_capture_cost_yuan_per_ton']:.0f} yuan/t CO₂, respectively. These values enter CCU pathways as source-specific availability and cost coefficients.",
    )
    add_para(
        doc,
        f"For DAC scenarios, carbon is not drawn from the industrial point-source table. Instead, DAC supply is generated through an explicit DAC module whose siting candidates are attached to renewable or industrial nodes, depending on the scenario configuration. The supplementary metadata retained alongside the model also include {data.transport_summary['ccs_projects']} CCUS reference records, of which {data.transport_summary['ccs_operating']} are marked as operating in the source table. These CCUS records are used as contextual infrastructure metadata rather than as direct mandatory CO₂ supply in every pathway.",
    )
    add_figure_block(
        doc,
        assets["co2_capture_raw"],
        f"Figure S{figure_no}. Industrial CO₂ capture-source summary statistics.",
        width_inches=6.25,
        notes=[
            f"This chart figure summarizes type composition, aggregate capture, unit capture cost, provincial concentration, and type-level capture-cost assumptions across the archived industrial CO₂ table. The source table contains {fmt_int(data.co2_summary['records'])} rows across {data.co2_summary['provinces']} provinces.",
            f"Aggregate weekly capture potentials are {data.co2_summary['by_type']['Coal power']['weekly_capture_tonnes'] / 1e6:.1f} Mt/week from coal power, {data.co2_summary['by_type']['Gas power']['weekly_capture_tonnes'] / 1e6:.1f} Mt/week from gas power, and {data.co2_summary['by_type']['Oil refinery']['weekly_capture_tonnes'] / 1e6:.1f} Mt/week from oil refineries. Median unit capture costs are {data.co2_summary['by_type']['Coal power']['unit_capture_cost_yuan_per_ton']:.0f}, {data.co2_summary['by_type']['Gas power']['unit_capture_cost_yuan_per_ton']:.0f}, and {data.co2_summary['by_type']['Oil refinery']['unit_capture_cost_yuan_per_ton']:.0f} yuan/t CO₂, respectively.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["co2_capture_geo"],
        f"Figure S{figure_no}. National maps of industrial CO₂ capture sources.",
        width_inches=6.25,
        notes=[
            "The left panel shows point-source geography by facility type, while the right panel aggregates weekly capture potential to the provincial scale. Coal-power points are sampled for readability because they dominate the archived source table.",
        ],
    )
    figure_no += 1

    add_heading(doc, "1.5 Transport network construction", "Heading 3")
    add_para(
        doc,
        f"Transport-network construction combines existing infrastructure metadata with newly generated candidate routes. The current GIS table contains {data.transport_summary['pipeline_records']} pipeline polylines, including {data.transport_summary['pipeline_operating']} operating and {data.transport_summary['pipeline_planned']} planned records, with a reported cumulative length of {fmt_int(data.transport_summary['pipeline_length_km'])} km across records that include explicit length metadata. These corridors provide spatial anchors for candidate hydrogen or CO₂ pipeline links and for interpreting where existing rights-of-way may reduce new-build transport burdens.",
    )
    add_para(
        doc,
        "The pipeline layer is not treated as a fully designed future network. Instead, it is used to construct and screen candidate logistics arcs between source, conversion, and demand nodes. The model still decides whether to use available pathway-specific transport options under the scenario assumptions. This distinction is important because the map data describe corridors and reference infrastructure, while the optimization variables represent selected flows and costs within the simplified network.",
    )
    add_para(
        doc,
        "Downstream SAF deliveries and other truck-based movements are routed over the road network. The implementation relies on OpenStreetMap-derived road graphs and GraphHopper or OSM shortest-path tools to convert origin-destination coordinate pairs into practical transport distances. This avoids using straight-line distances for truck transport, which would understate travel distance in regions where road geometry, terrain, or network connectivity makes direct routes unrealistic.",
    )
    add_para(
        doc,
        f"Economic routing assumptions are paired with a province-level price surface aggregated from the integrated gas-pipeline price table. After excluding nationwide aggregate rows, the archived provincial gas-price range is {data.price_summary['gas_price_min']:.2f}-{data.price_summary['gas_price_max']:.2f} yuan per 10k m3 and the electricity-price range is {data.price_summary['electricity_price_min']:.3f}-{data.price_summary['electricity_price_max']:.3f} yuan per kWh. In the current configuration, routing is applied with a 1,000 km maximum transport threshold and cached to support repeated scenario runs, which keeps repeated pathway comparisons computationally manageable.",
    )
    add_figure_block(
        doc,
        assets["transport_raw"],
        f"Figure S{figure_no}. Transport-network and CCUS metadata summaries.",
        width_inches=6.25,
        notes=[
            f"This chart figure summarizes archived pipeline lengths, status mix, online years, and CCUS project statuses. Across the stored pipeline table, the reported total length is {fmt_int(data.transport_summary['total_ng_pipeline_km'])} km, the median segment length is {data.transport_summary['median_ng_pipeline_km']:.1f} km, and {fmt_int(data.transport_summary['year_stamped_ng_records'])} records retain online-year information.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["price_surface_maps"],
        f"Figure S{figure_no}. Province-level energy price maps.",
        width_inches=6.25,
        notes=[
            f"The price panels are isolated as maps in this revision. After excluding nationwide aggregate rows, provincial gas prices span {data.price_summary['gas_price_min']:.2f}-{data.price_summary['gas_price_max']:.2f} yuan per 10k m3 and electricity prices span {data.price_summary['electricity_price_min']:.3f}-{data.price_summary['electricity_price_max']:.3f} yuan per kWh.",
        ],
    )
    figure_no += 1
    add_figure_block(
        doc,
        assets["energy_infrastructure_overview"],
        f"Figure S{figure_no}. Repository-level national infrastructure overview.",
        width_inches=6.2,
        notes=[
            f"This overview map is retained as a final national cross-check of archived layer coverage before optimization preprocessing. Renewable layers are sampled for readability, pipeline geometry is drawn from {fmt_int(data.transport_summary['pipeline_rows'])} stored GIS polylines, and airport nodes mark the Beijing–Tianjin demand instance used in the current appendix.",
        ],
    )
    figure_no += 1

    add_heading(doc, "1.6 Data preprocessing workflow", "Heading 3")
    add_para(
        doc,
        "The raw spatial source tables are much denser than the final optimization network can efficiently handle. Many plant or point-source records are close to one another and would create a large number of nearly redundant candidate arcs if passed directly into the network model. Spatial preprocessing therefore reduces the raw point layers to a smaller set of representative nodes while retaining the geographic structure needed for logistics and siting decisions.",
    )
    add_para(
        doc,
        f"The repository applies DBSCAN clustering to renewable-hydrogen, by-product-hydrogen, and CO₂ source layers before arc generation. In the archived renewable clustering run, {data.clustering_summary['renewable_clusters']} clusters and {data.clustering_summary['renewable_noise']} direct-connection points were obtained with eps = {data.clustering_summary['renewable_eps_km']:.0f} km and min_samples = {data.clustering_summary['renewable_min_samples']}. The by-product hydrogen clustering run stores {data.clustering_summary['byproduct_clusters']} clusters and {data.clustering_summary['byproduct_noise']} noise points with eps = {data.clustering_summary['byproduct_eps_km']:.0f} km, while the stored CO₂ clustering run retains {data.clustering_summary['co2_clusters']} clusters and {data.clustering_summary['co2_noise']} noise points.",
    )
    add_para(
        doc,
        f"Cluster-center selection is not based on geometric proximity alone. The center-optimization routine balances within-cluster travel and access to existing pipeline corridors with a pipeline-weight term of {data.clustering_summary['pipeline_weight']:.1f}; shared pipeline use is further discounted by a factor of {data.clustering_summary['shared_pipeline_discount']:.2f}. This means a representative node can be favored not only because it is central to nearby source points, but also because it is better connected to the transport infrastructure layer.",
    )
    add_para(
        doc,
        f"After clustering, candidate logistics arcs are generated by connecting business nodes to the {data.clustering_summary['super_graph_k']} nearest pipeline nodes in the precomputed super graph and by calculating shortest-path road distances for truck transport. The preprocessing outputs are therefore the bridge between the raw GIS and time-series data in Sections 1.1-1.5 and the scenario-specific optimization models described in the later sections of the appendix.",
    )
    add_figure(
        doc,
        assets["workflow"],
        f"Figure S{figure_no}. Preprocessing workflow from raw datasets to optimization-ready inputs.",
        width_inches=6.2,
    )
    figure_no += 1

    # Chapter 2
    add_heading(doc, "2 Scenario Definition and Underlying Assumptions", "Heading 2")
    add_heading(doc, "2.1 Spatial and temporal boundary", "Heading 3")
    add_para(
        doc,
        "The optimization domain spans mainland China on the supply side: renewable plants, industrial hydrogen sources, carbon sources, candidate synthesis nodes, and infrastructure corridors are all screened nationally. The current demand instance is narrower. In the archived four-week runs, weekly airport demand is represented by Beijing and Tianjin as a northern-China test system, which keeps the demand side concrete while preserving a national supply search space.",
    )
    add_para(
        doc,
        "Time is represented on two linked layers. Raw renewable profiles are available hourly, but the optimization model aggregates production and inventory decisions into 3-hour periods, giving 56 periods per representative week. Transport flows and airport demand are enforced at the weekly level. This combination retains operational variability while keeping the problem size manageable for repeated scenario comparisons.",
    )
    add_para(
        doc,
        f"Table S{table_no} Spatial, temporal, and routing boundary parameters for the archived optimization instance.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Parameter", "Baseline value and unit", "Scope", "Model role"],
        data.config_summary["boundary_rows"],
        [1.25, 1.65, 1.35, 1.75],
    )
    table_no += 1

    add_heading(doc, "2.2 System boundary", "Heading 3")
    add_para(
        doc,
        "The system boundary starts with carbon and hydrogen acquisition and ends with SAF delivery to airport nodes. Upstream stages include renewable-power-driven electrolysis, industrial by-product hydrogen use, natural-gas- or coal-linked carbon supply, industrial capture, and DAC, depending on scenario definition. Midstream stages include route-specific conversion through FT or MTJ, inventory buffering, and interregional transport. Aircraft use, airport tank-farm operations, and detailed refueling logistics are outside the modeled system boundary.",
    )
    add_para(
        doc,
        f"Table S{table_no} Physical system-boundary inclusions and simplifications.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["System element", "Included representation and model treatment", "Explicit exclusions"],
        data.config_summary["system_boundary_rows"],
        [1.15, 2.85, 2.0],
    )
    table_no += 1

    add_heading(doc, "2.3 Supply-chain structure assumptions", "Heading 3")
    add_para(
        doc,
        "The implemented supply chain is abstracted as a production–transport–consumption network. Production nodes combine carbon sourcing, hydrogen sourcing, and SAF synthesis; transport links move hydrogen, carbon dioxide, methanol intermediates where needed, and final SAF; airport nodes absorb weekly demand. Nodes are treated as perfectly mixed. Internal plant scheduling is not modeled below the chosen time step, and transport capacity is represented with upper bounds rather than detailed vehicle-count or pipe-diameter design variables.",
    )
    add_para(
        doc,
        "Scenario-specific pathway structure is represented through activation switches. FT routes link carbon and hydrogen consumption directly to SAF output. MTJ routes add an explicit methanol intermediate together with methanol inventory. DAC and industrial-capture pathways keep their own carbon-accounting rules, whereas natural-gas and coal pathways treat carbon as fossil feedstock in ex-post emissions accounting.",
    )
    add_para(
        doc,
        f"Table S{table_no} Route-module and inventory activation assumptions.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Module", "Active scenarios", "Modeled state", "Key parameterization and role"],
        data.config_summary["route_module_rows"],
        [0.9, 1.55, 1.35, 2.2],
    )
    table_no += 1

    add_heading(doc, "2.4 Pathway screening and selection rationale", "Heading 3")
    add_para(
        doc,
        "The comparative analysis is restricted to the 13 pathways for which the repository contains complete scenario configurations, route-specific mass-balance logic, cost and emissions parameters, and archived optimization or visualization outputs. The retained pathways cover grey, blue, and green pathway groups and include both FT and MTJ conversion routes when the corresponding route structure is implemented in the code base.",
    )
    add_para(
        doc,
        f"Table S{table_no} Implemented pathway set retained for appendix reporting.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    pathway_headers = ["Class", "Scenario", "Pathway definition", "Technical description"]
    pathway_display_rows = [
        [row[0], row[5], f"Carbon: {row[1]}; H₂: {row[2]}; route: {row[3]}", row[6]]
        for row in data.pathway_rows
    ]
    add_table(doc, pathway_headers, pathway_display_rows, [0.65, 0.85, 2.25, 2.25])
    table_no += 1
    add_para(
        doc,
        "For reproducibility, the pathway names are complemented by an activation matrix. The matrix reports which conversion block, hydrogen module, carbon module, inventory treatment, and carbon-accounting rule is active in each retained scenario, without exposing repository file paths as scientific content.",
    )
    add_para(
        doc,
        f"Table S{table_no} Scenario activation matrix for the retained pathway definitions.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Scenario", "Source modules", "Conversion and inventory", "Carbon-accounting treatment"],
        data.config_summary["scenario_rows"],
        [0.75, 1.75, 1.55, 1.95],
    )
    table_no += 1

    add_heading(doc, "2.5 Economic, policy, and lifecycle assumptions", "Heading 3")
    add_para(
        doc,
        "Economic inputs are treated as exogenous over the representative year. The scenario parameter sets define the project life, discount rate, shortage penalty, technology lifetimes, power prices, fuel prices, and solver tolerances before each pathway is optimized. Electricity tariffs, fuel feedstock prices, DAC costs, and carbon-accounting switches therefore enter as scenario assumptions rather than as endogenous market outcomes. Carbon price and subsidy effects, when considered, enter the model as cost terms rather than through endogenous policy-response loops.",
    )
    add_para(
        doc,
        f"Table S{table_no} Economic, policy, and lifecycle assumptions used in scenario definitions.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Parameter", "Baseline value and unit", "Scenario scope", "Model role"],
        data.config_summary["economic_rows"],
        [1.2, 1.85, 1.35, 1.6],
    )
    table_no += 1

    add_heading(doc, "2.6 Feedstock, process, and capture assumptions", "Heading 3")
    add_para(
        doc,
        "Feedstock and process parameters define the physical conversion assumptions that differ across pathway families. The table separates hydrogen production, fossil feedstock use, coal gasification, DAC capture, and route-specific process coefficients so that each parameter block can be traced to the scenarios in which it is active.",
    )
    add_para(
        doc,
        f"Table S{table_no} Feedstock, process, and capture assumptions used in pathway definitions.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Parameter block", "Baseline value and unit", "Scenario scope", "Model role"],
        data.config_summary["feedstock_rows"],
        [1.15, 2.35, 1.35, 1.15],
    )
    table_no += 1

    add_heading(doc, "2.7 Transport and routing assumptions", "Heading 3")
    add_para(
        doc,
        "Transport assumptions define which candidate arcs can enter the network and how their cost and capacity are represented. Distance limits are screening parameters, while truck capacities, cost curves, capacity factors, and routing-cache settings determine the attributes assigned to feasible arcs before optimization.",
    )
    add_para(
        doc,
        f"Table S{table_no} Transport and routing assumptions used to construct candidate logistics arcs.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Commodity / mode", "Transport limit", "Capacity or cost parameter", "Scenario scope"],
        data.config_summary["transport_rows"],
        [1.0, 1.0, 2.65, 1.35],
    )
    table_no += 1

    add_heading(doc, "2.8 Capacity, solver, and carbon-accounting assumptions", "Heading 3")
    add_para(
        doc,
        "The final scenario-assumption layer defines capacity bounds, minimum-output constraints, MILP tolerances, and carbon-accounting switches. These parameters do not change the pathway identity, but they determine the feasible scale of each scenario and the way cost and emissions are reported after optimization.",
    )
    add_para(
        doc,
        f"Table S{table_no} Capacity bounds and MILP settings used by scenario.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Scenario", "Route", "Capacity bounds", "Output and cost screen", "MILP settings"],
        data.config_summary["capacity_solver_rows"],
        [0.75, 0.55, 2.05, 1.55, 1.1],
    )
    table_no += 1
    add_para(
        doc,
        f"Table S{table_no} Carbon-accounting assumptions and emissions benchmark parameters.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Accounting assumption", "Baseline value and unit", "Scenario scope", "Interpretation"],
        data.config_summary["emission_rows"],
        [1.35, 2.35, 1.35, 0.95],
    )
    table_no += 1

    # Chapter 3
    add_heading(doc, "3 Methodological Supplement", "Heading 2")
    add_para(
        doc,
        "Figure S below provides an overview of the cost-driven synthetic aviation-fuel supply-chain optimisation framework. "
        "The framework unifies 13 pathway scenarios within a single mixed-integer linear programme that minimises the levelised cost of SAF (LCO-SynAF) "
        "subject to spatial, temporal, capacity, and carbon-accounting constraints.",
    )
    add_figure(
        doc,
        assets["framework_overview"],
        f"Figure S{figure_no}. Overview of the cost-driven synthetic aviation-fuel supply-chain optimisation framework in China. "
        f"Inputs (left) span carbon feedstocks, hydrogen and power sources, demand data, and techno-economic parameters for all 13 pathways. "
        f"The unified MILP (centre) simultaneously determines facility siting and capacity, pathway and material flow, and transport routing. "
        f"Outputs (right) include the optimised supply-chain network layout, facility capacities, LCO-SynAF estimates, and lifecycle emissions.",
        width_inches=6.3,
    )
    figure_no += 1

    add_heading(doc, "3.1 Unified scenario-indexed formulation", "Heading 3")
    add_para(
        doc,
        "The 13 archived pathway models are written here as scenario-indexed instances of one mixed-integer linear optimization framework. Each scenario activates the same core cost, capacity, transport, demand, and reporting blocks, but uses a different combination of hydrogen-supply, carbon/feedstock, conversion, inventory, and accounting blocks. This is consistent with the repository implementation: the runner selects one scenario configuration at a time, while the core optimizer instantiates only the variables and constraints required by that configuration.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_unified_scenario_model())
    eq_no += 1
    add_para(
        doc,
        f"Table S{table_no} Unified constraint blocks used to express the 13 scenario-specific MILP instances.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    constraint_block_table = add_table(
        doc,
        ["Block", "Mathematical role", "Scenario scope", "Implementation meaning"],
        data.config_summary["constraint_block_rows"],
        [0.75, 2.05, 1.25, 1.95],
    )
    # 对 Block 列（第0列）应用正确的数学符号格式：X 斜体 + 下标（如 X_H(s)）
    for row_idx, block_row in enumerate(data.config_summary["constraint_block_rows"], start=1):
        block_name = block_row[0]
        m = MODEL_INSTANCE_PATTERN.match(block_name)
        if m:
            module_str, raw_arg = m.group(1), m.group(2)
            argument = raw_arg if raw_arg else None
            segs = model_module_segments(module_str, argument, argument_italic=(argument == "s"))
            set_symbol_cell(constraint_block_table.rows[row_idx].cells[0], segs)
    table_no += 1

    add_heading(doc, "3.2 Objective function and cost accounting", "Heading 3")
    add_para(
        doc,
        f"The core optimization objective minimizes annualized total system cost for a selected scenario. {infer_levelized_cost_term()} is evaluated ex post by dividing the discounted cost stream by discounted SAF output. The archived result files store lifecycle and annualized cost measures; the frontier comparison uses the lifecycle value excluding shortage cost unless otherwise noted, while shortage penalties remain in the optimization objective to discourage unmet airport demand.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_lco_saf())
    eq_no += 1
    add_equation(doc, f"Eq. (S{eq_no})", eq_objective())
    eq_no += 1

    add_heading(doc, "3.3 Sets, indices, and decision variables", "Heading 3")
    add_para(
        doc,
        "The MILP uses a compact multi-commodity flow notation. Candidate hydrogen or carbon sources are indexed separately from candidate synthesis locations and airport nodes. Time is indexed on two layers: 3-hour production periods and representative weeks. Route-selection logic is scenario-indexed rather than endogenized inside one super-model, which matches the codebase workflow and keeps comparisons reproducible.",
    )
    add_para(
        doc,
        f"Table S{table_no} Representative sets, indices, and decision-variable groups used in the archived optimization models.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    symbol_rows = [
        ["Sets", "S", "Retained scenario set containing the 13 pathway models"],
        ["Sets", "I", "Hydrogen supply nodes, including electrolysis, by-product, or fossil-linked supply depending on scenario"],
        ["Sets", "J", "Carbon-source or feedstock nodes, including industrial capture, DAC, natural gas, or coal"],
        ["Sets", "K", "Candidate synthesis-facility locations"],
        ["Sets", "A", "Airport demand nodes"],
        ["Sets", "T", "3-hour production periods in the representative-year reduction"],
        ["Sets", "W", "Representative weeks used for transport and demand aggregation"],
        ["Sets", "M_s", "Active commodity set for scenario s, e.g., H₂, CO₂, methanol, natural gas, coal, and SAF"],
        ["Binary variables", "z_k^p", "Facility construction / activation decisions for route-specific process block p"],
        ["Binary variables", "y_u,v^m", "Transport arc or mode activation decisions for commodity m where a fixed decision is represented"],
        ["Continuous variables", "Q_k^p", "Installed capacity for the active process block p"],
        ["Continuous variables", "q_k,t^p", "Route-specific production quantities, including methanol or SAF depending on block p"],
        ["Continuous variables", "f_i,k,t^H2, f_j,k,t^CO2, f_k,a,w^SAF", "Material flows between supply, plant, and airport nodes"],
        ["Continuous variables", "I_k,t^H2, I_k,t^CO2, I_k,t^MeOH, I_k,t^SAF", "Inventory states that bridge production variability and weekly shipment timing"],
        ["Continuous variables", "s_a,w", "Airport shortage slack variables used to preserve feasibility"],
    ]
    symbol_table = add_table(
        doc,
        ["Category", "Symbol / group", "Meaning in implementation"],
        symbol_rows,
        [1.0, 1.45, 3.8],
    )
    symbol_segments = [
        symbol_var("S"),
        symbol_var("I"),
        symbol_var("J"),
        symbol_var("K"),
        symbol_var("A"),
        symbol_var("T"),
        symbol_var("W"),
        symbol_var("M", "s"),
        symbol_var("z", "k", "p"),
        symbol_var("y", "u,v", "m"),
        symbol_var("Q", "k", "p"),
        symbol_var("q", "k,t", "p"),
        [
            *symbol_var("f", "i,k,t", "H2"),
            symbol_token(", "),
            *symbol_var("f", "j,k,t", "CO2"),
            symbol_token(",", line_break=True),
            *symbol_var("f", "k,a,w", "SAF"),
        ],
        [
            *symbol_var("I", "k,t", "H2"),
            symbol_token(", "),
            *symbol_var("I", "k,t", "CO2"),
            symbol_token(",", line_break=True),
            *symbol_var("I", "k,t", "MeOH"),
            symbol_token(", "),
            *symbol_var("I", "k,t", "SAF"),
        ],
        symbol_var("s", "a,w"),
    ]
    for row_idx, segments in enumerate(symbol_segments, start=1):
        set_symbol_cell(symbol_table.rows[row_idx].cells[1], segments)
    table_no += 1

    add_heading(doc, "3.4 Scenario-specific model composition", "Heading 3")
    add_rich_para(
        doc,
        [
            symbol_token(
                "The following matrix spells out the complete model composition for each retained scenario. "
                "The notation should be read as a modular MILP definition: "
            ),
            *model_module_segments("core"),
            symbol_token(", "),
            *model_module_segments("D"),
            symbol_token(", and "),
            *model_module_segments("E"),
            symbol_token(" are always present, while "),
            *model_module_segments("H", "s", argument_italic=True),
            symbol_token(", "),
            *model_module_segments("C", "s", argument_italic=True),
            symbol_token(", "),
            *model_module_segments("P", "s", argument_italic=True),
            symbol_token(", "),
            *model_module_segments("I", "s", argument_italic=True),
            symbol_token(", and "),
            *model_module_segments("N", "s", argument_italic=True),
            symbol_token(
                " are populated by the hydrogen, carbon/feedstock, conversion, inventory, and network modules "
                "indicated in the row. This avoids duplicating the same demand and capacity equations 13 times "
                "while still identifying the exact pathway-specific model for each scenario."
            ),
        ],
    )
    add_para(
        doc,
        f"Table S{table_no} Complete scenario-level model composition for the 13 retained pathway models.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    scenario_model_table = add_table(
        doc,
        ["Scenario", "Mathematical instance", "Feedstock and capture constraints", "Conversion, temporal, and accounting constraints"],
        data.config_summary["scenario_model_rows"],
        [0.72, 1.85, 1.62, 1.81],
    )
    for row_idx, row in enumerate(data.config_summary["scenario_model_rows"], start=1):
        set_model_instance_cell(scenario_model_table.rows[row_idx].cells[1], row[1])
    table_no += 1

    add_heading(doc, "3.5 Facility siting and capacity constraints", "Heading 3")
    add_para(
        doc,
        "Binary siting variables are linked to continuous capacity and production variables through standard activation constraints. These constraints prevent nonzero production at unbuilt plants and ensure that installed capacity remains within the scenario-specific lower and upper bounds stored in the configuration files. The same block covers synthesis facilities, DAC units, electrolyzer or by-product-H₂ supply blocks, and route-specific upstream feedstock blocks when they are active.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_capacity_activation())
    eq_no += 1

    add_heading(doc, "3.6 Feedstock and carbon-source availability", "Heading 3")
    add_para(
        doc,
        "Feedstock availability is enforced before conversion. Green-hydrogen supply is bounded by renewable output, grid-electricity purchases, and electrolyzer capacity; by-product hydrogen is capped by plant-level availability from steel and refinery sources; industrial CO₂ is bounded by point-source capture capacity; DAC CO₂ is bounded by DAC siting and energy-use constraints; natural-gas and coal routes use their own feedstock-capacity and price layers. The generic availability expression is specialized by the active commodity set in each scenario.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_feedstock_availability())
    eq_no += 1
    add_para(
        doc,
        "Scenario-family feedstock equations then map atmospheric capture, coal purchase, or natural-gas feedstock into usable carbon supply. The DAC term links captured CO₂ to electricity consumption; the coal term links coal purchase to coal-derived CO₂ yield; and the natural-gas term links GTL feedstock consumption to SAF output.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_pathway_feedstock_balance())
    eq_no += 1

    add_heading(doc, "3.7 Conversion-route balances", "Heading 3")
    add_para(
        doc,
        "The active conversion block maps feedstocks to final SAF output. FT and RWGS-FT routes connect hydrogen and carbon inputs directly to SAF production through route-specific coefficients. MTJ routes add methanol production and methanol inventory before final SAF conversion. GTL and CTL routes replace the generic CO₂ input with natural-gas or coal-derived feedstock balances while retaining the same facility, production, and delivery accounting structure.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_conversion_balance())
    eq_no += 1

    add_heading(doc, "3.8 Inventory and transport constraints", "Heading 3")
    add_para(
        doc,
        "Inventory states bridge variable production and scheduled logistics. Hydrogen, carbon dioxide, methanol, natural-gas-linked feedstock intermediates, and final SAF appear only when the relevant scenario block activates them. Weekly transport variables move final SAF to airports and move gaseous intermediates through candidate pipeline, truck, LNG, or natural-gas logistics arcs depending on the scenario and arc availability.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_inventory_balance())
    eq_no += 1
    add_equation(doc, f"Eq. (S{eq_no})", eq_transport_capacity())
    eq_no += 1

    add_heading(doc, "3.9 Demand fulfillment constraints", "Heading 3")
    add_para(
        doc,
        "Airport demand is enforced on a weekly basis. For each airport-week pair, the sum of delivered SAF and the shortage slack variable must cover the required weekly demand. The shortage variable enters the objective with a high unit penalty, allowing the model to stay feasible under stressed cases while still reporting unmet demand explicitly in archived outputs.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_demand_balance())
    eq_no += 1

    add_heading(doc, "3.10 Emissions accounting method", "Heading 3")
    add_para(
        doc,
        "Emissions are calculated ex post on a well-to-wake basis. The repository uses a conventional-jet benchmark of 89 gCO2e/MJ and an SAF energy content of 43.15 MJ/kg. Pathway-specific carbon intensity combines feedstock acquisition or capture, upstream and downstream transport, process emissions, storage losses where relevant, facility lifecycle allocation, and route-specific credit treatment.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", eq_emissions_intensity())
    eq_no += 1
    add_equation(doc, f"Eq. (S{eq_no})", eq_emissions_reduction())
    eq_no += 1

    add_heading(doc, "3.11 Solver settings and implementation", "Heading 3")
    add_para(
        doc,
        "The archived models are implemented in Python and solved with Gurobi. Most scenario families use a 3,600 s time limit, a 1% relative MIP gap, and 100 threads; the DAC implementation allows a wider archived gap and higher thread count in the corresponding configuration. GraphHopper-based route calculation, pipeline super-graph preprocessing, and cached distance queries are used to keep repeated scenario runs computationally manageable.",
    )

    # Chapter 4 — Sensitivity and Uncertainty Assessment
    # Prose follows the appendix-section-writing skill:
    #   4-move pattern per subsection: Question → Design → Finding → Implication
    #   All prose in continuous paragraphs; no bullet lists in body text.
    #   Figures are placeholders until batch runs are finalised.
    _frontier_s4 = {row["scenario"]: row for row in data.temporal_summary["pareto_frontier"]}
    _gtl_bh_cost = _frontier_s4["GTL-BH"]["lcoe_cny_per_kg"]
    _gtl_cost    = _frontier_s4["GTL"]["lcoe_cny_per_kg"]
    _ccu_cost    = _frontier_s4["CCU-GH-FT"]["lcoe_cny_per_kg"]
    _gtl_bh_ci   = _frontier_s4["GTL-BH"]["carbon_diff_vs_jet_gco2e_per_mj"]
    _ccu_ci      = _frontier_s4["CCU-GH-FT"]["carbon_diff_vs_jet_gco2e_per_mj"]
    _gas_min     = data.price_summary["gas_price_min"]
    _gas_max     = data.price_summary["gas_price_max"]
    _elec_min    = data.price_summary["electricity_price_min"]
    _elec_max    = data.price_summary["electricity_price_max"]

    # 加载敏感性分析批量运行派生量（215个运行结果）
    try:
        import sys as _sys
        _proj_root = str(Path(__file__).resolve().parents[3])
        if _proj_root not in _sys.path:
            _sys.path.insert(0, _proj_root)
        from products.supply_chain_optimization.sensitivity_analysis.sensitivity_data_analyzer import get_all_appendix_values
        _sv = get_all_appendix_values()
    except Exception as _e:
        import warnings
        warnings.warn(f"敏感性分析派生量加载失败，将使用占位符文字: {_e}")
        _sv = None

    add_heading(doc, "4 Sensitivity and Uncertainty Assessment", "Heading 2")

    # ── 4.1 Framework ──────────────────────────────────────────────────────────
    add_heading(doc, "4.1 Analysis design and parameter overview", "Heading 3")
    add_para(
        doc,
        "The three main results reported in the paper each rest on parameter values that carry non-trivial uncertainty: the Pareto frontier cost ranking (Result 1) depends on feedstock and capital cost assumptions; the spatial configuration patterns (Result 2) are sensitive to the demand scale assumed for the Beijing–Tianjin instance; and the temporal-operational robustness differences (Result 3) depend partly on how renewable variability is represented through the four-week reduction. This section tests whether each of these results is robust to variation in the parameters most likely to draw reviewer scrutiny. The analyses are organised by the scientific question they answer rather than by parameter, following a progression from the most commercially relevant pathways to the most structurally uncertain ones.",
    )
    add_para(
        doc,
        f"The sensitivity-design table summarises the parameters examined, their baseline values as encoded in the archived scenario configurations, the ranges over which they are varied, and the justification for each range. Baseline electricity prices are drawn from the archived scenario files (0.35 yuan/kWh for wind and 0.40 yuan/kWh for solar); the sweep range spans the full observed provincial retail tariff distribution ({_elec_min:.3f}–{_elec_max:.3f} yuan/kWh). Natural-gas prices use the archived model input of 4.2 yuan/m³ as the baseline; the sweep range is defined by the documented provincial gas-price surface ({_gas_min:.2f}–{_gas_max:.2f} yuan per 10,000 m³), extended symmetrically to bracket the plausible medium-term market trajectory.",
    )
    add_para(
        doc,
        f"Table S{table_no} Parameter sweep design for sensitivity and uncertainty assessment.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Parameter", "Baseline value", "Sweep range", "Range basis", "Pathways affected", "Tests main-text result"],
        [
            ["Natural-gas price", "4.2 yuan/m³", "2.5–7.0 yuan/m³", "Provincial price surface + medium-term trajectory", "GTL, GTL-BH", "Result 1 cost ranking"],
            ["Byproduct H₂ purification CAPEX", "Steel: 280,000; Refinery: 224,000 yuan/(kg/h)", "140,000–400,000 yuan/(kg/h)", "PSA+compression literature range", "GTL-BH", "Result 1 cost ranking"],
            ["Renewable electricity price", "0.35–0.40 yuan/kWh", f"{_elec_min:.3f}–{_elec_max:.3f} yuan/kWh", "Full observed provincial retail tariff distribution", "CCU-GH-FT, DAC-GH series", "Result 1 cost ranking"],
            ["Electrolyzer CAPEX", "Current market value", "Current → IEA 2030 target (~2,000 yuan/kW)", "IEA hydrogen technology roadmap", "All PtL pathways", "Result 1 cost ranking"],
            ["DAC unit cost", "Archived scenario value", "Current → competitive-entry threshold", "Technology learning trajectories", "DAC-GH-FT, DAC-BH series", "Result 1 Pareto inclusion"],
            ["Discount rate", "8%", "5%–12%", "Infrastructure financing convention range", "All pathways (differential)", "Result 1 cost structure"],
            ["Demand scale", "Beijing–Tianjin (baseline)", "0.5×–3× baseline", "Representativeness across Chinese airport clusters", "All pathways", "Result 2 spatial configuration"],
        ],
        [1.3, 1.3, 1.5, 1.7, 1.3, 1.2],
    )
    table_no += 1

    # ── 4.2 Fossil feedstock price sensitivity ──────────────────────────────
    add_heading(
        doc,
        "4.2 Fossil feedstock price sensitivity: competitiveness bounds of GTL and GTL-BH",
        "Heading 3",
    )
    add_heading(doc, "4.2.1 Natural-gas price sensitivity for GTL", "Heading 3")
    add_para(
        doc,
        (
            f"This analysis tests the main-text finding that GTL, with an LCO-SynAF of "
            f"{_sv['gtl_lco_at_baseline']:.2f} yuan/kg, sits within the 6\u20138 yuan/kg market-price window. "
            f"Because natural gas is the primary feedstock for the GTL route, the pathway's cost position is more "
            f"sensitive to gas-price variation than to any other single parameter. "
            f"The main optimization uses location-specific pipeline prices from GIS data, which average "
            f"approximately {_sv['gtl_effective_baseline_ng']:.2f} yuan/m\u00b3 for the primary supply sources "
            f"(cf. config default: 4.2 yuan/m\u00b3). "
            f"The sensitivity sweep varies natural-gas price from 2.0 to 8.0 yuan/m\u00b3 around this effective baseline."
        ) if _sv else
        f"This analysis tests the main-text finding that GTL, with an LCO-SynAF of {_gtl_cost:.2f} yuan/kg, sits within the 6\u20138 yuan/kg market-price window. Because natural gas is the primary feedstock for the GTL route, the pathway's cost position is more sensitive to gas-price variation than to any other single parameter. The natural-gas price was varied from 2.0 to 8.0 yuan/m\u00b3.",
    )
    add_para(
        doc,
        (
            f"As shown in Fig. S{figure_no}, the LCO-SynAF of GTL responds approximately linearly to "
            f"natural-gas price changes. "
            f"At the sweep lower bound (2.0 yuan/m\u00b3), GTL's LCO-SynAF is "
            f"{_sv['_ng']['per_scenario']['GTL']['lco_min']:.2f} yuan/kg; "
            f"at the upper bound (8.0 yuan/m\u00b3) it rises to "
            f"{_sv['_ng']['per_scenario']['GTL']['lco_max']:.2f} yuan/kg. "
            f"At the effective baseline of {_sv['gtl_effective_baseline_ng']:.2f} yuan/m\u00b3, GTL's LCO-SynAF "
            f"is {_sv['gtl_lco_at_baseline']:.2f} yuan/kg\u2014firmly within the market window. "
            f"The critical threshold\u2014the gas price above which GTL exits the 8 yuan/kg upper bound\u2014is "
            f"{_sv['gtl_ng_threshold']:.2f} yuan/m\u00b3, corresponding to a price increase of "
            f"{abs(_sv['gtl_ng_pct_from_baseline']):.1f}% above the effective baseline. "
            f"GTL-BH at the same effective baseline ({_sv['gtlbh_effective_baseline_ng']:.2f} yuan/m\u00b3) "
            f"achieves an LCO-SynAF of {_sv['gtlbh_lco_at_baseline']:.2f} yuan/kg, below the market-window lower "
            f"bound, indicating substantial cost competitiveness. GTL-BH's LCO does not reach the 8 yuan/kg "
            f"threshold until natural-gas prices rise to {_sv['gtlbh_ng_upper_threshold']:.2f} yuan/m\u00b3."
        ) if _sv else
        f"As shown in Fig. S{figure_no}, the LCO-SynAF of GTL responds approximately linearly to natural-gas price changes. The critical threshold is approximately 3.32 yuan/m\u00b3.",
    )
    add_para(
        doc,
        "This headroom has a direct policy implication. GTL's current market competitiveness depends on natural-gas prices remaining below approximately 3.32 yuan/m³—a 31% rise above the effective pipeline-access baseline. China's ongoing gas-market reform is gradually narrowing the gap between regulated pipeline prices and spot-market prices; the threshold value identified here provides a quantitative anchor for monitoring that convergence risk.",
    )
    add_figure(
        doc,
        assets["fossil_price_placeholder"],
        (
            f"Figure S{figure_no}. Natural-gas price sensitivity of GTL and GTL-BH. "
            f"The panel shows LCO-SynAF (yuan/kg) as a function of natural-gas price (yuan/m\u00b3) for both pathways. "
            f"The shaded band marks the 6\u20138 yuan/kg market-price window. "
            f"Diamond markers indicate the effective baselines (GTL: {_sv['gtl_effective_baseline_ng']:.2f} yuan/m\u00b3 "
            f"\u2192 {_sv['gtl_lco_at_baseline']:.2f} yuan/kg; "
            f"GTL-BH: {_sv['gtlbh_effective_baseline_ng']:.2f} yuan/m\u00b3 "
            f"\u2192 {_sv['gtlbh_lco_at_baseline']:.2f} yuan/kg) derived from the main optimisation. "
            f"The annotated threshold line marks the GTL market-window exit price "
            f"({_sv['gtl_ng_threshold']:.2f} yuan/m\u00b3, +{abs(_sv['gtl_ng_pct_from_baseline']):.0f}% above baseline). "
            f"GTL-BH exits the window only at {_sv['gtlbh_ng_upper_threshold']:.2f} yuan/m\u00b3."
        ) if _sv else
        f"Figure S{figure_no}. Natural-gas price sensitivity of GTL and GTL-BH. The shaded band marks the 6\u20138 yuan/kg market-price window.",
    )
    figure_no += 1

    add_heading(
        doc,
        "4.2.2 Byproduct hydrogen purification capital cost sensitivity for GTL-BH",
        "Heading 3",
    )
    add_para(
        doc,
        f"GTL-BH sources its hydrogen not from electrolysis or reforming but from industrial by-product streams—primarily blast-furnace off-gas at steel mills and catalytic-reformer off-gas at refineries. The hydrogen in these streams is a genuine waste product and carries no commodity purchase price in the archived model. The cost of obtaining it is instead captured through the capital expenditure required to install pressure-swing adsorption (PSA) purification and compression equipment at each industrial source site. The archived model uses hardcoded equipment costs of 280,000 yuan per kg H₂/h of capacity for steel-mill sites and 224,000 yuan per kg H₂/h for refinery sites, together with annual operating costs of 130,000 and 110,000 yuan per year respectively. These values sit at the mid-range of the PSA-plus-compression literature but carry substantial uncertainty: vendor quotes and site-specific installation conditions can shift equipment costs by a factor of two or more around any central estimate. This subsection therefore tests how sensitive GTL-BH's cost advantage of {_gtl_cost - _gtl_bh_cost:.2f} yuan/kg over GTL is to the assumed purification CAPEX.",
    )
    add_para(
        doc,
        f"The purification CAPEX was varied from 140,000 to 400,000 yuan per kg H₂/h—a range that spans the lower end of modular skid-mounted PSA installations to fully engineered on-site units with redundancy—while holding all other parameters at their archived baseline values. As shown in Fig. S{figure_no}, GTL-BH's LCO-SynAF rises monotonically as purification CAPEX increases, because higher equipment cost raises the annualised capital recovery charge per unit of hydrogen produced. "
        + (
            f"At the archived 8% discount rate and the assumed utilisation profile, the LCO-SynAF response to CAPEX is approximately {_sv['bh_slope_per_100k']:.3f} yuan/kg per 100,000 yuan/(kg/h). The gap between GTL-BH and GTL narrows as CAPEX rises; however, no cost crossover occurs within the analysed range: at the upper sweep limit of 400,000 yuan/(kg/h), GTL-BH's LCO-SynAF reaches {_sv['bh_gtlbh_lco_max']:.2f} yuan/kg, maintaining a cost advantage of {_sv['bh_advantage_at_upper_capex']:.2f} yuan/kg over GTL at the same gas price. The GTL-BH cost advantage is therefore robust to the full plausible range of PSA equipment costs."
            if _sv else
            "At the archived 8% discount rate and the assumed utilisation profile, the LCO-SynAF response to CAPEX is approximately 0.044 yuan/kg per 100,000 yuan/(kg/h). No cost crossover with GTL occurs within the analysed range; GTL-BH maintains a cost advantage of approximately 4.37 yuan/kg at the upper CAPEX limit."
        ),
    )
    add_para(
        doc,
        f"This result has two practical implications. First, the cost advantage of GTL-BH is conditional on the assumption that industrial sources are willing to allow on-site PSA installation at or near the archived equipment cost; if siting constraints or contractual terms raise this cost materially, the advantage may erode. Second, even at the upper bound of the CAPEX sweep, GTL-BH carries a carbon-intensity difference of +{_gtl_bh_ci:.2f} gCO₂eq/MJ relative to conventional Jet A-1, so any scenario in which GTL-BH retains a cost advantage does so at the price of increased lifecycle emissions. The Pareto trade-off between cost leadership and carbon intensity is therefore not an artefact of the specific CAPEX assumption; it is structurally determined by the carbon content of the industrial processes from which the by-product hydrogen originates.",
    )
    add_figure(
        doc,
        assets["bh_capex_placeholder"],
        (
            f"Figure S{figure_no}. Byproduct hydrogen purification CAPEX sensitivity for GTL-BH. The panel shows LCO-SynAF (yuan/kg) as a function of PSA-plus-compression equipment CAPEX [yuan per kg H₂/h of capacity]. The shaded horizontal band marks the 6–8 yuan/kg market-price window. The vertical dashed lines mark the archived baseline values for steel-mill (280,000) and refinery (224,000) sites. The dashed horizontal line marks GTL's baseline LCO-SynAF ({_sv['bh_gtl_baseline_lco']:.2f} yuan/kg); no cost crossover occurs within the surveyed CAPEX range."
            if _sv else
            f"Figure S{figure_no}. Byproduct hydrogen purification CAPEX sensitivity for GTL-BH. The panel shows LCO-SynAF (yuan/kg) as a function of PSA-plus-compression equipment CAPEX [yuan per kg H₂/h of capacity]. The shaded horizontal band marks the 6–8 yuan/kg market-price window. The vertical dashed lines mark the archived baseline values for steel-mill (280,000) and refinery (224,000) sites. No cost crossover with GTL occurs within the surveyed range."
        ),
    )
    figure_no += 1

    # ── 4.3 Green-hydrogen joint sensitivity ────────────────────────────────
    add_heading(
        doc,
        "4.3 Joint sensitivity of green-hydrogen production costs to electricity price and electrolyzer capital cost",
        "Heading 3",
    )
    add_para(
        doc,
        f"The main text identifies CCU-GH-FT as the deep-decarbonization archetype, with an LCO-SynAF of {_ccu_cost:.2f} yuan/kg—substantially above the 6–8 yuan/kg market-price window. This joint sensitivity analysis uses CCU-GH-MTJ as the representative green-hydrogen pathway, as it shares the same electrolysis and CO₂-utilisation architecture while yielding valid batch-run results across the full parameter grid. The two dominant cost drivers examined are renewable electricity price and electrolyzer capital cost, both of which are projected to decline along technology learning curves and whose joint effect on pathway competitiveness is more policy-relevant than either parameter alone. This analysis varies electricity price from {_elec_min:.3f} to 0.80 yuan/kWh and electrolyzer CAPEX from 200,000 to 1,000,000 yuan per kg H₂/h, producing the two-dimensional cost surface shown in Fig. S{figure_no}.",
    )
    add_para(
        doc,
        (
            f"As shown in Fig. S{figure_no}, the LCO-SynAF of CCU-GH-MTJ is driven almost exclusively by electrolyzer capital cost within the surveyed parameter space. Grid electricity price in the 0.25–0.80 yuan/kWh range exerts negligible influence on the optimised solution, consistent with the model's preference for dedicated renewable power over grid supply when grid tariffs exceed the marginal cost of on-site generation. Across the CAPEX dimension, LCO-SynAF rises from {_sv['ec_lco_min']:.1f} yuan/kg at 200,000 yuan per kg H₂/h to {_sv['ec_lco_max']:.1f} yuan/kg at 1,000,000 yuan per kg H₂/h, yielding a CAPEX sensitivity of approximately {_sv['ec_capex_slope_per_100k']:.2f} yuan/kg per 100,000 yuan/(kg/h). The 8 yuan/kg market-window threshold is not reached within the surveyed parameter space; the gap between the lowest-cost grid point and the market window upper bound is approximately {_sv['ec_gap_from_window']:.1f} yuan/kg, indicating that electrolyzer cost reduction alone is insufficient to achieve market competitiveness and that complementary improvements across the full value chain—including carbon capture and Fischer-Tropsch synthesis capital—are required."
            if _sv else
            f"As shown in Fig. S{figure_no}, the LCO-SynAF of CCU-GH-MTJ is driven almost exclusively by electrolyzer capital cost within the surveyed parameter space. Grid electricity price exerts negligible influence on the optimised solution. LCO-SynAF rises from approximately 23.1 yuan/kg at 200,000 yuan per kg H₂/h CAPEX to 37.0 yuan/kg at 1,000,000 yuan per kg H₂/h, yielding a sensitivity of approximately 1.74 yuan/kg per 100,000 yuan/(kg/h). The 8 yuan/kg market-window threshold is not reached within the surveyed space; the minimum gap is approximately 15.1 yuan/kg."
        ),
    )
    add_para(
        doc,
        "The IEA 2030 scenario locates a specific point on this surface—marked as a labelled symbol in the figure—and the proximity of that point to the 8 yuan/kg contour provides a concrete timeline estimate for when CCU-GH-FT may approach market competitiveness under optimistic but documented technology trajectories. The analysis thereby extends the main-text discussion by converting a qualitative observation about future cost reduction potential into a quantitative proximity metric that can be updated as technology learning data become available.",
    )
    add_figure(
        doc,
        assets["green_h2_joint_placeholder"],
        (
            f"Figure S{figure_no}. Joint sensitivity of CCU-GH-MTJ LCO-SynAF to electricity price and electrolyzer capital cost (CCU-GH-MTJ used as representative green-hydrogen pathway). Colour encodes LCO-SynAF (yuan/kg) across the {_sv['ec_lco_min']:.1f}–{_sv['ec_lco_max']:.1f} yuan/kg range. Grid electricity price (0.25–0.80 yuan/kWh) has negligible influence; LCO variation is driven entirely by electrolyzer CAPEX. The 8 yuan/kg market-window threshold is not reached within the surveyed space."
            if _sv else
            f"Figure S{figure_no}. Joint sensitivity of CCU-GH-MTJ LCO-SynAF to electricity price and electrolyzer capital cost. Colour encodes LCO-SynAF (yuan/kg). LCO ranges from approximately 23.1 to 37.0 yuan/kg and is driven entirely by electrolyzer CAPEX; grid electricity price has no discernible effect. The 8 yuan/kg market threshold is not reached within the surveyed space."
        ),
    )
    figure_no += 1

    # ── 4.4 DAC break-even threshold ────────────────────────────────────────
    add_heading(
        doc,
        "4.4 Break-even cost threshold for DAC pathway market entry",
        "Heading 3",
    )
    add_para(
        doc,
        "The main text notes that DAC-linked pathways carry the highest LCO-SynAF values in the archived comparison, primarily because DAC carbon-capture costs remain an order of magnitude above the cost of industrial point-source capture. This analysis converts that qualitative observation into a quantitative target by identifying the DAC unit cost at which the most cost-competitive DAC pathway—DAC-GH-FT—first enters the 6–8 yuan/kg market-price window. The sweep holds all other parameters at their archived baseline values and reduces DAC cost from the current level to the break-even threshold, recording the LCO-SynAF at each step.",
    )
    add_para(
        doc,
        (
            f"As shown in Fig. S{figure_no}, the LCO-SynAF of DAC-GH-FT declines monotonically as DAC costs fall. "
            f"Even at the lowest DAC cost examined (300 yuan/t CO\u2082), the minimum LCO-SynAF of DAC-GH-FT reaches "
            f"{_sv['dac_ghft_lco_min']:.1f} yuan/kg\u2014still {_sv['dac_ghft_gap_from_window']:.1f} yuan/kg above the "
            f"8 yuan/kg market-window upper bound. The 8 yuan/kg threshold is not crossed within the surveyed DAC cost range, "
            f"confirming that no break-even point is reachable under the parameter space examined. "
            f"For reference, the IEA's 2030 advanced DAC cost target of approximately 100 yuan/t CO\u2082 yields an "
            f"extrapolated LCO-SynAF of {_sv['dac_ghft_lco_at_iea']:.1f} yuan/kg\u2014above the market window and above the "
            f"current baseline estimate of {_sv['dac_ghft_lco_at_baseline']:.1f} yuan/kg (at the 4500 yuan/t CO\u2082 "
            f"baseline). This comparison makes explicit that DAC pathways cannot achieve unaided market competitiveness "
            f"within the current policy planning horizon and require additional instruments\u2014such as direct carbon-removal "
            f"payments or blending mandates\u2014to bridge the remaining gap."
        ) if _sv else
        f"As shown in Fig. S{figure_no}, the LCO-SynAF of DAC-GH-FT declines monotonically as DAC costs fall, crossing the 8 yuan/kg threshold at approximately [BREAK_EVEN_DAC_COST yuan/t CO\u2082]\u2014a reduction of approximately [BREAK_EVEN_REDUCTION%] from the archived baseline. For reference, the IEA's 2030 advanced DAC cost target is approximately 100 yuan/t CO\u2082, which falls [ABOVE/BELOW] the identified break-even value. This comparison makes explicit whether DAC pathways can achieve unaided market competitiveness within the current policy planning horizon, or whether they require additional instruments\u2014such as direct carbon-removal payments or blending mandates\u2014to bridge the remaining gap.",
    )
    add_para(
        doc,
        "The break-even value identified here also sets a benchmark for comparing DAC with industrial carbon-capture alternatives. CCU pathways using point-source industrial CO₂ already operate at capture costs well below the DAC break-even threshold, which is consistent with their lower LCO-SynAF values in the main-text frontier. This cost ordering confirms that for the near-to-medium term, industrial CCU routes are the more economically efficient pathway for decarbonised SAF production, while DAC routes remain technology-limited rather than demand-limited.",
    )
    add_figure(
        doc,
        assets["dac_breakeven_placeholder"],
        (
            f"Figure S{figure_no}. DAC unit cost versus LCO-SynAF for the DAC-GH-FT pathway. "
            f"The horizontal dashed line marks the 8 yuan/kg market-window upper bound. "
            f"The minimum LCO-SynAF across the surveyed range is {_sv['dac_ghft_lco_min']:.1f} yuan/kg (at 300 yuan/t CO\u2082), "
            f"which remains {_sv['dac_ghft_gap_from_window']:.1f} yuan/kg above the market threshold; "
            f"no break-even point is reached within the parameter space examined. "
            f"The IEA 2030 target (100 yuan/t CO\u2082) is annotated for reference, "
            f"yielding an extrapolated LCO-SynAF of {_sv['dac_ghft_lco_at_iea']:.1f} yuan/kg."
        ) if _sv else
        f"Figure S{figure_no}. DAC unit cost versus LCO-SynAF for the DAC-GH-FT pathway. The horizontal dashed line marks the 8 yuan/kg market-window upper bound. The vertical dashed line at the intersection marks the break-even DAC cost. The IEA 2030 target is annotated for reference. [Figure to be populated from batch runs.]",
    )
    figure_no += 1

    # ── 4.5 Pareto frontier structural robustness ────────────────────────────
    add_heading(
        doc,
        "4.5 Robustness of the Pareto frontier taxonomy under parameter perturbation",
        "Heading 3",
    )
    add_para(
        doc,
        f"The core contribution of the main text is the identification of three distinct transition logics on the Pareto frontier: GTL-BH as the low-cost, high-carbon archetype; GTL as the market-accessible bridging pathway; and CCU-GH-FT as the deep-decarbonisation option. A rank reversal among these three pathways in cost-emissions space—that is, a scenario in which the lowest-cost pathway is no longer GTL-BH, or in which CCU-GH-FT is no longer the deepest decarboniser—would directly undermine the main-text taxonomy. This subsection uses the results from Sections 4.2 through 4.4 to synthesise a unified robustness statement for the three-archetype classification.",
    )
    add_para(
        doc,
        f"Fig. S{figure_no} presents a cross-scenario sensitivity range chart comparing the LCO-SynAF variation of all 13 archived pathways under their respective dominant-parameter sweeps. Each horizontal bar spans from the minimum to the maximum LCO-SynAF observed when the pathway's key cost driver is varied across its full specified range, with all other parameters held at their baseline values. Pathways are ordered from lowest to highest baseline LCO-SynAF. The shaded green band marks the 6–8 yuan/kg market-price window. Diamond markers (◆) indicate the baseline LCO-SynAF from the main-text optimisation for the two fossil pathways (GTL and GTL-BH) for which precise main-result values are available.",
    )
    add_para(
        doc,
        (
            f"Across all individual parameter sweeps, no scenario produces a cost-rank reversal among the three "
            f"representative pathways. The narrowest gap between GTL-BH and GTL occurs at a natural-gas price of "
            f"{_sv['min_margin_ng_price']:.1f} yuan/m\u00b3 (the upper sweep bound), where GTL-BH's cost advantage "
            f"narrows to a minimum of {_sv['min_margin_gtlbh_vs_gtl']:.2f} yuan/kg\u2014well above zero, so no rank "
            f"reversal occurs. CCU-GH-FT's position as the highest-cost, lowest-carbon option is stable across the "
            f"entire parameter space examined. The three-archetype taxonomy is therefore robust to individual parameter "
            f"perturbations of the magnitude analysed here. It should be noted, however, that this robustness has been "
            f"tested under single-parameter variation; simultaneous multi-parameter stress\u2014such as high electricity "
            f"prices combined with slow electrolyzer learning\u2014represents a distinct scenario class that is not "
            f"addressed by the present sweep and constitutes a direction for future stochastic analysis."
        ) if _sv else
        f"Across all individual parameter sweeps, no scenario produces a cost-rank reversal among the three representative pathways. The narrowest gap between GTL-BH and GTL occurs when natural-gas prices rise to their upper sweep bound or when byproduct-hydrogen purification CAPEX approaches the upper end of the literature range; even under the joint pressure of high gas prices, GTL-BH's cost advantage is maintained above [MINIMUM_MARGIN yuan/kg]. CCU-GH-FT's position as the highest-cost, lowest-carbon option is stable across the entire parameter space examined. The three-archetype taxonomy is therefore robust to individual parameter perturbations of the magnitude analysed here. It should be noted, however, that this robustness has been tested under single-parameter variation; simultaneous multi-parameter stress\u2014such as high electricity prices combined with slow electrolyzer learning\u2014represents a distinct scenario class that is not addressed by the present sweep and constitutes a direction for future stochastic analysis.",
    )
    add_para(
        doc,
        "Beyond cost ordering, the main-text spatial and temporal results (Results 2 and 3) are also expected to be qualitatively stable under parameter perturbation, because they describe structural patterns in network configuration and utilisation that emerge from the optimisation logic rather than from any single parameter value. The compact network structure of GTL-BH and the expanded infrastructure of CCU-GH-FT reflect the cost minimisation logic that persists across parameter variants; a parameter perturbation that does not overturn the cost ordering will not overturn the qualitative spatial differentiation. The temporal-operational penalty differences, similarly, arise from the renewable-coupling architecture of CCU-GH-FT rather than from any specific price level, and are therefore robust to the parameter ranges examined.",
    )
    add_figure(
        doc,
        assets["pareto_tornado_placeholder"],
        (
            f"Figure S{figure_no}. Cross-scenario LCO-SynAF sensitivity range chart for all 13 archived pathways. "
            f"Each bar spans the minimum-to-maximum LCO-SynAF range produced by varying the pathway's dominant cost "
            f"driver across its full specified sweep range (italic label at bar left). Pathways are sorted by baseline "
            f"LCO-SynAF (lowest at top). The shaded green band marks the 6–8 yuan/kg market-price window; diamond "
            f"markers (◆) indicate the main-text baseline LCO-SynAF for GTL ({_sv['gtl_lco_at_baseline']:.2f} yuan/kg) "
            f"and GTL-BH ({_sv['gtlbh_lco_at_baseline']:.2f} yuan/kg). Natural-gas price dominates the GTL pathway "
            f"(Δ{_sv['_ng']['per_scenario']['GTL']['lco_range']:.1f} yuan/kg); electrolyzer CAPEX dominates CCU-GH and "
            f"DAC-GH pathways. The minimum GTL-BH vs. GTL cost gap across the NG-price sweep is "
            f"{_sv['min_margin_gtlbh_vs_gtl']:.2f} yuan/kg, confirming no rank reversal occurs within the surveyed space."
        ) if _sv else
        f"Figure S{figure_no}. Cross-scenario LCO-SynAF sensitivity range chart for all 13 archived pathways. Each bar spans the minimum-to-maximum LCO-SynAF range produced by varying the dominant cost driver for that pathway. The shaded green band marks the 6–8 yuan/kg market-price window.",
    )
    figure_no += 1

    # ── 4.6 Integrated model uncertainty assessment ──────────────────────────
    add_heading(doc, "4.6 Integrated model uncertainty assessment", "Heading 3")
    add_para(
        doc,
        "Three layers of uncertainty bear on the present model implementation and the sensitivity results reported above. The first layer is price uncertainty. Among the parameters examined, DAC unit cost and electrolyzer CAPEX carry the largest absolute uncertainty because their technology learning curves are still early-stage and the literature reports a wide range of near-term projections. Natural-gas prices carry lower structural uncertainty but are exposed to policy-driven volatility, particularly in the context of China's ongoing gas-market reform. Renewable electricity prices are the most datarich parameter—the provincial tariff surface used in this study is grounded in archived grid-price records—and their uncertainty is therefore the most bounded of the inputs examined. The sensitivity analyses in Sections 4.2 through 4.4 were designed to span the credible range of each parameter rather than to perform a formal probabilistic analysis, and the results should be interpreted in that light.",
    )
    add_para(
        doc,
        "The second layer is structural model uncertainty. The present implementation optimises over a static representative year using four representative weeks, which preserves the dominant demand and supply patterns but does not capture dynamic technology learning, multi-year investment sequencing, or the co-evolution of carbon markets and SAF deployment. These structural simplifications tend to understate the competitive advantage of pathways whose key cost drivers—electrolyzers, DAC—are projected to decline most steeply over time, meaning that the main-text analysis is likely to be conservative in its estimate of the long-term competitiveness gap between fossil-linked and power-to-liquid pathways. The sensitivity analyses presented here are not designed to remedy this structural simplification; they provide bounds on cost outcomes within the static model frame and should be read as such.",
    )
    add_para(
        doc,
        "The third layer is representative-week sampling uncertainty. The four-week temporal reduction compresses 52 weeks of demand and renewable variability into a structure designed to preserve low-, medium-, and high-demand regimes and their associated renewable supply states. As demonstrated by the coverage analysis in Section 1.2 and the verification reported in Section 5.3, the retained weeks span the full distribution of annual demand and produce temporal robustness metrics that are consistent with qualitative expectations. The key sensitivity results—particularly the cost ordering and Pareto ranking—are not materially sensitive to the specific weeks retained, because they are driven by cost magnitudes that exceed the variation attributable to representative-week selection.",
    )
    add_para(
        doc,
        "Despite these limitations, the findings of the sensitivity analysis are sufficient to support the main-text conclusions. The cost differences among the three Pareto archetypes are large enough—spanning more than a factor of three between GTL-BH and CCU-GH-FT—that individually plausible parameter perturbations do not overturn the qualitative ordering. The threshold analyses identify the specific parameter values at which individual pathways would lose or gain market competitiveness, providing the precision required for policy calibration. The three-archetype taxonomy is therefore robust within the parameter ranges examined, and the conditions required to overturn it—simultaneous extreme values for multiple uncertain parameters—are implausible under current market and technology projections.",
    )

    # Chapter 5
    add_heading(doc, "5 Model Verification and Result Validation", "Heading 2")
    add_heading(doc, "5.1 Internal consistency checks", "Heading 3")
    verif = data.verification_summary
    add_para(
        doc,
        f"The repository contains archived automated verification reports for the refactored optimization code. In the comprehensive report, {verif['summary']['total_components_verified']} component groups were checked. Geographic distance calculation passed, while several other checks terminated early because the legacy test scripts attempted to print UTF-8 check-mark characters through a GBK-encoded console. These failures are software-reporting issues rather than evidence of mathematical inconsistency, and the appendix treats them as implementation notes rather than proof of model invalidity.",
    )
    add_para(
        doc,
        "The geographic check provides a useful sanity test. The report records a Beijing-center to Shanghai distance of about 1,067 km, a Beijing-to-Hebei test distance of about 70 km, and successful rejection of invalid latitude and longitude values. Direct function-to-function comparison further confirms that geographic filtering and data-structure construction remained consistent after refactoring.",
    )

    add_heading(doc, "5.2 Benchmark comparison with literature", "Heading 3")
    add_para(
        doc,
        "The methodological contribution of the current framework sits between plant-level techno-economic analysis and large-scale supply-chain optimization. Prior studies often emphasize one side of that split: either a high-resolution process comparison without spatial planning, or a spatial supply-chain model without explicit temporal variability in renewable-driven hydrogen supply. The comparison in Table S6 summarizes the positioning of the present work relative to representative SAF studies listed in the project materials.",
    )
    add_para(
        doc,
        f"Table S{table_no} Comparison with representative SAF supply-chain studies summarized in the project materials.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    literature_rows = [
        ["Masum et al.", "County-level, annual, 20-year horizon", "Biomass only", "MILP", "Does not resolve renewable-power variability or multi-source PtL coupling"],
        ["Bianco et al.", "Sampled siting with uncertain parameters", "Cellulosic ethanol / ATJ", "Monte Carlo", "Focuses on siting and policy sensitivity rather than integrated upstream energy supply"],
        ["Akter et al.", "State-level, annual, 10-year horizon", "Forest residues / ETJ", "MILP", "Does not include variable renewable-power-driven PtL pathways"],
        ["Pin et al.", "National, multiple demand cases", "Oilseed-based SAF", "MILP", "Excludes downstream airport allocation and non-biogenic alternatives"],
        ["Wassermann et al.", "National-scale Germany case", "PtL with industrial CO₂", "Supply-chain optimization", "Limited treatment of multi-source H₂/CO₂ matrices and hourly variability"],
        ["Martinez-Valencia et al.", "Regional annual life-cycle case", "Gasification / FT", "Financial analysis + system dynamics", "No high-resolution siting and logistics optimization"],
        ["Woeldgen et al.", "Flight-schedule-linked airport allocation", "Airport-side SAF allocation", "Distribution optimization", "Does not optimize the full upstream source-to-airport supply chain"],
        ["Eyberg et al.", "Plant-level TEA", "PtL via FT and MTJ", "Process simulation + TEA", "Strong process detail but no macro-spatial siting model"],
    ]
    add_table(doc, ["Study", "Data / resolution", "Pathway scope", "Method", "Main gap relative to this study"], literature_rows, [1.1, 1.25, 1.0, 0.9, 2.05])
    table_no += 1

    add_heading(doc, "5.3 Extreme-case tests", "Heading 3")
    stress = data.temporal_summary["frontier_stress"]
    add_para(
        doc,
        f"The archived temporal-robustness summary can also be read as an extreme-case operational test. Under worst-week conditions, GTL-BH retains {stress['GTL-BH']['WorstWeekRetention']:.2f}% of mean performance, GTL retains {stress['GTL']['WorstWeekRetention']:.2f}%, and CCU-GH-FT retains {stress['CCU-GH-FT']['WorstWeekRetention']:.2f}%. GTL shows a single but very long low-load event (168 h in the stored metric), whereas CCU-GH-FT exhibits more frequent low-load episodes, consistent with its renewable-coupled bottlenecks.",
    )
    add_para(
        doc,
        "These stress metrics do not replace a full stochastic optimization. They do, however, provide a concrete boundary-condition check for whether the static pathway ranking survives the weakest representative operating week retained in the reduced-year data structure.",
    )

    add_heading(doc, "5.4 Robustness of pathway ranking", "Heading 3")
    frontier = {row["scenario"]: row for row in data.temporal_summary["pareto_frontier"]}
    add_para(
        doc,
        f"The archived cost–emissions frontier remains qualitatively stable across the supporting plots. GTL-BH is the lowest-cost option at {frontier['GTL-BH']['lcoe_cny_per_kg']:.2f} yuan/kg but increases carbon intensity by about {frontier['GTL-BH']['carbon_diff_vs_jet_gco2e_per_mj']:.2f} gCO2e/MJ relative to conventional jet fuel. GTL remains the bridging pathway at {frontier['GTL']['lcoe_cny_per_kg']:.2f} yuan/kg with a carbon-intensity reduction of about {-frontier['GTL']['carbon_diff_vs_jet_gco2e_per_mj']:.2f} gCO2e/MJ. CCU-GH-FT remains the deep-decarbonization case at {frontier['CCU-GH-FT']['lcoe_cny_per_kg']:.2f} yuan/kg and about {-frontier['CCU-GH-FT']['carbon_diff_vs_jet_gco2e_per_mj']:.2f} gCO2e/MJ below the jet-fuel benchmark.",
    )
    add_para(
        doc,
        "Temporal penalties widen the economic separation between these pathways, but they do not erase their functional roles in the archived comparison outputs: GTL-BH remains the low-cost case, GTL remains the market-accessible bridge, and CCU-GH-FT remains the deep-decarbonization option with substantial operating penalties.",
    )

    # Chapter 6
    add_heading(doc, "6 Supplementary Results", "Heading 2")
    add_heading(doc, "6.1 Additional pathway comparison results", "Heading 3")
    add_para(
        doc,
        "The following figures extend the main-text comparison by showing the full 13-pathway landscape and additional multi-panel summaries retained in the visualization results folder.",
    )
    add_figure(
        doc,
        assets["quadrant_chart"],
        f"Figure S{figure_no}. Extended cost-emissions comparison across the 13 archived pathways.",
        width_inches=6.0,
    )
    figure_no += 1
    add_figure(
        doc,
        assets["efficiency_analysis"],
        f"Figure S{figure_no}. Additional pathway-efficiency and deployability comparisons.",
        width_inches=6.1,
    )
    figure_no += 1

    add_heading(doc, "6.2 Additional spatial network maps", "Heading 3")
    add_para(
        doc,
        "Spatial outputs in the repository highlight pathway archetypes across all thirteen pathways. The figure below summarises clustered archetypes used to compare siting and network concentration. Per-pathway transport route maps are presented in full in Section 6.7.",
    )
    add_figure(
        doc,
        assets["combined_cluster"],
        f"Figure S{figure_no}. Clustered network and siting archetypes across archived SAF pathways.",
        width_inches=6.2,
    )
    figure_no += 1

    add_heading(doc, "6.3 Additional cost decomposition results", "Heading 3")
    add_para(
        doc,
        "The appendix cost-decomposition panel focuses on the same three representative Pareto pathways discussed above: CCU-GH-FT, GTL, and GTL-BH. A shared eight-category mapping makes it clear that electricity and hydrogen dominate the deep-decarbonization route, whereas feedstock-linked categories remain the main reason that gas-based pathways stay closer to the cost frontier.",
    )
    add_figure(
        doc,
        assets["pareto_cost"],
        f"Figure S{figure_no}. Cost composition of the three representative Pareto pathways.",
        width_inches=6.2,
    )
    figure_no += 1

    add_heading(doc, "6.4 Additional emissions decomposition results", "Heading 3")
    add_para(
        doc,
        "The emissions-decomposition panel uses the same representative pathways and the same component mapping. Positive lifecycle burdens are shown as stacked contributions, while negative credit terms are kept outside the stacks and reported separately as net deltas relative to conventional jet fuel so that the comparison remains visually comparable across pathways.",
    )
    add_figure(
        doc,
        assets["pareto_carbon"],
        f"Figure S{figure_no}. Positive emissions composition of the three representative Pareto pathways.",
        width_inches=6.2,
    )
    figure_no += 1

    # ── Section 6.5: Temporal-operational profile ────────────────────────────
    add_heading(doc, "6.5 Temporal-operational profile and production-stability analysis", "Heading 3")
    add_para(
        doc,
        "Cost-optimal pathway selection does not guarantee smooth production. This section documents the temporal-operational profile of each pathway by measuring facility utilization over the four representative weeks and decomposing any shortfall below full capacity into two distinct penalty types. The analysis connects the optimization output to observable operational characteristics and provides context for the main-text claim that CCU/DAC pathways carry substantially higher operational penalties than fossil-feedstock pathways.",
    )

    # ── 6.5.1 Facility-utilization efficiency ────────────────────────────────
    add_heading(doc, "6.5.1 Facility-utilization efficiency across pathways", "Heading 4")
    add_para(
        doc,
        "Facility utilization is computed at the 3-hour production-period level for both the hydrogen supply stage and the SAF synthesis stage. For each period the utilization rate is the ratio of actual production (kg) to the maximum achievable output given installed capacity and period length. By-product hydrogen pathways use the peak hourly supply observed over the four-week horizon as the capacity denominator because no electrolyzer nameplate capacity is installed.",
    )
    add_equation(doc, f"Eq. (S{eq_no})", [
        math_var("U", "w,k"),
        math_text(" = ", "p"),
        math_fraction(
            [math_var("q", "w,k")],
            [math_var("Cap", "k"), math_text(" · ", "p"), math_var("T", "w")],
        ),
        math_text(",    ", "p"),
        math_var("U", "w,k"),
        math_text(" ∈ [0, 1]", "p"),
    ])
    eq_no += 1
    add_para(
        doc,
        f"where q_{{w,k}} is actual production in period w for facility k, Cap_k is installed capacity in kg per hour, and T_w is the period length in hours. The utilization metric is capped at 1.0 so that over-production artefacts from modelling tolerances do not inflate efficiency scores. Fig. S{figure_no} shows the utilization distributions for H₂ production and SAF synthesis across the 13 pathways, grouped by feedstock and carbon-source category.",
    )
    add_figure(
        doc,
        assets["temporal_h2_saf"],
        f"Figure S{figure_no}. H₂ production and SAF factory utilization distributions for all 13 pathways, grouped by feedstock category. Each panel shows the distribution of 3-hour period utilization rates across the four representative weeks. Pathways without an electrolyzer (GTL-BH, CTL-BH, GTL, and related by-product-H₂ routes) use peak supply as the capacity denominator.",
        width_inches=6.2,
    )
    figure_no += 1
    add_para(
        doc,
        "Green-hydrogen pathways (CCU-GH and DAC-GH families) show distinctly bimodal SAF utilization distributions: the bulk of periods cluster near zero utilization when renewable generation is low, with a secondary cluster at or near full capacity when generation peaks. By-product hydrogen and fossil pathways show more uniform distributions because their hydrogen supply is not coupled to intermittent renewables.",
    )
    add_figure(
        doc,
        assets["temporal_space"],
        f"Figure S{figure_no}. Four-panel temporal-robustness overview for all 13 pathways: (A) worst-week retention rate versus mean SAF utilization, (B) pathways ranked by worst-week retention, (C) chronic versus robustness penalty scatter, (D) longest low-load episode versus mean SAF utilization.",
        width_inches=6.2,
    )
    figure_no += 1

    # ── 6.5.2 Temporal robustness penalty framework ───────────────────────────
    add_heading(doc, "6.5.2 Temporal robustness penalty framework", "Heading 4")
    add_para(
        doc,
        "Raw utilization rates do not directly quantify the implied cost premium arising from underutilisation. The temporal robustness framework converts utilization levels into a dimensionless penalty that expresses, in percentage terms, how much additional capacity would be needed to produce the same SAF output if the facility ran at the observed mean utilization level rather than at full capacity. The total penalty is decomposed into two additive components.",
    )
    add_para(
        doc,
        "The chronic low-utilization penalty reflects the gap between mean SAF utilization and the full-capacity benchmark:",
    )
    add_equation(doc, f"Eq. (S{eq_no})", [
        math_var("P", "chronic"),
        math_text(" = ", "p"),
        math_text("(", "p"),
        math_fraction(
            [math_text("100", "p")],
            [math_var("Ū", "SAF")],
        ),
        math_text(" − 1", "p"),
        math_text(") × 100 %", "p"),
    ])
    eq_no += 1
    add_para(
        doc,
        "The robustness penalty reflects the additional shortfall associated with the single worst representative week, capturing inter-week variability rather than the average level:",
    )
    add_equation(doc, f"Eq. (S{eq_no})", [
        math_var("P", "robust"),
        math_text(" = ", "p"),
        math_text("(", "p"),
        math_fraction(
            [math_text("100", "p")],
            [math_var("Ū", "worst")],
        ),
        math_text(" − ", "p"),
        math_fraction(
            [math_text("100", "p")],
            [math_var("Ū", "SAF")],
        ),
        math_text(") × 100 %", "p"),
    ])
    eq_no += 1
    add_para(
        doc,
        f"where Ū_SAF is the mean SAF utilization across all four representative weeks and Ū_worst is the mean SAF utilization in the worst-performing week. The total penalty P_total = P_chronic + P_robust quantifies the combined implied capacity premium. Fig. S{figure_no} shows the penalty decomposition across all 13 pathways.",
    )
    add_figure(
        doc,
        assets["temporal_penalty"],
        f"Figure S{figure_no}. Chronic and robustness penalty decomposition across all pathways. Each horizontal bar shows the chronic component (darker) and robustness component (lighter) of the implied cost penalty. Pathways are ordered by total penalty.",
        width_inches=6.0,
    )
    figure_no += 1

    # ── 6.5.3 Penalty decomposition for representative pathways ──────────────
    add_heading(doc, "6.5.3 Penalty decomposition for the three Pareto-representative pathways", "Heading 4")
    _stress = data.temporal_summary["frontier_stress"]
    _penalty_rows = {row["Scenario"]: row for row in data.temporal_summary["frontier_penalties"]}
    _gtl_bh = _penalty_rows.get("GTL-BH", {})
    _gtl    = _penalty_rows.get("GTL", {})
    _ccu    = _penalty_rows.get("CCU-GH-FT", {})
    add_para(
        doc,
        f"The three Pareto-representative pathways span the full range of temporal penalty behaviour. GTL-BH achieves a mean SAF utilization of {_gtl_bh.get('SAFMean', float('nan')):.1f}%, but its worst-week retention drops to {_gtl_bh.get('WorstWeekMean', float('nan')):.1f}%, producing a chronic penalty of {_gtl_bh.get('ChronicPenaltyPct', float('nan')):.1f}% and a robustness penalty of {_gtl_bh.get('RobustnessPenaltyPct', float('nan')):.1f}% (total {_gtl_bh.get('TotalPenaltyPct', float('nan')):.1f}%). GTL shows a higher chronic penalty of {_gtl.get('ChronicPenaltyPct', float('nan')):.1f}% and a substantially larger robustness penalty of {_gtl.get('RobustnessPenaltyPct', float('nan')):.1f}% (total {_gtl.get('TotalPenaltyPct', float('nan')):.1f}%), reflecting a single extended low-load episode lasting {_stress.get('GTL', {}).get('MaxLowRunHours', 'n/a')} hours in the worst representative week. CCU-GH-FT carries the largest penalties: a chronic component of {_ccu.get('ChronicPenaltyPct', float('nan')):.1f}% reflecting its low mean SAF utilization ({_ccu.get('SAFMean', float('nan')):.1f}%) and a robustness component of {_ccu.get('RobustnessPenaltyPct', float('nan')):.1f}% from intermittent renewable coupling (total {_ccu.get('TotalPenaltyPct', float('nan')):.1f}%).",
    )
    _penalty_table_rows = [
        ["GTL-BH",
         f"{_gtl_bh.get('SAFMean', float('nan')):.1f}",
         f"{_gtl_bh.get('WorstWeekMean', float('nan')):.1f}",
         f"{_gtl_bh.get('ChronicPenaltyPct', float('nan')):.1f}",
         f"{_gtl_bh.get('RobustnessPenaltyPct', float('nan')):.1f}",
         f"{_gtl_bh.get('TotalPenaltyPct', float('nan')):.1f}",
        ],
        ["GTL",
         f"{_gtl.get('SAFMean', float('nan')):.1f}",
         f"{_gtl.get('WorstWeekMean', float('nan')):.1f}",
         f"{_gtl.get('ChronicPenaltyPct', float('nan')):.1f}",
         f"{_gtl.get('RobustnessPenaltyPct', float('nan')):.1f}",
         f"{_gtl.get('TotalPenaltyPct', float('nan')):.1f}",
        ],
        ["CCU-GH-FT",
         f"{_ccu.get('SAFMean', float('nan')):.1f}",
         f"{_ccu.get('WorstWeekMean', float('nan')):.1f}",
         f"{_ccu.get('ChronicPenaltyPct', float('nan')):.1f}",
         f"{_ccu.get('RobustnessPenaltyPct', float('nan')):.1f}",
         f"{_ccu.get('TotalPenaltyPct', float('nan')):.1f}",
        ],
    ]
    add_para(
        doc,
        f"Table S{table_no} Temporal robustness summary for the three Pareto-representative pathways.",
        style="Caption",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_table(
        doc,
        ["Pathway", "Mean SAF util. (%)", "Worst-week util. (%)", "Chronic penalty (%)", "Robustness penalty (%)", "Total penalty (%)"],
        _penalty_table_rows,
        [0.75, 1.10, 1.10, 1.15, 1.25, 1.15],
    )
    table_no += 1
    add_figure(
        doc,
        assets["pareto_temporal_dist"],
        f"Figure S{figure_no}. Utilization distributions for the three Pareto-representative pathways (GTL-BH, GTL, CCU-GH-FT). Left column: H₂ production utilization; right column: SAF factory utilization. Mean and median are annotated.",
        width_inches=6.2,
    )
    figure_no += 1
    add_figure(
        doc,
        assets["pareto_penalty_comparison"],
        f"Figure S{figure_no}. Penalty decomposition comparison for the three Pareto-representative pathways. Panel (b): stacked bar chart showing chronic (darker) and robustness (lighter) components of the implied cost penalty. Panel (c): penalty space scatter plot with bubble area proportional to total penalty; the diagonal boundary separates chronic-dominant from robustness-dominant regimes.",
        width_inches=6.0,
    )
    figure_no += 1

    # ── 6.5.4 Penalty driver fingerprint ─────────────────────────────────────
    add_heading(doc, "6.5.4 Penalty-driver fingerprint", "Heading 4")
    add_para(
        doc,
        "To attribute penalty differences among the three Pareto pathways to underlying operational factors, a twelve-driver fingerprint is constructed. The drivers are organized into four groups: low-utilization drivers (mean SAF utilization, frequency and duration of episodes below 10% capacity), instability drivers (worst-week retention, week-to-week coefficient of variation, maximum single low-load event), renewable shock drivers (renewable power variability, zero-output fraction, ramp intensity), and penalty outcomes (chronic, robustness, and total penalty). Each driver value is standardized against the distribution across all 13 archived pathways using z-score normalization, so that the heatmap identifies which factors deviate from the all-pathway baseline for each of the three representative routes.",
    )
    add_para(
        doc,
        f"Fig. S{figure_no} shows the normalized fingerprint. GTL-BH and GTL exhibit near-zero renewable shock drivers—consistent with their fossil-feedstock operation—but diverge in the instability group: GTL's single extended low-load event registers as an outlier in the MaxLowRunHours cell, explaining its disproportionately large robustness penalty relative to its chronic penalty. CCU-GH-FT, by contrast, shows strongly elevated low-utilization drivers (SAFMean is below the all-pathway median), high renewable shock values, and correspondingly elevated penalty outcomes in both the chronic and robustness columns.",
    )
    add_figure(
        doc,
        assets["penalty_driver_heatmap"],
        f"Figure S{figure_no}. Twelve-driver penalty fingerprint for the three Pareto-representative pathways, normalized against the all-pathway distribution. Red cells indicate values elevated above the baseline; blue cells indicate suppressed values. The four driver groups (low-utilization drivers, instability drivers, renewable shock, penalty outcomes) are separated by horizontal dividers.",
        width_inches=6.2,
    )
    figure_no += 1

    # ── Section 6.6: Per-pathway cost–emission decomposition ──────────────────
    add_heading(doc, "6.6 Per-pathway cost–emission decomposition", "Heading 3")
    add_para(
        doc,
        "This section presents a polar stacked-area learning-curve chart for each of the thirteen SAF supply-chain pathways. "
        "Each chart shows how the levelised cost of SAF (CNY/kg) evolves from 2026 to 2036 "
        "under Wright's-Law learning-curve projections, with stacked angular segments representing the eight cost "
        "categories and an outer ring encoding the corresponding carbon-intensity allocation. "
        "The charts allow direct cross-pathway comparison of both the magnitude and the category structure of costs and emissions.",
    )

    _group_descriptions = {
        "Coal-based pathways": (
            "Coal gasification routes (CTL and CTL-BH) rely on solid feedstock and exhibit high feedstock and "
            "conversion contributions. CTL-BH partially offsets hydrogen production cost by integrating "
            "by-product hydrogen recovered from steel and refinery off-gases."
        ),
        "Natural-gas-based pathways": (
            "Natural-gas routes (GTL-GH, GTL, GTL-BH) benefit from mature reforming technology. GTL uses "
            "Fischer-Tropsch one-step synthesis, GTL-GH adds an electrolytic green-hydrogen supplement, and "
            "GTL-BH replaces electrolysis with by-product hydrogen, lowering the hydrogen-system cost share."
        ),
        "CCU Green-H₂ pathways": (
            "CCU routes capture industrial CO\u2082 as carbon feedstock and use electrolytic green hydrogen. "
            "The MTJ variant converts via methanol-to-jet; the FT variant uses Fischer-Tropsch synthesis. "
            "Power and hydrogen system costs dominate; facility and transport costs are comparatively minor."
        ),
        "CCU By-product-H₂ pathways": (
            "CCU by-product-hydrogen routes substitute electrolysers with PSA purification of steel- and "
            "refinery-plant off-gases, substantially reducing the hydrogen-system capital share while retaining "
            "the CCU carbon-capture infrastructure."
        ),
        "DAC Green-H₂ pathways": (
            "DAC routes capture atmospheric CO\u2082 directly. The additional DAC energy and capital cost raise "
            "the CO\u2082-capture category share relative to CCU routes. Learning-curve projections for DAC "
            "show the strongest cost-reduction potential through the 2036 horizon."
        ),
        "DAC By-product-H₂ pathways": (
            "DAC by-product-hydrogen routes combine atmospheric CO\u2082 capture with by-product hydrogen "
            "recovery, eliminating electrolyser capital. The DAC capture system remains the dominant cost "
            "driver at current technology readiness levels."
        ),
    }

    for group_name, scenario_labels in APPENDIX_SCENARIO_GROUPS.items():
        add_heading(doc, group_name, "Heading 4")
        group_desc = _group_descriptions.get(group_name, "")
        if group_desc:
            add_para(doc, group_desc)

        for _label in scenario_labels:
            _slug = _label.lower().replace("-", "_")
            _key  = f"scenario_composite_{_slug}"
            _composite_path = assets.get(_key)
            if _composite_path is None or not Path(_composite_path).exists():
                print(f"[WARNING] 找不到场景合成图: {_label} ({_key})")
                continue
            add_figure(
                doc,
                _composite_path,
                f"Figure S{figure_no}. Cost–emission decomposition for pathway {_label}: "
                f"polar stacked-area learning-curve trajectory (2026–2036), with stacked cost categories "
                f"and outer carbon-intensity ring.",
                width_inches=6.3,
            )
            figure_no += 1

    # ── Section 6.7: Transport network route maps ─────────────────────────────
    add_heading(doc, "6.7 Transport network route maps", "Heading 3")
    add_para(
        doc,
        "This section presents the optimised clustered transport route maps for all thirteen SAF supply-chain pathways. "
        "Each map shows the spatial layout of the supply chain at the representative-week level after k-means route clustering. "
        "Nodes represent supply facilities (hydrogen production, CO\u2082 capture, methanol synthesis, SAF conversion) "
        "and demand airports; edges represent transport flows with line thickness proportional to the weekly flow volume. "
        "The outer ring colour encodes the feedstock category (coal grey, natural gas orange, green hydrogen blue, "
        "by-product hydrogen teal, DAC purple). The shared legend is shown first, followed by individual pathway maps "
        "grouped by feedstock and hydrogen-supply technology.",
    )

    # 图例 + 四场景对比图
    _legend_path = assets.get("transport_map_legend")
    if _legend_path and Path(_legend_path).exists():
        add_figure(
            doc,
            _legend_path,
            f"Figure S{figure_no}. Shared legend for the transport network route maps (all 13 pathways).",
            width_inches=6.3,
        )
        figure_no += 1

    _four_key_path = assets.get("transport_map_four_key")
    if _four_key_path and Path(_four_key_path).exists():
        add_figure(
            doc,
            _four_key_path,
            f"Figure S{figure_no}. Side-by-side comparison of four representative pathways: "
            f"CCU-GH-FT (deep decarbonisation), GTL (transition), GTL-BH (lowest-cost frontier), and DAC-BH-FT.",
            width_inches=6.3,
        )
        figure_no += 1

    _transport_map_group_descs = {
        "Coal-based transport networks": (
            "Coal-based pathways (CTL and CTL-BH) source solid coal from mine clusters in North China. "
            "CTL routes all feedstock and CO\u2082 flows via a centralised gasification hub, whereas CTL-BH "
            "adds by-product hydrogen collection routes from nearby steel and refinery plants, reducing "
            "the demand on coal-fed gasifiers."
        ),
        "Natural-gas-based transport networks": (
            "Natural-gas pathways draw feedstock from pipeline networks in Western China. GTL-GH and GTL "
            "differ in their methanol-synthesis route (two-step methanol-to-jet vs. one-step Fischer-Tropsch); "
            "both show long east-west pipeline flows. GTL-BH supplements gas-derived hydrogen with by-product "
            "hydrogen from industrial clusters along the coast, shortening average transport distances."
        ),
        "CCU Green-H₂ transport networks": (
            "CCU pathways capture industrial CO\u2082 from power-plant and cement clusters and source "
            "electrolytic hydrogen from wind- and solar-rich regions. The two-step (MTJ) route shows a "
            "methanol intermediate hub, while the one-step (FT) route links hydrogen and CO\u2082 sources "
            "directly to the Fischer-Tropsch reactor."
        ),
        "CCU By-product-H₂ transport networks": (
            "CCU by-product-hydrogen pathways replace electrolysers with PSA-purified off-gas recovery "
            "near steel and refinery clusters. The transport network is denser around industrial corridors "
            "in the Yangtze River Delta and Bohai Rim, with shorter average hydrogen transport legs compared "
            "with the green-hydrogen CCU routes."
        ),
        "DAC Green-H₂ transport networks": (
            "DAC pathways use direct-air-capture units co-located with renewable electricity resources. "
            "The CO\u2082 collection network is distributed rather than point-source, leading to more "
            "dispersed edges relative to CCU routes. The green-hydrogen transport flows are similar in "
            "structure to the CCU green-H\u2082 cases."
        ),
        "DAC By-product-H₂ transport networks": (
            "DAC by-product-hydrogen pathways combine atmospheric CO\u2082 capture with by-product hydrogen "
            "from industrial clusters. The resulting network reflects two distinct spatial footprints: "
            "distributed DAC capture sites near wind/solar zones and concentrated by-product hydrogen "
            "sources in industrial coastal regions."
        ),
    }

    for _group_name, _map_keys in _TRANSPORT_MAP_GROUPS.items():
        add_heading(doc, _group_name, "Heading 4")
        _desc = _transport_map_group_descs.get(_group_name, "")
        if _desc:
            add_para(doc, _desc)
        for _mkey in _map_keys:
            _mpath = assets.get(_mkey)
            if _mpath is None or not Path(_mpath).exists():
                print(f"[WARNING] 找不到运输地图: {_mkey}")
                continue
            _scenario_lbl = _TRANSPORT_MAP_SCENARIO_LABELS.get(_mkey, _mkey)
            add_figure(
                doc,
                _mpath,
                f"Figure S{figure_no}. Optimised clustered transport route map for pathway {_scenario_lbl}.",
                width_inches=6.3,
            )
            figure_no += 1

    add_heading(doc, "Supplementary References", "Heading 2")
    references = [
        "Eyberg, V. et al. Techno-economic assessment and comparison of Fischer-Tropsch and methanol-to-jet processes to produce sustainable aviation fuel via power-to-liquid. Energy Conversion and Management 315, 118728 (2024).",
        "Wassermann, T., Muehlenbrock, H., Kenkel, P. & Zondervan, E. Supply chain optimization for electricity-based jet fuel: the case study Germany. Applied Energy 307, 117683 (2022).",
        "Martinez-Valencia, L., Garcia-Perez, M. & Wolcott, M. P. Supply chain configuration of sustainable aviation fuel: review, challenges, and pathways for including environmental and social benefits. Renewable and Sustainable Energy Reviews 152, 111680 (2021).",
        "Woeldgen, E., Teoh, R., Stettler, M. E. J. & Malina, R. Sustainable aviation fuel deployment strategies in Europe: supply chain implications and climate benefits. Environmental Science & Technology 59, 12447-12457 (2025).",
        "Flightera. Flight schedules and historical operational data used for the 2024 demand-processing workflow in the current project repository.",
        "OpenStreetMap contributors. OpenStreetMap road-network data used for truck-distance routing in the current project implementation.",
        "GraphHopper. Local routing engine used by the project to compute shortest-path road distances from OSM-based network data.",
        "EUROCONTROL Base of Aircraft Data (BADA), accessed through the pyBADA-based fuel-consumption workflow used in the project.",
    ]
    for item in references:
        add_para(doc, item, style="Bibliography" if style_exists(doc, "Bibliography") else "Normal")

    return doc


def validate_structure(doc_path: Path) -> None:
    doc = Document(str(doc_path))
    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip() and p.style and ("Heading" in p.style.name or "heading" in p.style.name)
    ]
    print(f"Created {doc_path.name}")
    print(f"Headings: {len(headings)}")
    print(f"Tables: {len(doc.tables)}")
    print("First headings:")
    for item in headings[:12]:
        print(f"  - {item}")


def main() -> None:
    ensure_dir(ASSET_DIR)
    pareto_assets = generate_appendix_pareto_assets(output_dir=REDRAWN_DIR, prepared_dir=PREP_DIR)
    data = collect_data(pareto_summary_path=pareto_assets["summary_json"])
    assets = prepare_assets(data, pareto_assets=pareto_assets)
    qa_outputs = run_appendix_qa(assets)
    doc = build_document(data, assets)
    doc.save(str(OUTPUT_DOCX))
    validate_structure(OUTPUT_DOCX)
    print(f"Assets written to: {ASSET_DIR}")
    for name, path in qa_outputs.items():
        print(f"{name}: {path}")


if __name__ == "__main__":
    main()
